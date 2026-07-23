#!/usr/bin/env python3
"""IJKM compliance pass on Paper2_R2_Full_submission_IJKM_v2.docx:
(1) remove title/authors/abstract/keywords (and Acknowledgments — double-blind);
(2) tables renumbered Table 1-7 (Arabic) with an in-text callout for each;
(3) all bracketed citations converted to APA author-date; reference list
    rewritten in APA 7, alphabetical. Also: the two unverified Kim IJERPH
    entries are replaced by the thesis-verified FRA validation sources
    (Park et al., 2019, Ann Rehabil Med; Kim et al., 2018, AGMR).
Also produces a separate title-page file for the submission system."""
import re
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = "/tmp/claude-0/-home-user-gamification/3d8a50e8-9258-51ce-a7d6-194da50f9a04/scratchpad/"
d = docx.Document(BASE + "P2_IJKM_v2.docx")
log, fail = [], []

# ---------------- citation label map ----------------
LBL = {
 1: "NICE, 2013", 2: "WHO, 2007", 3: "Guner & Acarturk, 2020",
 4: "Wildenbos et al., 2018", 5: "Hawthorn, 2007", 6: "Tsai et al., 2020",
 7: "Eysenbach, 2005", 8: "Davis, 1989", 9: "Venkatesh et al., 2003",
 10: "Deterding et al., 2011", 11: "An et al., 2024",
 12: "Kelders et al., 2012", 13: "Christensen et al., 2009",
 14: "Lee & Coughlin, 2015", 15: "Grossi et al., 2024",
 16: "Tinetti & Powell, 1993", 17: "Yardley et al., 2005",
 18: "Scheffer et al., 2008", 19: "Koivisto & Malik, 2021",
 20: "Chen & Chan, 2014", 21: "Chou, 2015", 22: "Bandura, 1997",
 23: "Compeau & Higgins, 1995", 24: "Kahneman & Tversky, 1979",
 25: "Tversky & Kahneman, 1992", 26: "Frey & Jegen, 2001",
 27: "Deci & Ryan, 1985", 28: "Michie et al., 2011", 29: "Shannon, 1948",
 30: "Stringer, 2014", 31: "Lewin, 1946",
 32: "Standing Committee of the National People's Congress, 2021",
 33: "Park et al., 2019", 34: "Kim et al., 2018", 35: "Festinger, 1954",
 36: "Cohen, 1988", 37: "Thaler, 1980", 38: "Kahneman et al., 1990",
 39: "Dionigi, 2015", 40: "Barber, 2017", 41: "McGlone & Aronson, 2006",
 42: "von Elm et al., 2007", 43: "Liu et al., 2020", 44: "WHO, 2021",
 45: "FDA, Health Canada, & MHRA, 2021", 46: "Tam & Cheung, in press",
}

# narrative conversions: (old phrase, new phrase, bracket to drop)
NARR = [
 ("Guner and Acarturk found", "Guner and Acarturk (2020) found", " [3]"),
 ("Wildenbos et al. organised", "Wildenbos et al. (2018) organised", " [4]"),
 ("Hawthorn observed", "Hawthorn (2007) observed", " [5]"),
 ("Tsai et al. found technology anxiety",
  "Tsai et al. (2020) found technology anxiety", " [6]"),
 ("Lee and Coughlin noted", "Lee and Coughlin (2015) noted", " [14]"),
 ("Eysenbach argued", "Eysenbach (2005) argued", " [7]"),
 ("Eysenbach's 'law of attrition' distinguished",
  "Eysenbach's (2005) 'law of attrition' distinguished", " [7]"),
 ("Kelders et al. found", "Kelders et al. (2012) found", " [12]"),
 ("Christensen et al. noted", "Christensen et al. (2009) noted", " [13]"),
 ("Yardley et al. developed", "Yardley et al. (2005) developed", " [17]"),
 ("An, Cheung and Willoughby reported",
  "An, Cheung and Willoughby (2024) reported", " [11]"),
 ("An et al. is the closest precedent",
  "An et al. (2024) is the closest precedent", " [11]"),
 ("An et al.'s conclusion", "An et al.'s (2024) conclusion", " [11]"),
 ("Bandura's work indicates", "Bandura's (1997) work indicates", " [22]"),
 ("Frey and Jegen's motivation-crowding effect",
  "Frey and Jegen's (2001) motivation-crowding effect", " [26]"),
 ("Deci and Ryan's basic needs",
  "Deci and Ryan's (1985) basic needs", " [27]"),
]

BR = re.compile(r"\s*\[\d+\](?:,\s*\[\d+\])*")


def conv(m):
    nums = re.findall(r"\d+", m.group(0))
    labels = sorted({LBL[int(n)] for n in nums})
    return " (" + "; ".join(labels) + ")"


def set_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in list(p.runs[1:]):
            r._element.getparent().remove(r._element)
    else:
        p.add_run(text)


# ---------------- (1) title block + acknowledgments removal ----------------
to_delete = []
paras = d.paragraphs
# title block: everything before the "I. Introduction" heading
for p in paras:
    if p.text.strip().startswith("I. Introduction"):
        break
    to_delete.append(p)
# acknowledgments: heading + its paragraph(s) up to References heading
in_ack = False
for p in paras:
    t = p.text.strip()
    if t == "Acknowledgments":
        in_ack = True
    elif t == "References":
        in_ack = False
    if in_ack:
        to_delete.append(p)
for p in to_delete:
    p._element.getparent().remove(p._element)
log.append(f"removed title block + acknowledgments ({len(to_delete)} paragraphs)")

# ---------------- (3) references: find and replace the numbered list ----------------
ref_heading = None
for p in d.paragraphs:
    if p.text.strip() == "References" and p.style.name.startswith("Heading"):
        ref_heading = p
        break
assert ref_heading is not None
old_refs = []
seen = False
for p in d.paragraphs:
    if p is ref_heading:
        seen = True
        continue
    if seen and re.match(r"\[\d+\]", p.text.strip()):
        old_refs.append(p)
for p in old_refs:
    p._element.getparent().remove(p._element)
log.append(f"removed {len(old_refs)} numbered reference entries")

APA = [
 "An, S., Cheung, C. F., & Willoughby, K. W. (2024). A gamification approach for enhancing older adults' technology adoption and knowledge transfer. Technological Forecasting and Social Change, 205.",
 "Bandura, A. (1997). Self-efficacy: The exercise of control. W. H. Freeman.",
 "Barber, S. J. (2017). An examination of age-based stereotype threat about cognitive decline. Perspectives on Psychological Science, 12(1), 62–90.",
 "Chen, K., & Chan, A. H. S. (2014). Gerontechnology acceptance by elderly Hong Kong Chinese: A senior technology acceptance model (STAM). Ergonomics, 57(5), 635–652.",
 "Chou, Y.-K. (2015). Actionable gamification: Beyond points, badges, and leaderboards. Octalysis Media.",
 "Christensen, H., Griffiths, K. M., & Farrer, L. (2009). Adherence in internet interventions for anxiety and depression. Journal of Medical Internet Research, 11(2), e13.",
 "Cohen, J. (1988). Statistical power analysis for the behavioral sciences (2nd ed.). Erlbaum.",
 "Compeau, D. R., & Higgins, C. A. (1995). Computer self-efficacy: Development of a measure and initial test. MIS Quarterly, 19(2), 189–211.",
 "Davis, F. D. (1989). Perceived usefulness, perceived ease of use, and user acceptance of information technology. MIS Quarterly, 13(3), 319–340.",
 "Deci, E. L., & Ryan, R. M. (1985). Intrinsic motivation and self-determination in human behavior. Plenum.",
 "Deterding, S., Dixon, D., Khaled, R., & Nacke, L. (2011). From game design elements to gamefulness: Defining gamification. Proceedings of the 15th International Academic MindTrek Conference, 9–15.",
 "Dionigi, R. A. (2015). Stereotypes of aging: Their effects on the health of older adults. Journal of Geriatrics, 2015, Article 954027.",
 "Eysenbach, G. (2005). The law of attrition. Journal of Medical Internet Research, 7(1), e11.",
 "Festinger, L. (1954). A theory of social comparison processes. Human Relations, 7(2), 117–140.",
 "Frey, B. S., & Jegen, R. (2001). Motivation crowding theory. Journal of Economic Surveys, 15(5), 589–611.",
 "Grossi, G., Ferrario, R., & Tosi, V. (2024). Barriers and facilitators to health technology adoption by older adults with chronic diseases: An integrative systematic review. BMC Geriatrics, 24, Article 189.",
 "Guner, H., & Acarturk, C. (2020). The use and acceptance of ICT by senior citizens: A comparison of TAM for elderly and young adults. Universal Access in the Information Society, 19, 311–330.",
 "Hawthorn, D. (2007). Interface design and engagement with older people. Behaviour & Information Technology, 26(4), 333–341.",
 "Kahneman, D., Knetsch, J. L., & Thaler, R. H. (1990). Experimental tests of the endowment effect and the Coase theorem. Journal of Political Economy, 98(6), 1325–1348.",
 "Kahneman, D., & Tversky, A. (1979). Prospect theory: An analysis of decision under risk. Econometrica, 47(2), 263–291.",
 "Kelders, S. M., Kok, R. N., Ossebaard, H. C., & Van Gemert-Pijnen, J. E. W. C. (2012). Persuasive system design does matter: A systematic review of adherence to web-based interventions. Journal of Medical Internet Research, 14(6), e152.",
 "Kim, M., Kim, S., & Won, C. W. (2018). Test–retest reliability of the Fall Risk Assessment system. Annals of Geriatric Medicine and Research, 22(2), 80–87.",
 "Koivisto, J., & Malik, A. (2021). Gamification for older adults: A systematic literature review. The Gerontologist, 61(7), 360–372.",
 "Lee, C., & Coughlin, J. F. (2015). Older adults' adoption of technology: An integrated approach. Geriatrics & Gerontology International, 15(5), 519–525.",
 "Lewin, K. (1946). Action research and minority problems. Journal of Social Issues, 2(4), 34–46.",
 "Liu, X., Cruz Rivera, S., Moher, D., Calvert, M. J., & Denniston, A. K. (2020). Reporting guidelines for clinical trials evaluating artificial intelligence interventions. Nature Medicine, 26, 1364–1374.",
 "McGlone, M. S., & Aronson, J. (2006). Stereotype threat, identity salience, and spatial reasoning. Journal of Applied Developmental Psychology, 27(5), 486–493.",
 "Michie, S., van Stralen, M. M., & West, R. (2011). The behaviour change wheel: A new method for characterising and designing behaviour change interventions. Implementation Science, 6, Article 42.",
 "National Institute for Health and Care Excellence [NICE]. (2013). Falls in older people: Assessing risk and prevention (Clinical Guideline CG161).",
 "Park, W.-C., Kim, M., & Kim, S. (2019). Introduction of the Fall Risk Assessment system and cross-sectional validation among community-dwelling older adults. Annals of Rehabilitation Medicine, 43(1), 87–95.",
 "Scheffer, A. C., Schuurmans, M. J., van Dijk, N., van der Hooft, T., & de Rooij, S. E. (2008). Fear of falling: Measurement strategy, prevalence, risk factors and consequences among older persons. Age and Ageing, 37(1), 19–24.",
 "Shannon, C. E. (1948). A mathematical theory of communication. Bell System Technical Journal, 27, 379–423, 623–656.",
 "Standing Committee of the National People's Congress. (2021). Personal Information Protection Law of the People's Republic of China.",
 "Stringer, E. T. (2014). Action research (4th ed.). SAGE.",
 "Tam, S. K. T., & Cheung, C. F. (in press). A study of utilisation of Fall Risk Assessment among elderly users: A critical review through TAM and Octalysis frameworks with gamification approach.",
 "Thaler, R. (1980). Toward a positive theory of consumer choice. Journal of Economic Behavior & Organization, 1(1), 39–60.",
 "Tinetti, M. E., & Powell, L. (1993). Fear of falling and low self-efficacy: A cause of dependence in elderly persons. Journal of Gerontology, 48(Special issue), 35–38.",
 "Tsai, H. S., Shillair, R., & Cotten, S. R. (2020). Exploring the role of technology anxiety and self-efficacy in older adults' information seeking. iConference 2020 Proceedings.",
 "Tversky, A., & Kahneman, D. (1992). Advances in prospect theory: Cumulative representation of uncertainty. Journal of Risk and Uncertainty, 5(4), 297–323.",
 "U.S. Food and Drug Administration, Health Canada, & Medicines and Healthcare products Regulatory Agency. (2021). Good machine learning practice for medical device development: Guiding principles.",
 "Venkatesh, V., Morris, M. G., Davis, G. B., & Davis, F. D. (2003). User acceptance of information technology: Toward a unified view. MIS Quarterly, 27(3), 425–478.",
 "von Elm, E., Altman, D. G., Egger, M., Pocock, S. J., Gøtzsche, P. C., & Vandenbroucke, J. P. (2007). The STROBE statement: Guidelines for reporting observational studies. Annals of Internal Medicine, 147(8), 573–577.",
 "Wildenbos, G. A., Peute, L., & Jaspers, M. (2018). Aging barriers influencing mobile health usability for older adults: MOLD-US framework. International Journal of Medical Informatics, 114, 66–75.",
 "World Health Organization [WHO]. (2007). WHO global report on falls prevention in older age.",
 "World Health Organization. (2021). Ethics and governance of artificial intelligence for health.",
 "Yardley, L., Beyer, N., Hauer, K., Kempen, G., Piot-Ziegler, C., & Todd, C. (2005). Development and initial validation of the Falls Efficacy Scale-International (FES-I). Age and Ageing, 34(6), 614–619.",
]
from docx.oxml import OxmlElement
prev = ref_heading
for entry in APA:
    newp = OxmlElement("w:p")
    prev._p.addnext(newp)
    np_ = docx.text.paragraph.Paragraph(newp, ref_heading._parent)
    np_.add_run(entry)
    prev = np_
log.append(f"inserted {len(APA)} APA reference entries (alphabetical)")

# ---------------- (3) in-text conversion ----------------
n_narr = n_gen = 0
for p in d.paragraphs:
    t = p.text
    if not t.strip() or "–" in t[:4]:
        pass
    orig = t
    for old, new, br in NARR:
        if old in t:
            t = t.replace(old, new)
            t = t.replace(br + ",", ",", 1) if (br + ",") in t else t.replace(br, "", 1)
            n_narr += 1
    if re.search(r"\[\d+\]", t):
        t = BR.sub(conv, t)
        t = t.replace("  (", " (").replace(" ,", ",").replace(" .", ".")
        n_gen += 1
    if t != orig:
        set_text(p, t)
log.append(f"narrative conversions: {n_narr}; paragraphs with generic conversion: {n_gen}")

# ---------------- (2) table renumbering + callouts ----------------
ROMAN = [("Table IIA", "Table 3"), ("Table III", "Table 4"),
         ("Table VI", "Table 7"), ("Table IV", "Table 5"),
         ("Table II", "Table 2"), ("Table V", "Table 6"),
         ("Table I", "Table 1")]
n_ren = 0
for p in d.paragraphs:
    t = p.text
    new = t
    for old, nw in ROMAN:
        new = re.sub(re.escape(old) + r"(?![A-Z])", nw, new)
    if new != t:
        set_text(p, new)
        n_ren += 1
log.append(f"table numbering: {n_ren} paragraphs updated to Arabic")

def sub1(needle, old, new, label):
    for p in d.paragraphs:
        if needle in p.text:
            set_text(p, p.text.replace(old, new))
            log.append("callout: " + label)
            return
    fail.append("callout fail: " + label)

sub1("Despite minimal engagement, the low-engagement group",
     "Despite minimal engagement, the low-engagement group shows",
     "Table 3 reports the low-engagement group: despite minimal "
     "engagement, it shows", "Table 3")
sub1("All 30 participants were coded with thirteen binary attributes",
     "thirteen binary attributes drawn from STAM constructs",
     "thirteen binary attributes (Table 5) drawn from STAM constructs",
     "Table 5")
sub1("The empty interval between scores 5 and 9 is not read",
     "The empty interval between scores 5 and 9 is not read",
     "The empty interval between scores 5 and 9 (Table 6) is not read",
     "Table 6")

cp = d.core_properties
cp.title = "Anonymous manuscript (IJKM submission)"
cp.author = ""
cp.comments = ("IJKM compliance: title block/abstract/keywords/"
               "acknowledgments removed; tables 1-7 Arabic with callouts; "
               "APA author-date citations and alphabetical reference list.")
try:
    cp.last_modified_by = ""
except Exception:
    pass
OUT = BASE + "Paper2_IJKM_Anonymous_Manuscript.docx"
d.save(OUT)
print("saved", OUT)
for l in log:
    print(" -", l)
for f in fail:
    print(" !", f)

# ---------------- title page file ----------------
t = docx.Document()
t.styles["Normal"].font.name = "Times New Roman"
t.styles["Normal"].font.size = Pt(11)
p = t.add_paragraph()
r = p.add_run("Gamification as a Proactive Approach: A Study of Elderly "
              "Failure-Induced Abandonment in Gamified Fall Risk Assessment")
r.bold = True; r.font.size = Pt(14)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
a = t.add_paragraph("Sidney K. T. Tam and Chi Fai Cheung\n"
                    "Behaviour and Knowledge Engineering Research Centre, "
                    "Department of Industrial and Systems Engineering,\n"
                    "The Hong Kong Polytechnic University, Hong Kong, China\n"
                    "Corresponding author: sidney.tam@connect.polyu.hk")
a.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.add_heading("Abstract", level=1)
t.add_paragraph("Fall Risk Assessment (FRA) systems assess balance and "
    "support fall prevention, but rarely sustain elderly engagement. In an "
    "eight-week action-research deployment at a Southern China care "
    "facility, thirty residents (aged 62-84) used an AI-assisted skiing "
    "game on the FRA's pressure panel: one weight shift both steers the "
    "game and yields the balance measurement. Seventy per cent disengaged "
    "within three cycles, 26.7% after one session, and none returned. Yet "
    "the recorded FRA index improved (+6.9 points among sustained "
    "participants; +4.0 across all thirty), and three who improved in "
    "their only session never returned. Satisfaction was strongly "
    "associated with sustained engagement (p = 0.013). We propose two "
    "working hypotheses: Failure-Induced Dropout (FIDS), self-efficacy "
    "collapse at first failure, and Failure-Induced Abandonment (FIAS), "
    "loss aversion when failure reads as evidence of decline. Five "
    "falsifiable predictions are specified for a pre-registered three-arm "
    "confirmatory study.")
t.add_heading("Keywords", level=1)
t.add_paragraph("fall risk assessment; gamification; elderly technology "
                "adoption; failure-induced abandonment; loss aversion; "
                "eHealth attrition; action research")
t.add_heading("Acknowledgments", level=1)
t.add_paragraph("The authors thank the management and nursing staff, and "
    "the residents of the participating residential care facility in "
    "Southern China, for their generous cooperation. Supervision by "
    "Prof. Chi Fai Cheung at The Hong Kong Polytechnic University is "
    "gratefully acknowledged. AI tooling (Anthropic Claude Code and API "
    "services) was used as an enabling technology for game development "
    "and difficulty adaptation; the contribution of this paper is the "
    "behavioural analysis and the design framework. The authors declare "
    "no competing financial or non-financial interests; the lead author "
    "holds no financial relationship with Anthropic; the facility was "
    "not involved in data analysis or manuscript preparation; the study "
    "received no external funding.")
t.save(BASE + "Paper2_IJKM_Title_Page.docx")
print("saved title page file")

# ---------------- verification ----------------
d2 = docx.Document(OUT)
full = "\n".join(p.text for p in d2.paragraphs)
for tb in d2.tables:
    for r_ in tb.rows:
        full += "\n" + " | ".join(c.text for c in r_.cells)
print("\nbrackets remaining:", len(re.findall(r"\[\d+\]", full)))
print("Roman 'Table I/V' remaining:",
      len(re.findall(r"Table [IVX]+\b", full)))
for i in range(1, 8):
    print(f"Table {i} occurrences:", len(re.findall(rf"Table {i}\b", full)))
print("Abstract present:", "Abstract" in full,
      "| Keywords present:", "Keywords" in full,
      "| author names present:", "Tam and Cheung" in full or "Sidney" in full,
      "| supervisor named:", "Chi Fai" in full)
print("'Tam & Cheung, in press' citations:", full.count("Tam & Cheung, in press"))
EOF_ = None
