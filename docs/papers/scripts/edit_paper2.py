#!/usr/bin/env python3
"""Paper 2 alignment with thesis V20.
Removes the statistics withdrawn from the thesis (Cox HR 3.24 + K-M figure:
no time variable in the archive; motivational-orientation OR 31.5: coding
records not producible; pooled d 0.78 / ITT d 0.57: not reproducible from
the archive) and replaces them with the verified V20 statistics. Applies the
V20 composite-circularity treatment, calibrated language, and factual fixes
(site, age SD, Stage-1 characterisation, session-based definitions).
"""
import shutil
import docx
from docx.oxml import OxmlElement

BASE = "/tmp/claude-0/-home-user-gamification/3d8a50e8-9258-51ce-a7d6-194da50f9a04/scratchpad/"
SRC = BASE + "Paper2.docx"
DST = BASE + "Paper2_R1_aligned_V20.docx"
shutil.copy(SRC, DST)
d = docx.Document(DST)
log, fail = [], []


def set_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in list(p.runs[1:]):
            r._element.getparent().remove(r._element)
    else:
        p.add_run(text)


def find(needle, nth=1):
    n = 0
    for p in d.paragraphs:
        if needle in p.text:
            n += 1
            if n == nth:
                return p
    return None


def repl(needle, text, label):
    p = find(needle)
    if p is None:
        fail.append("FIND-FAIL: " + label)
        return None
    set_text(p, text)
    log.append("rewrote: " + label)
    return p


def sub(needle, pairs, label):
    p = find(needle)
    if p is None:
        fail.append("FIND-FAIL: " + label)
        return
    t = p.text
    for old, new in pairs:
        if old not in t:
            fail.append(f"SUB-FAIL ({label}): {old[:50]}")
        t = t.replace(old, new)
    set_text(p, t)
    log.append("edited: " + label)


def delete_para(needle, label):
    p = find(needle)
    if p is None:
        fail.append("DEL-FAIL: " + label)
        return
    p._element.getparent().remove(p._element)
    log.append("deleted: " + label)


# ---------------- Abstract ----------------
repl("Fall Risk Assessment (FRA) system can assess balance",
    "Fall Risk Assessment (FRA) systems can assess balance and assist with "
    "fall prevention as clinical devices but rarely achieve sustained elderly "
    "user engagement. This study examines the factors associated with low "
    "engagement and rapid permanent dropout in a gamified, AI-assisted FRA "
    "among elderly residents after their first adverse user experience. An "
    "action research approach was employed in a residential care facility in "
    "Southern China. An AI-assisted skiing game was developed on a "
    "weight-shifting foot pressure panel — a standard accessory of the FRA "
    "system — allowing elderly individuals to undergo balance measurement "
    "while controlling a skier on a display. Thirty residents (aged 62-84; "
    "mean 71.1 ± 5.9) participated over eight weeks. Seventy per cent "
    "(21 of 30; 95% CI 52.1-83.3%) disengaged at or before three game "
    "cycles, 26.7% (8 of 30; 95% CI 14.2-44.4%) exited after a single "
    "session, and no disengaged participant re-engaged during the study "
    "window (0 of 21; 95% CI 0.0-15.5%). Among the nine sustained "
    "participants, the recorded FRA index improved by a mean of +6.9 points "
    "(95% CI 4.5-9.3; paired dz = 2.22, 95% CI 0.95-3.46); across all "
    "thirty participants the paired change was +4.0 points (95% CI 0.5-7.6; "
    "dz = 0.42, 95% CI 0.05-0.79). Satisfaction was strongly associated "
    "with sustained engagement (9/9 sustained participants satisfied versus "
    "11/21 among the low-engagement group; Fisher's exact p = 0.013). The "
    "research did not seek engineering excellence in game design; the "
    "findings concern the behavioural context beneath the game play. We "
    "propose two working hypotheses of failure response: Failure-Induced "
    "Dropout (FIDS), associated with loss of self-efficacy, and "
    "Failure-Induced Abandonment (FIAS), associated with loss aversion when "
    "negative feedback is read as evidence of declining function. The study "
    "also provides directions for confirmatory testing using the "
    "domain-reframing hypothesis.", "Abstract")

# ---------------- Introduction: Stage-1 characterisation ----------------
sub("A survey at the study site found that 61%",
    [("A survey at the study site found that 61% of nearly 200 residents "
      "had never tried the available FRA equipment. The remaining 39% had "
      "tried it once or twice before giving up.",
      "A survey of 200 residents at the study site (of a resident "
      "population of 280) found that only 39% had ever used the available "
      "FRA equipment; among those users, approximately 74% failed on their "
      "first attempt and 82% did not return for subsequent use [46].")],
    "Intro survey figures aligned with SS-293")

# ---------------- Section V definitions text ----------------
sub("The distinction between FIDS and FIAS is operationalized here",
    [("FIDS is coded as dropout within 48 hours of a first failure event; "
      "FIAS additionally requires zero re-engagement and a composite score "
      "of five or above.",
      "FIDS is coded as initial dropout at, or immediately following, the "
      "session containing the first failure event (the archived dataset "
      "records cycle counts rather than timestamps, so all codings are "
      "session-based); FIAS additionally requires zero re-engagement "
      "through the study end. The descriptive composite of Section IX is "
      "not part of the evidential definition.")],
    "Section V operationalisation")

# ---------------- Table I rows ----------------
for t in d.tables:
    head = t.rows[0].cells[0].text.strip()
    if head == "Term":
        for row in t.rows[1:]:
            term = row.cells[0].text.strip()
            if term.startswith("FIDS"):
                c = row.cells[2].paragraphs[0]
                set_text(c, "Initial dropout at/immediately after the "
                            "session with the first failure (session-based; "
                            "no timestamps retained)")
            elif term.startswith("FIAS"):
                c = row.cells[2].paragraphs[0]
                set_text(c, "FIDS coding + zero re-engagement through "
                            "study end (composite is descriptive only)")
            elif term.startswith("Re-engagement"):
                c = row.cells[1].paragraphs[0]
                set_text(c, "At least one additional completed session "
                            "after a disengagement, in a later session "
                            "block")
                c2 = row.cells[2].paragraphs[0]
                set_text(c2, "Later-session completion after disengagement "
                             "(session-based)")
        log.append("Table I: FIDS/FIAS/re-engagement rows made session-based")

# ---------------- Leaderboard regression (unverifiable) ----------------
repl("The team experimented with and included a leaderboard",
    "The team experimented with a leaderboard mapped after Octalysis Drive "
    "5 (Social Influence), and the strategy backfired. Staff observation "
    "and field notes recorded that participants who attended to the "
    "leaderboard compared themselves unfavourably with other users and "
    "were quicker to judge themselves negatively; session satisfaction "
    "appeared lower among those participants. Because the archived dataset "
    "does not include leaderboard-viewing as a recorded variable, this "
    "observation is reported qualitatively and no coefficient is claimed "
    "for it. The pattern is consistent with Festinger's (1954) social "
    "comparison theory [37] and with An et al.'s (2024) conclusion that "
    "building confidence may matter more than creating competition in this "
    "population [11]. The leaderboard will be removed in the confirmatory "
    "study.", "Leaderboard paragraph (regression removed)")

# ---------------- Data verification statement ----------------
repl("All statistics presented were recalculated from the raw 30-participant dataset",
    "All statistics presented were recalculated from the raw "
    "30-participant dataset, and this version of the paper reports only "
    "the statistics that reproduce from that archive. The verification "
    "steps were as follows. First, the direction of FRA index change for "
    "each participant was compared with self-reported improvement to "
    "confirm scale directionality. Second, all means, standard deviations "
    "and effect sizes were recalculated directly from individual records: "
    "sustained participants (N = 9) improved by +6.9 ± 3.1 points (95% CI "
    "4.5-9.3; paired dz = 2.22, 95% CI 0.95-3.46), and the all-participant "
    "paired change was +4.0 points (95% CI 0.5-7.6; dz = 0.42, 95% CI "
    "0.05-0.79). Third, the low-engagement rate was confirmed at 70.0% "
    "(21 of 30; 95% CI 52.1-83.3%) and the single-session exit rate at "
    "26.7% (8 of 30; 95% CI 14.2-44.4%). Fourth, the satisfaction "
    "association with sustained engagement was confirmed (9/9 versus "
    "11/21; Fisher's exact p = 0.013). Withdrawal statement: three "
    "statistics reported in earlier drafts have been withdrawn and do not "
    "appear in this paper. A Cox proportional-hazards model and "
    "Kaplan-Meier analysis were withdrawn because the archived dataset "
    "preserves no time variable and therefore cannot support time-to-event "
    "modelling. A week-2 motivational-orientation classification and its "
    "odds ratio were withdrawn because the classification variable does "
    "not appear in the archived dataset and its coding records could not "
    "be produced. Pooled and zero-imputation effect sizes were withdrawn "
    "because they did not reproduce from the archive.",
    "Data verification + withdrawal statement")

# ---------------- Statistical methods ----------------
repl("Four methods are employed, each addressing a specific question.",
    "Three methods are employed, matched to a small single-site sample. "
    "Proportions are reported with Wilson score 95% confidence intervals. "
    "Within-participant change in the recorded FRA index is analysed with "
    "paired t and Wilcoxon tests, with exact noncentral-t confidence "
    "intervals for the paired effect size dz. Associations between binary "
    "participant characteristics and sustained engagement are tested with "
    "Fisher's exact test. Time-to-event modelling is not used: the "
    "archived dataset records cycle counts and session outcomes but no "
    "timestamps, so no survival analysis is supportable and none is "
    "reported.", "Statistical methods")

delete_para("Cox proportional hazards model.", "Cox model heading para")
delete_para("h(t|X) = h₀(t) × exp(βX)", "Cox formula")
delete_para("where t is time (weeks), h₀(t) is the baseline hazard",
            "Cox formula explanation")
delete_para("Model assumptions. A minimum of 10 events per predictor",
            "Cox assumptions para")

repl("Cohen’s d. Two versions are reported in Table II",
    "Cohen's dz. Two paired effect sizes are reported in Table II: for the "
    "nine sustained participants, dz = 2.22 (95% CI 0.95-3.46; mean change "
    "+6.9 points, 95% CI 4.5-9.3); for all thirty participants, dz = 0.42 "
    "(95% CI 0.05-0.79; mean change +4.0 points, 95% CI 0.5-7.6). Both are "
    "computed from archived paired pre/post records. No pooled or "
    "zero-imputation estimate is reported (see the withdrawal statement in "
    "Section VII-A).", "Cohen's d paragraph")

delete_para("Odds ratio (OR) for motivational orientation.",
            "OR paragraph")
repl("Note on motivational classification:",
    "Note on motivational classification: earlier drafts reported a "
    "week-2 motivational-orientation classification (learning-focused "
    "versus health-deficit-focused) with a large but very imprecise "
    "association with sustained engagement. The classification variable "
    "does not appear in the archived dataset and its coding records could "
    "not be produced, so the classification and every statistic derived "
    "from it have been withdrawn from this paper. The confirmatory study "
    "will administer a validated motivational measure at baseline, before "
    "any engagement outcome is observed.", "Motivational withdrawal note")

# ---------------- Table II: drop pooled/ITT rows ----------------
for t in d.tables:
    rows = [r.cells[0].text.strip() for r in t.rows]
    if any(r.startswith("Cohen’s d (paired") for r in rows):
        for row in list(t.rows):
            k = row.cells[0].text.strip()
            if k.startswith("Cohen’s d (pooled") or k.startswith("Cohen’s d (ITT"):
                row._element.getparent().remove(row._element)
            elif k.startswith("Cohen’s d (paired"):
                set_text(row.cells[0].paragraphs[0],
                         "Cohen's dz (paired, 95% CI)")
                set_text(row.cells[1].paragraphs[0], "2.22 (0.95–3.46)")
            elif k.startswith("Change (mean"):
                set_text(row.cells[1].paragraphs[0], "+6.9 ± 3.1 (95% CI 4.5–9.3)")
        log.append("Table II: pooled/ITT rows removed, CIs added")

sub("A higher FRA index denotes better balance. The three d values differ",
    [("The three d values differ because they measure change from "
      "different baselines. The 95% CI around the pooled d at n = 9 ranges "
      "from 0.20 to 1.76. Table II reports clinical balance outcomes for "
      "the nine participants who sustained engagement across four or more "
      "session cycles. Three effect size estimates are provided to reflect "
      "different analytical perspectives: the paired Cohen's d (2.22) "
      "captures within-subject change for completers; the pooled d (0.78) "
      "compares pre- and post-intervention group means; and the "
      "intention-to-treat d (0.57) assigns a change score of zero to all "
      "21 dropouts, providing a conservative estimate of the "
      "intervention's population-level effect.",
      "Table II reports clinical balance outcomes for the nine "
      "participants who sustained engagement across four or more session "
      "cycles: paired dz = 2.22 (95% CI 0.95-3.46) with a mean change of "
      "+6.9 points (95% CI 4.5-9.3). For the full sample of thirty, the "
      "paired change was +4.0 points (95% CI 0.5-7.6; dz = 0.42, 95% CI "
      "0.05-0.79). A single-session change of this size may include "
      "measurement variation, familiarisation and practice effects, and "
      "is read as a promising signal rather than a demonstrated clinical "
      "effect.")],
    "Table II commentary")

# ---------------- Abandonment pattern: replace Cox/KM block ----------------
sub("In the Table III summarizes the overall engagement pattern",
    [("In the Table III summarizes", "Table III summarizes")],
    "Table III intro fragment")

repl("Table IV presents the result of the Cox proportional hazards regression",
    "No time-to-event model is reported: the archived dataset preserves "
    "cycle counts and session outcomes but no timestamps, so survival "
    "modelling is not supportable from this archive. The evidential core "
    "of the abandonment pattern is instead carried by three directly "
    "observed quantities: the low-engagement proportion (70.0%, 95% CI "
    "52.1-83.3%), the single-session exit proportion (26.7%, 95% CI "
    "14.2-44.4%), and the zero re-engagement count among all 21 "
    "disengaged participants (95% CI 0.0-15.5%). Satisfaction was "
    "strongly associated with sustained engagement: all nine sustained "
    "participants were satisfied or very satisfied, against eleven of "
    "twenty-one in the low-engagement group (Fisher's exact p = 0.013).",
    "Cox commentary replaced")

delete_para("In Fig. 3. Kaplan–Meier survival curves showing time to permanent dropout",
            "KM figure commentary")
delete_para("Fig. 3. Kaplan–Meier survival curves: time to permanent dropout",
            "KM figure caption")
delete_para("TABLE IV. COX PROPORTIONAL HAZARDS MODEL (N = 30)",
            "Table IV caption")
for t in list(d.tables):
    if t.rows[0].cells[0].text.strip() == "Predictor" and "HR" in t.rows[0].cells[1].text:
        t._tbl.getparent().remove(t._tbl)
        log.append("Table IV (Cox) deleted")

# ---------------- Motivational orientation section ----------------
p = find("TABLE V. OUTCOMES BY MOTIVATIONAL ORIENTATION")
if p is not None:
    set_text(p, "Withdrawn analysis. Earlier drafts presented Table V, "
        "contrasting sustained engagement between motivational-orientation "
        "groups identified through week-2 interviews (odds ratio 31.5). "
        "The classification variable does not appear in the archived "
        "dataset and its coding records could not be produced, so the "
        "classification, the table and the odds ratio have been withdrawn. "
        "The direction of the idea — that a learning-oriented reading of "
        "the activity accompanies sustained engagement — remains a "
        "hypothesis for the confirmatory study, where motivational "
        "orientation will be measured with a validated instrument at "
        "baseline.")
    p.style = "Normal"
    log.append("Table V caption replaced with withdrawal note")
delete_para("Table V contrasts sustained engagement and clinical improvement rates",
            "Table V commentary")
for t in list(d.tables):
    if t.rows[0].cells[0].text.strip() == "Orientation":
        t._tbl.getparent().remove(t._tbl)
        log.append("Table V (motivational) deleted")

sub("E. Motivational Orientation",
    [("E. Motivational Orientation", "E. Withdrawn Motivational-Orientation Analysis")],
    "Section VIII-E heading")

# ---------------- Composite sections ----------------
sub("All 30 participants were coded according to a scale",
    [("Scores are aggregated into a single FIAS composite that ranges from 0 to 10.",
      "Scores are aggregated into a single FIAS composite that ranges from "
      "0 to 10. The composite is a descriptive case-identification device "
      "and carries no evidential weight in this paper: several attributes "
      "record engagement and exit behaviour, so participants who abandoned "
      "early necessarily score high, and the separation of the F4 cluster "
      "is partly true by construction.")],
    "IX composite status")

repl("Table VII shows the distribution of FIAS composite scores",
    "Table VII shows the distribution of FIAS composite scores across the "
    "five severity levels. The absence of scores between 6 and 8 is not "
    "read as evidence of a behavioural discontinuity: given the "
    "construction of the composite, the gap may be a mechanical effect of "
    "the scoring thresholds and the outcome-linked items, or sampling "
    "variation in a cohort of thirty. The six F4-level participants "
    "(scores 9-10) remain the most theoretically interesting subgroup for "
    "a directly observed reason that does not depend on the composite: "
    "three of them recorded a positive FRA index change during their "
    "single session and did not return. That divergence between recorded "
    "score direction and return behaviour challenges a purely clinical "
    "explanation of dropout and is consistent with the loss-aversion "
    "mechanism at the core of the FIAS hypothesis.",
    "Table VII commentary")

delete_para("Fig. 4. FIAS score distribution (N = 30). The gap between scores 5 and 9 is consistent",
            "Fig 4 duplicate caption")
repl("Fig. 4. Distribution of FIAS composite scores across all 30 participants. Scores range",
    "Fig. 4. Distribution of FIAS composite scores across all 30 "
    "participants. Scores range from 0 (F0) to 10 (F4). The composite "
    "includes engagement-outcome attributes, so the separation of the F4 "
    "cluster is expected by construction; the figure is descriptive case "
    "identification, not evidence. The empty interval between scores 5 "
    "and 9 is not read as a behavioural threshold.",
    "Fig 4 caption")

sub("Three of the F4-level participants did improve in FRA",
    [("This demonstrates that a clinical improvement measure that "
      "indicates greater health doesn't override the negative psychic "
      "force that drive F4 behaviours.",
      "The evidence suggests that a recorded improvement, by itself, did "
      "not bring these participants back — the observation that most "
      "clearly motivates the identity-level account examined in the "
      "Discussion.")],
    "F4 improvement sentence calibrated")

sub("The participants in F4 were asked about whether the decision to stop",
    [("Table VIII shows the distribution of FIAS composite scores across "
      "the five severity levels. The absence of any scores between 6 and 8 "
      "produces a visible gap in the distribution, consistent with the "
      "threshold model proposed in the FIDS–FIAS hypothesis. The six "
      "F4-level participants (scores 9–10) represent the most "
      "theoretically significant subgroup: three of these participants "
      "recorded a positive FRA index change during their single session, "
      "yet none returned. This finding challenges a purely clinical "
      "explanation of dropout and supports the loss aversion mechanism at "
      "the core of the FIAS hypothesis.", "")],
    "Section X duplicated text removed")

# ---------------- Discussion ----------------
repl("In the subset who stuck with the therapy, gamification worked well",
    "In the subset who sustained engagement, gamification worked well: "
    "the nine sustained participants improved by a mean of +6.9 FRA index "
    "points (95% CI 4.5-9.3; paired dz = 2.22, 95% CI 0.95-3.46), and all "
    "nine reported being satisfied or very satisfied. Across all thirty "
    "participants the paired improvement was +4.0 points (95% CI 0.5-7.6). "
    "These within-deployment changes are promising signals rather than "
    "demonstrated clinical effects.", "RQ1 discussion")

repl("Of those who stayed with the program but quit before cycle 3",
    "The abandonment pattern has three defining features. First, it is "
    "early: 70.0% of participants disengaged at or before three cycles, "
    "and 26.7% exited after a single session. Second, it is absolute "
    "within the observation window: none of the 21 disengaged "
    "participants returned (95% CI 0.0-15.5%), which distinguishes the "
    "pattern from ordinary usage attrition, where some users return [7]. "
    "Third, it is not explained by dissatisfaction alone, although "
    "satisfaction and sustained engagement are strongly associated "
    "(Fisher's exact p = 0.013): among the six F4-cluster participants, "
    "three recorded a positive FRA index change in their only session and "
    "still did not return. Practical obstacles — usability, staffing, "
    "scheduling — do not account for a pattern in which the technology "
    "was free, supervised, on site, and in three documented cases "
    "measurably working for the participant who walked away. The evidence "
    "is consistent with a psychological mechanism attached to the meaning "
    "of failure, which the FIDS-FIAS account names.", "RQ2 discussion")

repl("We hypothesize that the drop-off in the F3/F0-F2 groups",
    "The proposed mechanism has two steps. FIDS names the initial "
    "dropout: a first failure event collapses task self-efficacy and the "
    "participant withdraws from the session. FIAS names the consolidation "
    "into permanence: loss aversion makes re-engagement feel like a "
    "gamble on further unwelcome evidence about one's own body, so the "
    "participant never returns. The four falsifiable predictions H1-H4 "
    "(Section II-F) operationalise this account for the confirmatory "
    "study. If first-failure response ceases to predict abandonment once "
    "fear of falling is controlled for (H1), or if participants describe "
    "deliberate rather than impulsive exits (H2), or game-language rather "
    "than health-language accounts of performance (H3), or if technical "
    "ease alone prevents abandonment (H4), the account will be "
    "contradicted in whole or in part.", "RQ3 discussion")

sub("The findings showed that the high (87.5% engagement) subject subset",
    [("The findings showed that the high (87.5% engagement) subject "
      "subset, the learning subjects, all treated FRA as a training "
      "subject, never a fixed condition, whereas the low (18.2% "
      "engagement) subject subset, the health subjects, saw FRA as a "
      "fixed condition and because of the endowment effect, would never "
      "re-entry into the service again.",
      "Field observation during the deployment was consistent with this "
      "account: participants who spoke about the activity in learning "
      "terms tended to sustain engagement, while participants who spoke "
      "about it in health-deficit terms tended to exit early. Because the "
      "week-2 classification that formalised this observation has been "
      "withdrawn (Section VII-A), the pattern is reported as a field "
      "observation and a hypothesis for the confirmatory study, not as a "
      "measured effect.")],
    "Endowment observation calibrated")

# ---------------- Limitations ----------------
repl("The sample of 30 participants was relatively small",
    "The sample of 30 participants was small (aged 62-84) and drawn from "
    "a single residential care facility in Southern China, so the "
    "findings apply to this relatively homogeneous cohort and require "
    "replication in culturally diverse populations. The archived dataset "
    "preserves cycle counts and session outcomes but no timestamps, so "
    "the temporal texture of disengagement — including whether an exit "
    "followed the first or second cycle within a session — cannot be "
    "reconstructed, and no time-to-event claims are made. Three "
    "statistics from earlier drafts (a Cox hazard ratio, a "
    "motivational-orientation odds ratio, and pooled/zero-imputation "
    "effect sizes) have been withdrawn for the reasons stated in Section "
    "VII-A; the withdrawal is disclosed rather than silent. The FES-I was "
    "not administered. Prospect Theory offers a plausible framework, but "
    "no direct test of it was performed in this study, and extending the "
    "endowment effect to health self-concept moves beyond the domains in "
    "which that effect has been documented; both remain propositions for "
    "the confirmatory study. Neither the domain-reframing hypothesis (H5) "
    "nor the coaching layer was tested here. The investigator's dual role "
    "as observer and participant creates a risk of bias that reflexive "
    "field notes mitigate but do not eliminate; no independent observer "
    "monitored interpretations.", "Limitations")

# ---------------- Conclusion ----------------
repl("This study built and tested a gamified FRA system using dual-purpose input design",
    "This study built and tested a gamified FRA system using a "
    "dual-purpose input design, deployed at a residential care facility "
    "in Southern China. For participants who sustained engagement, "
    "gamification helped: the nine sustained participants improved by a "
    "mean of +6.9 FRA index points (95% CI 4.5-9.3; paired dz = 2.22), "
    "and all nine were satisfied or very satisfied. But 70.0% of "
    "participants never passed three cycles, 26.7% stopped after their "
    "first session, and no participant returned after disengaging.",
    "Conclusion para 1")

repl("We propose that initial dropout (FIDS) is induced by self-efficacy collapse",
    "We propose that initial dropout (FIDS) is associated with "
    "self-efficacy collapse — the user loses the belief that they can "
    "improve, and withdraws — and that loss aversion (FIAS) consolidates "
    "the withdrawal when the game failure is read as evidence of physical "
    "decline. Satisfaction was strongly associated with sustained "
    "engagement (Fisher's exact p = 0.013), and field observation "
    "suggested that a learning-oriented reading of the activity "
    "accompanied persistence; the formal test of that suggestion belongs "
    "to the confirmatory study.", "Conclusion para 2")

sub("Drawing upon research in the endowment effect",
    [("Drawing upon research in the endowment effect (Kahneman & Tversky "
      "1979) and stereotype threat reframing (Ben-Ami & Ruef 2008), we "
      "propose the domain reframing hypothesis:",
      "Drawing upon research on the endowment effect (Thaler, 1980; "
      "Kahneman, Knetsch & Thaler, 1990) and stereotype-threat reframing "
      "(McGlone & Aronson, 2006; Barber, 2017), we propose the domain "
      "reframing hypothesis:")],
    "Conclusion citations fixed")

sub("The three-armed, Phase 2 clinical trial will be conducted",
    [("All study hypotheses (H1–H5) and primary end points are registered.",
      "All study hypotheses (H1-H5) and primary end points will be "
      "pre-registered before data collection begins. An interim "
      "between-cohort deployment of a decoupled version of the game with "
      "120 community-dwelling adults, reported in a companion manuscript, "
      "provides preliminary quasi-experimental support for the "
      "decoupling direction while leaving the randomised within-cohort "
      "test to the study described here.")],
    "Conclusion Phase-2 bridge")

# ---------------- SS-293 reference ----------------
p = find("[45] FDA, Health Canada, and MHRA")
if p is not None:
    new = OxmlElement("w:p")
    p._p.addnext(new)
    np_ = docx.text.paragraph.Paragraph(new, p._parent)
    np_.style = p.style
    np_.add_run("[46] S. K. T. Tam and C. F. Cheung, “A study of "
                "utilisation of Fall Risk Assessment among elderly users: "
                "A critical review through TAM and Octalysis frameworks "
                "with gamification approach,” accepted for "
                "publication (SS-293).")
    log.append("[46] SS-293 reference added")
else:
    fail.append("reference [45] anchor")

cp = d.core_properties
cp.title = "Paper 2 R1 - aligned with thesis V20"
cp.comments = ("R1 (19 Jul 2026): withdrawn statistics removed (Cox/KM, "
               "OR 31.5, pooled/ITT d), verified V20 statistics inserted, "
               "composite made descriptive-only, session-based "
               "definitions, calibrated language, factual fixes.")

d.save(DST)
print("saved", DST)
for l in log:
    print(" -", l)
for f in fail:
    print(" !", f)

# verification
d2 = docx.Document(DST)
full = "\n".join(p.text for p in d2.paragraphs)
for tb in d2.tables:
    for r in tb.rows:
        full += "\n" + " | ".join(c.text for c in r.cells)
for s in ("3.24", "31.5", "0.78", "0.57", "Kaplan", "hazard", "48 hours",
          "48h", "Shanghai", "6.75", "87.5%", "18.2%", "61% of nearly",
          "Ben-Ami", "beta = -0.38", "$\\beta"):
    c = full.count(s)
    if c:
        print(f"RESIDUAL '{s}': {c}")
        for line in full.split("\n"):
            if s in line:
                print("   >", line[:120])
