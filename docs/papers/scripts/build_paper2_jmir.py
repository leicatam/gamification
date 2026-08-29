#!/usr/bin/env python3
"""Convert Paper2_R2_Full_Rewrite.docx into a JMIR Serious Games
submission: IMRD structure, structured abstract (<=450 words), JMIR
back-matter sections (Data Availability, Authors' Contributions,
Conflicts of Interest, Abbreviations), Arabic table numbers, Figure N
captions with callouts, and a new difficulty-levels figure."""
import re
import shutil
import docx
from docx.oxml import OxmlElement
from docx.shared import Inches

BASE = "/home/user/gamification/docs/papers/"
SRC = BASE + "Paper2_R2_Full_Rewrite.docx"
DST = BASE + "Paper2_JMIR_SeriousGames_Manuscript.docx"
shutil.copy(SRC, DST)
d = docx.Document(DST)
log, fail = [], []


def set_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in list(p.runs[1:]):
            r._element.getparent().remove(r._element)
    else:
        p.add_run(text)


def find(needle, style_prefix=None):
    for p in d.paragraphs:
        if needle in p.text and (style_prefix is None
                                 or p.style.name.startswith(style_prefix)):
            return p
    return None


def insert_after(p, text="", style=None):
    new = OxmlElement("w:p")
    p._p.addnext(new)
    np_ = docx.text.paragraph.Paragraph(new, p._parent)
    if style:
        np_.style = style
    if text:
        np_.add_run(text)
    return np_


# ---------- 1. Title & byline ----------
set_text(d.paragraphs[0],
    "Failure-Induced Abandonment of a Gamified, Artificial "
    "Intelligence-Assisted Fall Risk Assessment System Among Older "
    "Adults in Residential Care: Action Research Study")
log.append("title")

# ---------- 2. Structured abstract ----------
ABS = [
 ("Background: ", "Fall Risk Assessment (FRA) systems assess balance and "
  "support fall prevention as clinical devices, but they rarely achieve "
  "sustained engagement among older adults. At the study facility, a "
  "survey of 200 residents found that only 39% had ever used the "
  "installed FRA system; approximately 74% of users failed at their "
  "first attempt, and 82% of those never returned. Established "
  "acceptance models and gamification frameworks predict whether people "
  "start using a system and how engagement is motivated within a "
  "session; none represents what happens after a user fails."),
 ("Objective: ", "This study aimed to characterise engagement and "
  "abandonment behaviour in a gamified, artificial intelligence "
  "(AI)-assisted FRA deployment among elderly residents, and to propose "
  "a testable behavioural account of rapid, permanent post-failure "
  "abandonment."),
 ("Methods: ", "An action-research deployment ran for eight weeks at a "
  "residential care facility in Southern China. An AI-assisted skiing "
  "game was built on the FRA system's weight-shifting foot pressure "
  "panel, so that one physical action steers the on-screen skier while "
  "yielding the centre-of-pressure stream from which the clinical "
  "balance index is computed. First-session difficulty was set low "
  "enough for the weakest-balance participant to succeed, and an AI "
  "layer adapted difficulty to individual performance. Thirty residents "
  "(aged 62-84 years; mean 71.1, SD 5.9; 13 men, 17 women) participated "
  "with informed consent. Behavioural outcomes (disengagement, "
  "re-engagement), the recorded FRA index and satisfaction were "
  "analysed with Wilson 95% confidence intervals, paired t and Wilcoxon "
  "tests, and Fisher's exact test."),
 ("Results: ", "Seventy per cent of participants (21/30; 95% CI "
  "52.1-83.3) disengaged at or before three game cycles, 26.7% (8/30; "
  "95% CI 14.2-44.4) exited after a single session, and no disengaged "
  "participant re-engaged within the study window (0/21; 95% CI "
  "0.0-15.5). Among the nine sustained participants, the recorded FRA "
  "index improved by a mean of +6.9 points (95% CI 4.5-9.3; paired "
  "dz=2.22, 95% CI 0.95-3.46); across all thirty participants the "
  "paired change was +4.0 points (95% CI 0.5-7.6; dz=0.42, 95% CI "
  "0.05-0.79). Satisfaction was strongly associated with sustained "
  "engagement (9/9 vs 11/21; P=.013). Three participants whose recorded "
  "balance index improved during their only session still never "
  "returned."),
 ("Conclusions: ", "The findings motivate two working hypotheses: "
  "failure-induced dropout (FIDS), in which a first failure collapses "
  "task self-efficacy, and failure-induced abandonment (FIAS), in which "
  "loss aversion converts negative feedback into evidence of bodily "
  "decline, so that non-return becomes permanent. Four falsifiable "
  "predictions (H1-H4) and a domain-reframing hypothesis (H5) are "
  "specified for a preregistered three-arm confirmatory study. "
  "Practitioner guidance follows directly: guarantee success in the "
  "first session, monitor observable abandonment-risk signals, coach "
  "with mastery-oriented language, and never display output the user "
  "could read as a health verdict."),
]
absp = d.paragraphs[3]
set_text(absp, "")
absp.runs[0].text = ""
lead = absp.add_run(ABS[0][0]); lead.bold = True
absp.add_run(ABS[0][1])
anchor = absp
for label, body in ABS[1:]:
    np_ = insert_after(anchor)
    r = np_.add_run(label); r.bold = True
    np_.add_run(body)
    anchor = np_
log.append("structured abstract")

# ---------- 3. Keywords ----------
kw = find("Keywords: fall risk assessment")
set_text(kw, "Keywords: gamification; serious games; older adults; fall "
    "risk assessment; technology abandonment; eHealth attrition; "
    "self-efficacy; loss aversion; balance; action research")
log.append("keywords")

# ---------- 4. Heading map ----------
H = {
 "I. Introduction": ("Introduction", "Heading 1"),
 "A. Industrial Problem": ("Background", "Heading 2"),
 "B. Research Questions": ("Research Questions", "Heading 2"),
 "II. Literature Review": ("Prior Work", "Heading 2"),
 "A. Elderly Technology Adoption": ("Elderly Technology Adoption", "Heading 3"),
 "B. Positioning Against the Science of Attrition": ("Positioning Against the Science of Attrition", "Heading 3"),
 "C. Fear of Falling": ("Fear of Falling", "Heading 3"),
 "D. Gamification and Elderly Health Technology": ("Gamification and Elderly Health Technology", "Heading 3"),
 "E. Theoretical Frameworks": ("Theoretical Frameworks", "Heading 3"),
 "F. Self-Efficacy, Prospect Theory, and the FIDS-FIAS Transition": ("Self-Efficacy, Prospect Theory, and the FIDS-FIAS Transition", "Heading 3"),
 "G. Motivation Crowding and the FRA-Game Coupling": ("Motivation Crowding and the FRA-Game Coupling", "Heading 3"),
 "H. Behaviour-Change Lens: COM-B": ("Behaviour-Change Lens: COM-B", "Heading 3"),
 "I. Information Entropy as a Conceptual Lens": ("Information Entropy as a Conceptual Lens", "Heading 3"),
 "J. Research Gap": ("Research Gap", "Heading 3"),
 "III. Original Contribution Statement": ("Objectives and Contribution", "Heading 2"),
 "IV. Research Design, Ethics and Participants": ("Methods", "Heading 1"),
 "A. Action-Research Methodology": ("Study Design", "Heading 2"),
 "B. Ethics, Consent and Safeguarding": ("Ethical Considerations", "Heading 2"),
 "C. Data Governance and Protection": ("Data Governance and Protection", "Heading 2"),
 "D. Study Site and Participants": ("Study Site and Participants", "Heading 2"),
 "V. Definitions and Operationalisation": ("Definitions and Operationalisation", "Heading 2"),
 "VI. The Gamified Intervention": ("The Gamified Intervention", "Heading 2"),
 "A. Dual-Purpose Input": ("Dual-Purpose Input", "Heading 3"),
 "B. Game Mechanics": ("Game Mechanics", "Heading 3"),
 "C. Five-Layer Architecture": ("Five-Layer Architecture", "Heading 3"),
 "D. Proposed Coaching Layer (Confirmatory Study)": ("Proposed Coaching Layer (Confirmatory Study)", "Heading 3"),
 "VII. Data Collection and Measures": ("Data Collection and Measures", "Heading 2"),
 "A. Data Verification and Withdrawal Statement": ("Data Verification and Withdrawal Statement", "Heading 3"),
 "A. Statistical Methods": ("Statistical Analysis", "Heading 2"),
 "VIII. Quantitative Findings": ("Results", "Heading 1"),
 "B. Clinical Outcomes": ("Clinical Outcomes", "Heading 2"),
 "C. The Abandonment Pattern": ("The Abandonment Pattern", "Heading 2"),
 "IX. Qualitative Coding and the FIAS Composite": ("Qualitative Coding and the FIAS Composite", "Heading 2"),
 "X. F4 Interview Protocol": ("F4 Interviews", "Heading 2"),
 "XI. Discussion": ("Discussion", "Heading 1"),
 "A. RQ1: Did Gamification Help?": ("Principal Findings: Did Gamification Help? (RQ1)", "Heading 2"),
 "B. RQ2: Why Does Abandonment Persist?": ("Why Does Abandonment Persist? (RQ2)", "Heading 2"),
 "C. RQ3: The FIDS-FIAS Mechanism": ("The FIDS-FIAS Mechanism (RQ3)", "Heading 2"),
 "D. Extended Theory: Endowment, Stereotype Threat, Domain Reframing": ("Extended Theory: Endowment, Stereotype Threat, and Domain Reframing", "Heading 2"),
 "XII. Limitations": ("Limitations", "Heading 2"),
 "XIII. Design Implications for Practitioners": ("Design Implications for Practitioners", "Heading 2"),
 "XIV. The Entropy Constraint": ("The Entropy Constraint", "Heading 2"),
 "XV. Industrial Context and Confirmatory Design": ("Future Work: Industrial Context and the Confirmatory Study", "Heading 2"),
 "A. Sponsor Requirements and Cost-Effectiveness": ("Sponsor Requirements and Cost-Effectiveness", "Heading 3"),
 "B. Three-Arm Confirmatory Design": ("Three-Arm Confirmatory Design", "Heading 3"),
 "XVI. Conclusion": ("Conclusions", "Heading 2"),
}
renamed = 0
for p in d.paragraphs:
    t = p.text.strip()
    if t in H and p.style.name.startswith("Heading"):
        new, sty = H[t]
        set_text(p, new)
        p.style = sty
        renamed += 1
log.append(f"headings renamed: {renamed}/{len(H)}")
if renamed != len(H):
    seen = {p.text.strip() for p in d.paragraphs if p.style.name.startswith("Heading")}
    fail.append("headings missed: " + str([k for k in H if H[k][0] not in seen]))

# ---------- 5. Move Statistical Analysis into Methods ----------
stat_h = find("Statistical Analysis", "Heading")
res_h = find("Results", "Heading 1")
if stat_h is None or res_h is None:
    fail.append("stat move")
else:
    stat_body = stat_h._p.getnext()  # the single methods paragraph
    res_h._p.addprevious(stat_h._p)
    stat_h._p.addnext(stat_body)
    log.append("Statistical Analysis moved into Methods")

# ---------- 6. Figure renumbering (Fig. N -> Figure M) ----------
FIGMAP = {5: 6, 4: 5, 3: 4, 2: 3, 1: 1}
count = 0
for p in d.paragraphs:
    if re.search(r"Fig\. \d", p.text):
        t = p.text
        for old in (5, 4, 3, 2, 1):
            t = t.replace(f"Fig. {old}", f"Figure {FIGMAP[old]}")
        set_text(p, t)
        count += 1
log.append(f"Fig->Figure conversions in {count} paragraphs")

# ---------- 7. New Figure 2 (difficulty levels) ----------
p57 = find("The system comprises five layers.")
if p57 is None:
    fail.append("five-layer anchor")
else:
    set_text(p57, p57.text.rstrip() + " Figure 2 shows the deployed game "
             "at five staged difficulty levels.")
    img = insert_after(p57)
    img.add_run().add_picture(
        BASE + "figures_submission/Figure3_Game_Screenshots_Difficulty_Levels.png",
        width=Inches(6.0))
    insert_after(img,
        "Figure 2. The skiing game at five staged difficulty multipliers "
        "(x0.40-x1.40). Obstacle density and scroll speed scale with the "
        "AI-set difficulty multiplier; screenshots are from the archived "
        "build in the deployed design lineage.")
    log.append("Figure 2 inserted")

# ---------- 8. Callouts for figures 3, 4, 6 ----------
for anchor_text, sentence in [
    ("The cycle distribution was", " Figure 3 shows the participant flow."),
    ("The empty interval between scores 5 and 9 is not read as evidence",
     " Figure 4 shows the distribution."),
    ("The confirmatory study marks the shift from exploratory action research",
     " Figure 6 summarises the design.")]:
    p = find(anchor_text)
    if p is None:
        fail.append(f"callout: {anchor_text[:30]}")
    else:
        set_text(p, p.text.rstrip() + sentence)
log.append("figure callouts added")

# ---------- 9. Roman -> Arabic table numbers ----------
ROMAN = {"I": 1, "II": 2, "III": 3, "IV": 4, "V": 5, "VI": 6, "VII": 7}
tconv = 0
for p in d.paragraphs:
    m = re.search(r"\bTable ([IVX]+)\b", p.text)
    if m:
        t = re.sub(r"\bTable ([IVX]+)\b",
                   lambda mm: f"Table {ROMAN[mm.group(1)]}", p.text)
        set_text(p, t)
        tconv += 1
log.append(f"Roman table numbers converted in {tconv} paragraphs")

# ---------- 10. Split COI out of Acknowledgments ----------
ack = find("The authors thank the management and nursing staff")
coi_text = "None declared."
if ack and "The authors declare no competing" in ack.text:
    idx = ack.text.index("The authors declare no competing")
    coi_text = ack.text[idx:].strip()
    set_text(ack, ack.text[:idx].rstrip())
    log.append("COI split from Acknowledgments")

# ---------- 11. Back matter before References ----------
ref_h = None
for p in d.paragraphs:
    if p.text.strip() == "References" and p.style.name.startswith("Heading"):
        ref_h = p
        break
if ref_h is None:
    fail.append("References heading")
else:
    prev = ref_h._p.getprevious()
    anchor = docx.text.paragraph.Paragraph(prev, ref_h._parent)
    sections = [
        ("Data Availability",
         "The de-identified Stage-2 dataset (N=30), the derived-variable "
         "code and the analysis scripts that reproduce every statistic "
         "reported in this paper are held in the study archive and are "
         "available from the corresponding author on reasonable request."),
        ("Authors' Contributions",
         "SKTT conceived the study, developed the gamified system, "
         "conducted the fieldwork, performed the analysis and drafted the "
         "manuscript. CFC supervised the research programme, contributed "
         "to the methodology and critically revised the manuscript. Both "
         "authors read and approved the final manuscript."),
        ("Conflicts of Interest", coi_text),
        ("Abbreviations",
         "AI: artificial intelligence; CI: confidence interval; COM-B: "
         "capability, opportunity, motivation-behaviour; COP: centre of "
         "pressure; FES-I: Falls Efficacy Scale-International; FIAS: "
         "failure-induced abandonment syndrome; FIDS: failure-induced "
         "dropout syndrome; FRA: fall risk assessment; RQ: research "
         "question; STAM: senior technology acceptance model; TAM: "
         "technology acceptance model; WTA: willingness to accept; WTP: "
         "willingness to pay."),
    ]
    for title, body in sections:
        h = insert_after(anchor, title, style="Heading 1")
        anchor = insert_after(h, body)
    log.append("back matter inserted")

cp = d.core_properties
cp.title = "Failure-Induced Abandonment of a Gamified AI-Assisted FRA: Action Research Study"
cp.comments = "JMIR Serious Games submission build from Paper2 R2."
d.save(DST)
print("LOG:", *log, sep="\n  ")
print("FAIL:", fail if fail else "none")
