#!/usr/bin/env python3
"""Paper 3 rewritten as the Stage-3 manuscript: the between-cohort decoupled
deployment actually conducted (N = 120, register 38/41/41), replacing the
'Key Contexts' scaffold that presented an unrun three-arm study as its
evidence. Theory sections (design logic of the response-challenge framing,
WTA/WTP, adaptive-learning/endowment mutual exclusion) are retained from the
scaffold, condensed and calibrated. All figures follow thesis V20."""
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

d = docx.Document()
style = d.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(11)


def H1(t):
    return d.add_heading(t, level=1)


def H2(t):
    return d.add_heading(t, level=2)


def P(t):
    return d.add_paragraph(t)


def table(rows, widths=None):
    t = d.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri].cells[ci]
            cell.paragraphs[0].add_run(str(val))
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
                if ri == 0:
                    run.bold = True
    return t


def caption(t):
    p = d.add_paragraph(t)
    for r in p.runs:
        r.font.size = Pt(9)
        r.italic = True
    return p


# ---------------- front matter ----------------
t = d.add_paragraph()
r = t.add_run("Decoupling the Clinical Attachment from Gamified Fall Risk "
              "Assessment: Between-Cohort Evidence for a Frame-Sensitive "
              "Account of Elderly Abandonment")
r.bold = True
r.font.size = Pt(14)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
a = d.add_paragraph("Sidney K. T. Tam and Chi Fai Cheung\n"
                    "Behaviour and Knowledge Engineering Research Centre, "
                    "Department of Industrial and Systems Engineering,\n"
                    "The Hong Kong Polytechnic University, Hong Kong, China\n"
                    "Corresponding author: sidney.tam@connect.polyu.hk")
a.alignment = WD_ALIGN_PARAGRAPH.CENTER

H1("Abstract")
P("Earlier work in this programme documented rapid, permanent abandonment "
  "of a gamified, AI-assisted Fall Risk Assessment (FRA) by elderly "
  "residents after first failure: 70.0% disengaged at or before three game "
  "cycles and none returned, even where the recorded balance index had "
  "improved. Two rival readings of that result make opposite predictions — "
  "a dispositional reading, on which gamification does not transfer to "
  "frail older adults, and a coupling reading, on which the attachment of "
  "the game to a clinical verdict on the user's body carried the failure. "
  "This paper reports the evaluation that separates them: a between-cohort "
  "quasi-experimental deployment in which the same broad apparatus, skiing "
  "mechanic, five-attempt rule and session length were retained while the "
  "participant-facing clinical attachment was removed as a package — "
  "recruitment language, pre-session clinical scan, clinical labels and "
  "participant-facing clinical outputs — and the activity was framed as an "
  "eye-leg coordination challenge. One hundred and twenty community-"
  "dwelling adults aged 25-80 were recruited in Hong Kong across three age "
  "bands (38 aged 25-44, 41 aged 45-64, 41 aged 65-80). In the preliminary "
  "aggregate record, retention at the end of the study window was 93.3% "
  "overall and 85.4% in the 65-80 band; 15 of 23 initial dropouts returned "
  "voluntarily (65.2%; 10 of 16, or 62.5%, in the 65-80 band); "
  "coordination-challenge endorsement was 71.7%; and stated regular-use "
  "intention was 63.3% overall and 56.1% in the 65-80 band. Initial "
  "dropout climbed steeply with age (0 of 38 versus 16 of 41; Fisher's "
  "exact p < 0.0001), while observed non-return converged to low levels in "
  "every band (0.0%, 4.9%, 14.6%). The pattern is consistent with the "
  "coupling account and inconsistent with the dispositional account, "
  "within the limits of a between-cohort, multi-component design. A "
  "residual 29.3-percentage-point gap between session retention and "
  "regular-use intention in the oldest band motivates adaptive learning — "
  "trajectory-oriented feedback with calibrated challenge — as the design "
  "response, and a pre-specified within-cohort randomised study as the "
  "confirmatory test.")
P("Keywords: fall risk assessment, gamification, elderly technology "
  "abandonment, FIAS, decoupling, framing, endowment effect, adaptive "
  "learning.")

# ---------------- I ----------------
H1("I. Introduction")
P("A survey of 200 residents at a Guangzhou care facility (of a resident "
  "population of 280) found that only 39% had ever used an installed FRA "
  "system; among users, approximately 74% failed on their first attempt "
  "and 82% did not return [1]. A subsequent eight-week action-research "
  "deployment with thirty residents tested whether a gamified, AI-assisted "
  "FRA — engineered to satisfy the acceptance and gamification design "
  "criteria of the technology-acceptance and Octalysis literatures — would "
  "sustain engagement. It did not: 70.0% of participants (95% CI "
  "52.1-83.3%) disengaged at or before three cycles, 26.7% exited after a "
  "single session, and none of the 21 disengaged participants re-engaged "
  "during the study window (95% CI 0.0-15.5%) [2]. Most strikingly, three "
  "of the six participants in the strongest-abandonment cluster recorded a "
  "positive change in the balance index during their single session and "
  "still did not return.")
P("That rejection is ambiguous between two explanations. Under the "
  "dispositional explanation, gamification simply does not work for frail "
  "older adults: the literature's successes with non-clinical technologies "
  "do not transfer. Under the coupling explanation, the game was never the "
  "problem: attaching the gameplay to a clinical verdict on the user's own "
  "body changed the meaning of failure, and the failure-induced "
  "abandonment (FIAS) documented in [1], [2] is a property of that "
  "attachment. The two explanations make opposite predictions for a "
  "decoupled version of the same game. This paper reports that evaluation. "
  "It evaluates a participant-facing redesign package in a separate "
  "cohort; it does not isolate a single causal variable, and its claims "
  "are calibrated accordingly.")

# ---------------- II ----------------
H1("II. Theoretical Background")
H2("A. The FIAS Account and the Endowment Mechanism")
P("FIAS (Failure-Induced Abandonment Syndrome) is a provisional "
  "behavioural construct: active, rapid and persistent non-return after an "
  "adverse first experience under identity-relevant conditions [1]. The "
  "candidate mechanism developed in [2] runs through the endowment effect "
  "[3], [4]. An elderly participant arrives carrying an identity-level "
  "asset — the belief in their own physical capability. A clinically "
  "framed failure forces a devaluation of that asset. The "
  "willingness-to-accept (WTA) such a devaluation is very high, because "
  "the asset is neither fungible nor separable from the person's sense of "
  "self; the willingness-to-pay (WTP) for re-engagement — bounded by "
  "finite expected benefits — cannot match it. On this account the "
  "transaction never clears and the participant never returns, which is "
  "consistent with the observed zero re-engagement. Extending the "
  "endowment effect to health self-concept goes beyond the domains in "
  "which it has been demonstrated experimentally; it is carried here as a "
  "candidate mechanism, not a measured one.")
H2("B. Two Appraisal Regimes")
P("The same account implies a structural proposition: an endowment-"
  "protective appraisal and a learning-trajectory appraisal are competing "
  "readings of the same feedback event. Read against a fixed self-concept, "
  "a failure is a verdict on a possessed asset; read against a learning "
  "curve, it is a data point on a trajectory. The two regimes route "
  "identical information — a missed obstacle on a pressure panel — to "
  "different targets: the self, or the task. Which regime a participant "
  "enters is hypothesised to depend on the frame the deployment "
  "establishes. This proposition motivates both the decoupling evaluated "
  "here (remove the verdict frame) and the adaptive-learning design "
  "direction discussed in Section VI (install the trajectory frame); as a "
  "claim of mutual exclusion it is theoretical, and its direct test is "
  "specified as future work.")

# ---------------- III ----------------
H1("III. Method: A Participant-Facing Decoupling Package in a Separate Cohort")
H2("A. Design Logic of the Coordination-Challenge Framing")
P("The design problem was practical. A pure entertainment framing would "
  "not recruit meaningfully; a health-assessment framing re-activates the "
  "very mechanism under study. The deployment therefore framed the "
  "activity as an eye-leg coordination challenge: participants were told "
  "the game tests how quickly the eyes can spot obstacles and how fast the "
  "legs can respond, and that response speed improves with practice. The "
  "framing was chosen for three properties: it avoids the physical-decline "
  "schema that 'balance' and 'fall' language activates; it avoids "
  "cognitive-decline connotations; and it presents performance as "
  "inherently variable rather than as a stable capacity, so that no "
  "identity-level reference point is established for a failure to "
  "threaten.")
H2("B. The Decoupling Package")
P("Relative to the coupled deployment of [2], four participant-facing "
  "components changed as a package: recruitment language (a coordination "
  "challenge, with no medical or fall-prevention vocabulary); removal of "
  "the pre-session clinical scan; removal of clinical labels from the "
  "station and screens; and removal of participant-facing clinical "
  "outputs, with no balance score or risk index shown at any point. The "
  "broad pressure-panel apparatus, skiing mechanic, five-attempt rule and "
  "session length were retained. Participant-level clinical balance "
  "telemetry was deliberately not collected. Because cohort, setting and "
  "recruitment channel also differed, the design estimates the combined "
  "association of the decoupled condition and its context with behaviour.")
H2("C. Participants")
P("One hundred and twenty community-dwelling adults aged 25-80 were "
  "recruited through Hong Kong community channels: 38 aged 25-44, 41 aged "
  "45-64 and 41 aged 65-80; 48 men and 72 women; all reported high-school "
  "education or above. Recruiting three age bands builds the comparison "
  "group into the study and licenses a second null hypothesis, H0-age: "
  "that under the decoupled condition there is no difference between age "
  "groups. Data collection is complete; this paper reports the "
  "preliminary aggregate record, with the participant-level records held "
  "in the study archive supporting analysis beyond it.")
H2("D. Outcomes and Analysis")
P("Outcomes are behavioural and stated: initial dropout (exit before "
  "completing the five-attempt protocol), voluntary return after initial "
  "dropout, observed non-return at the end of the study window, "
  "coordination-challenge endorsement (a manipulation check on the "
  "intended frame), and stated regular-use intention. Proportions carry "
  "Wilson score 95% confidence intervals; band contrasts use Fisher's "
  "exact and chi-square tests. No time-to-event modelling is used.")

# ---------------- IV ----------------
H1("IV. Results (Preliminary Aggregate Record)")
P("Overall, 23 of 120 participants (19.2%) dropped out before completing "
  "the protocol; 15 of the 23 (65.2%; 95% CI 44.9-81.2%) returned "
  "voluntarily; 8 (6.7%) did not return, giving end-of-window retention "
  "of 93.3% (Table 1).")
caption("Table 1. Overall engagement and dropout (N = 120).")
table([
    ("Metric", "n", "%"),
    ("Initial dropout before completing the protocol", "23", "19.2"),
    ("Voluntary return after initial dropout", "15", "65.2 of dropouts"),
    ("Observed non-return at end of window", "8", "6.7"),
    ("Retention at end of window", "112", "93.3"),
])
P("Engagement was strongly age-stratified (Table 2). Initial dropout rose "
  "from 0.0% (0 of 38) in the 25-44 band through 17.1% (7 of 41) in the "
  "45-64 band to 39.0% (16 of 41) in the 65-80 band (chi-square across "
  "bands p < 0.001; 25-44 versus 65-80, Fisher's exact p < 0.0001), "
  "nullifying H0-age for initial engagement. Observed non-return, by "
  "contrast, converged to low levels in every band: 0.0%, 4.9% and 14.6% "
  "(95% CI 6.9-28.4% in the 65-80 band), because most older dropouts came "
  "back: 10 of 16 in the 65-80 band (62.5%; 95% CI 38.6-81.5%) against 0 "
  "of 21 observed re-engagements in the coupled deployment of [2].")
caption("Table 2. Age-stratified engagement.")
table([
    ("Age band", "n", "Initial dropout", "Returned", "Non-return", "Retention"),
    ("25–44", "38", "0 (0.0%)", "0", "0 (0.0%)", "100.0%"),
    ("45–64", "41", "7 (17.1%)", "5", "2 (4.9%)", "95.1%"),
    ("65–80", "41", "16 (39.0%)", "10", "6 (14.6%)", "85.4%"),
    ("Total", "120", "23 (19.2%)", "15", "8 (6.7%)", "93.3%"),
])
P("Coordination-challenge endorsement reached 71.7% (86 of 120; 95% CI "
  "63.0-79.0%), indicating that most participants perceived the intended "
  "frame; the per-band breakdown was not retained in the preliminary "
  "summary and no per-band endorsement claim is made. Stated regular-use "
  "intention was 63.3% overall (76 of 120) — 65.8%, 68.3% and 56.1% by "
  "band (Table 3). In the 65-80 band, the difference between observed "
  "session retention (85.4%) and stated regular-use intention (56.1%; 95% "
  "CI 41.0-70.1%) is 29.3 percentage points; the two are distinct "
  "constructs, and the difference is a design diagnostic rather than a "
  "measured causal effect.")
caption("Table 3. Endorsement and stated regular-use intention.")
table([
    ("Item", "25–44", "45–64", "65–80", "Total"),
    ("Coordination-challenge endorsement", "—", "—", "—", "71.7%"),
    ("Would use regularly if accessible", "65.8%", "68.3%", "56.1%", "63.3%"),
])

# ---------------- V ----------------
H1("V. Discussion")
H2("A. Frame-Sensitive Versus Dispositional Accounts")
P("The pattern is one-sided between the two rival readings. The decoupled "
  "game was accepted by the elderly cohort at levels consistent with what "
  "the gamification literature predicts for non-clinical technologies — "
  "85.4% in-session retention in the 65-80 band, voluntary return in ten "
  "of sixteen cases, majority stated intention — none of which the "
  "dispositional account predicted. The result is therefore consistent "
  "with gamification theory holding for this population, and with the "
  "failure documented in [2] residing in the coupling rather than in the "
  "game or the participants. The causal interpretation remains bounded: "
  "the cohorts differed in setting, recruitment, residential status, age "
  "distribution, health profile and prior exposure, and the decoupling "
  "was a package rather than a single variable. The decisive confirmation "
  "is a within-cohort randomised comparison, which is specified as future "
  "work.")
H2("B. The Age Gradient Read Through the WTA/WTP Account")
P("The age-stratified pattern is what an identity-level account predicts. "
  "The burden that the coupling imposes is age-linked — initial dropout "
  "climbs steeply with age even under decoupling — but under the "
  "decoupled condition that burden no longer converts into abandonment, "
  "because most older dropouts return. In WTA/WTP terms: with no clinical "
  "verdict attached, a slow round does not devalue an identity-level "
  "asset, the WTA for accepting failure stays low, and re-engagement "
  "clears. The account remains inferential — no WTA/WTP elicitation was "
  "performed in this deployment — and its direct measurement is specified "
  "for the confirmatory study.")
H2("C. The Commitment Gap and the Adaptive-Learning Direction")
P("Decoupling did not answer every engagement question. Half of the "
  "older participants who happily completed the session saw no place for "
  "the game in their weekly lives: the 29.3-point gap between retention "
  "and stated intention marks the boundary of what removing the threat "
  "can achieve. Decoupling removes a reason to leave; it does not add a "
  "reason to return. The design response this motivates is adaptive "
  "learning: presenting performance as movement on a personal trajectory, "
  "with challenge calibrated to evolving capability, so that a rising "
  "curve becomes the reason to return — and eventually becomes the frame "
  "within which the clinical purpose can be re-attached without "
  "re-activating the identity threat. The gap does not identify adaptive "
  "learning as its sole remedy; fatigue, access, burden, relevance and "
  "enjoyment are plausible contributors, and the confirmatory study "
  "pairs behavioural follow-up with measures of each.")

# ---------------- VI ----------------
H1("VI. Limitations")
P("The comparison with [2] is between-cycle and between-cohort, without "
  "randomisation; the decoupling is a multi-component package; the "
  "proposed mechanisms (endowment, mental accounting, stereotype threat) "
  "were not directly measured; regular-use intention is stated, not "
  "longitudinal behaviour; this paper reports the preliminary aggregate "
  "record, with participant-level analysis continuing beyond it; "
  "participants were volunteers in Chinese-cultural settings, and "
  "generalisation beyond these contexts is untested; and no clinical "
  "telemetry was collected, so no claim about clinical outcomes under "
  "decoupling is made.")

# ---------------- VII ----------------
H1("VII. Future Work: The Within-Cohort Confirmatory Study")
P("A three-arm randomised within-cohort study is pre-specified to convert "
  "the present associations into causal tests: Arm 1, the coupled game "
  "(control); Arm 2, the decoupled coordination-challenge framing; Arm 3, "
  "the coupled game with a structured coaching layer. The primary "
  "endpoint is re-engagement within seven days of first failure. The "
  "study adds instruments this deployment lacked: baseline FES-I, a "
  "validated motivational measure administered before any outcome is "
  "observed, WTA/WTP elicitation after the first session, and "
  "harmonised session-level logging with timestamps. Hypotheses and "
  "endpoints will be pre-registered before data collection begins.")

# ---------------- VIII ----------------
H1("VIII. Conclusion")
P("Offered the same game without the clinical attachment, a separate "
  "cohort of 120 adults — including 41 aged 65-80 — accepted it at the "
  "levels the gamification literature predicts, and most older "
  "participants who initially dropped out came back. The evidence is "
  "consistent with a frame-sensitive account of elderly abandonment: the "
  "problem documented in this programme lives at the attachment between "
  "gameplay and clinical verdict, not in gamification and not in the "
  "population. What decoupling leaves open — the commitment gap — points "
  "the engineering programme toward adaptive learning as the driver that "
  "can make returning worthwhile, and the pre-specified within-cohort "
  "study as the test that can make the causal claim this design cannot.")

H1("Acknowledgments")
P("The authors thank the Hong Kong community partners and participants, "
  "and the staff who supported the deployment. Supervision by Prof. Chi "
  "Fai Cheung at The Hong Kong Polytechnic University is gratefully "
  "acknowledged. AI tooling used in the programme's game development is "
  "an enabling technology; the contribution of this paper is the "
  "behavioural evaluation and the design framework it informs.")

H1("References")
refs = [
 "[1] S. K. T. Tam and C. F. Cheung, “A study of utilisation of Fall "
 "Risk Assessment among elderly users: A critical review through TAM and "
 "Octalysis frameworks with gamification approach,” accepted for "
 "publication (SS-293).",
 "[2] S. K. T. Tam and C. F. Cheung, “Gamification as proactive "
 "approach: A study of elderly failure-induced abandonment in Fall Risk "
 "Assessment,” manuscript under submission (Paper 2 of this "
 "programme).",
 "[3] R. Thaler, “Toward a positive theory of consumer choice,” "
 "J. Econ. Behav. Organ., vol. 1, no. 1, pp. 39-60, 1980.",
 "[4] D. Kahneman, J. L. Knetsch, and R. H. Thaler, “Experimental "
 "tests of the endowment effect and the Coase theorem,” J. Political "
 "Econ., vol. 98, no. 6, pp. 1325-1348, 1990.",
 "[5] S. An, C. F. Cheung, and K. W. Willoughby, “A gamification "
 "approach for enhancing older adults' technology adoption and knowledge "
 "transfer,” Technol. Forecast. Soc. Change, vol. 205, 2024.",
 "[6] K. Chen and A. H. S. Chan, “Gerontechnology acceptance by "
 "elderly Hong Kong Chinese: STAM,” Ergonomics, vol. 57, no. 5, "
 "pp. 635-652, 2014.",
 "[7] Y.-K. Chou, Actionable Gamification: Beyond Points, Badges, and "
 "Leaderboards. Fremont, CA: Octalysis Media, 2015.",
 "[8] G. Eysenbach, “The law of attrition,” J. Med. Internet "
 "Res., vol. 7, no. 1, e11, 2005.",
 "[9] S. J. Barber, “An examination of age-based stereotype threat "
 "about cognitive decline,” Perspect. Psychol. Sci., vol. 12, no. 1, "
 "pp. 62-90, 2017.",
 "[10] M. S. McGlone and J. Aronson, “Stereotype threat, identity "
 "salience, and spatial reasoning,” J. Appl. Dev. Psychol., vol. 27, "
 "no. 5, pp. 486-493, 2006.",
 "[11] D. Kahneman and A. Tversky, “Prospect theory: An analysis of "
 "decision under risk,” Econometrica, vol. 47, no. 2, pp. 263-291, "
 "1979.",
 "[12] A. Bandura, Self-Efficacy: The Exercise of Control. New York: "
 "W. H. Freeman, 1997.",
 "[13] L. Yardley et al., “Development and initial validation of the "
 "Falls Efficacy Scale-International (FES-I),” Age Ageing, vol. 34, "
 "no. 6, pp. 614-619, 2005.",
]
for r_ in refs:
    P(r_)

cp = d.core_properties
cp.title = "Paper 3 R1 - Stage-3 decoupling manuscript (aligned with thesis V20)"
cp.comments = ("R1 (19 Jul 2026): full rewrite of the Key Contexts scaffold "
               "as the manuscript of the conducted between-cohort decoupled "
               "deployment (N=120, register 38/41/41). Three-arm study is "
               "future work. All figures follow thesis V20.")

OUT = ("/tmp/claude-0/-home-user-gamification/3d8a50e8-9258-51ce-a7d6-"
       "194da50f9a04/scratchpad/Paper3_R1_Stage3_Manuscript.docx")
d.save(OUT)
print("saved", OUT)

# verification
d2 = docx.Document(OUT)
full = "\n".join(p.text for p in d2.paragraphs)
for tb in d2.tables:
    for r in tb.rows:
        full += "\n" + " | ".join(c.text for c in r.cells)
import re
words = len(re.findall(r"\S+", full))
print("words:", words)
for s in ("3.24", "31.5", "35/40/45", "35.6", "86.7", "51.1", "LODF",
          "confirmatory evidence", "Phase 2 data", "HR ="):
    c = full.count(s)
    if c:
        print(f"RESIDUAL '{s}': {c}")
for s in ("38", "29.3", "65.2", "62.5", "71.7", "56.1", "85.4", "39.0"):
    print(f"'{s}':", full.count(s))
