#!/usr/bin/env python3
"""Paper 2 — full rewrite, submission-ready, aligned with thesis V20 and the
accepted SS-293. Every statistic is a V20-verified value; withdrawn analyses
are disclosed once, in the verification section. Reference numbering follows
the original submitted list (1-45) plus [46] SS-293, and in-text citations
are written against that list consistently."""
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

d = docx.Document()
st = d.styles["Normal"]
st.font.name = "Times New Roman"
st.font.size = Pt(11)

H1 = lambda t: d.add_heading(t, level=1)
H2 = lambda t: d.add_heading(t, level=2)
P = lambda t: d.add_paragraph(t)


def table(rows):
    t = d.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri].cells[ci]
            c.paragraphs[0].add_run(str(val))
            for r in c.paragraphs[0].runs:
                r.font.size = Pt(9)
                if ri == 0:
                    r.bold = True
    return t


def caption(t):
    p = d.add_paragraph(t)
    for r in p.runs:
        r.font.size = Pt(9)
        r.italic = True


# ================= front matter =================
t = d.add_paragraph()
r = t.add_run("Gamification as a Proactive Approach: A Study of Elderly "
              "Failure-Induced Abandonment in Gamified Fall Risk Assessment")
r.bold = True
r.font.size = Pt(14)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
a = d.add_paragraph("Sidney K. T. Tam and Chi Fai Cheung\n"
                    "Behaviour and Knowledge Engineering Research Centre, "
                    "Department of Industrial and Systems Engineering,\n"
                    "The Hong Kong Polytechnic University, Hong Kong, China\n"
                    "Corresponding author: sidney.tam@connect.polyu.hk")
a.alignment = WD_ALIGN_PARAGRAPH.CENTER

H1("Abstract")
P("Fall Risk Assessment (FRA) systems can assess balance and support fall "
  "prevention as clinical devices, but rarely achieve sustained elderly "
  "user engagement. This study examines the factors associated with low "
  "engagement and rapid, permanent dropout in a gamified, AI-assisted FRA "
  "among elderly residents after their first adverse user experience. An "
  "action-research approach was employed at a residential care facility in "
  "Southern China. An AI-assisted skiing game was developed on the FRA "
  "system's weight-shifting foot pressure panel, allowing elderly "
  "individuals to undergo balance measurement while steering a skier on a "
  "display. Thirty residents (aged 62-84; mean 71.1 ± 5.9) participated "
  "over eight weeks. Seventy per cent (21 of 30; 95% CI 52.1-83.3%) "
  "disengaged at or before three game cycles, 26.7% (8 of 30; 95% CI "
  "14.2-44.4%) exited after a single session, and no disengaged "
  "participant re-engaged within the study window (0 of 21; 95% CI "
  "0.0-15.5%). Among the nine sustained participants the recorded FRA "
  "index improved by +6.9 points (95% CI 4.5-9.3; paired dz = 2.22, 95% "
  "CI 0.95-3.46); across all thirty, the paired change was +4.0 points "
  "(95% CI 0.5-7.6; dz = 0.42, 95% CI 0.05-0.79). Satisfaction was "
  "strongly associated with sustained engagement (9/9 versus 11/21; "
  "Fisher's exact p = 0.013). Most strikingly, three participants whose "
  "recorded index improved during their only session still never "
  "returned. The study did not seek engineering excellence in game "
  "design; its findings concern the behavioural context beneath the play. "
  "We carry two working hypotheses of failure response: Failure-Induced "
  "Dropout (FIDS), associated with self-efficacy collapse at first "
  "failure, and Failure-Induced Abandonment (FIAS), associated with loss "
  "aversion when negative feedback is read as evidence of declining "
  "function. Four falsifiable predictions (H1-H4) and a domain-reframing "
  "hypothesis (H5) are specified for a pre-registered three-arm "
  "confirmatory study.")
P("Keywords: fall risk assessment, gamification, elderly technology "
  "adoption, failure-induced abandonment, loss aversion, eHealth "
  "attrition, action research.")

# ================= I =================
H1("I. Introduction")
P("Falls are the leading cause of injury-related hospitalisation in "
  "adults aged 65 or older [1], [2]. FRA systems assess balance and "
  "postural control and recommend prevention strategies, yet elderly "
  "users struggle with ongoing participation in these technologies. A "
  "survey of 200 residents at the study site (of a resident population "
  "of 280) found that only 39% had ever used the installed FRA "
  "equipment; among those users, approximately 74% failed on their "
  "first attempt and 82% did not return [46]. The facility had invested "
  "in a clinically validated system that the great majority of its "
  "residents were not using.")
P("When older adults encounter a barrier during first use of a new "
  "technology, they respond differently from younger adults: they tend "
  "to give up rather than persist [3], [4], partly through low "
  "self-efficacy and a growing perception of age-related limitation "
  "[5], [6]. Eysenbach argued that abandonment is the most common "
  "outcome of digital health interventions and called for a 'science of "
  "attrition' [7]. The adoption frameworks in wide use — TAM and its "
  "extensions — concern what makes people start using a system and have "
  "little to say about what happens after a user fails [8], [9].")
P("Gamification has been proposed as one route to better engagement "
  "among older adults [10]. An, Cheung and Willoughby reported that a "
  "gamified learning system lifted perceived usefulness, ease of use and "
  "behavioural intention among older participants within a single "
  "session [11], while noting that their design could not show whether "
  "people return after experiencing failure, and calling for work on "
  "post-adoption obstacles. That call motivated the present study.")
H2("A. Industrial Problem")
P("The study was conducted in partnership with a large residential care "
  "provider in Southern China, at a facility of several hundred beds "
  "that acquired an InBody FRA system and found it largely unused. Two "
  "circumstances make the setting a strong test: staffing is sufficient "
  "for one-on-one instruction, and use is free with experienced "
  "assistance, so the standard practical barriers — cost, access, "
  "support — are largely absent. The research therefore focused not on "
  "engineering excellence in the FRA or novel game mechanisms, but on "
  "the behavioural context of a health-technology user following a "
  "failure incident with the technology.")
H2("B. Research Questions")
P("RQ1: Does gamified exercise enhance both clinical outcomes measured "
  "via FRA and user acceptance among frail elderly residents in "
  "residential care? RQ2: What explains rapid and permanent attrition "
  "despite a gamified deployment with high user acceptance and "
  "potential health benefit? RQ3: What psychological mechanisms lead to "
  "continued abandonment of FRA after the first adverse experience, and "
  "what principles of game design might counteract them? The objectives "
  "are correspondingly: to design and pilot-test an exercise gamified "
  "to double as a clinical FRA measurement (RQ1); to characterise the "
  "failure-response pattern (RQ2); and to propose a model of failure "
  "abandonment with falsifiable predictions for confirmatory testing "
  "(RQ3).")

# ================= II =================
H1("II. Literature Review")
H2("A. Elderly Technology Adoption")
P("Guner and Acarturk found that TAM remains useful in predicting which "
  "technologies older adults choose to use [3]. Wildenbos et al. "
  "organised ageing barriers into cognitive, physical, perceptual and "
  "motivational types, with motivational aspects the least addressed "
  "[4]. Hawthorn observed that many elderly people read their age as "
  "evidence against their ability to learn new technologies [5], and "
  "Tsai et al. found technology anxiety peaking at first use [6]. Lee "
  "and Coughlin noted that older adults respond better to training "
  "structured around explicit goals [14], and reviews report that "
  "reluctance toward health technologies is often driven by perceived "
  "threats to autonomy and independence [15].")
H2("B. Positioning Against the Science of Attrition")
P("Eysenbach's 'law of attrition' distinguished usage attrition from "
  "research dropout [7]; the users in this study fall into the former "
  "group. Kelders et al. found persuasive features and human "
  "involvement the strongest indicators of long-term adherence [12]; "
  "Christensen et al. noted that the psychological mechanisms beneath "
  "adherence remain poorly understood [13]. None of this literature "
  "explains why, specifically among frail older adults using FRA "
  "systems, abandonment after first failure is rapid, persistent and "
  "immune to redesign. FIAS is intended to fill this gap.")
H2("C. Fear of Falling")
P("Fear of falling denotes persistent fear that results in activity "
  "avoidance even when the person is physically capable [16], [18]. "
  "Yardley et al. developed the Falls Efficacy Scale-International "
  "(FES-I), capturing falls-related self-efficacy across sixteen "
  "everyday activities [17]. A participant who enters an FRA session "
  "already carrying high fear may interpret any difficulty as proof of "
  "decline rather than as a momentary game challenge. The confirmatory "
  "study will administer the FES-I at baseline to test whether fear of "
  "falling mediates the link between first failure and permanent "
  "abandonment (H1).")
H2("D. Gamification and Elderly Health Technology")
P("Gamification applies game-design elements in non-game contexts [10]; "
  "work with older adults has concentrated on physical activity and "
  "cognitive training [11], [19]. An et al. is the closest precedent: "
  "improvements in perceived usefulness, ease of use and behavioural "
  "intention among 115 participants aged 50-75, with gamification "
  "effectiveness the strongest predictor of ease of use [11]. Their "
  "design covered a single session and could not capture what happens "
  "when a person fails; the present eight-week longitudinal study "
  "responds directly to that limitation.")
H2("E. Theoretical Frameworks")
P("TAM predicts adoption from perceived usefulness and ease of use [8]; "
  "UTAUT adds social influence and facilitating conditions [9]; STAM "
  "extends these for older users with gerontechnology self-efficacy and "
  "anxiety [20]; and the Octalysis framework describes eight core "
  "drives of engagement with game designs [21]. None of these "
  "frameworks addresses the response to failure: all treat engagement "
  "as continuous or increasing under appropriate design.")
H2("F. Self-Efficacy, Prospect Theory, and the FIDS-FIAS Transition")
P("Bandura's work indicates that self-efficacy determines effort and "
  "persistence in the face of adversity [22], and computer "
  "self-efficacy shapes system use [23]. A frail older adult who fails "
  "an FRA task may experience a collapse of technology self-efficacy "
  "leading to disengagement — the mechanism denoted FIDS. Prospect "
  "Theory explains why re-engagement may not follow [24], [25]: people "
  "evaluate outcomes relative to a reference point, and losses loom "
  "roughly twice as large as commensurate gains. For an elderly person, "
  "the reference point is their sense of their own physical capability; "
  "failure delivers unwelcome news about it, and returning risks more. "
  "The psychological cost outweighs the perceived health benefit of "
  "replaying. This, the study hypothesises, constitutes the transition "
  "from FIDS to FIAS: users do not quit because they doubt they could "
  "improve — they quit because playing again might force them to face "
  "evidence that their body is failing. Four falsifiable predictions "
  "follow, to be tested in the confirmatory study: (H1) participants "
  "scoring higher on the FES-I will abandon more rapidly after first "
  "failure; (H2) participants will report the decision to quit as "
  "impulsive rather than deliberate; (H3) participants will describe "
  "their performance in health-related rather than game-related terms; "
  "(H4) technical ease of use alone will not prevent abandonment while "
  "the health-identity framing remains prominent.")
H2("G. Motivation Crowding and the FRA-Game Coupling")
P("Frey and Jegen's motivation-crowding effect names the erosion of "
  "intrinsic motivation by external evaluation [26]; Deci and Ryan's "
  "basic needs — autonomy, competence, relatedness — ground intrinsic "
  "motivation [27]. When a game is linked directly to health "
  "evaluation, play becomes assessment. For older adults already "
  "carrying low self-efficacy and high health anxiety, the coupling "
  "not only lowers interest but can produce an identity-driven negative "
  "reaction.")
H2("H. Behaviour-Change Lens: COM-B")
P("In COM-B terms [28], the adaptive difficulty engine manages "
  "capability and on-site deployment provides opportunity; the FIAS "
  "pattern lives in motivation — specifically reflective motivation, "
  "comprising beliefs about consequences and future identity. "
  "Enablement and persuasion are the relevant intervention functions, "
  "and the coaching layer proposed for the confirmatory study employs "
  "both: mastery-oriented framing for reflective motivation, and care "
  "staff for social opportunity.")
H2("I. Information Entropy as a Conceptual Lens")
P("Shannon entropy, H(p) = −Σ pᵢ log₂ pᵢ [29], is used here as a "
  "conceptual lens, not a quantitative model. At the start of a "
  "session an elderly participant typically holds an open belief state "
  "about what the outcome might mean (high entropy). A failure can "
  "collapse that state toward a single conviction — 'this technology "
  "is not for me' (low entropy) — after which positive feedback must "
  "travel through a channel already dominated by negative information, "
  "including negative accounts circulating among residents. No entropy "
  "quantities are computed in this study; the confirmatory study will "
  "elicit outcome probability distributions before and after sessions "
  "so that entropy over those distributions can serve as a measure of "
  "belief openness.")
H2("J. Research Gap")
P("No existing framework covers the full arc from initial adoption "
  "through failure to permanent abandonment among older users of "
  "health technology. TAM explains initial adoption; Octalysis "
  "explains within-session engagement; self-efficacy theory explains "
  "initial dropout; Prospect Theory explains non-return; the attrition "
  "literature describes abandonment without a mechanism for this "
  "population; the fear-of-falling literature supplies instruments not "
  "yet applied to gamified FRA; and COM-B maps design choices without "
  "the underlying theory. This study addresses that gap.")

# ================= III =================
H1("III. Original Contribution Statement")
P("This paper makes three original contributions. First, it documents "
  "empirical evidence of a pattern in which older users permanently "
  "abandon a gamified FRA after first failure: 70.0% low engagement "
  "(95% CI 52.1-83.3%), 26.7% single-session exit (95% CI 14.2-44.4%), "
  "and zero re-engagement among all 21 disengaged participants (95% CI "
  "0.0-15.5%) over eight weeks — including three participants whose "
  "recorded balance index improved in their only session. Second, it "
  "offers a mechanistic account linking self-efficacy collapse at "
  "first failure with subsequent permanent abandonment driven by loss "
  "aversion, with four falsifiable predictions (H1-H4). Third, it "
  "presents a framework for failure-aware coaching in health-"
  "assessment gamification, including an entropy-constraint concept, "
  "observable risk signals and an intervention specification congruent "
  "with COM-B. The Discussion develops an extended argument from the "
  "endowment-effect and stereotype-threat literatures, generating the "
  "domain-reframing hypothesis (H5). The extended claims go beyond the "
  "present evidence and are offered as directions for the "
  "pre-registered three-arm confirmatory study, not as findings.")

# ================= IV =================
H1("IV. Research Design, Ethics and Participants")
H2("A. Action-Research Methodology")
P("An action-research methodology was followed throughout [30], [31], "
  "chosen on the view that the environment cannot be understood or "
  "changed without sustained immersion. The observer-participant's "
  "presence in the care facility provides intimate familiarity with "
  "older adults' reality, alongside an attendant risk of observer bias; "
  "reflective field notes were kept to monitor evolving observations "
  "and limit personal bias, and the residual risk is stated in the "
  "limitations.")
H2("B. Ethics, Consent and Safeguarding")
P("Ethical approval was obtained from the relevant institutional "
  "ethics sub-committee, and the study was conducted in accordance "
  "with the Declaration of Helsinki. Participants received written "
  "information sheets in Mandarin and Cantonese covering the research "
  "objective, session routine, data protection and the unconditional "
  "right to refuse or withdraw. For participants unable to read, a "
  "non-involved caregiver relayed the information verbally; verbal "
  "consent was recorded and signed for by the caregiver witness. "
  "Trained personnel remained present throughout sessions, with "
  "non-slip flooring and grab rails in reach. Eligibility required "
  "being fall-free for the preceding two weeks and having no medical "
  "condition deemed inappropriate by care staff; cognitive impairment "
  "affecting comprehension was grounds for exclusion. Any fall, "
  "near-fall or dizziness would trigger immediate stoppage, referral "
  "and adverse-event recording; no adverse events occurred.")
H2("C. Data Governance and Protection")
P("Data collected comprised FRA assessment scores, play-session logs, "
  "satisfaction ratings and demographics. All data were pseudonymised "
  "at collection, with the code key stored separately and encrypted. "
  "Compliance with the Personal Information Protection Law (PIPL) "
  "[32] was maintained throughout. The AI difficulty-adaptation "
  "service received only non-identifying aggregate game parameters "
  "(obstacle density, speed, timing); no participant identities, "
  "demographics, FRA scores or other health information were "
  "transmitted to external servers, and all API requests were logged. "
  "Throughout the paper, the AI tooling is an enabling technology; "
  "the contribution is the behavioural analysis and design framework.")
H2("D. Study Site and Participants")
P("Thirty elderly residents of the facility participated without "
  "remuneration: ages 62-84 (mean 71.1, SD 5.9), 13 men and 17 women, "
  "all with fall-risk markers identified at initial assessment. Over "
  "eight weeks, each resident was exposed to both the standard FRA "
  "and the gamified FRA. This was a within-participant observational "
  "study of dropout patterns, not a randomised controlled trial.")

# ================= V =================
H1("V. Definitions and Operationalisation")
P("Table I defines the core terms with their operational criteria as "
  "recorded in the study dataset. FIDS and FIAS are operationalised "
  "as working hypotheses only. Because the archived dataset records "
  "cycle counts and session outcomes rather than timestamps, all "
  "codings are session-based: FIDS is coded as initial dropout at, or "
  "immediately following, the session containing the first failure "
  "event; FIAS additionally requires zero re-engagement through the "
  "study end. The descriptive composite of Section IX is not part of "
  "the evidential definition. The term 'syndrome' is deliberately "
  "avoided pending confirmatory evidence.")
caption("Table I. Key definitions and operationalisation.")
table([
    ("Term", "Definition", "Operationalisation"),
    ("Failure event", "A session in which the participant does not clear "
     "the minimum obstacle threshold from the competency baseline",
     "Binary: below baseline = 1"),
    ("Low engagement", "Three or fewer session cycles completed",
     "Cycle count ≤ 3"),
    ("Sustained engagement", "Four or more session cycles completed",
     "Cycle count ≥ 4"),
    ("Abandonment", "Permanent cessation with zero re-engagement through "
     "study end", "No completed session after disengagement"),
    ("Re-engagement", "At least one additional completed session after "
     "disengagement, in a later session block",
     "Later-session completion (session-based; no timestamps retained)"),
    ("Session cycle", "One complete game run from start to endpoint",
     "System-logged completion"),
    ("FIDS (hypothesis)", "Failure-Induced Dropout: initial dropout after "
     "self-efficacy collapse",
     "Initial dropout at/immediately after first-failure session"),
    ("FIAS (hypothesis)", "Failure-Induced Abandonment: permanent "
     "withdrawal driven by loss aversion when failure is read as health "
     "decline", "FIDS coding + zero re-engagement through study end"),
    ("Completers", "Participants completing ≥ 4 cycles across 8 weeks",
     "System-logged cycle count ≥ 4"),
])

# ================= VI =================
H1("VI. The Gamified Intervention")
H2("A. Dual-Purpose Input")
P("The system was engineered so that one physical action performs two "
  "functions. Standing on the FRA's pressure panel and tilting from "
  "side to side steers an on-screen skier while, in the background, "
  "clinical balance data are acquired from the same centre-of-pressure "
  "(COP) stream. Independent studies have validated the FRA platform "
  "against reference balance batteries [33] and established its "
  "test-retest reliability [34]; whether gameplay-derived COP metrics "
  "are clinically interchangeable with validated FRA outputs has not "
  "been established and is not assumed. The FRA alone presents as a "
  "clinical test; the game version casts the same measurement as play, "
  "so the meaning of difficulty depends on whether the activity is "
  "experienced as a medical procedure or as a game.")
H2("B. Game Mechanics")
P("Users navigate a downhill ski run filled with obstacles, steering "
  "by weight shift; each full run counts as one session cycle. A "
  "pre-session assessment set the initial difficulty low enough for "
  "the weakest-balance participant to complete safely, so that the "
  "first session could produce an experience of success rather than a "
  "de facto fall-risk test. Achievements followed Octalysis Drive 2 "
  "(Accomplishment) [21]. Staff were instructed to use phrases such "
  "as 'you are practising a new balance skill' in place of 'you are "
  "being tested for fall risk'.")
P("A leaderboard mapped to Octalysis Drive 5 (Social Influence) was "
  "trialled, and the strategy backfired. Staff observation and field "
  "notes recorded that participants who attended to the leaderboard "
  "compared themselves unfavourably with other users and judged "
  "themselves more quickly and more negatively; satisfaction appeared "
  "lower among those participants. Because the archived dataset does "
  "not record leaderboard viewing as a variable, this observation is "
  "reported qualitatively and no coefficient is claimed. The pattern "
  "is consistent with social comparison theory [35] and with An et "
  "al.'s conclusion that building confidence matters more than "
  "competition in this population [11]. The leaderboard is removed "
  "from the confirmatory design.")
H2("C. Five-Layer Architecture")
P("The system comprises five layers. Layer 1 (sensor interface): the "
  "pressure panel streams COP coordinates in real time, which are "
  "simultaneously logged for clinical review. Layer 2 (game engine): "
  "the skiing game consumes the COP stream as its control input; the "
  "engine was built with AI-assisted code generation from "
  "natural-language specifications, with design feedback from players "
  "and clinicians folded back into the specification. Layer 3 (AI "
  "adaptation): a language-model service, orchestrated through "
  "LangChain, reviews performance at five-minute intervals and "
  "adjusts obstacle density (one step up above 80% success, one step "
  "down below 40%) and speed banding. Layer 4 (logging): completed "
  "cycles, per-cycle timing, COP traces, success accuracy, per-cycle "
  "affect and symptoms. Layer 5 (clinical output): pre/post FRA "
  "indices and balance measures exported for therapist review. On "
  "API delay or loss, the game continues at the last-selected "
  "difficulty; all calls are logged with parameters and responses; "
  "the AI service has no access to participant identity or health "
  "status; staff can manually override at any time. The AI tooling "
  "here is an enabling technology — it lowered the cost of iterating "
  "engineering decisions that remain the developers' own.")
H2("D. Proposed Coaching Layer (Confirmatory Study)")
P("The present AI layer manages technical difficulty; failure's "
  "larger effect may be psychological. The confirmatory study "
  "therefore adds a coaching layer governing language and affective "
  "tone. Three real-time FIAS risk signals are specified: "
  "satisfaction dropping below neutral after the first cycle; no "
  "improvement across three consecutive cycles; and participants "
  "attributing difficulty to age or physical deterioration. In COM-B "
  "terms, coaching addresses reflective motivation through "
  "reframing, and staff activation adds social opportunity.")

# ================= VII =================
H1("VII. Data Collection and Measures")
P("Each session involved: baseline FRA index on the InBody machine; "
  "performance on the gamified task; a five-point satisfaction "
  "rating; and self-reported symptom change at follow-up. Five "
  "variables were collected per participant: V1 cycles completed; V2 "
  "pre/post FRA index; V3 symptom change; V4 satisfaction; V5 age "
  "and gender. Individual Octalysis and STAM questionnaires could "
  "not be administered as initially planned. Session-level temporal "
  "data (inter-session intervals, durations) were not captured. A "
  "higher FRA index signifies better balance; participants with "
  "increased indices between sessions reported improved well-being.")
H2("A. Data Verification and Withdrawal Statement")
P("All statistics in this paper were recalculated from the raw "
  "30-participant dataset, and only statistics that reproduce from "
  "that archive are reported. Verification steps: (1) the direction "
  "of FRA index change was checked against self-reported improvement "
  "to confirm scale directionality; (2) all means, standard "
  "deviations and effect sizes were recalculated from individual "
  "records — sustained participants (N = 9): +6.9 ± 3.1 points (95% "
  "CI 4.5-9.3; paired dz = 2.22, 95% CI 0.95-3.46); all participants "
  "(N = 30): +4.0 points (95% CI 0.5-7.6; dz = 0.42, 95% CI "
  "0.05-0.79); (3) the low-engagement rate was confirmed at 70.0% "
  "(95% CI 52.1-83.3%) and the single-session exit rate at 26.7% "
  "(95% CI 14.2-44.4%); (4) the satisfaction association with "
  "sustained engagement was confirmed (9/9 versus 11/21; Fisher's "
  "exact p = 0.013).")
P("Withdrawal statement. Three statistics reported in earlier drafts "
  "have been withdrawn and do not appear in this paper. A Cox "
  "proportional-hazards model and Kaplan-Meier analysis were "
  "withdrawn because the archived dataset preserves no time variable "
  "and cannot support time-to-event modelling. A week-2 "
  "motivational-orientation classification and its odds ratio were "
  "withdrawn because the classification variable does not appear in "
  "the archived dataset and its coding records could not be "
  "produced. Pooled and zero-imputation effect sizes were withdrawn "
  "because they did not reproduce from the archive. The confirmatory "
  "study replaces these with pre-registered instruments administered "
  "at baseline and session-level logging with timestamps.")

# ================= VIII =================
H1("VIII. Quantitative Findings")
H2("A. Statistical Methods")
P("Three methods are used, matched to a small single-site sample. "
  "Proportions carry Wilson score 95% confidence intervals. "
  "Within-participant change in the recorded FRA index is analysed "
  "with paired t and Wilcoxon tests, with exact noncentral-t "
  "confidence intervals for the paired effect size dz [36]. "
  "Associations between binary characteristics and sustained "
  "engagement use Fisher's exact test. No time-to-event modelling is "
  "used or reported (Section VII-A).")
H2("B. Clinical Outcomes")
caption("Table II. FRA index outcomes, sustained participants (N = 9).")
table([
    ("Measure", "Gamified FRA"),
    ("Baseline FRA index (mean ± SD)", "42.0 ± 8.3"),
    ("Post-intervention (mean ± SD)", "48.9 ± 9.4"),
    ("Change (mean ± SD; 95% CI)", "+6.9 ± 3.1 (4.5–9.3)"),
    ("Cohen's dz (paired; 95% CI)", "2.22 (0.95–3.46)"),
    ("Self-reported improvement", "7/9 (78%)"),
    ("Satisfaction (Satisfied/Very Satisfied)", "9/9 (100%)"),
])
P("Table II reports outcomes for the nine participants who sustained "
  "engagement (≥ 4 cycles). Across the full sample of thirty, the "
  "paired change was +4.0 points (95% CI 0.5-7.6; dz = 0.42, 95% CI "
  "0.05-0.79). A within-deployment change of this size may include "
  "measurement variation, familiarisation and practice effects, and "
  "is read as a promising signal rather than a demonstrated clinical "
  "effect.")
caption("Table IIA. FRA index outcomes, low-engagement group (cycles ≤ 3, N = 21).")
table([
    ("Measure", "Gamified FRA"),
    ("Baseline FRA index (mean ± SD)", "44.7 ± 12.0"),
    ("Post-intervention (mean ± SD)", "47.5 ± 11.5"),
    ("Change (mean ± SD)", "+2.8 ± 11.1"),
    ("Self-reported improvement", "11/21 (52%)"),
    ("Satisfaction (Satisfied/Very Satisfied)", "11/21 (52%)"),
])
P("Despite minimal engagement, the low-engagement group shows a mean "
  "change of +2.8 points with high variability. Only 52% reported "
  "satisfaction or perceived improvement, against 100% of sustained "
  "participants — satisfaction and sustained engagement are strongly "
  "associated (Fisher's exact p = 0.013).")
H2("C. The Abandonment Pattern")
caption("Table III. Engagement pattern (N = 30).")
table([
    ("Measure", "Value (95% CI)"),
    ("Low engagement (≤ 3 cycles)", "21/30, 70.0% (52.1–83.3%)"),
    ("Sustained (≥ 4 cycles)", "9/30, 30.0%"),
    ("Single-session exits (1 cycle)", "8/30, 26.7% (14.2–44.4%)"),
    ("Re-engagement after disengagement", "0/21, 0.0% (0.0–15.5%)"),
])
P("The cycle distribution was: eight participants completed one "
  "cycle, seven completed two, six reached three, one reached four, "
  "five reached five, and three reached six. The most telling figure "
  "is the zero re-engagement count: of the 21 participants who "
  "disengaged, not one returned within the eight-week window. This "
  "absolute permanence distinguishes the pattern from typical usage "
  "attrition, in which some users return [7]. The 26.7% single-"
  "session exit rate marks the first session as the highest-risk "
  "point. No time-to-event model is reported: the archive preserves "
  "cycle counts, not timestamps. The evidential core of the pattern "
  "is carried by the three directly observed quantities in Table III "
  "and by the satisfaction association above.")

# ================= IX =================
H1("IX. Qualitative Coding and the FIAS Composite")
P("All 30 participants were coded with thirteen binary attributes "
  "drawn from STAM constructs and selected Octalysis drives, "
  "collapsed onto a 0-10 composite. The composite is a descriptive "
  "case-identification device and carries no evidential weight: "
  "several attributes record engagement and exit behaviour, so a "
  "participant who abandoned early necessarily scores high, and the "
  "clustering of the F4 group at the top of the scale is partly true "
  "by construction. Every evidential statement in this paper rests "
  "on directly observed behaviour — single-session exit, non-return, "
  "and the divergence between recorded score direction and return.")
caption("Table IV. Coding scheme (abbreviated; 8 of 13 attributes).")
table([
    ("Attribute", "Coding rule", "Rationale"),
    ("STAM-PU", "Symptom improvement AND FRA index improved",
     "Perceived clinical usefulness"),
    ("STAM-PEOU", "Satisfaction ≥ Satisfied AND cycles ≥ 2",
     "System manageable"),
    ("STAM-BI", "Satisfaction ≥ Satisfied", "Continuation-intention proxy"),
    ("STAM-GSE", "Cycles ≥ 3 AND verified",
     "Self-efficacy through repeated success"),
    ("STAM-GA", "Satisfaction ≤ Dissatisfied OR (1 cycle AND no "
     "improvement)", "Anxiety signal"),
    ("D2 (Accomplishment)", "Cycles ≥ 4 AND improvement",
     "Progress achieved"),
    ("D5− (Neg. social)", "Cycles ≤ 2 AND no improvement AND "
     "satisfaction ≤ Neutral", "Social withdrawal"),
    ("D8 (Loss/avoidance)", "1 cycle AND no improvement",
     "Loss-aversion exit"),
])
caption("Table V. Composite score distribution (descriptive).")
table([
    ("Level", "n", "Score", "Age range"),
    ("F0 (none)", "4", "0", "63–69"),
    ("F1 (low)", "14", "1–2", "62–84"),
    ("F2 (moderate)", "3", "3–4", "67–73"),
    ("F3 (high)", "3", "5", "62–70"),
    ("F4 (strong)", "6", "9–10", "65–80"),
])
P("The empty interval between scores 5 and 9 is not read as evidence "
  "of a behavioural threshold: given the construction, it may be a "
  "mechanical effect of the scoring thresholds and outcome-linked "
  "items, or sampling variation in a cohort of thirty. What matters "
  "about the F4 group does not depend on the composite: three of its "
  "six members recorded a positive FRA index change in their single "
  "session and did not return. That divergence between recorded "
  "improvement and return behaviour challenges a purely clinical "
  "explanation of dropout and is the observation that most clearly "
  "motivates the identity-level account of the Discussion.")

# ================= X =================
H1("X. F4 Interview Protocol")
P("F4-cluster participants were interviewed within three to five days "
  "of the abandoned session about whether the decision to stop was "
  "immediate or deliberate. Table VI shows the nine-question "
  "protocol; questions 3 and 4 test H2 and H3, question 7 tests H4, "
  "and question 9 probes domain reframing.")
caption("Table VI. F4 interview protocol.")
table([
    ("Q", "Question", "Tests"),
    ("1", "Describe what happened in your last session", "Account"),
    ("2", "How did you feel when the session ended?", "Immediate anxiety"),
    ("3", "When did you decide not to come back?", "H2: impulsive vs deliberate"),
    ("4", "What were you thinking about your performance?", "H3: health identity"),
    ("5", "Did others' opinions matter?", "Social comparison"),
    ("6", "Did you connect difficulty to your age?", "Age attribution"),
    ("7", "If easier, would you have continued?", "H4: technical vs psychological"),
    ("8", "Anything that would make you try again?", "Permanence"),
    ("9", "If described as a reaction-speed game instead of a health "
     "test, would that change how you felt?", "Domain-reframing probe"),
])

# ================= XI =================
H1("XI. Discussion")
H2("A. RQ1: Did Gamification Help?")
P("For the subset who sustained engagement, yes: +6.9 FRA index "
  "points (95% CI 4.5-9.3; dz = 2.22), with all nine satisfied or "
  "very satisfied; across all thirty, +4.0 points (95% CI 0.5-7.6). "
  "These are promising within-deployment signals rather than "
  "demonstrated clinical effects.")
H2("B. RQ2: Why Does Abandonment Persist?")
P("The pattern has three defining features. It is early: 70.0% "
  "disengaged at or before three cycles, 26.7% after one session. It "
  "is absolute within the window: none of 21 disengaged participants "
  "returned (95% CI 0.0-15.5%), unlike ordinary usage attrition [7]. "
  "And it is not accounted for by dissatisfaction alone: although "
  "satisfaction and engagement are strongly associated (p = 0.013), "
  "three F4 participants improved on the recorded index in their "
  "only session and still never returned. Practical obstacles — "
  "usability, staffing, cost, scheduling — do not explain a pattern "
  "in which the technology was free, supervised, on site, and in "
  "documented cases measurably working for the person who walked "
  "away. The evidence is consistent with a psychological mechanism "
  "attached to the meaning of failure.")
H2("C. RQ3: The FIDS-FIAS Mechanism")
P("The proposed mechanism has two steps: FIDS names the initial "
  "dropout — first failure collapses task self-efficacy and the "
  "participant withdraws; FIAS names the consolidation — loss "
  "aversion makes re-engagement a gamble on further unwelcome "
  "evidence about one's own body, so the participant never returns. "
  "H1-H4 operationalise the account. Field observation during the "
  "deployment was consistent with it: participants who spoke about "
  "the activity in learning terms tended to sustain engagement, "
  "while those who spoke in health-deficit terms tended to exit "
  "early. Because the week-2 classification that formalised this "
  "observation was withdrawn (Section VII-A), the pattern is "
  "reported as a field observation and a hypothesis for the "
  "confirmatory study, not as a measured effect.")
H2("D. Extended Theory: Endowment, Stereotype Threat, Domain Reframing")
P("Three further perspectives are offered as theory for confirmatory "
  "testing, not as findings. The endowment effect [37], [38]: "
  "people value what they possess above what they might acquire, "
  "with willingness-to-accept (WTA) a loss typically exceeding "
  "willingness-to-pay (WTP) to obtain the same item. Applied here: "
  "an elderly participant arrives carrying an endowed, identity-"
  "level asset — their health self-concept. A clinically framed "
  "failure threatens to devalue it; the WTA for accepting that "
  "devaluation is extremely high, and the WTP for re-engagement — "
  "bounded by finite expected benefits — cannot match it. "
  "Identity-protective cognition research suggests people do treat "
  "core self-beliefs as possessions to defend. Extending endowment "
  "to an intangible self-concept is a theoretical step requiring "
  "validation. Stereotype threat [39], [40]: activating negative "
  "ageing stereotypes degrades older adults' performance, "
  "predominantly as self-concept threat [40]; reframing a task from "
  "trait measurement to skill activity reduces the threat [41]. "
  "Domain reframing draws the threads together: frame task "
  "performance as a trainable skill with variable outcomes, not a "
  "measurement of fixed capacity. With no fixed reference point, "
  "there is no endowed asset; with nothing to lose, the FIAS "
  "cascade has nothing to initiate. Prediction (H5): participants "
  "under an instant-response-test framing will show lower permanent "
  "dropout than under FRA framing, performing the identical "
  "physical task.")
P("Ethical note. If the game is framed as a response test while FRA "
  "data are captured in the background, participants must still be "
  "informed that balance patterns are recorded. The recommended "
  "consent wording leads with the response-test framing and "
  "discloses the dual purpose: 'this game measures how quickly you "
  "can react to obstacles — it also helps us understand your "
  "balance patterns.' Disclosure is preserved; the identity-"
  "threatening frame is not imposed.")

# ================= XII =================
H1("XII. Limitations")
P("The sample of 30 was small (aged 62-84) and drawn from a single "
  "residential facility in Southern China; findings apply to this "
  "relatively homogeneous cohort and require replication elsewhere. "
  "The archive preserves cycle counts and session outcomes but no "
  "timestamps, so the temporal texture of disengagement cannot be "
  "reconstructed and no time-to-event claims are made. Three "
  "statistics from earlier drafts (a Cox hazard ratio, a "
  "motivational-orientation odds ratio, and pooled/zero-imputation "
  "effect sizes) have been withdrawn for the reasons in Section "
  "VII-A — disclosed rather than silent. The FES-I was not "
  "administered. Prospect Theory is a plausible frame but was not "
  "directly tested, and extending the endowment effect to health "
  "self-concept goes beyond its documented domains; both remain "
  "propositions. Neither H5 nor the coaching layer was tested here. "
  "The investigator's dual role as observer and participant creates "
  "a bias risk that reflexive notes mitigate but do not eliminate.")

# ================= XIII =================
H1("XIII. Design Implications for Practitioners")
P("Four principles follow, three grounded in the present findings "
  "and one awaiting confirmation. First, set first-session "
  "difficulty low enough that every user succeeds; the first "
  "session's purpose is success, not measurement. Second, make "
  "participants comfortable before play — let them talk with "
  "residents who completed the process, and have staff accompany "
  "them in. Third, watch for the three FIAS risk signals during "
  "sessions (satisfaction below neutral after the first cycle; no "
  "improvement over three cycles; age-attribution talk) and "
  "intervene — adjust difficulty downward, reframe with "
  "mastery-oriented language ('your reaction time sped up') rather "
  "than performance language ('you missed again'). Fourth — the H5 "
  "hypothesis, pending confirmation — describe the session as a "
  "response challenge rather than a health assessment, avoid "
  "displaying any output the participant could read as a health "
  "verdict, and train staff in skill-development language.")

# ================= XIV =================
H1("XIV. The Entropy Constraint")
P("A game with no real-world stakes confines the meaning of failure "
  "to the game domain: one loses a life and tries again. When game "
  "performance and clinical assessment are coupled to the same "
  "physical action, a failure event reduces belief entropy in both "
  "domains at once — the game and the participant's beliefs about "
  "their own health. This is the entropy constraint: a structural "
  "property of any system coupling play and assessment. Beliefs "
  "about health and self are far more resistant to reopening than "
  "game beliefs, and if the endowment analysis above is correct, a "
  "newly settled negative conviction ('this technology is not for "
  "me') will itself be defended as a possession. Domain reframing "
  "aims to confine failure's meaning to the game domain in one "
  "step. This framework guides the confirmatory design; no "
  "quantitative entropy claim is made here.")

# ================= XV =================
H1("XV. Industrial Context and Confirmatory Design")
H2("A. Sponsor Requirements and Cost-Effectiveness")
P("The sponsor requires FRA usage above 39% of eligible users, "
  "retention of existing infrastructure, PIPL compliance, and no "
  "increase in fall risk. The gamified system runs on the existing "
  "InBody equipment at negligible added cost (AI-service calls "
  "below US$0.50 per session; staff training under half a day per "
  "person). Reducing the single-session exit rate from 26.7% "
  "toward 10% would substantially reduce cost per continuing user; "
  "the full cost-effectiveness analysis belongs to the confirmatory "
  "phase.")
H2("B. Three-Arm Confirmatory Design")
P("The confirmatory study marks the shift from exploratory action "
  "research to a pre-registered design. Arm 1 (control): the "
  "FRA-coupled game, framed explicitly as fall risk assessment, "
  "replicating the present conditions. Arm 2 (domain reframing): "
  "identical game, hardware and adaptation, framed as an instant "
  "response test with no health-assessment language; background FRA "
  "capture disclosed as in Section XI-D. Arm 3 (coaching): the "
  "FRA-coupled condition plus the coaching layer of Section VI-D. "
  "Arm 1 versus Arm 2 tests whether preventing the identity threat "
  "works (H5); Arm 1 versus Arm 3 tests whether repairing it after "
  "activation works (H1-H4). A pilot (n = 15-25) precedes a "
  "two-site study (n = 60-80) with stratified sampling by age band "
  "and baseline FES-I, randomised allocation, and baseline "
  "measures including cognitive status (MoCA), PHQ-4, technology "
  "familiarity, fall history and a validated motivational measure "
  "administered before any outcome is observed. The primary "
  "outcome is re-engagement attempt within seven days of first "
  "failure. Secondary outcomes include the FRA index, FES-I, "
  "satisfaction, three-month fall rate, EQ-5D-5L and staff time. "
  "Reporting follows STROBE [42] and SPIRIT-AI [43]; WHO and "
  "regulator guidance on AI in health informs the AI component "
  "[44], [45]. All hypotheses (H1-H5) and endpoints will be "
  "pre-registered before data collection.")
P("An interim between-cohort deployment of a decoupled version of "
  "the game with 120 community-dwelling adults, reported in a "
  "companion manuscript, provides preliminary quasi-experimental "
  "support for the decoupling direction — post-failure return where "
  "the present study observed none — while leaving the randomised "
  "within-cohort test to the design above.")

# ================= XVI =================
H1("XVI. Conclusion")
P("This study built and tested a gamified FRA using a dual-purpose "
  "input design at a residential care facility in Southern China. "
  "For participants who sustained engagement, gamification helped: "
  "+6.9 FRA index points (95% CI 4.5-9.3; dz = 2.22), with all "
  "nine satisfied or very satisfied. But 70.0% never passed three "
  "cycles, 26.7% stopped after their first session, no participant "
  "returned after disengaging — and three participants whose "
  "recorded index improved in their only session were among those "
  "who never came back. We propose that initial dropout (FIDS) is "
  "associated with self-efficacy collapse and that loss aversion "
  "(FIAS) consolidates the withdrawal when game failure is read as "
  "evidence of physical decline. Satisfaction was strongly "
  "associated with sustained engagement (p = 0.013), and field "
  "observation suggested that a learning-oriented reading of the "
  "activity accompanied persistence — a suggestion whose formal "
  "test belongs to the pre-registered three-arm study, alongside "
  "the domain-reframing hypothesis drawn from the endowment-effect "
  "[37], [38] and stereotype-threat [40], [41] literatures. The "
  "framing of a gamified health technology may matter as much as "
  "its engineering; the confirmatory programme is designed to "
  "find out.")

H1("Acknowledgments")
P("The authors thank the management and nursing staff, and the "
  "residents of the participating residential care facility in "
  "Southern China, for their generous cooperation. Supervision by "
  "Prof. Chi Fai Cheung at The Hong Kong Polytechnic University is "
  "gratefully acknowledged. AI tooling (Anthropic Claude Code and "
  "API services) was used as an enabling technology for game "
  "development and difficulty adaptation; the contribution of this "
  "paper is the behavioural analysis and the design framework. The "
  "authors declare no competing financial or non-financial "
  "interests; the lead author holds no financial relationship with "
  "Anthropic; the facility was not involved in data analysis or "
  "manuscript preparation; the study received no external funding.")

H1("References")
refs = [
 "[1] NICE, 'Falls in older people: assessing risk and prevention,' "
 "Clinical Guideline CG161, London, UK, 2013.",
 "[2] WHO, 'WHO global report on falls prevention in older age,' Geneva, 2007.",
 "[3] H. Guner and C. Acarturk, 'The use and acceptance of ICT by senior "
 "citizens: A comparison of TAM for elderly and young adults,' Universal "
 "Access Inf. Soc., vol. 19, pp. 311-330, 2020.",
 "[4] G. A. Wildenbos, L. Peute, and M. Jaspers, 'Aging barriers "
 "influencing mobile health usability for older adults: MOLD-US,' Int. J. "
 "Med. Inform., vol. 114, pp. 66-75, 2018.",
 "[5] D. Hawthorn, 'Interface design and engagement with older people,' "
 "Behav. Inf. Technol., vol. 26, no. 4, pp. 333-341, 2007.",
 "[6] H. S. Tsai, R. Shillair, and S. R. Cotten, 'Exploring the role of "
 "technology anxiety and self-efficacy in older adults' information "
 "seeking,' in Proc. iConference, 2020.",
 "[7] G. Eysenbach, 'The law of attrition,' J. Med. Internet Res., "
 "vol. 7, no. 1, e11, 2005.",
 "[8] F. D. Davis, 'Perceived usefulness, perceived ease of use, and "
 "user acceptance of information technology,' MIS Quarterly, vol. 13, "
 "no. 3, pp. 319-340, 1989.",
 "[9] V. Venkatesh et al., 'User acceptance of information technology: "
 "Toward a unified view,' MIS Quarterly, vol. 27, no. 3, pp. 425-478, 2003.",
 "[10] S. Deterding et al., 'From game design elements to gamefulness: "
 "Defining gamification,' in Proc. 15th Int. Acad. MindTrek Conf., 2011.",
 "[11] S. An, C. F. Cheung, and K. W. Willoughby, 'A gamification "
 "approach for enhancing older adults' technology adoption and knowledge "
 "transfer,' Technol. Forecast. Soc. Change, vol. 205, 2024.",
 "[12] S. M. Kelders et al., 'Persuasive system design does matter: A "
 "systematic review of adherence to web-based interventions,' J. Med. "
 "Internet Res., vol. 14, no. 6, e152, 2012.",
 "[13] H. Christensen, K. M. Griffiths, and L. Farrer, 'Adherence in "
 "internet interventions for anxiety and depression,' J. Med. Internet "
 "Res., vol. 11, no. 2, e13, 2009.",
 "[14] C. Lee and J. F. Coughlin, 'Older adults' adoption of technology: "
 "An integrated approach,' Geriatr. Gerontol. Int., vol. 15, no. 5, "
 "pp. 519-525, 2015.",
 "[15] G. Grossi, R. Ferrario, and V. Tosi, 'Barriers and facilitators "
 "to health technology adoption by older adults with chronic diseases,' "
 "BMC Geriatrics, vol. 24, no. 189, 2024.",
 "[16] M. E. Tinetti and L. Powell, 'Fear of falling and low "
 "self-efficacy: A cause of dependence in elderly persons,' J. "
 "Gerontol., vol. 48, pp. 35-38, 1993.",
 "[17] L. Yardley et al., 'Development and initial validation of the "
 "Falls Efficacy Scale-International (FES-I),' Age Ageing, vol. 34, "
 "no. 6, pp. 614-619, 2005.",
 "[18] A. C. Scheffer et al., 'Fear of falling: Measurement strategy, "
 "prevalence, risk factors and consequences among older persons,' Age "
 "Ageing, vol. 37, no. 1, pp. 19-24, 2008.",
 "[19] J. Koivisto and A. Malik, 'Gamification for older adults: A "
 "systematic literature review,' Gerontologist, vol. 61, no. 7, "
 "pp. 360-372, 2021.",
 "[20] K. Chen and A. H. S. Chan, 'Gerontechnology acceptance by "
 "elderly Hong Kong Chinese: STAM,' Ergonomics, vol. 57, no. 5, "
 "pp. 635-652, 2014.",
 "[21] Y.-K. Chou, Actionable Gamification: Beyond Points, Badges, and "
 "Leaderboards. Fremont, CA: Octalysis Media, 2015.",
 "[22] A. Bandura, Self-Efficacy: The Exercise of Control. New York: "
 "W. H. Freeman, 1997.",
 "[23] D. R. Compeau and C. A. Higgins, 'Computer self-efficacy: "
 "Development of a measure and initial test,' MIS Quarterly, vol. 19, "
 "no. 2, pp. 189-211, 1995.",
 "[24] D. Kahneman and A. Tversky, 'Prospect theory: An analysis of "
 "decision under risk,' Econometrica, vol. 47, no. 2, pp. 263-291, 1979.",
 "[25] A. Tversky and D. Kahneman, 'Advances in prospect theory: "
 "Cumulative representation of uncertainty,' J. Risk Uncertain., "
 "vol. 5, no. 4, pp. 297-323, 1992.",
 "[26] B. S. Frey and R. Jegen, 'Motivation crowding theory,' J. Econ. "
 "Surv., vol. 15, no. 5, pp. 589-611, 2001.",
 "[27] E. L. Deci and R. M. Ryan, Intrinsic Motivation and "
 "Self-Determination in Human Behavior. New York: Plenum, 1985.",
 "[28] S. Michie, M. M. van Stralen, and R. West, 'The behaviour "
 "change wheel,' Implement. Sci., vol. 6, no. 42, 2011.",
 "[29] C. E. Shannon, 'A mathematical theory of communication,' Bell "
 "Syst. Tech. J., vol. 27, pp. 379-423, 623-656, 1948.",
 "[30] E. T. Stringer, Action Research, 4th ed. Thousand Oaks, CA: "
 "SAGE, 2014.",
 "[31] K. Lewin, 'Action research and minority problems,' J. Soc. "
 "Issues, vol. 2, no. 4, pp. 34-46, 1946.",
 "[32] Standing Committee NPC, 'Personal Information Protection Law of "
 "the People's Republic of China,' Beijing, effective 1 November 2021.",
 "[33] S. H. Kim et al., 'Introduction of Fall Risk Assessment system "
 "and cross-sectional validation among community-dwelling older "
 "adults,' Int. J. Environ. Res. Public Health, vol. 16, no. 5, "
 "p. 740, 2019.",
 "[34] S. H. Kim et al., 'Test-retest reliability and sensitivity to "
 "change of a new Fall Risk Assessment system,' Int. J. Environ. Res. "
 "Public Health, vol. 17, no. 14, p. 5082, 2020.",
 "[35] L. Festinger, 'A theory of social comparison processes,' Hum. "
 "Relat., vol. 7, no. 2, pp. 117-140, 1954.",
 "[36] J. Cohen, Statistical Power Analysis for the Behavioral "
 "Sciences, 2nd ed. Hillsdale, NJ: Erlbaum, 1988.",
 "[37] R. Thaler, 'Toward a positive theory of consumer choice,' J. "
 "Econ. Behav. Organ., vol. 1, no. 1, pp. 39-60, 1980.",
 "[38] D. Kahneman, J. L. Knetsch, and R. H. Thaler, 'Experimental "
 "tests of the endowment effect and the Coase theorem,' J. Political "
 "Econ., vol. 98, no. 6, pp. 1325-1348, 1990.",
 "[39] R. A. Dionigi, 'Stereotypes of aging: Their effects on the "
 "health of older adults,' J. Geriatr., vol. 2015, art. 954027, 2015.",
 "[40] S. J. Barber, 'An examination of age-based stereotype threat "
 "about cognitive decline,' Perspect. Psychol. Sci., vol. 12, no. 1, "
 "pp. 62-90, 2017.",
 "[41] M. S. McGlone and J. Aronson, 'Stereotype threat, identity "
 "salience, and spatial reasoning,' J. Appl. Dev. Psychol., vol. 27, "
 "no. 5, pp. 486-493, 2006.",
 "[42] E. von Elm et al., 'STROBE statement: Guidelines for reporting "
 "observational studies,' Ann. Intern. Med., vol. 147, no. 8, "
 "pp. 573-577, 2007.",
 "[43] X. Liu et al., 'Reporting guidelines for clinical trials "
 "evaluating artificial intelligence interventions (SPIRIT-AI/"
 "CONSORT-AI),' Nature Med., vol. 26, pp. 1364-1374, 2020.",
 "[44] WHO, 'Ethics and governance of artificial intelligence for "
 "health,' Geneva, 2021.",
 "[45] FDA, Health Canada, and MHRA, 'Good Machine Learning Practice "
 "for Medical Device Development: Guiding Principles,' 2021.",
 "[46] S. K. T. Tam and C. F. Cheung, 'A study of utilisation of Fall "
 "Risk Assessment among elderly users: A critical review through TAM "
 "and Octalysis frameworks with gamification approach,' accepted for "
 "publication (SS-293).",
]
for r_ in refs:
    P(r_)

cp = d.core_properties
cp.title = "Paper 2 R2 - full rewrite aligned with thesis V20"
cp.comments = ("R2 (19 Jul 2026): complete rewrite for submission. V20 "
               "statistics only; withdrawn analyses disclosed once; "
               "composite descriptive-only; session-based definitions; "
               "renumbered references incl. SS-293 as [46].")

OUT = ("/tmp/claude-0/-home-user-gamification/3d8a50e8-9258-51ce-a7d6-"
       "194da50f9a04/scratchpad/Paper2_R2_Full_Rewrite.docx")
d.save(OUT)
print("saved", OUT)

d2 = docx.Document(OUT)
full = "\n".join(p.text for p in d2.paragraphs)
for tb in d2.tables:
    for r in tb.rows:
        full += "\n" + " | ".join(c.text for c in r.cells)
import re
print("words:", len(re.findall(r"\S+", full)))
for s in ("3.24", "31.5", "0.78", "0.57", "Kaplan", "hazard", "48 hours",
          "Shanghai", "6.75", "87.5", "18.2", "Ben-Ami", "LODF", "61% of"):
    c = full.count(s)
    if c:
        # allowed only inside the withdrawal statement/limitations
        lines = [l for l in full.split("\n") if s in l]
        flagged = [l for l in lines if "withdraw" not in l.lower()]
        if flagged:
            print(f"RESIDUAL '{s}':", flagged[0][:120])
        else:
            print(f"'{s}' only in withdrawal disclosure - OK")
