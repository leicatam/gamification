#!/usr/bin/env python3
"""Paper 3 R3 — aligned with thesis V24: Stage 3 described as the actual
remote deployment (HTML build distributed by e-mail; unsupervised play on
participants' own computers; trackpad/arrow-keys/foot-pad input; returns
compiled from participants' reply e-mails). Corrects the R2 claims that
the pressure-panel apparatus and session setting were retained and that
participant-level records sit in a physical study archive."""
import shutil
import docx

BASE = "/home/user/gamification/docs/papers/"
SRC = BASE + "Paper3_R2_Full_Rewrite.docx"
DST = BASE + "Paper3_R3_Aligned_V24.docx"
shutil.copy(SRC, DST)
d = docx.Document(DST)
log, fail = [], []


def set_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in list(p.runs[1:]):
            r._element.getparent().remove(r._element)
    else:
        p.add_run(text)


def replace(tag, needle, new_text):
    for p in d.paragraphs:
        if needle in p.text:
            set_text(p, new_text)
            log.append(tag)
            return
    fail.append(tag)


# (1) Abstract
replace("abstract",
    "the same broad apparatus, skiing mechanic, five-attempt rule and session length were retained",
    "Earlier work in this programme documented rapid, permanent abandonment "
    "of a gamified, AI-assisted Fall Risk Assessment (FRA) by elderly "
    "residents after first failure: 70.0% disengaged at or before three "
    "game cycles and none returned, even where the recorded balance index "
    "had improved. Two rival readings of that result make opposite "
    "predictions — a dispositional reading, on which gamification does not "
    "transfer to frail older adults, and a coupling reading, on which the "
    "attachment of the game to a clinical verdict on the user's body "
    "carried the failure. This paper reports the evaluation that separates "
    "them: a between-cohort quasi-experimental remote deployment in which "
    "the game was ported to an HTML build and distributed by e-mail. "
    "Participants played unsupervised on their own computers, using a "
    "trackpad, arrow keys or, where available, a foot-pad sensor; the "
    "skiing mechanic, five-attempt rule and session structure were "
    "retained in the port, while the participant-facing clinical "
    "attachment was removed as a package — the clinical apparatus and "
    "setting, recruitment language, pre-session clinical scan, clinical "
    "labels and participant-facing clinical outputs — and the activity was "
    "framed as an eye-leg coordination challenge. One hundred and twenty "
    "community-dwelling adults aged 25-80 were recruited through the "
    "authors' Hong Kong-based networks across three age bands (38 aged "
    "25-44, 41 aged 45-64, 41 aged 65-80). In the preliminary aggregate "
    "record, compiled from the results files and questionnaire responses "
    "each participant returned by e-mail, retention at the end of the "
    "study window was 93.3% overall and 85.4% in the 65-80 band; 15 of 23 "
    "initial dropouts returned voluntarily (65.2%; 10 of 16, or 62.5%, in "
    "the 65-80 band); coordination-challenge endorsement was 71.7%; and "
    "stated regular-use intention was 63.3% overall and 56.1% in the "
    "65-80 band. Initial dropout climbed steeply with age (0 of 38 versus "
    "16 of 41; Fisher's exact p < 0.0001), while observed non-return "
    "converged to low levels in every band (0.0%, 4.9%, 14.6%). The "
    "pattern is consistent with the coupling account and inconsistent "
    "with the dispositional account, within the limits of a "
    "between-cohort, multi-component design in which platform, input "
    "device, setting and supervision changed alongside framing. A "
    "residual 29.3-percentage-point gap between session retention and "
    "regular-use intention in the oldest band motivates adaptive "
    "learning — trajectory-oriented feedback with calibrated challenge — "
    "as the design response, and a pre-specified within-cohort randomised "
    "study as the confirmatory test.")

# (2) III.B decoupling package
replace("III.B-package",
    "The broad pressure-panel apparatus, skiing mechanic, five-attempt rule and session length were retained.",
    "Relative to the coupled deployment of [2], the participant-facing "
    "clinical attachment was removed as a package, and the removal was "
    "implemented by re-platforming: the game was ported to an HTML build "
    "and sent to participants by e-mail, so that play occurred away from "
    "the clinical station altogether. The package comprised recruitment "
    "language (a coordination challenge, with no medical or "
    "fall-prevention vocabulary); the absence of any pre-session clinical "
    "scan; the absence of clinical labels, and of any clinical station or "
    "setting; and the absence of participant-facing clinical outputs, "
    "with no balance score or risk index shown at any point. The skiing "
    "mechanic, five-attempt rule and session structure were retained in "
    "the ported build; the pressure-panel apparatus and in-person "
    "supervision were not — participants played unsupervised on their own "
    "computers using a trackpad, arrow keys or, where available, a "
    "foot-pad sensor. Participant-level clinical balance telemetry was "
    "not collected. Because cohort, setting, platform, input device, "
    "supervision and recruitment channel also differed, the design "
    "estimates the combined association of the decoupled condition and "
    "its context with behaviour.")

# (3) III.C participants
replace("III.C-participants",
    "recruited through Hong Kong community channels: 38 aged 25-44",
    "One hundred and twenty community-dwelling adults aged 25-80 were "
    "recruited through the authors' Hong Kong-based networks — work "
    "contacts, a charity community network, university support channels "
    "and working peers in medical centres: 38 aged 25-44, 41 aged 45-64 "
    "and 41 aged 65-80; 48 men and 72 women; all reported high-school "
    "education or above. Recruitment materials described a coordination "
    "challenge and used no medical or fall-prevention vocabulary. "
    "Recruiting three age bands builds the comparison group into the "
    "study and licenses a second null hypothesis, H0-age: that under the "
    "decoupled condition there is no difference between age groups. Data "
    "collection is complete: each participant returned, by e-mail, the "
    "results file logged by the HTML build (including the session score "
    "and game-logged data) together with their questionnaire responses. "
    "This paper reports the preliminary aggregate record compiled from "
    "those returns; participant-level compilation continues beyond it.")

# (4) III.D outcomes
replace("III.D-outcomes",
    "Outcomes are behavioural and stated: initial dropout",
    "Outcomes are behavioural and stated, observed from the returned game "
    "logs and questionnaire responses: initial dropout (exit before "
    "completing the five-attempt protocol), voluntary return after "
    "initial dropout, observed non-return at the end of the study window, "
    "coordination-challenge endorsement (a manipulation check on the "
    "intended frame), and stated regular-use intention. Proportions carry "
    "Wilson score 95% confidence intervals; band contrasts use Fisher's "
    "exact and chi-square tests. No time-to-event modelling is used.")

# (5) V.A discussion — widen the confound list
replace("V.A-discussion",
    "the cohorts differed in setting, recruitment, residential status, age distribution, health profile and prior exposure",
    "The pattern is one-sided between the two rival readings. The "
    "decoupled game was accepted by the elderly cohort at levels "
    "consistent with what the gamification literature predicts for "
    "non-clinical technologies — 85.4% in-session retention in the 65-80 "
    "band, voluntary return in ten of sixteen cases, majority stated "
    "intention — none of which the dispositional account predicted. The "
    "result is therefore consistent with gamification theory holding for "
    "this population, and with the failure documented in [2] residing in "
    "the coupling rather than in the game or the participants. The causal "
    "interpretation remains bounded: the cohorts differed in setting, "
    "platform, input device, supervision, recruitment, residential "
    "status, age distribution, health profile and prior exposure, and the "
    "decoupling was a package rather than a single variable — indeed, the "
    "decoupled condition was delivered by re-platforming the game "
    "entirely. The decisive confirmation is a within-cohort randomised "
    "comparison, which is specified as future work.")

# (6) VII limitations
replace("VII-limitations",
    "The comparison with [2] is between-cycle and between-cohort",
    "The comparison with [2] is between-cycle and between-cohort, without "
    "randomisation; the decoupling is a multi-component package that "
    "includes a change of platform, input device, setting and supervision "
    "(an e-mailed HTML build played unsupervised on participants' own "
    "computers, rather than the supervised pressure-panel station of "
    "[2]); the proposed mechanisms (endowment, mental accounting, "
    "stereotype threat) were not directly measured; regular-use intention "
    "is stated, not longitudinal behaviour; this paper reports the "
    "preliminary aggregate record compiled from participants' e-mailed "
    "returns, with participant-level compilation continuing beyond it; "
    "input device varied by participant and was not centrally "
    "standardised; remote unsupervised play means session conditions were "
    "not controlled; participants were volunteers in Chinese-cultural "
    "settings, and generalisation beyond these contexts is untested; and "
    "no clinical telemetry was collected, so no claim about clinical "
    "outcomes under decoupling is made.")

# (7) IX conclusion
replace("IX-conclusion",
    "Offered the same game without the clinical attachment",
    "Offered the same game mechanics without the clinical attachment — an "
    "HTML port of the game, played on their own computers — a separate "
    "cohort of 120 adults, including 41 aged 65-80, accepted it at the "
    "levels the gamification literature predicts, and most older "
    "participants who initially dropped out came back. The evidence is "
    "consistent with a frame-sensitive account of elderly abandonment: "
    "the problem documented in this programme lives at the attachment "
    "between gameplay and clinical verdict, not in gamification and not "
    "in the population. What decoupling leaves open — the commitment "
    "gap — points the engineering programme toward adaptive learning as "
    "the driver that can make returning worthwhile, and the pre-specified "
    "within-cohort study as the test that can make the causal claim this "
    "design cannot.")

# (8) Acknowledgments
replace("acknowledgments",
    "the staff who supported the deployment",
    "The authors thank the Hong Kong network partners — work contacts, "
    "the charity community network, university supporters and "
    "medical-centre colleagues — who distributed the game build, and all "
    "participants who returned their game logs and questionnaire "
    "responses. Supervision by Prof. Chi Fai Cheung at The Hong Kong "
    "Polytechnic University is gratefully acknowledged. AI tooling used "
    "in the programme's game development is an enabling technology; the "
    "contribution of this paper is the behavioural evaluation and the "
    "design framework it informs.")

cp = d.core_properties
cp.comments = "R3: aligned with thesis V24 — remote e-mailed HTML deployment; archive claims corrected."
d.save(DST)
print("LOG:", *log, sep="\n  ")
print("FAIL:", fail if fail else "none")
