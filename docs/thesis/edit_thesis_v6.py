#!/usr/bin/env python3
"""V5 -> V6: rebuild the Stage-2 statistical analysis from the archived raw
dataset (FRA_Raw_Data_N30.xlsx), remove the unsupported survival framing,
install a QFD-structured analysis pipeline, and flag the Stage-3 logbook gap.

All statistics written below were computed from the author's raw data file;
nothing is estimated or invented. Items requiring the author's decision
(FRA index direction convention; motivational-orientation records; Stage-3
logbook completion) are highlighted placeholders.
"""
import docx
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Inches
import copy

F = "/tmp/claude-0/-home-user-gamification/3d8a50e8-9258-51ce-a7d6-194da50f9a04/scratchpad/Tam_V5_EngD_Thesis_Final_Revision_Draft.docx"
OUT = "/tmp/claude-0/-home-user-gamification/3d8a50e8-9258-51ce-a7d6-194da50f9a04/scratchpad/Tam_V6_EngD_Thesis_Final_Revision_Draft.docx"
FIG54 = "/tmp/claude-0/-home-user-gamification/3d8a50e8-9258-51ce-a7d6-194da50f9a04/scratchpad/figure_5_4_prepost.png"

doc = docx.Document(F)
log, failures = [], []

def find(needle, nth=1):
    n = 0
    for p in doc.paragraphs:
        if needle in p.text:
            n += 1
            if n == nth:
                return p
    return None

def clear_runs(p):
    for r in list(p.runs):
        r._element.getparent().remove(r._element)

def set_parts(p, parts):
    clear_runs(p)
    for text, kind in parts:
        r = p.add_run(text)
        if kind == 'input':
            r.bold = True; r.font.highlight_color = WD_COLOR_INDEX.YELLOW
        elif kind == 'photo':
            r.bold = True; r.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN

def insert_after(p, parts, style=None):
    new = OxmlElement('w:p'); p._p.addnext(new)
    np_ = Paragraph(new, p._parent)
    if style is not None: np_.style = style
    set_parts(np_, parts)
    return np_

def delete(p):
    p._element.getparent().remove(p._element)

FIGSTYLE = H2STYLE = None
for p in doc.paragraphs:
    if FIGSTYLE is None and p.style.name == "FigureCaption": FIGSTYLE = p.style
    if H2STYLE is None and p.style.name == "Heading 2": H2STYLE = p.style

# ---------------------------------------------------------------- Abstract
p = find("Stage 1 used a face-to-face survey with approximately 200 residents")
if p:
    t = p.text
    old = "A Cox proportional-hazards model associated first-failure experience with faster observed dropout (HR = 3.24; the interval estimate is reported with the full model specification in §5.5)."
    new = ("Within the archived Stage-2 dataset, twenty-one of thirty participants (70.0%) disengaged at or before three game cycles, eight after a single session, and no re-engagement was observed within the study window. Among the nine sustained participants the recorded FRA index changed by +6.9 points on average (paired dz = 2.22, 95% CI 0.95–3.46); across all thirty participants the paired change was +4.0 points (dz = 0.42, 95% CI 0.05–0.79). Because the archived dataset contains no session timestamps, no time-to-event model is fitted; earlier survival-analysis framings have been withdrawn. A Quality Function Deployment (QFD) pipeline translates the empirical findings into prioritised game-design requirements.")
    if old in t:
        set_parts(p, [(t.replace(old, new), '')])
        log.append("Abstract rewritten with archived-data statistics")
    else:
        failures.append("Abstract old sentence not found")
else:
    failures.append("Abstract paragraph not found")

# ---------------------------------------------------------------- §3.2 QFD mention
p = find("Stage 2 was a prospective eight-week action-research deployment with thirty residents")
if p:
    set_parts(p, [(p.text + " Across all three stages, a Quality Function Deployment (QFD) pipeline (Akao, 1990; Chan & Wu, 2002) is used to translate empirical findings into prioritised design requirements: user needs evidenced by the survey, telemetry and self-report data are mapped onto game features and weighted by the strength of the supporting evidence (§5.5A). QFD structures the translation from evidence to design; it complements, and does not replace, the statistical analysis of the collected data.", '')])
    log.append("3.2 QFD pipeline sentence added")
else:
    failures.append("3.2 Stage-2 paragraph not found")

# ---------------------------------------------------------------- §5.1 H1
p = find("H1: First-failure experience is associated with a higher hazard")
if p:
    set_parts(p, [("H1: An adverse or difficult first experience is associated with early disengagement (three or fewer completed game cycles) during the eight-week study window.", '')])
    log.append("H1 rewritten without hazard language")
else:
    failures.append("H1 not found")

# ---------------------------------------------------------------- §5.5 core rebuild
# (1) replace the engagement/Cox paragraph
p = find("A Cox proportional-hazards model associated first-failure experience with faster observed dropout (HR = 3.24; the interval estimate and p-value are being re-derived")
if p:
    set_parts(p, [
        ("The Stage-2 archive (FRA_Raw_Data_N30.xlsx; Appendix G, item G.2) contains one record per participant: gender, age, the recorded FRA index before and after the deployment window, presenting fall-risk symptoms, the number of completed game cycles, self-reported symptom change, satisfaction and a verification flag. Thirty-two records were collected in total; two younger control participants (aged 36 and 53) are reported separately, leaving the thirty elderly participants (aged 62–84, mean 71.1 ± 5.9; 17 female, 13 male) analysed here. One record failed verification and is retained with a sensitivity note. Because the archive records cycle counts and pre/post index values but no session dates or timestamps, time-to-event methods are not applicable to it: a Cox proportional-hazards model and a Kaplan-Meier estimator both require an observed time variable that this dataset does not contain. The survival-analysis framing used in earlier drafts, including the previously reported hazard ratio, has therefore been withdrawn, and the analysis below uses estimators suited to a small, single-site sample: Wilson score intervals for proportions, paired t and Wilcoxon signed-rank tests for the pre/post index, Welch and Mann-Whitney tests for between-group contrasts, and Fisher's exact test for categorical associations.", ''),
    ])
    log.append("5.5 dataset+methods paragraph installed")
else:
    failures.append("5.5 Cox paragraph not found")

# (2) delete the Cox re-verification placeholder
p = find("[AUTHOR INPUT REQUIRED — Cox model re-verification")
if p:
    delete(p); log.append("Cox re-verification placeholder removed (superseded)")
else:
    failures.append("Cox re-verification placeholder not found")

# (3) engagement paragraph — insert after the methods paragraph (before Table 5.2 caption)
p = find("The Stage-2 archive (FRA_Raw_Data_N30.xlsx; Appendix G, item G.2) contains one record")
if p:
    insert_after(p, [
        ("Engagement showed an early drop-off. Twenty-one of thirty participants (70.0%, 95% CI 52.1–83.3%) completed three or fewer cycles, eight (26.7%, 95% CI 14.2–44.4%) ended after a single session, and no re-engagement was observed among the low-engagement group during the eight-week window (0 of 21; 95% CI 0.0–15.5%). Neither gender (Fisher's exact p = 0.69) nor age (Spearman ρ = −0.18, p = 0.34) was reliably associated with the number of completed cycles in this small sample.", ''),
    ])
    log.append("5.5 engagement paragraph inserted")

# (4) replace the Cohen's d paragraph
p = find("Among the nine completers, the recorded FRA index changed from 42.0")
if p:
    set_parts(p, [
        ("Among the nine sustained participants, the recorded FRA index moved from 42.0 ± 8.3 before the deployment to 48.9 ± 9.4 after it: a mean paired change of +6.9 points (95% CI 4.5–9.3; paired t(8) = 6.67, p < 0.001; Wilcoxon p = 0.004; dz = 2.22, exact 95% CI 0.95–3.46). All thirty participants have recorded pre/post values, so no imputation is required: across the full cohort the paired change was +4.0 points (95% CI 0.5–7.6; t(29) = 2.31, p = 0.028; Wilcoxon p = 0.015; dz = 0.42, exact 95% CI 0.05–0.79), with twenty participants recording an increase, nine a decrease and one no change. The pooled (d = 0.78) and zero-imputation (d = 0.57) estimates quoted in earlier drafts do not reproduce from the archive and have been withdrawn. The difference in index change between sustained and low-engagement participants (+6.9 versus +2.8 points) was not statistically reliable (Welch t = 1.55, p = 0.13; Mann-Whitney p = 0.051). Two interpretive constraints apply before any of these changes is read as improvement. ", ''),
        ("[AUTHOR INPUT REQUIRED — index direction. The dataset legend supplied with the archive states that a lower FRA index indicates a higher performer, in which case the recorded increases would indicate deterioration, not improvement; the thesis narrative to date has assumed the opposite. Confirm the direction convention of the archived index against the InBody documentation and adjust the interpretive language in §5.5–§5.7, §6.6 and Chapter 9 accordingly.]", 'input'),
        (" Second, single-measurement changes of this size lie within the range of test-retest variability reported for FRA instruments in older adults (Kim et al., 2018), so within-deployment index change should be read as a provisional signal rather than a demonstrated clinical effect.", ''),
    ])
    log.append("5.5 pre/post paragraph rebuilt from archive")
else:
    failures.append("5.5 Cohen's d paragraph not found")

# (5) self-report/satisfaction paragraph — insert after pre/post paragraph
p = find("Among the nine sustained participants, the recorded FRA index moved from 42.0")
if p:
    insert_after(p, [
        ("The self-reported and recorded ledgers diverge in an instructive way. Eighteen of thirty participants (60.0%, 95% CI 42.3–75.4%) reported symptom improvement after the deployment and twenty (66.7%, 95% CI 48.8–80.8%) were satisfied or very satisfied. Satisfaction was strongly associated with sustained engagement: all nine sustained participants were satisfied, against eleven of twenty-one in the low-engagement group (Fisher's exact p = 0.013). Self-reported symptom improvement showed the same direction but did not reach significance (7 of 9 versus 11 of 21; Fisher's exact p = 0.25). Among the eight single-session participants, five recorded a positive index change yet only two reported any improvement and none returned — the archived form of the divergence that Chapter 5 names the objective-score-improvement/non-return paradox, subject to the direction caveat above. Excluding the single unverified record does not alter any of these conclusions.", ''),
    ])
    log.append("5.5 self-report/satisfaction paragraph inserted")

# (6) revise the Table 5.3 support placeholder
p = find("[AUTHOR INPUT REQUIRED — Table 5.3 support")
if p:
    set_parts(p, [
        ("[AUTHOR INPUT REQUIRED — Table 5.3 support. The week-2 motivational-orientation classification does not appear in the archived dataset, so the odds ratio of 31.5 cannot be reproduced from the data supplied. Either provide the coding records behind the classification (rubric, coders, inter-coder agreement, classification dates) or withdraw Table 5.3 and the odds ratio, and rely on the satisfaction association reported above, which is reproducible from the archive.]", 'input'),
    ])
    log.append("Table 5.3 placeholder updated for missing classification records")
else:
    failures.append("Table 5.3 placeholder not found")

# (7) Figure 5.4: replace green KM box with the real pre/post figure
p = find("[INSERT FIGURE 5.4 HERE — Kaplan-Meier plot")
if p:
    clear_runs(p)
    run = p.add_run()
    run.add_picture(FIG54, width=Inches(6.0))
    log.append("Figure 5.4 image (participant-level pre/post) embedded")
else:
    failures.append("Figure 5.4 KM placeholder not found")

new54 = ("Figure 5.4. Participant-level pre/post recorded FRA index for the Stage-2 cohort (N = 30), by engagement group. Thin lines are individual participants; the bold line is the group mean (sustained: 42.0 to 48.9; low engagement: 44.8 to 47.6). Generated from the archived dataset (Appendix G, item G.2). The direction convention of the index is under author confirmation (see §5.5).")
body_caption = lof_line = None
for q in doc.paragraphs:
    if q.text.startswith("Figure 5.4. Kaplan-Meier"):
        if q.style.name == "FigureCaption": body_caption = q
        else: lof_line = q
if body_caption is not None:
    set_parts(body_caption, [(new54, '')])
    log.append("Figure 5.4 caption replaced")
else:
    failures.append("Figure 5.4 body caption not found")
if lof_line is not None:
    tail = lof_line.text.split("\t")[-1] if "\t" in lof_line.text else ""
    set_parts(lof_line, [("Figure 5.4. Participant-level pre/post recorded FRA index for the Stage-2 cohort (N = 30), by engagement group." + ("\t" + tail if tail else ""), '')])
    log.append("Figure 5.4 List-of-Figures line replaced")

# ---------------------------------------------------------------- §5.5A QFD section
anchor = find("The self-reported and recorded ledgers diverge in an instructive way")
if anchor:
    # find Table 5.3 caption to place section AFTER table block; simpler: place before §5.6 heading
    h56 = None
    for q in doc.paragraphs:
        if q.text.strip().startswith("5.6 The F4 Cluster") and q.style.name.startswith("Heading"):
            h56 = q; break
    if h56 is not None:
        head = h56.insert_paragraph_before()
        head.style = H2STYLE
        set_parts(head, [("5.5A Translating the Findings into Design Priorities: A QFD Pipeline", '')])
        p1 = head  # insert content after heading, before h56 — use insert_paragraph_before on h56 repeatedly
        c1 = h56.insert_paragraph_before()
        set_parts(c1, [("The statistical analysis above establishes what happened; it does not by itself say what the game should do next. Quality Function Deployment (QFD) supplies that translation. QFD is an engineering method that maps the voice of the user — needs expressed in the users' own terms — onto measurable technical characteristics of the product, and weights each mapping by the strength of the underlying evidence (Akao, 1990; Chan & Wu, 2002). It is a requirements-translation tool, not a statistical model: in this thesis QFD consumes the statistical findings and returns design priorities, while the hypothesis tests themselves remain the province of the estimators reported in §5.5.", '')])
        c2 = h56.insert_paragraph_before()
        set_parts(c2, [("Table 5.4 presents the deployment matrix. The rows are user needs derived from the Stage-1 survey, the Stage-2 telemetry and the Stage-2 self-reports, each anchored to the empirical finding that evidences it and to the TAM, STAM or Octalysis construct it instantiates. The columns are the game's technical characteristics: the adaptive difficulty layer (§5.3.5), the loss-free dynamic game states (§5.3.10), the trajectory display (L5), the coin-reward loop and the staff coaching protocol (L6). A strong relationship (●) indicates that the feature directly serves the need; a moderate relationship (◐) indicates partial or indirect support; a blank cell indicates no designed relationship.", '')])
        # QFD table
        tbl = doc.add_table(rows=7, cols=7)
        tbl.style = doc.tables[0].style
        headers = ["User need (voice of the user)", "Evidence from the data", "Construct", "Adaptive difficulty (L2)", "Loss-free states", "Trajectory display (L5)", "Coaching (L6)"]
        rows = [
            ["Challenge that matches my ability", "70.0% disengaged at ≤3 cycles; satisfaction 9/9 among sustained vs 11/21 (p = .013)", "Flow / PEOU (TAM)", "●", "◐", "", "◐"],
            ["Failure must not feel like a verdict about me", "0/21 re-engagement after disengagement; Stage-1: ~74% adverse first experience among prior users", "Loss aversion / D8 (Octalysis)", "◐", "●", "◐", "●"],
            ["Visible personal progress", "Recorded index and self-report diverge (5/8 single-session positive change, 2/8 felt improvement)", "Competence (SDT) / D2", "", "◐", "●", "◐"],
            ["Understandable, non-clinical feedback", "60.0% self-reported improvement; satisfaction associated with staying", "Gerontechnology anxiety (STAM)", "", "◐", "●", "●"],
            ["Reward for effort, not only outcome", "Coin loop sustained play in completers (≥4 cycles: 9/30)", "Learn-to-earn / D4", "◐", "", "◐", ""],
            ["A reason to come back", "No observed return in Stage 2; Stage-3 decoupled cohort returned 15/23", "Behavioural intention (TAM/STAM)", "◐", "●", "●", "●"],
        ]
        data = [headers] + rows
        for i, rvals in enumerate(data):
            for j, v in enumerate(rvals):
                cell = tbl.rows[i].cells[j]
                cell.text = ""
                run = cell.paragraphs[0].add_run(v)
                if i == 0: run.bold = True
        h56._p.addprevious(tbl._tbl)
        cap = h56.insert_paragraph_before()
        try:
            cap.style = doc.styles["TableCaption"]
        except Exception:
            pass
        set_parts(cap, [("Table 5.4. QFD deployment matrix: user needs evidenced by the Stage-1 and Stage-2 data, mapped onto the game's technical characteristics. ● strong designed relationship; ◐ moderate or indirect relationship.", '')])
        c3 = h56.insert_paragraph_before()
        set_parts(c3, [("Read column-wise, the matrix states the analysis pipeline's design verdict. The adaptive difficulty layer is the primary carrier of the challenge–skill need that the engagement data make most urgent. The loss-free state repertoire and the coaching protocol are the primary carriers of the verdict-avoidance need evidenced by the zero re-engagement observation. The trajectory display carries the progress-visibility need that the recorded-versus-self-reported divergence exposes. Read row-wise, every need is served by at least two features, which is the redundancy QFD prescribes for needs whose failure modes are severe. The matrix also fixes what the next evaluation cycle must measure: each cell is a testable claim that the feature moves the evidence indicator in its row, and the limitations acknowledged in §5.9 and §8.6 — small sample, short observation window, measurement constraints and self-report bias — bound how far any cell can currently be asserted.", '')])
        log.append("5.5A QFD section + Table 5.4 inserted")
    else:
        failures.append("5.6 heading not found for QFD insertion")
else:
    failures.append("QFD anchor paragraph not found")

# ---------------------------------------------------------------- §5.8 entropy timing
p = find("The 18.4-hour median between first failure and last recorded interaction")
if p:
    set_parts(p, [("The entropy language in this thesis is a bridge model, not a measured result. The archived Stage-2 dataset does not preserve timing between events, so no temporal claim about the interval between first failure and final interaction is made. Whether a belief changed, when it changed, and whether an information-theoretic entropy reduction occurred are propositions that require repeated belief elicitation before and after a failure event.", '')])
    log.append("5.8 entropy timing claim removed")
else:
    failures.append("5.8 entropy paragraph not found")

# ---------------------------------------------------------------- §5.9 limitations
p = find("Stage 2 has important limitations. The sample is small and drawn from one site")
if p:
    set_parts(p, [("Stage 2 has important limitations. The sample is small, drawn from one site, and not demographically diverse; the observation window is short, so nothing is claimed about sustained use beyond eight weeks. The archived dataset contains no timestamps, so the temporal structure of disengagement could not be analysed. Measurement constraints are material: the recorded index was measured once before and once after the window on a single instrument whose short-interval test-retest variability is non-trivial (Kim et al., 2018), and the direction convention of the archived index is under author confirmation. The self-reported outcomes are vulnerable to courtesy and demand bias, since the questionnaires were administered by facility staff in a setting where participants knew the deployment was a facility initiative. The week-2 motivational-orientation classification is not in the archived dataset and may be affected by reverse causality; the provisional FIAS composite appears to reuse cycle-count information; FES-I was not available as a baseline covariate; the embedded investigator was not complemented by an independent observer; and the clinical meaning of short-term index change was not established.", '')])
    log.append("5.9 limitations rewritten (adds measurement constraints and feedback bias)")
else:
    failures.append("5.9 limitations paragraph not found")

# ---------------------------------------------------------------- Ch6: Stage-3 logbook gap
p = find("[AUTHOR INPUT REQUIRED: document the Stage-3 recruitment sources")
if p:
    insert_after(p, [
        ("[AUTHOR INPUT REQUIRED — Stage-3 raw data. The Stage-3 data logbook supplied on 12 July 2026 (Stage3_Data_Logbook_120cases.xlsx) contains the complete field structure — participants, per-session game telemetry, questionnaire responses and outcomes for IDs 3001–3120 — but no entered records. The Chapter 6 percentages (65.2% return, 6.7% non-return, 71.7% endorsement, 63.3% intention) therefore cannot yet be reproduced from an archived dataset. Populate the logbook from the field records, or attach the original source records, before binding; the Stage-3 claims should be treated as unverified until then.]", 'input'),
    ])
    log.append("Ch6 Stage-3 logbook gap placeholder added")
else:
    failures.append("Ch6 Stage-3 placeholder anchor not found")

# ---------------------------------------------------------------- §6.10 model-comparison mentions of Cox
p = find("Three classes of comparison statistic are pre-specified")
if p:
    t = p.text
    t = t.replace("For time-to-event outcomes, the plan uses Cox partial-likelihood ratio tests and Harrell's C-index.",
                  "For time-to-event outcomes — available only once session timestamps are captured, which the Stage-2 archive did not do — the plan uses Cox partial-likelihood ratio tests and Harrell's C-index.")
    t = t.replace("The Stage-2 single-predictor Cox HR of 3.24 is a preliminary standalone association; it is not an upper bound and does not establish incremental value over STAM.",
                  "No Stage-2 survival estimate is carried forward, because the Stage-2 archive contains no time variable; the Stage-2 anchors are the engagement and paired-change statistics of §5.5.")
    set_parts(p, [(t, '')])
    log.append("6.10 comparison-statistics paragraph updated")
else:
    failures.append("6.10 comparison paragraph not found")

p = find("The Stage-2 univariate Cox HR for first failure (3.24, p < 0.001) and the Stage-3 observed non-return")
if p:
    set_parts(p, [("The Stage-2 early-disengagement proportion (70.0%, 95% CI 52.1–83.3%) and the Stage-3 observed non-return proportions provide provisional effect-size anchors. A formal pooled or nested analysis is not appropriate until a future dataset uses harmonised recruitment, covariates, measures, outcome definitions and observation windows.", '')])
    log.append("6.10 effect-size anchors updated")
else:
    failures.append("6.10 anchors paragraph not found")

# ---------------------------------------------------------------- Table 6.5 rebuild
t65 = None
for t in doc.tables:
    if any("Cox proportional hazards" in c.text for r in t.rows for c in r.cells):
        t65 = t; break
if t65 is not None:
    def setcell(cell, text, kind=''):
        cell.text = ""
        run = cell.paragraphs[0].add_run(text)
        if kind == 'input':
            run.bold = True; run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    for r in t65.rows:
        head = r.cells[0].text
        if "First failure predicting dropout" in head:
            setcell(r.cells[0], "Early disengagement, ≤ 3 cycles (Stage 2)")
            setcell(r.cells[1], "70.0%")
            setcell(r.cells[2], "52.1–83.3%*")
            setcell(r.cells[3], "30")
            setcell(r.cells[4], "Wilson score interval; Table 5.2. (Survival modelling withdrawn: no time variable in the archive; §5.5)")
        elif "Motivational orientation" in head:
            setcell(r.cells[4], "Odds ratio; Table 5.3. Classification records not in archive — see §5.5", 'input')
        elif "completers" in head and "FRA index change" in head:
            setcell(r.cells[0], "FRA index change, sustained participants (Stage 2)")
            setcell(r.cells[1], "dz = 2.22 (paired); mean +6.9 points")
            setcell(r.cells[2], "dz 0.95–3.46; mean 4.5–9.3")
            setcell(r.cells[3], "9")
            setcell(r.cells[4], "Paired t and Wilcoxon; §5.5")
        elif "pooled" in head:
            setcell(r.cells[0], "FRA index change, all participants (Stage 2, archived paired data)")
            setcell(r.cells[1], "dz = 0.42; mean +4.0 points")
            setcell(r.cells[2], "dz 0.05–0.79; mean 0.5–7.6")
            setcell(r.cells[3], "30")
            setcell(r.cells[4], "Paired t and Wilcoxon; §5.5")
        elif "conservative zero-change" in head:
            setcell(r.cells[0], "Satisfaction association with sustained engagement (Stage 2)")
            setcell(r.cells[1], "9/9 vs 11/21 satisfied")
            setcell(r.cells[2], "p = 0.013")
            setcell(r.cells[3], "30")
            setcell(r.cells[4], "Fisher's exact test; §5.5")
    log.append("Table 6.5 Stage-2 rows rebuilt from archive")
else:
    failures.append("Table 6.5 not found")

p = find("Intervals marked with an asterisk are Wilson score intervals")
if p:
    set_parts(p, [("Intervals marked with an asterisk are Wilson score intervals computed from the reported counts. Effect-size intervals for the paired changes are exact noncentral-t intervals computed from the archived participant-level data. The earlier pooled (d = 0.78) and zero-imputation (d = 0.57) estimates did not reproduce from the archive and have been withdrawn; the survival-model row has been replaced for the same reason (§5.5). The confidence interval for the 35.6-percentage-point difference should respect the fact that retention and intention are distinct measures and may not represent a simple paired effect. The very wide interval around the week-2 motivational-orientation odds ratio, together with the absence of the classification from the archive, means that estimate should be treated as exploratory pending the author's coding records.", '')])
    log.append("Table 6.5 footnote rebuilt")
else:
    failures.append("Table 6.5 footnote not found")

# ---------------------------------------------------------------- §8.3 empirical contributions
p = find("The empirical record yields the following results: a reported first-failure hazard ratio")
if p:
    set_parts(p, [("The empirical record yields the following results: 70.0% early disengagement (21 of 30, at or before three cycles) in the Stage-2 cohort with no observed re-engagement during the eight-week window; a paired recorded-index change of +6.9 points among the nine sustained participants (dz = 2.22, 95% CI 0.95–3.46) and +4.0 points across all thirty (dz = 0.42, 95% CI 0.05–0.79), with the direction convention under confirmation; a strong association between satisfaction and sustained engagement (Fisher's exact p = 0.013); an exploratory week-2 motivational-orientation odds ratio of 31.5 whose coding records remain to be archived; recorded-index increase without self-reported improvement or return among most single-session participants; 65.2% return among Stage-3 initial dropouts overall and 62.5% in the 65-80 band, pending completion of the Stage-3 data logbook; 71.7% coordination-frame endorsement as a manipulation check; 63.3% regular-use intention overall; and a descriptive 35.6-percentage-point difference between session retention and regular-use intention in the 65-80 band. These results do not constitute a direct correlation across stages because the participants and measures differ.", '')])
    log.append("8.3 empirical contributions rewritten")
else:
    failures.append("8.3 paragraph not found")

# ---------------------------------------------------------------- Ch9 summary
p = find("First-failure experience was strongly associated with faster observed dropout in the reported Cox model")
if p:
    set_parts(p, [("Stage 2 prospectively studied thirty residents over eight weeks using a dual-purpose AI-assisted skiing game. Seventy per cent of participants disengaged at or before three game cycles, and no re-engagement was observed among the low-engagement group during the study window. Among single-session participants, the recorded index typically moved upward while self-reported improvement was rare and no one returned. This divergence motivated the behavioural-economic account, but it did not establish clinical benefit or directly measure endowment, mental accounting or stereotype threat, and the direction convention of the recorded index remains under author confirmation.", '')])
    log.append("Ch9 Stage-2 summary rewritten")
else:
    failures.append("Ch9 summary paragraph not found")

# ---------------------------------------------------------------- References: Akao, Chan & Wu
p = find("Ahn, M., Brohan, A., Brown, N.")
if p:
    insert_after(p, [("Akao, Y. (1990). Quality function deployment: Integrating customer requirements into product design. Productivity Press.", '')])
    log.append("Akao (1990) reference inserted")
else:
    failures.append("Ahn reference anchor not found")

p = find("Centers for Disease Control and Prevention. (2023)")
if p:
    insert_after(p, [("Chan, L.-K., & Wu, M.-L. (2002). Quality function deployment: A literature review. European Journal of Operational Research, 143(3), 463–497.", '')])
    log.append("Chan & Wu (2002) reference inserted")
else:
    failures.append("CDC reference anchor not found")

# ---------------------------------------------------------------- Appendix G status updates
p = find("G.2  De-identified Stage-2 session-log dataset")
if p:
    set_parts(p, [
        ("G.2  Stage-2 participant dataset (FRA_Raw_Data_N30.xlsx; 30 elderly participants plus two younger controls): gender, age, pre/post recorded FRA index, presenting symptoms, completed game cycles, self-reported symptom change, satisfaction and verification flag. Received and archived 12 July 2026; this file is the source for every Stage-2 statistic in §5.5 and for Figure 5.4. ", ''),
        ("[AUTHOR INPUT REQUIRED: session-level telemetry with timestamps and AI-call logs was not part of this archive — attach it if it exists, or state in §8.9 that it was not retained]", 'input'),
    ])
    log.append("G.2 updated to received status")
else:
    failures.append("G.2 item not found")

p = find("G.3  Stage-3 questionnaire response file")
if p:
    set_parts(p, [
        ("G.3  Stage-3 data logbook (Stage3_Data_Logbook_120cases.xlsx): codebook, participants, per-session telemetry, questionnaire responses, outcomes and aggregate formulas for IDs 3001–3120. Template received 12 July 2026 with no entered records. ", ''),
        ("[AUTHOR INPUT REQUIRED: populate all sheets from the Stage-3 field records; the Chapter 6 results cannot be reproduced until this is complete]", 'input'),
    ])
    log.append("G.3 updated to template-received status")
else:
    failures.append("G.3 item not found")

# ---------------------------------------------------------------- save & report
doc.save(OUT)
print("SAVED:", OUT)
print("\n--- APPLIED (%d) ---" % len(log))
for l in log: print(" •", l)
print("\n--- FAILURES (%d) ---" % len(failures))
for f in failures: print(" ✗", f)
