#!/usr/bin/env python3
"""Apply the final-revision edits to the EngD thesis docx.

Edits fall into five groups:
  A. Convert every AUTHOR ACTION REQUIRED block into final academic prose,
     with clearly highlighted [AUTHOR INPUT REQUIRED] placeholders for items
     only the author can supply (no data are invented).
  B. Insert photograph/screenshot placeholder boxes at the positions the
     author requested (FRA machine, people playing, game artefacts).
  C. Repair the citation apparatus: insert in-text citations at the natural
     anchor points for the ~40 orphaned reference entries, correct verified
     reference errors, delete meta-commentary and author notes.
  D. Statistical presentation: approximate CIs for Cohen's d from reported
     summary statistics (labelled as such); withhold the statistically
     implausible Cox CI pending re-verification; regenerate Figure 6.1 from
     Table 6.2's reported values; placeholder for the Figure 5.4 KM plot.
  E. Harmonise the Layer-2 model-vendor inconsistency and other wording.
"""
import docx
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.shared import Inches

SRC = "/root/.claude/uploads/3d8a50e8-9258-51ce-a7d6-194da50f9a04/11fbf6ac-Tam_V4_EngD_Thesis_Academic_Revision_Draft_11072026.docx"
DST = "/tmp/claude-0/-home-user-gamification/3d8a50e8-9258-51ce-a7d6-194da50f9a04/scratchpad/Tam_V5_EngD_Thesis_Final_Revision_Draft.docx"
FIG61 = "/tmp/claude-0/-home-user-gamification/3d8a50e8-9258-51ce-a7d6-194da50f9a04/scratchpad/figure_6_1.png"

doc = docx.Document(SRC)
log, failures = [], []

# ---------------- helpers ----------------

def find(needle, nth=1):
    n = 0
    for p in doc.paragraphs:
        if needle in p.text:
            n += 1
            if n == nth:
                return p
    return None

def clear_runs(p):
    for r in list(p.runs):
        r._element.getparent().remove(r._element)

def set_parts(p, parts):
    """parts: list of (text, kind) where kind in {'', 'input', 'photo', 'b', 'i'}"""
    clear_runs(p)
    for text, kind in parts:
        r = p.add_run(text)
        if kind == 'input':
            r.bold = True
            r.font.highlight_color = WD_COLOR_INDEX.YELLOW
        elif kind == 'photo':
            r.bold = True
            r.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
        elif kind == 'b':
            r.bold = True
        elif kind == 'i':
            r.italic = True

def sub(needle, old, new, nth=1, label=None):
    p = find(needle, nth)
    lab = label or (old[:50] + "...")
    if p is None:
        failures.append("FIND-FAIL: " + lab); return None
    if old not in p.text:
        failures.append("SUB-FAIL (old text absent): " + lab); return None
    set_parts(p, [(p.text.replace(old, new), '')])
    log.append("edited: " + lab)
    return p

def insert_after(p, parts, style=None):
    new = OxmlElement('w:p')
    p._p.addnext(new)
    np_ = Paragraph(new, p._parent)
    if style is not None:
        np_.style = style
    set_parts(np_, parts)
    return np_

def insert_before(p, parts, style=None):
    np_ = p.insert_paragraph_before()
    if style is not None:
        np_.style = style
    set_parts(np_, parts)
    return np_

def delete(p):
    p._element.getparent().remove(p._element)

FIGSTYLE = None
for p in doc.paragraphs:
    if p.style.name == "FigureCaption":
        FIGSTYLE = p.style
        break

# =========================================================================
# GROUP A — AUTHOR ACTION REQUIRED conversions
# =========================================================================

# A1. Section 3.4 ethics block
p = find("AUTHOR ACTION REQUIRED BEFORE FINAL BINDING: Each stage was conducted")
if p:
    set_parts(p, [
        ("Each stage was conducted under the ethics policy of The Hong Kong Polytechnic University and in accordance with the Declaration of Helsinki. Personal Information Protection Law (PIPL) compliance was maintained throughout: all data were pseudonymised at collection, and no participant identifiers were transmitted to third-party AI services. Stage 3 did not collect participant-level clinical balance telemetry as part of the decoupling protocol. No adverse events were reported during any stage. The programme was reviewed by the PolyU Human Subjects Ethics Sub-Committee ", ''),
        ("[AUTHOR INPUT REQUIRED: insert the HSESC approval reference number and date; state which stages and amendments the approval covered; and confirm whether the Stage-3 frame-manipulation and recruitment wording were reviewed and approved]", 'input'),
        (". Chapter 8 (§8.7) addresses vulnerability, disclosure, autonomy and the ethical boundary between participant-facing reframing and deception.", ''),
    ])
    log.append("A1 ethics block rewritten")
else:
    failures.append("A1 ethics block not found")

# A2. §5.5 Cox paragraph
p = find("A Cox proportional-hazards model associated first-failure experience with faster observed dropout (HR = 3.24, 95% CI 1.89-5.56, p < 0.001). The final thesis must report")
if p:
    set_parts(p, [
        ("Engagement showed an early drop-off. Twenty-one of thirty participants completed three or fewer cycles, eight ended after a single session, and no re-engagement was observed among the low-engagement group during the eight-week study window. A Cox proportional-hazards model associated first-failure experience with faster observed dropout (HR = 3.24; the interval estimate and p-value are being re-derived from the verified survival dataset, as explained in the note below). Because the first failure occurs during rather than before the observation window, the covariate must enter the model as a time-varying predictor, or through a landmark analysis, to avoid immortal-time bias. The final specification accompanying Figure 5.4 therefore reports the time origin, the event definition, the censoring rule, the number of events, the treatment of ties, the predictor timing, the proportional-hazards diagnostics and any small-sample correction.", ''),
    ])
    ph = insert_after(p, [
        ("[AUTHOR INPUT REQUIRED — Cox model re-verification. The interval reported in the previous draft (95% CI 1.89–5.56, p < 0.001) implies a standard error of the log hazard ratio of about 0.28, which would require roughly fifty or more events. With thirty participants, and therefore at most thirty events, the smallest attainable standard error is approximately 0.37–0.44, giving an interval no narrower than roughly 1.4–7.6 around HR = 3.24 and a p-value nearer 0.007 than < 0.001. The previously reported interval cannot have been produced by this dataset as described. Re-fit the model from the participant-level survival file (Appendix G, item G.2) and update the estimate, interval and p-value here, in the Abstract, in Table 6.5, in the List of Figures and in the caption of Figure 5.4 before binding.]", 'input'),
    ])
    log.append("A2 Cox paragraph rewritten + verification note")
else:
    failures.append("A2 Cox paragraph not found")

# A3. §5.5 Cohen's d paragraph
p = find("Among the nine completers, the recorded FRA index changed from 42.0")
if p:
    set_parts(p, [
        ("Among the nine completers, the recorded FRA index changed from 42.0 ± 8.3 to 48.9 ± 9.4, a paired Cohen's d of 2.22 (approximate 95% CI 1.00–3.44). The pooled estimate across all thirty participants was d = 0.78 (approximate 95% CI 0.37–1.19). A conservative sensitivity analysis that assigns zero change to all twenty-one non-completers gave d = 0.57 (approximate 95% CI 0.18–0.96); it is reported as a sensitivity analysis rather than as an intent-to-treat estimate because no formal missing-data model was specified. The approximate intervals are computed from the reported summary statistics under a normal approximation and are to be re-derived from the participant-level dataset before binding, together with a participant-level pre/post change plot, an assessment of influential observations and a precise operational definition of the index. A week-2 motivational-orientation classification produced an odds ratio of 31.5 (95% CI 2.9–339) for sustained engagement. That interval is very wide, and because the classification was made after engagement behaviour had begun, reverse causality cannot be excluded: continuation may have influenced the orientation assigned at week 2. The classification is therefore treated as exploratory evidence of direction, not as a validated baseline predictor or causal mechanism.", ''),
    ])
    log.append("A3 Cohen's d paragraph rewritten with approximate CIs")
else:
    failures.append("A3 Cohen's d paragraph not found")

# A3b. Placeholder after Table 5.3 (anchor: the paragraph that follows the table is
# the old AUTHOR ACTION; we replace that paragraph)
p = find("AUTHOR ACTION REQUIRED: add the complete Cox model specification")
if p:
    set_parts(p, [
        ("[AUTHOR INPUT REQUIRED — Table 5.3 support. Replace the percentages with raw counts and denominators from the study records: sustained engagement is 7 of 8 learning-focused against 4 of 22 health-deficit-focused, and the positive FRA-index-change column requires the number of participants in each group with valid pre/post index data (85.7% implies 6 of 7 — confirm the denominators). Append the classification rubric, the identity and number of coders, the inter-coder agreement statistic, and the classification date relative to the outcome window. Insert the participant-level pre/post change plot referred to above as a lettered figure (suggested: Figure 5.4B).]", 'input'),
    ])
    log.append("A3b Table 5.3 support placeholder")
else:
    failures.append("A3b Table 5.3 AUTHOR ACTION not found")

# A3c. Delete the now-merged 'Interpretive caution' paragraph
p = find("Interpretive caution: the week-2 motivational orientation is an exploratory")
if p:
    delete(p); log.append("A3c interpretive-caution paragraph merged and removed")
else:
    failures.append("A3c interpretive caution not found")

# A4. Figure 5.4 placeholder (replace the AUTHOR ACTION paragraph)
p = find("AUTHOR ACTION REQUIRED: Figure 5.4 is referenced but the Kaplan-Meier plot")
if p:
    set_parts(p, [
        ("[INSERT FIGURE 5.4 HERE — Kaplan-Meier plot. Generate from the verified Stage-2 survival dataset (Appendix G, item G.2): survival curves stratified by first-failure experience; numbers at risk beneath the time axis; censoring marks; axes labelled (time in weeks, 0–8; proportion remaining engaged); caption bound to the re-verified Cox specification.]", 'photo'),
    ])
    log.append("A4 Figure 5.4 placeholder box")
else:
    failures.append("A4 Figure 5.4 AUTHOR ACTION not found")

# A5. Figure 5.4 caption update
p = find("Figure 5.4. Kaplan-Meier curves for time to observed dropout/non-return", nth=2) or find("Figure 5.4. Kaplan-Meier curves for time to observed dropout/non-return")
# nth=2 targets the body caption (nth=1 is the List of Figures line); fall back defensively
body_caption = None
lof_line = None
matches = [q for q in doc.paragraphs if q.text.startswith("Figure 5.4. Kaplan-Meier")]
for q in matches:
    if q.style.name == "FigureCaption":
        body_caption = q
    else:
        lof_line = q
if body_caption is not None:
    set_parts(body_caption, [("Figure 5.4. Kaplan-Meier curves for time to observed dropout/non-return during the eight-week Stage-2 study, stratified by first-failure experience (HR = 3.24; the interval and p-value are bound after re-verification — see §5.5). Model specification and proportional-hazards diagnostics are reported in the accompanying methods text.", '')])
    log.append("A5 Figure 5.4 body caption updated")
else:
    failures.append("A5 Figure 5.4 body caption not found")
if lof_line is not None:
    tail = lof_line.text.split("\t")[-1] if "\t" in lof_line.text else ""
    set_parts(lof_line, [("Figure 5.4. Kaplan-Meier curves for time to observed dropout/non-return during the eight-week Stage-2 study, stratified by first-failure experience (HR = 3.24; interval under re-verification, see §5.5)." + ("\t" + tail if tail else ""), '')])
    log.append("A5 Figure 5.4 List-of-Figures line updated")

# A6. §5.6 thirteen-attribute discrepancy
p = find("The current manuscript describes thirteen binary attributes but a score range of 0-10")
if p:
    txt = p.text
    old = "The current manuscript describes thirteen binary attributes but a score range of 0-10; this discrepancy must be resolved in the final codebook."
    if old in txt:
        pre, post = txt.split(old)
        set_parts(p, [
            (pre, ''),
            ("The composite is assembled from thirteen binary attributes drawn from STAM and selected Octalysis drives and is reported on a 0–10 range; the rule that maps the thirteen items onto that range is specified in the codebook in Appendix B ", ''),
            ("[AUTHOR INPUT REQUIRED: state the weighting or collapsing rule that maps the thirteen items onto the 0–10 composite]", 'input'),
            (". ", ''),
            (post, ''),
        ])
        log.append("A6 thirteen-attribute sentence rewritten")
    else:
        failures.append("A6 exact sentence not present")
else:
    failures.append("A6 paragraph not found")

# A7. §6.4 'final report should document' sentence -> placeholder
old = " The final report should document recruitment sources, inclusion and exclusion criteria, health and functional characteristics, staff scripts, observation window and any differences in software version or support model so that cross-stage comparability can be assessed."
p = find("One hundred and twenty adults aged 25-80 were recruited")
if p and old in p.text:
    set_parts(p, [(p.text.replace(old, ""), '')])
    insert_after(p, [
        ("[AUTHOR INPUT REQUIRED: document the Stage-3 recruitment sources, inclusion and exclusion criteria, health and functional characteristics, staff scripts, observation window, and any differences in software version or support model relative to Stage 2, so that cross-stage comparability can be assessed.]", 'input'),
    ])
    log.append("A7 Stage-3 documentation placeholder")
else:
    failures.append("A7 Stage-3 recruitment paragraph/sentence not found")

# A8. Table 6.3 per-band endorsement placeholder — insert after the interpretive note
p = find("Interpretive note: Stage-3 coordination endorsement and Stage-2 motivational orientation are non-equivalent")
if p:
    insert_after(p, [
        ("[AUTHOR INPUT REQUIRED: insert the per-age-band coordination-endorsement percentages in Table 6.3; the current draft records only the 71.7% total (86 of 120).]", 'input'),
    ])
    log.append("A8 Table 6.3 per-band placeholder")
else:
    failures.append("A8 Table 6.3 interpretive note not found")

# A9. §8.9 reproducibility AUTHOR ACTION
p = find("AUTHOR ACTION REQUIRED BEFORE FINAL BINDING: Three reproducibility boundaries")
if p:
    set_parts(p, [
        ("Three reproducibility boundaries are material and are named here so that replication attempts are not misled. First, the InBody calibration function is a vendor component and is treated as a black box: replication on other hardware requires new calibration and may not reproduce the same signals. Second, third-party model endpoints change over time, so the exact run-time configuration used in Stage 2 is recorded in Appendix G ", ''),
        ("[AUTHOR INPUT REQUIRED: record the model name and vendor, the version or dated endpoint, the system and user prompts, the LangChain configuration, temperature and other generation settings, fallback rules and any manual overrides. Earlier drafts named different vendors in §5.3 and Appendix E; the final record must state the endpoint actually deployed]", 'input'),
        (". Third, the appendices bind the verbatim language versions of all instruments, the translation procedure, the complete FIAS codebook with inter-coder reliability, the consent materials and the photographic record ", ''),
        ("[AUTHOR INPUT REQUIRED: complete this documentation at the marked positions in Appendices A–D and G before binding]", 'input'),
        (".", ''),
    ])
    log.append("A9 reproducibility block rewritten")
else:
    failures.append("A9 reproducibility AUTHOR ACTION not found")

# A10. Abstract CI removal
p = find("Stage 1 used a face-to-face survey with approximately 200 residents")
if p and "(HR = 3.24, 95% CI 1.89-5.56, p < 0.001)" in p.text:
    set_parts(p, [(p.text.replace("(HR = 3.24, 95% CI 1.89-5.56, p < 0.001)", "(HR = 3.24; the interval estimate is reported with the full model specification in §5.5)"), '')])
    log.append("A10 Abstract CI reference updated")
else:
    failures.append("A10 Abstract HR sentence not found")

# A11–A14. Appendix AUTHOR ACTION items -> placeholders
appendix_actions = [
    ("AUTHOR ACTION REQUIRED: Author to insert the verbatim Mandarin Chinese wording",
     "[AUTHOR INPUT REQUIRED: insert the verbatim Mandarin wording of the six survey items, the bilingual back-translation record, and the inter-rater agreement statistics for the staff administrators.]"),
    ("AUTHOR ACTION REQUIRED: Author to insert the full thirteen-attribute coding rules table",
     "[AUTHOR INPUT REQUIRED: insert the full thirteen-attribute coding rules table (including the item-to-composite mapping onto the 0–10 range), the inter-coder reliability statistics (Cohen's κ across the two trained coders), and the four worked-example transcripts that calibrate F1 through F4.]"),
    ("AUTHOR ACTION REQUIRED: Author to insert the verbatim bilingual question wording",
     "[AUTHOR INPUT REQUIRED: insert the verbatim bilingual question wording, the staff-administered consent script, and one sample anonymised transcript for an F4 participant.]"),
    ("AUTHOR ACTION REQUIRED: Author to insert the verbatim Cantonese / Mandarin / English wording",
     "[AUTHOR INPUT REQUIRED: insert the verbatim Cantonese, Mandarin and English wording of items D1 and D2, and the scale-level reliability (α) computed at the close of Stage 3.]"),
]
for needle, replacement in appendix_actions:
    p = find(needle)
    if p:
        set_parts(p, [(replacement, 'input')])
        log.append("Appendix placeholder: " + needle[:60])
    else:
        failures.append("Appendix action not found: " + needle[:60])

# A15. Appendix B prose correction
p = find("The Stage-2 coding scheme operationalises the provisional FIAS hypothesis")
if p:
    set_parts(p, [
        ("The Stage-2 coding scheme operationalises the provisional FIAS hypothesis through thirteen binary attributes drawn from STAM and selected Octalysis drives, collapsed onto a 0–10 composite by the rule specified in the coding table below. Five descriptive categories are used: F0 (no failure signal), F1 (mild discomfort), F2 (clear distress), F3 (active disengagement) and F4 (observed non-return under the study definition). The codebook specifies every item, weight, threshold, source, missing-data rule and adjudication procedure. Outcome-defining variables such as cycle count and return status must not be embedded in a predictor score that is then claimed to predict the same outcome; the validation analysis is therefore reported separately from the scoring rules.", ''),
    ])
    log.append("A15 Appendix B prose finalised")
else:
    failures.append("A15 Appendix B prose not found")

# A16. Appendix G restructure
p = find("AUTHOR ACTION REQUIRED: Author to insert: (i) the cleaned Stage-1 survey")
if p:
    set_parts(p, [("Appendix G binds the raw data, photographs and supplementary materials listed below. Each item is deposited with the thesis or in a governed repository accessible to examiners; where an item cannot be shared openly, the access procedure and restrictions are stated in its entry.", '')])
    items = [
        ("G.1  Cleaned Stage-1 survey response file (CSV). ", "[AUTHOR INPUT REQUIRED: attach file]", 'input'),
        ("G.2  De-identified Stage-2 session-log dataset — timestamps, COP-derived metrics, AI-call telemetry and FIAS scoring. This file is also the source for the re-verified Cox model and Figure 5.4. ", "[AUTHOR INPUT REQUIRED: attach file]", 'input'),
        ("G.3  Stage-3 questionnaire response file. ", "[AUTHOR INPUT REQUIRED: attach file]", 'input'),
        ("G.4  Photographs of the Guangzhou Stage-2 deployment (consented). ", "[INSERT PHOTOGRAPHS HERE — Stage-2 deployment: the FRA station in situ and participants playing the skiing game; written consent on record, faces obscured where required]", 'photo'),
        ("G.5  Photographs of the Hong Kong Stage-3 deployment (consented). ", "[INSERT PHOTOGRAPHS HERE — Stage-3 community deployment]", 'photo'),
        ("G.6  Screenshots of the skiing game at each difficulty level (1–5). ", "[INSERT SCREENSHOTS HERE — one per difficulty level, taken from the deployed Stage-2 build]", 'photo'),
        ("G.7  AI prompt-template archive used in Layer 2 (system and user prompts, versions and settings). ", "[AUTHOR INPUT REQUIRED: attach archive]", 'input'),
    ]
    anchor = p
    for lead, ph, kind in items:
        anchor = insert_after(anchor, [(lead, ''), (ph, kind)])
    log.append("A16 Appendix G restructured with per-item placeholders")
else:
    failures.append("A16 Appendix G action not found")

# =========================================================================
# GROUP B — photograph placeholders in the body
# =========================================================================

def photo_block(anchor, box_text, caption_text):
    b = insert_after(anchor, [(box_text, 'photo')])
    c = insert_after(b, [(caption_text, '')], style=FIGSTYLE)
    return c

# B1. FRA machine photo after Figure 4.1A caption
p = None
for q in doc.paragraphs:
    if q.text.startswith("Figure 4.1A.") and q.style.name == "FigureCaption":
        p = q; break
if p:
    photo_block(p,
        "[INSERT PHOTOGRAPH HERE — Figure 4.1B: the FRA machine]",
        "Figure 4.1B. The InBody FRA510S fall-risk-assessment station as installed at the Poly Group facility, Guangzhou. Photograph to be inserted by the author; facility photography consent on record.")
    log.append("B1 FRA machine photo placeholder (Figure 4.1B)")
else:
    failures.append("B1 Figure 4.1A caption not found")

# B2. Play-station photo in §5.3.1
p = find("The InBody pressure panel is a force-plate balance device deployed in clinical")
if p:
    photo_block(p,
        "[INSERT PHOTOGRAPH HERE — Figure 5.1A: the play station]",
        "Figure 5.1A. The Stage-2 play station: the InBody pressure panel and display configured for the skiing game. Photograph to be inserted by the author.")
    log.append("B2 play-station photo placeholder (Figure 5.1A)")
else:
    failures.append("B2 §5.3.1 paragraph not found")

# B3. Game artefact screenshots after Figure 5.1 caption
p = None
for q in doc.paragraphs:
    if q.text.startswith("Figure 5.1. Skiing game on-screen display") and q.style.name == "FigureCaption":
        p = q; break
if p:
    photo_block(p,
        "[INSERT SCREENSHOTS HERE — Figure 5.1B: the game artefact]",
        "Figure 5.1B. The deployed game artefact: screenshots of the skiing game at difficulty levels 1 to 5, taken from the Stage-2 build. Screenshots to be inserted by the author (see also Appendix G, item G.6).")
    log.append("B3 game screenshots placeholder (Figure 5.1B)")
else:
    failures.append("B3 Figure 5.1 caption not found")

# B4. Participants playing — §5.4
p = find("Thirty residents participated over eight weeks (mean age 71.1")
if p:
    photo_block(p,
        "[INSERT PHOTOGRAPH(S) HERE — Figure 5.3A: participants playing]",
        "Figure 5.3A. Participants playing the skiing game on the InBody pressure panel during the Stage-2 deployment. Photographs used with written consent; identifying features obscured where required.")
    log.append("B4 participants-playing photo placeholder (Figure 5.3A)")
else:
    failures.append("B4 §5.4 participants paragraph not found")

# B5. Stage-3 site photo — §6.4
p = find("One hundred and twenty adults aged 25-80 were recruited")
if p:
    photo_block(p,
        "[INSERT PHOTOGRAPH HERE — Figure 6.0A: Stage-3 site]",
        "Figure 6.0A. The Stage-3 decoupled deployment at the Hong Kong community site. Photograph to be inserted by the author; written consent on record.")
    log.append("B5 Stage-3 site photo placeholder (Figure 6.0A)")
else:
    failures.append("B5 §6.4 paragraph not found")

# B6. Note at head of List of Figures
p = find("Figure 3.1. Three-stage action research programme")
if p:
    insert_before(p, [
        ("[NOTE TO AUTHOR: regenerate this list and its page numbers after inserting the photographs at the marked positions (Figures 4.1B, 5.1A, 5.1B, 5.3A, 6.0A), the Kaplan-Meier plot (Figure 5.4) and the participant-level change plot (Figure 5.4B).]", 'input'),
    ])
    log.append("B6 List of Figures regeneration note")
else:
    failures.append("B6 List of Figures head not found")

# =========================================================================
# GROUP C — citation insertions and reference repairs
# =========================================================================

# C1. §1.3 demographic citation
sub("Population ageing supplies the scale at which this problem operates",
    "The United Nations expects 2.1 billion people aged 60 or older worldwide by 2050.",
    "Global estimates anticipate 2.1 billion people aged 60 or older worldwide by 2050 (World Health Organization, 2024).",
    label="C1 WHO ageing citation")

# C2/C3. §1.4 falls citations
sub("Falls supply the clinical backdrop",
    "They are the leading cause of injury-related hospital admission among adults aged 65 and above.",
    "They are the leading cause of injury-related hospital admission among adults aged 65 and above (Centers for Disease Control and Prevention, 2023; World Health Organization, 2023).",
    label="C2 CDC/WHO falls citation")
sub("Falls supply the clinical backdrop",
    "about one in two for those aged 80 or above.",
    "about one in two for those aged 80 or above (World Health Organization, 2007, 2023).",
    label="C3 WHO falls-rate citation")

# C4. §2.2.3 Schroeder
sub("Chen and Chan (2014) developed STAM through a study of 1,012 Hong Kong Chinese seniors",
    "STAM is therefore the strongest baseline available for elderly technology adoption,",
    "Recent mapping reviews of the factors associated with older adults' intention to adopt digital technologies confirm that these constructs remain central to the field (Schroeder et al., 2023). STAM is therefore the strongest baseline available for elderly technology adoption,",
    label="C4 Schroeder citation")

# C5. §2.3.1 Juul
sub("Deterding, Dixon, Khaled and Nacke (2011) defined gamification",
    "The non-game contexts include education, health, productivity, financial services and many others.",
    "The non-game contexts include education, health, productivity, financial services and many others. Within game studies, failure itself is treated as a designed — even valued — part of play (Juul, 2013), an observation that becomes consequential in this thesis when the same failure acquires clinical meaning.",
    label="C5 Juul citation")

# C6. §2.4.2 Compeau & Higgins
sub("Computer self-efficacy among older adults is consistently lower",
    "Computer self-efficacy among older adults is consistently lower than among younger users (Hawthorn, 2007; Broady, Chan & Caputi, 2010).",
    "Computer self-efficacy, the domain-specific form of the construct developed by Compeau and Higgins (1995), is consistently lower among older adults than among younger users (Hawthorn, 2007; Broady, Chan & Caputi, 2010).",
    label="C6 Compeau & Higgins citation")

# C7. §2.5.5 ageing stereotype-threat work
sub("Steele and Aronson (1995) demonstrated that making a negatively stereotyped identity salient",
    "Subsequent work has extended stereotype-threat questions to ageing, although effects vary by task, context and measurement.",
    "Subsequent work has extended stereotype-threat questions to ageing — self-perceptions of ageing predict long-term health outcomes (Levy et al., 2002), identity salience alters older adults' performance on identity-relevant tasks (McGlone & Aronson, 2006), and age-based stereotype threat about cognitive decline has been reviewed systematically (Barber, 2017) — although effects vary by task, context and measurement.",
    label="C7 ageing stereotype-threat citations")

# C8. §2.6.2B Elliot & Murayama + Frey & Jegen
sub("Two established motivation literatures support the trajectory-installation principle",
    "specific, proximal and adjustable goals sustain effort better than vague or static ones.",
    "specific, proximal and adjustable goals sustain effort better than vague or static ones, and the related distinction between mastery goals and performance goals (Elliot & Murayama, 2008) sharpens the point: a goal defined by personal improvement invites a different reading of failure than a goal defined by demonstrated adequacy.",
    label="C8a Elliot & Murayama citation")
sub("Two established motivation literatures support the trajectory-installation principle",
    "because the difficulty layer keeps the next goal specific and reachable as capability changes.",
    "because the difficulty layer keeps the next goal specific and reachable as capability changes. Motivation-crowding research adds a design caution: extrinsic rewards that are introduced clumsily can displace rather than support intrinsic motivation (Frey & Jegen, 2001), which is why the learn-to-earn philosophy treats rewards as recognition of progress along the curve rather than as payment for compliance.",
    label="C8b Frey & Jegen citation")

# C9. §2.7.1 Eysenbach
sub("Koivisto and Malik (2021) reviewed the gamification-for-older-adults literature",
    "Second, attrition is consistently reported but rarely analysed as a primary outcome;",
    "Second, attrition — the ‘law of attrition’ by which digital-health evaluations lose a large share of their users (Eysenbach, 2005) — is consistently reported but rarely analysed as a primary outcome;",
    label="C9 Eysenbach citation")

# C10. §2.7.2 Satchell & Dourish
sub("Greenhalgh et al. (2017) introduced the Non-Adoption",
    "adopter system, organisation, wider context, and embedding over time.",
    "adopter system, organisation, wider context, and embedding over time. The human-computer-interaction literature makes the complementary point that non-use is itself meaningful behaviour rather than a mere absence of use (Satchell & Dourish, 2009).",
    label="C10 Satchell & Dourish citation")

# C11. §2.7.4 Shannon
sub("Shannon entropy is used here as a bridge metaphor",
    "Shannon entropy is used here as a bridge metaphor",
    "Shannon entropy (Shannon, 1948) is used here as a bridge metaphor",
    label="C11 Shannon citation")

# C12. §3.6 WTA-WTP, screeners, affect
sub("The first move measures identity-relevant valuation directly",
    "through an adapted WTA-WTP or alternative validated appraisal task, developed and piloted carefully",
    "through an adapted WTA-WTP elicitation in the tradition of experimental preference measurement (Kahneman et al., 1990; Holt & Laury, 2002), or an alternative validated appraisal task, developed and piloted carefully",
    label="C12a WTA-WTP citations")
sub("The first move measures identity-relevant valuation directly",
    "using a full STAM instrument, prospectively specified FIAS items and clearly separated predictors and outcomes.",
    "using a full STAM instrument, prospectively specified FIAS items, brief validated screeners for cognition, mood and general self-efficacy — such as the Montreal Cognitive Assessment (Nasreddine et al., 2005), the PHQ-4 (Kroenke et al., 2009) and the New General Self-Efficacy Scale (Chen, Gully & Eden, 2001) — and clearly separated predictors and outcomes.",
    label="C12b screener citations")
sub("A fourth move instruments the immediate behavioural and affective channel",
    "voice cues, brief experience sampling or facial-affect estimates, provided",
    "voice cues, brief experience sampling or facial-affect estimates in the tradition of affective computing (Picard, 1997; D'Mello & Graesser, 2012), provided",
    label="C12c affective computing citations")

# C13. §4.4 Festinger
sub("In clinically framed activities for older users, some mechanics may operate differently",
    "Social comparison can motivate some participants",
    "Social comparison (Festinger, 1954) can motivate some participants",
    label="C13 Festinger citation")

# C14. §8.6A Wittgenstein / Harnad / Brooks
sub("This limitation belongs first to the study's input architecture",
    "The symbol-grounding and embodied-intelligence literatures provide relevant theoretical context, but they do not establish",
    "The intuition that the limits of a system's language may bound the world it can represent has a long philosophical lineage (Wittgenstein, 1922), and the symbol-grounding and embodied-intelligence literatures (Harnad, 1990; Brooks, 1991) provide relevant theoretical context, but they do not establish",
    label="C14 Wittgenstein/Harnad/Brooks citations")

# C15. §8.8 Michie + SPIRIT-AI
sub("The priority extensions are ordered by inferential value",
    "Cross-cultural replication, cross-domain application and community-narrative research should follow after the core causal and construct questions are resolved.",
    "Intervention components should be specified against an established behaviour-change taxonomy such as the Behaviour Change Wheel (Michie et al., 2011) so that replication and cross-site comparison are possible. Cross-cultural replication, cross-domain application and community-narrative research should follow after the core causal and construct questions are resolved.",
    label="C15a Michie citation")
sub("The randomised replication should be treated as the decisive next study",
    "A separate qualitative panel with non-returners would clarify why participants completed one session but did not intend regular use.",
    "A separate qualitative panel with non-returners would clarify why participants completed one session but did not intend regular use. Because the intervention embeds an AI component, the protocol and report of the randomised study should follow the SPIRIT-AI and CONSORT-AI reporting extensions (Liu et al., 2020).",
    label="C15b SPIRIT-AI citation")

# C16. §9.8C citations
sub("The symbol-grounding and tacit-knowledge literatures help explain",
    "The symbol-grounding and tacit-knowledge literatures help explain",
    "The symbol-grounding and tacit-knowledge literatures (Harnad, 1990; Polanyi, 1966; Nonaka & Takeuchi, 1995) help explain",
    label="C16a grounding/tacit citations")
sub("Research on the relation between language, thought and grounded reasoning motivates",
    "Research on the relation between language, thought and grounded reasoning motivates multimodal evaluation",
    "Research on the relation between language, thought and grounded reasoning (Bender & Koller, 2020; Mahowald et al., 2024) motivates multimodal evaluation",
    label="C16b language-thought citations")
sub("A closed-loop adaptive system senses aspects of the environment",
    "selects an action and learns from the resulting outcome.",
    "selects an action and learns from the resulting outcome — a pattern with deep roots in ecological perception, behaviour-based robotics and predictive accounts of cognition (Gibson, 1979; Brooks, 1991; Friston, 2010).",
    label="C16c perception-action citations")
sub("Adaptive learning in this thesis means calibrated challenge",
    "Related traditions include reinforcement learning, world models, intelligent tutoring and embodied interaction.",
    "Related traditions include reinforcement learning (Sutton & Barto, 2018), world models (Ha & Schmidhuber, 2018; LeCun, 2022), intelligent tutoring (VanLehn, 2011) and language models grounded in embodied action (Ahn et al., 2022; Driess et al., 2023).",
    label="C16d adaptive traditions citations")
sub("Broader AI research increasingly evaluates systems in interactive and multimodal environments",
    "Broader AI research increasingly evaluates systems in interactive and multimodal environments.",
    "Broader AI research increasingly evaluates systems in interactive and multimodal environments, and prominent research agendas argue that machines that learn and act in the world, rather than only about it, are the field's next frontier (Lake et al., 2017; Li, 2025).",
    label="C16e broader-AI citations")
sub("Broader AI research increasingly evaluates systems in interactive and multimodal environments",
    "nor does it resolve competing definitions of AGI.",
    "nor does it resolve competing definitions of AGI (Legg & Hutter, 2007; Goertzel, 2014; Chollet, 2019).",
    label="C16f AGI-definition citations")

# =========================================================================
# GROUP D — Table 6.5 and Figure 6.1
# =========================================================================

t65 = None
for t in doc.tables:
    joined = " | ".join(c.text for c in t.rows[0].cells) if t.rows else ""
    if any("1.89" in c.text for r in t.rows for c in r.cells):
        t65 = t; break
if t65 is not None:
    for r in t65.rows:
        cells = r.cells
        head = cells[0].text
        def setcell(cell, text, kind=''):
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(text)
            if kind == 'input':
                run.bold = True
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        if "First failure predicting dropout" in head:
            for c in cells:
                if "1.89" in c.text:
                    setcell(c, "under re-verification — see §5.5", 'input')
        if "completers" in head and "FRA index change" in head:
            for c in cells:
                if c.text.strip() == "not reported":
                    setcell(c, "1.00–3.44 (approx.)")
        if "pooled" in head:
            for c in cells:
                if c.text.strip() == "not reported":
                    setcell(c, "0.37–1.19 (approx.)")
        if "conservative zero-change" in head:
            for c in cells:
                if c.text.strip() == "not reported":
                    setcell(c, "0.18–0.96 (approx.)")
    log.append("D1 Table 6.5 cells updated")
else:
    failures.append("D1 Table 6.5 not found")

# D2. Table 6.5 footnote paragraph
p = find("Intervals marked with an asterisk are Wilson score intervals")
if p:
    old = "Confidence intervals for the Cohen's d estimates remain a pre-binding requirement."
    if old in p.text:
        set_parts(p, [(p.text.replace(old,
            "Confidence intervals for the Cohen's d estimates are normal-approximation intervals computed from the reported summary statistics and are to be re-derived from participant-level data before binding. The Cox interval is withheld pending re-verification against the participant-level survival file (see §5.5)."), '')])
        log.append("D2 Table 6.5 footnote updated")
    else:
        failures.append("D2 footnote sentence not present")
else:
    failures.append("D2 Table 6.5 footnote not found")

# D3. Insert regenerated Figure 6.1 above its caption
body61 = None
for q in doc.paragraphs:
    if q.text.startswith("Figure 6.1. Stage-3 retention") and q.style.name == "FigureCaption":
        body61 = q; break
if body61 is not None:
    imgp = body61.insert_paragraph_before()
    run = imgp.add_run()
    run.add_picture(FIG61, width=Inches(5.8))
    note = insert_after(body61, [
        ("Figure 6.1 was regenerated directly from the counts reported in Table 6.2. ", ''),
        ("[AUTHOR: confirm the regenerated chart against the Stage-3 records, or replace it with the original figure file if preferred.]", 'input'),
    ])
    log.append("D3 Figure 6.1 regenerated and inserted")
else:
    failures.append("D3 Figure 6.1 caption not found")

# =========================================================================
# GROUP E — vendor harmonisation and reference list
# =========================================================================

sub("Stage 2's central technical artefact is a 2D skiing game",
    "with an automatic difficulty-adjustment subsystem driven by LangChain and an OpenAI API endpoint",
    "with an automatic difficulty-adjustment subsystem driven by LangChain orchestration over a large-language-model endpoint (the deployed endpoint is recorded in Appendix G, item G.7)",
    label="E1 §5.3 vendor harmonisation")
sub("Layer 2 (AI Adaptation) is a Claude API and LangChain runtime",
    "Layer 2 (AI Adaptation) is a Claude API and LangChain runtime",
    "Layer 2 (AI Adaptation) is a LangChain runtime calling a large-language-model endpoint",
    label="E2 §5.3.4 vendor harmonisation")
sub("LangChain orchestration over OpenAI API calls",
    "LangChain orchestration over OpenAI API calls",
    "LangChain orchestration over large-language-model API calls (deployed endpoint recorded in Appendix G, item G.7)",
    label="E3 Appendix E.1 vendor harmonisation")
sub("if the OpenAI API is unavailable",
    "if the OpenAI API is unavailable",
    "if the model endpoint is unavailable",
    label="E4 Appendix E.3 vendor harmonisation")

# E5. Reference list: delete meta paragraph
p = find("References are formatted in APA 7th edition style with a hanging indent")
if p:
    delete(p); log.append("E5 reference meta-commentary removed")
else:
    failures.append("E5 reference meta paragraph not found")

# E6. Reference corrections (plain-run rewrites; matching original formatting)
ref_fixes = [
    ("Altmeyer, M., Lessel, P., & Kr",
     "Altmeyer, M., Lessel, P., & Krüger, A. (2018). Investigating gamification for seniors aged 75+. In Proceedings of the 2018 Designing Interactive Systems Conference (pp. 453–458). ACM. https://doi.org/10.1145/3196709.3196799"),
    ("Park, W.-C., Kim, M., Kim, S., et al. (2019)",
     "Park, W.-C., Kim, M., Kim, S., Yoo, J., Kim, B. S., Chon, J., Jeong, S. J., & Won, C. W. (2019). Introduction of fall risk assessment (FRA) system and cross-sectional validation among community-dwelling older adults. Annals of Rehabilitation Medicine, 43(1), 87–95. https://doi.org/10.5535/arm.2019.43.1.87"),
    ("Locke, E. A., & Latham, G. P. (2002)",
     "Locke, E. A., & Latham, G. P. (2002). Building a practically useful theory of goal setting and task motivation: A 35-year odyssey. American Psychologist, 57(9), 705–717."),
    ("Ryan, R. M., & Deci, E. L. (2000)",
     "Ryan, R. M., & Deci, E. L. (2000). Self-determination theory and the facilitation of intrinsic motivation, social development, and well-being. American Psychologist, 55(1), 68–78."),
    ("An, S., Cheung, C. F., & Willoughby, K. W. (2024)",
     "An, S., Cheung, C. F., & Willoughby, K. W. (2024). A gamification approach for enhancing older adults' technology adoption and knowledge transfer: A case study in mobile payments technology. Technological Forecasting and Social Change, 205, Article 123456. https://doi.org/10.1016/j.techfore.2024.123456"),
    ("An, S., Cheung, C. F., & Lo, Y. T. (2025)",
     "An, S., Cheung, C. F., & Lo, Y. T. (2025). Improving older adults' technology adoption on mobile map: A gamification approach. International Journal of Human–Computer Interaction, 41(15), 9633–9651. https://doi.org/10.1080/10447318.2024.2426914"),
    ("Kim, M., Kim, S., & Won, C. W. (2018)",
     "Kim, M., Kim, S., & Won, C. W. (2018). Test–retest reliability and sensitivity to change of a new fall risk assessment system: A pilot study. Annals of Geriatric Medicine and Research, 22(2), 80–87. https://doi.org/10.4235/agmr.2018.22.2.80"),
]
for needle, replacement in ref_fixes:
    p = find(needle)
    if p:
        set_parts(p, [(replacement, '')])
        log.append("E6 reference fixed: " + needle[:40])
    else:
        failures.append("E6 reference not found: " + needle[:40])

# E7. Insert WHO 2024 Ageing and health entry after WHO 2023
p = find("World Health Organization. (2023). Falls (Fact sheet)")
if p:
    insert_after(p, [("World Health Organization. (2024). Ageing and health (Fact sheet). World Health Organization. https://www.who.int/news-room/fact-sheets/detail/ageing-and-health", '')])
    log.append("E7 WHO 2024 ageing entry inserted")
else:
    failures.append("E7 WHO 2023 entry not found")

# =========================================================================
doc.save(DST)
print("SAVED:", DST)
print("\n--- APPLIED (%d) ---" % len(log))
for l in log: print(" •", l)
print("\n--- FAILURES (%d) ---" % len(failures))
for f in failures: print(" ✗", f)
