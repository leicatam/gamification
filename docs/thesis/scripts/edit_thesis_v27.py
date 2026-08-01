#!/usr/bin/env python3
"""V27 — two Stage-3 design rationales supplied by the author (1 Aug 2026):
(1) the Stage-2 facility was deliberately not re-used for the decoupled
deployment, to avoid carryover of residents' prior experience of the
coupled system; (2) distribution went beyond the author's immediate
circle via onward forwarding (snowball distribution), so a portion of
participants were unknown to the author."""
import shutil
import docx
from docx.oxml import OxmlElement

REPO = "/home/user/gamification/docs/thesis/"
SRC = REPO + "Full_version_V26_EngD_Thesis_Submission_Draft.docx"
DST = REPO + "Full_version_V27_EngD_Thesis_Submission_Draft.docx"
shutil.copy(SRC, DST)
d = docx.Document(DST)

anchor = None
for p in d.paragraphs:
    if "48 were male and 72 female, and all reported high-school education or above" in p.text:
        anchor = p
        break
assert anchor is not None, "6.4 cohort anchor missing"

new = OxmlElement("w:p")
anchor._p.addnext(new)
np_ = docx.text.paragraph.Paragraph(new, anchor._parent)
np_.add_run(
    "Two recruitment decisions were deliberate. First, the Stage-2 "
    "facility was not re-used for the decoupled deployment: its "
    "residents already carried direct experience of the coupled system, "
    "and their prior exposure and attitudes would have contaminated a "
    "test of how the decoupled frame is received on first contact. A "
    "separate cohort was therefore a design choice, not only a "
    "constraint. Second, distribution deliberately extended beyond the "
    "author's immediate circle: the organisations and contacts who "
    "received the build forwarded it onward, so that a portion of "
    "participants were personally unknown to the author. This onward "
    "(snowball) distribution reduces, though it does not eliminate, the "
    "risk that participation reflected personal obligation to the "
    "researcher.")

cp = d.core_properties
cp.comments = "V27: Stage-3 recruitment rationale (no facility re-use; snowball distribution) added to 6.4."
d.save(DST)
print("V27 saved")
