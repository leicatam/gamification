#!/usr/bin/env python3
"""V25 — §3.4: Stage-3 consent procedure stated per the author's account
(25 Jul 2026): informed, implied consent in the remote deployment — the
distribution e-mails explained the purpose of the game and questionnaire
and the use of replies; participation was voluntary and unpaid; returning
the results file and responses constituted consent; no written consent
instrument was administered (acknowledged as a limitation)."""
import shutil
import docx
from docx.oxml import OxmlElement

REPO = "/home/user/gamification/docs/thesis/"
SRC = REPO + "Full_version_V24_EngD_Thesis_Submission_Draft.docx"
DST = REPO + "Full_version_V25_EngD_Thesis_Submission_Draft.docx"
shutil.copy(SRC, DST)
d = docx.Document(DST)

anchor = None
for p in d.paragraphs:
    if "Each stage was conducted under the ethics policy" in p.text:
        anchor = p
        break
assert anchor is not None, "3.4 anchor not found"

new = OxmlElement("w:p")
anchor._p.addnext(new)
np_ = docx.text.paragraph.Paragraph(new, anchor._parent)
np_.add_run(
    "Consent in the Stage-3 remote deployment operated by informed, "
    "implied consent, and the thesis states the procedure exactly as it "
    "was. The distribution e-mails explained the purpose of the game and "
    "of the accompanying questionnaire and the use that would be made of "
    "replies; participation was voluntary and unpaid, and recipients "
    "could decline by simply not taking part; returning the game's "
    "results file and the questionnaire responses constituted consent to "
    "their use in the research. No written consent instrument was "
    "administered in Stage 3, and this is acknowledged as a limitation "
    "of the remote design. Identities were separated from the analysed "
    "records at compilation: the analysis sheets retain only a "
    "participant code, age band, gender and education, and the "
    "distribution and return e-mails are retained in the author's "
    "records, with a de-identified index forming Appendix G, item G.5.")
print("3.4 consent paragraph inserted")

cp = d.core_properties
cp.comments = "V25: Stage-3 implied-consent procedure stated in 3.4."
d.save(DST)
print("saved", DST)
