#!/usr/bin/env python3
"""V29 — Appendix H gains H.5: the semi-structured interview guide
actually used for the expert consultation (author's file
Expert_Interview_List_2026.docx, archived in evidence/), reproduced in
compact form. Closes the outstanding 'question list' item for
Appendix H."""
import shutil
import docx
from docx.oxml import OxmlElement

REPO = "/home/user/gamification/docs/thesis/"
SRC = REPO + "Full_version_V28_EngD_Thesis_Submission_Draft.docx"
DST = REPO + "Full_version_V29_EngD_Thesis_Submission_Draft.docx"
shutil.copy(SRC, DST)
d = docx.Document(DST)

anchor = None
for p in d.paragraphs:
    if p.text.strip().startswith("H.4  Use in the thesis."):
        anchor = p
        break
assert anchor is not None, "H.4 anchor missing"


def insert_after(p, text):
    new = OxmlElement("w:p")
    p._p.addnext(new)
    np_ = docx.text.paragraph.Paragraph(new, p._parent)
    np_.add_run(text)
    return np_


anchor = insert_after(anchor,
    "H.5  Interview guide (as used). The consultation followed a "
    "semi-structured guide (45-60 minutes; archived as "
    "Expert_Interview_Guide_2026.docx, Appendix G item G.8), designed to "
    "elicit the expert's independent reading before study specifics were "
    "revealed. Its sections and questions, in compact form: "
    "A. Background — the expert's role and what their discipline sees "
    "about falls that other fields miss. "
    "B. The engagement problem — after hearing the field observation "
    "(adverse first experience, non-return even where measured "
    "performance improved): does the pattern match their experience; "
    "what is happening inside the non-return; why does the first "
    "experience carry so much weight. "
    "C. Framing and purpose — after hearing the decoupled-deployment "
    "result: why would removing the clinical frame change behaviour; is "
    "a \"test\" socially different from a \"game\" for this generation. "
    "D. The commitment gap — what would explain high session completion "
    "with low regular-use intention; what builds the habit rather than "
    "the visit. "
    "E. The social dimension — asked open-endedly BEFORE revealing that "
    "study play was individual: does alone-versus-together change the "
    "experience; would group play change engagement and in which "
    "direction; when would social play backfire; how does culture shape "
    "the answer. "
    "F. Adaptive learning and design — reaction to the adaptive-learning "
    "proposal; through what channel adaptation would actually operate "
    "(private mastery or between people); one measure they would add to "
    "a follow-up study. "
    "G. Closing — anything not covered; other perspectives to seek; "
    "confirmation of attribution preference and quote-approval contact. "
    "The guide's post-interview checklist requires the record to be "
    "written within 48 hours, the consent filed, and attribution "
    "sentences approved before binding.")

cp = d.core_properties
cp.comments = "V29: Appendix H.5 interview guide added (from the archived guide file)."
d.save(DST)
print("V29 saved with H.5")
