#!/usr/bin/env python3
"""V30 — commercial cross-reference: GrandPad as an industrial comparator
(Ch8, after the unit-of-innovation paragraph), with verified citations:
Perissinotto et al. (2019, Innovation in Aging conference abstract;
GrandPad telemedicine feasibility, N=21), Tha et al. (2025, Journal of
Applied Gerontology; two GrandPad implementation case studies, N=48),
and Gadbois et al. (2022, Innovation in Aging; supported tablet
onboarding for homebound elders — adjacent evidence, device not named).
References added alphabetically (incl. vendor material)."""
import shutil
import docx
from docx.oxml import OxmlElement

REPO = "/home/user/gamification/docs/thesis/"
SRC = REPO + "Full_version_V29_EngD_Thesis_Submission_Draft.docx"
DST = REPO + "Full_version_V30_EngD_Thesis_Submission_Draft.docx"
shutil.copy(SRC, DST)
d = docx.Document(DST)
log, fail = [], []


def insert_after(p, text, style=None):
    new = OxmlElement("w:p")
    p._p.addnext(new)
    np_ = docx.text.paragraph.Paragraph(new, p._parent)
    if style:
        np_.style = style
    np_.add_run(text)
    return np_


# (1) comparator paragraph after the unit-of-innovation paragraph
anchor = None
for p in d.paragraphs:
    if "The unit of innovation is the interface between an existing technology and the person who decides whether to return" in p.text:
        anchor = p
        break
if anchor is None:
    fail.append("unit-of-innovation anchor")
else:
    insert_after(anchor,
        "A commercial cross-reference suggests the market has "
        "independently converged on parts of this design space. GrandPad "
        "— a purpose-built tablet for older adults, sold in the United "
        "States on a subscription basis on the order of US$500 per year, "
        "with cellular connectivity built in so that no home broadband or "
        "technical set-up is required — packages games, video calling and "
        "photo sharing behind a radically simplified interface, with "
        "family members and a human support service connected around the "
        "user (GrandPad, n.d.). Peer-reviewed evaluations of its "
        "deployments report high perceived ease of use and connectedness "
        "among predominantly homebound users where the device filled a "
        "genuine access gap, and markedly weaker adoption where "
        "mainstream alternatives already existed — implementation "
        "context, not device quality, decided uptake (Perissinotto et "
        "al., 2019; Tha et al., 2025); adjacent trials of supported "
        "tablet onboarding for homebound elders likewise report increased "
        "technology use and directional loneliness reductions when a "
        "human scaffold accompanies the device (Gadbois et al., 2022). "
        "Three features of the model align with this programme's "
        "findings: access and set-up barriers are engineered away "
        "wholesale, consistent with the Stage-1 observation that their "
        "absence alone does not secure engagement; the games are framed "
        "as leisure and connection, never as assessment — a decoupled "
        "frame arrived at by commercial instinct; and the engagement "
        "scaffold is explicitly social, anticipating the group-dynamics "
        "mechanism raised in the §6.8A consultation. What the model does "
        "not contain is the layer this thesis contributes: GrandPad "
        "avoids threatening frames by carrying no clinical measurement at "
        "all, whereas the present programme concerns technologies that "
        "must measure while remaining bearable. The comparator therefore "
        "marks a boundary in the design space: engagement without "
        "clinical purpose is commercially solved; engagement with "
        "clinical purpose intact is the open problem the "
        "adaptive-learning framework addresses.")
    log.append("comparator paragraph inserted")

# (2) references, alphabetical
REFS = [
 ("Gadbois",
  "Gadbois, E. A., Jimenez, F., Brazier, J. F., Davoodi, N. M., Nunn, "
  "A. S., Mills, W. L., Dosa, D., & Thomas, K. S. (2022). Findings from "
  "Talking Tech: A technology training pilot intervention to reduce "
  "loneliness and social isolation among homebound older adults. "
  "Innovation in Aging, 6(5), igac040. "
  "https://doi.org/10.1093/geroni/igac040"),
 ("GrandPad",
  "GrandPad. (n.d.). GrandPad: The tablet for seniors [Product and "
  "subscription information]. GrandPad, Inc. https://www.grandpad.net"),
 ("Perissinotto",
  "Perissinotto, C., Zhang, C., Oseau, T., Balik, D., Sou, C., "
  "Burnight, C., & Burnight, K. (2019). Feasibility of a tablet "
  "designed for older adults to facilitate telemedicine visits "
  "[Conference abstract]. Innovation in Aging, 3(Suppl. 1), S975. "
  "https://doi.org/10.1093/geroni/igz038.3534"),
 ("Tha",
  "Tha, S. H., Hough, K., Bui, N., Dulaney, S., & Perissinotto, C. "
  "(2025). The use of age-friendly technology in the care of older "
  "adults: Two implementation case studies. Journal of Applied "
  "Gerontology. https://doi.org/10.1177/07334648251343513"),
]
paras = d.paragraphs
ref_i = next(i for i, p in enumerate(paras)
             if p.text.strip() == "References"
             and p.style.name.startswith("Heading") and i > len(paras)//2)
end_i = next(j for j in range(ref_i+1, len(paras))
             if paras[j].style.name.startswith("Heading"))
for key, entry in REFS:
    target = None
    for j in range(ref_i+1, end_i):
        t = paras[j].text.strip()
        if t and t > key:
            target = j
            break
    prevp = paras[target-1] if target else paras[end_i-1]
    new = OxmlElement("w:p")
    prevp._p.addnext(new)
    np_ = docx.text.paragraph.Paragraph(new, prevp._parent)
    np_.style = paras[target].style if target else prevp.style
    np_.add_run(entry)
    # refresh paragraph list after each insertion
    paras = d.paragraphs
    end_i = next(j for j in range(ref_i+1, len(paras))
                 if paras[j].style.name.startswith("Heading"))
    log.append(f"ref added: {key}")

cp = d.core_properties
cp.comments = "V30: GrandPad commercial cross-reference (Ch8) with verified citations."
d.save(DST)
print("LOG:", *log, sep="\n  ")
print("FAIL:", fail if fail else "none")
