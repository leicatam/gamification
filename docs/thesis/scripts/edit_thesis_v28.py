#!/usr/bin/env python3
"""V28 — Stage-2 qualitative record corrected per the author (1 Aug 2026):
the nine-question F4 interview protocol was DESIGNED but NOT administered
(the F4 residents' negative reaction made a research interview
inappropriate at the time). What actually happened: (i) an after-game
paper questionnaire completed by all thirty participants; (ii) intensive
post-deployment interviews with the persistent participants (those who
stayed the eight weeks, failed and returned, with positive recorded FRA
change). Appendix C, §5.6 and the ToC entry are corrected."""
import shutil
import docx

REPO = "/home/user/gamification/docs/thesis/"
SRC = REPO + "Full_version_V27_EngD_Thesis_Submission_Draft.docx"
DST = REPO + "Full_version_V28_EngD_Thesis_Submission_Draft.docx"
shutil.copy(SRC, DST)
d = docx.Document(DST)
log, fail = [], []


def set_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in list(p.runs[1:]):
            r._element.getparent().remove(r._element)
    else:
        p.add_run(text)


def replace(tag, needle, new_text, style_prefix=None):
    for p in d.paragraphs:
        if needle in p.text and (style_prefix is None
                                 or p.style.name.startswith(style_prefix)):
            set_text(p, new_text)
            log.append(tag)
            return
    fail.append(tag)


# (1) Appendix C body
replace("appC-body",
    "A nine-question semi-structured interview was administered to F4 participants",
    "The Stage-2 qualitative record has two genuine components. First, an "
    "after-game paper questionnaire was completed by every participant at "
    "the end of play; it is the source of the satisfaction data reported "
    "in Chapter 5. Second, intensive post-deployment interviews were "
    "conducted with the persistent participants — those who remained "
    "through the eight-week window, experienced failure and returned, and "
    "whose recorded FRA index improved; these conversations informed the "
    "learn-to-earn design orientation of §5.2. In addition, a "
    "nine-question semi-structured interview protocol was designed for "
    "the F4 cluster, to be administered within days of an abandoned "
    "session; it was NOT administered in Stage 2 — the F4 residents' "
    "negative reaction to the experience made a research interview "
    "inappropriate at the time — and no F4 interview data exist. The "
    "protocol below is therefore a designed instrument retained for the "
    "confirmatory programme (§8.8), not a record of interviews conducted.")

# (2) §5.6 sentence
replace("5.6-fix",
    "The F4 interview protocol (Appendix C) was administered to support design reflection",
    "The empty interval between scores 5 and 9 is not read as evidence of "
    "a behavioural discontinuity. Given the construction above, it may be "
    "a mechanical effect of the scoring thresholds and the outcome-linked "
    "items (criterion contamination), or sampling variation in a cohort "
    "of thirty. The item-level codebook — all items, weights, decision "
    "rules and missing-data handling, with cycle count and non-return "
    "excluded from any predictor score that would then be claimed to "
    "predict the same outcome — is specified as a deliverable of the "
    "confirmatory study (§8.8); pending it, the composite serves only as "
    "the descriptive case-identification device described above and "
    "carries no evidential weight. The F4 cluster rests on six "
    "participants, so every mechanism claim based on it is provisional "
    "until two independent coders apply a fixed rubric with reported "
    "agreement and the construct is evaluated in an external sample. A "
    "nine-question F4 interview protocol was designed (Appendix C) but "
    "not administered in Stage 2; no F4 interview data exist, and the "
    "protocol is retained for the confirmatory programme (§8.8). The "
    "Stage-2 qualitative record instead comprises the after-game "
    "questionnaire completed by all participants and post-deployment "
    "interviews with the sustained participants (Appendix C).")

# (3) Appendix C heading + ToC entry
replace("appC-heading",
    "Appendix C · Stage-2 F4 Interview Protocol",
    "Appendix C · Stage-2 Qualitative Record and the F4 Interview "
    "Protocol (Designed, Not Administered)", "Heading")
replace("appC-toc",
    "Appendix C · Stage-2 F4 Interview Protocol",
    "Appendix C · Stage-2 Qualitative Record and the F4 Interview "
    "Protocol (Designed, Not Administered)\t104", "toc")

cp = d.core_properties
cp.comments = "V28: F4 interviews corrected to designed-not-administered; real qualitative record described."
d.save(DST)
print("LOG:", *log, sep="\n  ")
print("FAIL:", fail if fail else "none")
