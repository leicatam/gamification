#!/usr/bin/env python3
"""V18 — final-readiness amendments per reviewer:
(1) Ch1 problem statement (behavioural, not clinical, failure);
(2) FIAS explicitly proposed as a provisional behavioural construct (Ch1);
(3) §7.4 justification of why exactly three principles (appraisal-chain logic);
(4) Stage 3 described as evaluating a participant-facing redesign package
    in a separate cohort (§3.2, §6.3);
(5) Ch8 opening: four contribution registers distinguished + Table 8.0;
(6) Figure 1.1 conceptual hierarchy at end of §1.12 (+ LoF entry).
"""
import shutil
import docx
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt

BASE = "/tmp/claude-0/-home-user-gamification/3d8a50e8-9258-51ce-a7d6-194da50f9a04/scratchpad/"
SRC = BASE + "Full_version_V17_EngD_Thesis_Submission_Draft.docx"
DST = BASE + "Full_version_V18_EngD_Thesis_Submission_Draft.docx"
shutil.copy(SRC, DST)
d = docx.Document(DST)
log, fail = [], []


def set_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in list(p.runs[1:]):
            r._element.getparent().remove(r._element)
    else:
        p.add_run(text)


def find(needle, nth=1):
    n = 0
    for p in d.paragraphs:
        if needle in p.text:
            n += 1
            if n == nth:
                return p
    return None


def insert_after(p, text="", style=None):
    new = OxmlElement("w:p")
    p._p.addnext(new)
    np_ = docx.text.paragraph.Paragraph(new, p._parent)
    if style:
        np_.style = style
    if text:
        np_.add_run(text)
    return np_


# ---------- (1) Ch1 problem statement (end of 1.4) ----------
p = find("Its intervention boundary is behavioural and engineering-focused")
if p is None:
    fail.append("1.4 anchor")
else:
    set_text(p, p.text.rstrip() + " Existing FRA systems fail not because "
        "they are clinically ineffective, but because the behavioural "
        "consequences of early failure are insufficiently considered in "
        "system design. This thesis investigates how adaptive interaction "
        "design can reduce disengagement while preserving clinical validity.")
    log.append("1.4: problem statement appended")

# ---------- (2) FIAS provisional-construct sentence (1.5) ----------
p = find("The later term Failure-Induced Abandonment Syndrome (FIAS) was adopted")
if p is None:
    fail.append("1.5 FIAS anchor")
else:
    set_text(p, p.text.rstrip() + " This thesis proposes FIAS as a "
        "provisional behavioural construct developed from the observed "
        "interaction patterns in this research programme; its definition, "
        "evidential status and limits are developed in Chapters 4 and 5.")
    log.append("1.5: FIAS provisional-construct sentence appended")

# ---------- (3) why exactly three principles (7.4) ----------
anchor = find("The framework therefore combines one empirically supported design direction")
if anchor is None:
    fail.append("7.4 anchor")
else:
    insert_after(anchor,
        "Why exactly three principles, and not two or four? The number and "
        "identity of the principles follow from the structure of the "
        "failure-appraisal chain that the FIAS account identifies, which is "
        "what makes them one framework rather than three recommendations. A "
        "failure event acquires behavioural force in three steps: an event "
        "is generated (a fall in the game, a low score); the event is read "
        "within a frame (as a verdict on a fixed state, or as a point on a "
        "moving curve); and the reading lands on a target (an identity-level "
        "asset, or a revisable skill estimate). Each principle intercepts "
        "exactly one link. Adaptive calibration (P3) governs the event link: "
        "it controls the rate, size and timing of failure events so that "
        "difficulty tracks capability. Trajectory installation (P2) governs "
        "the frame link: it supplies the moving-curve reading under which a "
        "failure is informative rather than conclusive. Identity "
        "disentanglement (P1) governs the target link: it detaches the "
        "reading from the identity-level asset that endowment protects. "
        "Fewer than three principles would leave a link untreated, and the "
        "empirical record indicates what happens then: the Stage-2 build "
        "already carried a partial calibration implementation, yet "
        "non-return occurred while the frame and target links were "
        "untreated; and Stage 3 treated the target link alone, which was "
        "associated with restored post-failure return but left the "
        "commitment gap open. More than three would duplicate a lever "
        "rather than add one: the candidate additions the programme "
        "examined — zero-failure presentation, social scaffolding, reward "
        "scheduling — decompose into presentation-layer or parameter "
        "choices within the three links, which is why zero-failure "
        "presentation appears above as the element P1 and P3 share rather "
        "than as a fourth principle. The framework is therefore jointly "
        "exhaustive over the appraisal chain, individually non-redundant, "
        "and anchored link by link to the observed failure modes.")
    log.append("7.4: why-three paragraph inserted")

# ---------- (4) Stage-3 descriptor ----------
p = find("it does not isolate a single causal variable. Its methodological role is attribution")
if p is None:
    fail.append("3.2 Stage-3 descriptor")
else:
    set_text(p, p.text.replace(
        "it does not isolate a single causal variable.",
        "it evaluates a participant-facing redesign package in a separate "
        "cohort, and it does not isolate a single causal variable."))
    log.append("3.2: redesign-package descriptor adopted")

p = find("The design reduces variation in hardware and game mechanics but does not isolate")
if p is None:
    fail.append("6.3 Stage-3 descriptor")
else:
    set_text(p, "Stage 3 evaluates a participant-facing redesign package in "
        "a separate cohort. " + p.text)
    log.append("6.3: redesign-package descriptor adopted")

# ---------- (5) Ch8 four registers + Table 8.0 ----------
h81 = None
for p in d.paragraphs:
    if p.style.name == "Heading 2" and p.text.strip().startswith("8.1 "):
        h81 = p
        break
if h81 is None:
    fail.append("8.1 heading")
else:
    # insert intro paragraph and caption BEFORE the 8.1 heading
    intro = OxmlElement("w:p")
    h81._p.addprevious(intro)
    pi = docx.text.paragraph.Paragraph(intro, h81._parent)
    pi.add_run(
        "This chapter states the contributions in four distinct registers "
        "and is explicit about which is which, because the four answer to "
        "different tests. A scientific contribution is judged by whether it "
        "explains observed behaviour and survives attempts to refute it; an "
        "engineering contribution by whether a competent team can build from "
        "it; an industrial contribution by whether it changes a deployment "
        "decision; and a methodological contribution by whether other "
        "researchers can reuse the method. Table 8.0 maps the thesis's "
        "contributions onto the four registers and the sections that state "
        "them; the sections that follow then present each in the hierarchy "
        "the evidence supports.")
    cap = OxmlElement("w:p")
    h81._p.addprevious(cap)
    pc = docx.text.paragraph.Paragraph(cap, h81._parent)
    try:
        pc.style = "TableCaption"
    except Exception:
        pass
    pc.add_run("Table 8.0. The four contribution registers and where each is stated.")

    rows = [
        ("Register", "Contribution", "Stated in"),
        ("Scientific",
         "FIAS as a provisional behavioural construct; the coupling boundary "
         "condition for gamification theory; the competing-appraisal account "
         "(endowment protection versus learning trajectory)",
         "§8.1, §8.3; Chs 2, 9"),
        ("Engineering",
         "The adaptive-learning framework — three principles over the "
         "failure-appraisal chain — and the dynamic game-state engine that "
         "operationalises it in real time",
         "Ch 7 (§7.4, §7.5, §7.5A); §8.4"),
        ("Industrial",
         "A deployment template on existing FRA infrastructure (decoupled "
         "presentation, trajectory layer, staff protocol) and "
         "addressable-population evidence for planning",
         "§8.4; §9.8"),
        ("Methodological",
         "The null-hypothesis-plus-attribution architecture; the "
         "dual-purpose mechanic; the QFD evidence-to-design pipeline; the "
         "commitment-gap measurement",
         "§8.2; Table 3.1"),
    ]
    tbl = d.add_table(rows=len(rows), cols=3)
    tbl.style = d.tables[2].style  # reuse the style of an existing inserted table
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri].cells[ci]
            cell.paragraphs[0].add_run(val)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
                if ri == 0:
                    run.bold = True
    h81._p.addprevious(tbl._tbl)
    log.append("Ch8: registers intro + Table 8.0 inserted")

# List of Tables entry
lot = find("Table 7.1. adaptive-learning principles")
if lot is not None and lot.style.name.startswith("toc"):
    t = insert_after(lot, "Table 8.0. The four contribution registers and where each is stated.\t72")
    t.style = lot.style
    log.append("LoT entry for Table 8.0 inserted")
else:
    # LoT line may have slightly different text; try by prefix
    cand = None
    for p in d.paragraphs:
        if p.style.name.startswith("toc") and p.text.strip().startswith("Table 7.1"):
            cand = p
            break
    if cand is not None:
        t = insert_after(cand, "Table 8.0. The four contribution registers and where each is stated.\t72")
        t.style = cand.style
        log.append("LoT entry for Table 8.0 inserted (prefix match)")
    else:
        fail.append("LoT Table 7.1 line not found")

# ---------- (6) Figure 1.1 at end of 1.12 ----------
p112 = find("These outputs form one architecture: FIAS explains why users disengage")
if p112 is None:
    fail.append("1.12 paragraph")
else:
    set_text(p112, p112.text.rstrip() + " Figure 1.1 places the "
        "contributions in a single chain from industrial problem to "
        "industrial outcome.")
    imgp = insert_after(p112)
    imgp.add_run().add_picture(BASE + "fig_1_1.png", width=Inches(5.6))
    capp = insert_after(imgp,
        "Figure 1.1. The intellectual structure of the thesis: an industrial "
        "problem produces an observed behavioural phenomenon (FIAS); the "
        "phenomenon motivates engineering design principles (the "
        "adaptive-learning framework); the principles are operationalised by "
        "the dynamic game-state engine; and the intended industrial outcome "
        "is sustained engagement with clinical validity preserved. The three "
        "middle layers are stated as one architecture in §7.5A.")
    try:
        capp.style = "FigureCaption"
    except Exception:
        pass
    log.append("1.12: Figure 1.1 + caption inserted")

# List of Figures entry (before Figure 3.1 line)
lof = None
for p in d.paragraphs:
    if p.style.name.startswith("toc") and p.text.strip().startswith("Figure 3.1"):
        lof = p
        break
if lof is None:
    fail.append("LoF Figure 3.1 line")
else:
    newp = OxmlElement("w:p")
    lof._p.addprevious(newp)
    nl = docx.text.paragraph.Paragraph(newp, lof._parent)
    nl.style = lof.style
    nl.add_run("Figure 1.1. The intellectual structure of the thesis: from "
               "industrial problem through FIAS, the adaptive-learning "
               "framework and the dynamic game-state engine to the "
               "industrial outcome.\t9")
    log.append("LoF entry for Figure 1.1 inserted")

# ---------- metadata ----------
cp = d.core_properties
cp.title = "From FIDS to FIAS - Submission Draft v18"
cp.comments = ("Submission draft v18 (19 Jul 2026). Reviewer pass: Ch1 "
               "problem statement; FIAS provisional construct; 7.4 why-three; "
               "Stage 3 as redesign package; Ch8 registers + Table 8.0; "
               "Figure 1.1. Pending: photos, Stage-3 data, interviews.")

d.save(DST)
print("saved", DST)
for l in log:
    print(" -", l)
for f in fail:
    print(" !", f)

# ---------- verification ----------
d2 = docx.Document(DST)
full = "\n".join(p.text for p in d2.paragraphs)
for tb in d2.tables:
    for r in tb.rows:
        full += "\n" + " | ".join(c.text for c in r.cells)
for s in ("Why exactly three principles", "provisional behavioural construct",
          "insufficiently considered in system design",
          "evaluates a participant-facing redesign package", "Table 8.0",
          "Figure 1.1", "LODF", "DSM"):
    print(f"'{s}':", full.count(s))
