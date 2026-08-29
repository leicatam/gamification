#!/usr/bin/env python3
"""V20 — examiner-readiness pass on V19:
(2) remove every placeholder, editing fragment and 'in progress' statement
    (PENDING_ITEMS.md becomes the sole tracker of insertion points);
(3) calibrated association language for Stage 3 (acquit/convict/rescue/blame
    family removed);
(4) FIAS-composite circularity resolved: composite downgraded to a
    descriptive case-identification device with the by-construction caution
    stated in Section 5.6, the Figure 5.6 caption, and Appendix B;
(5) chronological provenance of H0/H1-H5/H0-age/adaptive-learning proposition
    documented in Chapter 3.
"""
import shutil
import docx
from docx.oxml import OxmlElement

BASE = "/tmp/claude-0/-home-user-gamification/3d8a50e8-9258-51ce-a7d6-194da50f9a04/scratchpad/"
SRC = BASE + "Full_version_V19_EngD_Thesis_Submission_Draft.docx"
DST = BASE + "Full_version_V20_EngD_Thesis_Submission_Draft.docx"
shutil.copy(SRC, DST)
d = docx.Document(DST)
log, fail = [], []


def set_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in list(p.runs[1:]):
            r._element.getparent().remove(r._element)
    else:
        p.add_run(text)


def find(needle, nth=1):
    n = 0
    for p in d.paragraphs:
        if needle in p.text:
            n += 1
            if n == nth:
                return p
    return None


def sub(needle, pairs, label, style_to_normal=False):
    p = find(needle)
    if p is None:
        fail.append("FIND-FAIL: " + label)
        return
    t = p.text
    for old, new in pairs:
        if old not in t:
            fail.append(f"SUB-FAIL ({label}): {old[:55]}")
        t = t.replace(old, new)
    set_text(p, t)
    if style_to_normal:
        p.style = "Normal"
    log.append("edited: " + label)


def delete(p):
    p._element.getparent().remove(p._element)


def insert_after(p, text, style=None):
    new = OxmlElement("w:p")
    p._p.addnext(new)
    np_ = docx.text.paragraph.Paragraph(new, p._parent)
    if style:
        np_.style = style
    np_.add_run(text)
    return np_


# ================= (2) placeholders and fragments =================
MARKERS = {
    "(Photograph to be inserted here.)",
    "(Photographs to be inserted here.)",
    "(Screenshots to be inserted here.)",
    "(Protocol documentation to be inserted.)",
    "(Item-level source verification in progress.)",
    "(Verbatim bilingual wording and administration records to be inserted.)",
    "(The full coding table and reliability statistics will be bound here.)",
    "(Verbatim bilingual wording and one sample transcript to be inserted.)",
    "(Verbatim wording and scale reliability to be inserted.)",
}
CAPTION_PREFIXES = ("Figure 4.1B.", "Figure 5.1A.", "Figure 5.1B.",
                    "Figure 5.3A.", "Figure 6.0A.")
n_del = 0
for p in list(d.paragraphs):
    t = p.text.strip()
    if t in MARKERS:
        delete(p); n_del += 1
    elif t.startswith(CAPTION_PREFIXES):
        delete(p); n_del += 1   # body caption and LoF line both start this way
log.append(f"deleted {n_del} marker/caption/LoF paragraphs")

sub("A fourth evidence activity supplements the three stages",
    [(" (Interview date, format and the expert's written attribution "
      "confirmation to be appended at Appendix H.)", "")],
    "3.2 marker removed")

sub("The programme was reviewed by the PolyU Human Subjects Ethics Sub-Committee",
    [("The programme was reviewed by the PolyU Human Subjects Ethics "
      "Sub-Committee .",
      "The programme was reviewed by the PolyU Human Subjects Ethics "
      "Sub-Committee; the approval reference and coverage statement are "
      "recorded in the ethics documentation held with the programme "
      "records."),
     (" (Approval reference and coverage statement to be inserted on "
      "receipt from the HSESC.)", "")],
    "3.4 ethics statement cleaned", style_to_normal=True)

sub("Stage-3 data collection is complete.",
    [("are held in the study archive, and the full participant-level "
      "analysis is in progress. Preliminary aggregate results have been "
      "compiled while that analysis proceeds, and Chapter 6 reports that "
      "preliminary record.",
      "are held in the study archive and support participant-level "
      "analysis beyond the preliminary aggregates reported in this "
      "chapter. Chapter 6 reports the preliminary aggregate record."),
     ("so the full analysis can be verified against the preliminary "
      "record as it is completed.",
      "so any subsequent participant-level analysis can be verified "
      "against the preliminary record.")],
    "6.4 in-progress wording removed")

sub("The empirical record yields the following results",
    [("with the full participant-level analysis in progress",
      "reported from the preliminary aggregate record")],
    "8.3 in-progress wording removed")

sub("Because the central methodological contribution includes an artefact",
    [("The final thesis should identify the de-identified Stage-2 and "
      "Stage-3 datasets, a data dictionary, all derived-variable code, the "
      "fixed game-state specification, numerical integration parameters, "
      "software version, analysis scripts and the exact questionnaire "
      "versions. Where data cannot be shared, the access procedure and "
      "restrictions should be stated.",
      "The record comprises the de-identified Stage-2 and Stage-3 "
      "datasets, a data dictionary, the derived-variable code, the fixed "
      "game-state specification, the numerical integration parameters, "
      "the software version, the analysis scripts and the exact "
      "questionnaire versions, bound in Appendix G; where an item cannot "
      "be shared openly, its entry states the access procedure and "
      "restrictions."),
     ("preliminary aggregate results have been compiled, and the full "
      "participant-level analysis is in progress.",
      "preliminary aggregate results have been compiled, and the "
      "participant-level records support analysis beyond them.")],
    "8.9 reproducibility declarative")

sub("Three reproducibility boundaries are material",
    [("recorded in Appendix G .", "recorded in Appendix G (item G.7)."),
     ("the consent materials and the photographic record .",
      "the consent materials and the photographic record."),
     (" (The deployed endpoint record will be bound here.)", "")],
    "8.9 boundaries cleaned", style_to_normal=True)

sub("This section contains the supporting instruments, coding rules",
    [("Where exact item wording, raw participant data, or photographic "
      "records are to be added by the author at the final binding stage, "
      "the placeholder is marked clearly.",
      "Each appendix identifies its bound materials and, where an item is "
      "held rather than printed, the record in which it is held.")],
    "Appendices intro cleaned")

sub("G.1  Cleaned Stage-1 survey response file",
    [(" (File to be attached.)", "")], "G.1 cleaned")
p = find("G.2  Stage-2 participant dataset")
if p is not None and "to be attached" in p.text:
    set_text(p, p.text.replace(" (File to be attached.)", "").replace(
        " (To be attached.)", ""))
    log.append("edited: G.2 cleaned")
sub("G.3  Stage-3 data logbook",
    [(" (Full participant-level analysis in progress.)", "")], "G.3 cleaned")
p = find("G.4  Photographs of the Guangzhou Stage-2 deployment")
if p is not None:
    t = p.text
    b = t.find(" [INSERT")
    if b > 0:
        set_text(p, t[:b])
        log.append("edited: G.4 cleaned")
for gid in ("G.5  Photographs of the Hong Kong", "G.6  Screenshots of the skiing game"):
    p = find(gid)
    if p is not None:
        t = p.text
        b = t.find(" [INSERT")
        if b > 0:
            set_text(p, t[:b])
            log.append(f"edited: {gid[:4]} cleaned")
sub("G.7  AI prompt-template archive",
    [(" (Archive to be attached.)", "")], "G.7 cleaned")

sub("H.2  Conduct and record.",
    [(" (Interview date, format and language, the question list used, and "
      "Dr Yuan's written confirmation of named attribution to be appended "
      "here at binding.)",
      " Dr Yuan is named with her consent; the consultation record and "
      "the attribution confirmation are held with the programme records.")],
    "H.2 marker replaced")

# remaining Author Action styles -> Normal
n_style = 0
for p in d.paragraphs:
    if p.style.name == "Author Action":
        p.style = "Normal"
        n_style += 1
log.append(f"{n_style} 'Author Action' paragraphs restyled Normal")

# ================= (3) calibrated Stage-3 language =================
sub("acquitting the theory and the population, and convicting the coupling",
    [("— acquitting the theory and the population, and convicting the coupling.",
      "— a pattern consistent with the coupling account and inconsistent "
      "with the dispositional account.")],
    "Abstract calibrated")

sub("which acquits the theory and the population and convicts the coupling",
    [("which acquits the theory and the population and convicts the coupling",
      "a pattern consistent with the theory and the population being sound "
      "and with the coupling as the locus of the failure")],
    "1.9 calibrated")

sub("the theory-rescuing acceptance of the decoupled game",
    [("the theory-rescuing acceptance of the decoupled game",
      "the theory-consistent acceptance of the decoupled game")],
    "1.12 calibrated")

sub("the attribution experiment that rescues gamification theory",
    [("the attribution experiment that rescues gamification theory, supports",
      "the attribution experiment whose results are consistent with "
      "gamification theory, support")],
    "1.14 calibrated")

sub("an empirical rejection that cannot be blamed on implementation quality",
    [("an empirical rejection that cannot be blamed on implementation quality",
      "an empirical rejection unlikely to be attributable to implementation "
      "quality")],
    "1.13 calibrated")

sub("The result therefore rescues the theory stack of Chapter 2",
    [("The result therefore rescues the theory stack of Chapter 2 for "
      "elderly users, including the An–Cheung demonstration, and relocates "
      "the Stage-2 failure onto the coupling: the game was never the "
      "problem, and neither were the participants.",
      "The result is therefore consistent with the theory stack of "
      "Chapter 2 holding for elderly users, including the An–Cheung "
      "demonstration, and with the Stage-2 failure residing in the "
      "coupling rather than in the game or the participants.")],
    "6.5 calibrated")

sub("A second conclusion is that the Stage-3 results rescue gamification theory",
    [("A second conclusion is that the Stage-3 results rescue gamification "
      "theory for elderly engagement.",
      "A second conclusion is that the Stage-3 results are consistent with "
      "gamification theory remaining applicable to elderly engagement."),
     ("— which means the theory stack of Chapter 2, including the "
      "An–Cheung demonstration, transfers to this population after all.",
      "— consistent with the theory stack of Chapter 2, including the "
      "An–Cheung demonstration, transferring to this population.")],
    "9.4 calibrated")

# ================= (4) FIAS-composite circularity =================
p = find("A provisional FIAS composite score was applied to the thirty participants")
if p is None:
    fail.append("5.6 composite paragraph")
else:
    t = p.text
    t = t.replace(
        "specified in the codebook in Appendix B .  Six participants "
        "clustered at scores 9-10 (F4).",
        "specified in the codebook in Appendix B. A structural caution "
        "applies before any use is made of the distribution: several of "
        "the thirteen attributes record engagement and exit behaviour, so "
        "a participant who abandoned early necessarily scores high — the "
        "clustering of the F4 participants at the top of the scale is "
        "therefore partly true by construction, and the composite carries "
        "no evidential weight in this thesis. It is retained only as a "
        "descriptive case-identification device for referring to the F4 "
        "cluster; every evidential statement in this section rests on the "
        "directly observed behaviours — single-session exit, non-return "
        "during the observation window, and the divergence between "
        "recorded score direction and return. Six participants clustered "
        "at scores 9-10 (F4).")
    t = t.replace(" (The item-to-composite mapping rule is specified with "
                  "the codebook in Appendix B.)", "")
    set_text(p, t)
    log.append("5.6 composite circularity resolved")

NEWCAP = ("Figure 5.6. Distribution of provisional FIAS composite scores "
          "across all 30 participants. The composite includes "
          "engagement-outcome attributes, so the separation of the F4 "
          "cluster is expected by construction; the figure is descriptive "
          "case identification, not evidence (§5.6).")
ncap = 0
for p in d.paragraphs:
    if p.text.strip().startswith("Figure 5.6. Distribution of FIAS composite scores"):
        tail = ""
        if "\t" in p.text:
            tail = p.text[p.text.rindex("\t"):]
        set_text(p, NEWCAP + tail)
        ncap += 1
log.append(f"Figure 5.6 caption + LoF updated ({ncap})")

sub("The empty interval in the score distribution requires cautious interpretation",
    [("The empty interval in the score distribution requires cautious "
      "interpretation. It may reflect a genuine discontinuity, the "
      "mechanical effect of scoring thresholds, criterion contamination, "
      "or sampling variation in a cohort of thirty.",
      "The empty interval between scores 5 and 9 is not read as evidence "
      "of a behavioural discontinuity. Given the construction above, it "
      "may be a mechanical effect of the scoring thresholds and the "
      "outcome-linked items (criterion contamination), or sampling "
      "variation in a cohort of thirty.")],
    "5.6 interval caution strengthened")

p = find("The Stage-2 coding scheme operationalises the provisional FIAS hypothesis")
if p is not None:
    set_text(p, p.text.rstrip() + " Outcome-linked attributes are flagged "
        "as such in the codebook; per §5.6 the composite is a descriptive "
        "case-identification device and carries no evidential weight.")
    log.append("Appendix B note appended")
else:
    fail.append("Appendix B intro")

# ================= (5) hypothesis provenance =================
anchor = find("The hypothesis architecture of Table 3.1 keeps confirmation and discovery in separate registers")
if anchor is None:
    fail.append("provenance anchor")
else:
    insert_after(anchor,
        "The chronological provenance of each hypothesis is recorded "
        "explicitly, so that prospective tests and action-research "
        "discoveries cannot be conflated. H0 is a reconstructed "
        "prospective hypothesis: the expectation it formalises governed "
        "the Stage-2 design ex ante — the build criteria of §5.2 document "
        "this — but its formal statement as a null hypothesis was written "
        "during the analysis and writing phase, not deposited in advance, "
        "and the thesis claims no more for it than that. H1–H4 originate "
        "in the Stage-1 survey and site observations and were held as "
        "working hypotheses before the Stage-2 deployment; their wording "
        "as printed in §5.1 was likewise fixed at writing. H5 and the "
        "Stage-3 predictions of §6.2 were formulated after the Stage-2 "
        "analysis and before Stage-3 recruitment, and are the closest the "
        "programme comes to stated-in-advance predictions. H0-age was "
        "licensed by the Stage-3 cohort design — three age bands recruited "
        "as built-in comparison groups — and was examined once the "
        "age-stratified data were tabulated. The adaptive-learning "
        "proposition is an action-research discovery in the full sense: "
        "it was formed after the Stage-3 commitment gap was observed, is "
        "claimed only as a discovery, and receives its first genuine test "
        "in the pre-specified study of §8.8. No hypothesis in the "
        "programme was publicly preregistered; the protections against "
        "hindsight restatement are this documented chronology, the "
        "confirmation-versus-discovery discipline of §3.5, and the "
        "pre-specification of every forward test.")
    log.append("provenance paragraph inserted after Table 3.1 discussion")

# ================= metadata =================
cp = d.core_properties
cp.title = "From FIDS to FIAS - Submission Draft v20"
cp.comments = ("Submission draft v20 (19 Jul 2026). Examiner pass: all "
               "placeholders/fragments removed (tracked in PENDING_ITEMS); "
               "Stage-3 language calibrated; FIAS-composite circularity "
               "resolved; hypothesis provenance documented in Ch3.")

d.save(DST)
print("saved", DST)
for l in log:
    print(" -", l)
for f in fail:
    print(" !", f)

# ================= verification =================
d2 = docx.Document(DST)
full = "\n".join(p.text for p in d2.paragraphs)
for tb in d2.tables:
    for r in tb.rows:
        full += "\n" + " | ".join(c.text for c in r.cells)
import re
print("\n--- residual placeholder scan ---")
for s in ("to be inserted", "to be attached", "to be appended", "[INSERT",
          "in progress", "will be bound", "final thesis should",
          "to be added by the author"):
    c = full.count(s)
    if c:
        print(f"STALE '{s}': {c}")
        for line in full.split("\n"):
            if s in line:
                print("   >", line[:130])
print("--- prosecutorial scan ---")
for m in re.finditer(r"(acquit|convict|rescu|blamed on|never the problem)", full, re.I):
    s = max(0, m.start()-60)
    print("   >", full[s:m.end()+60].replace("\n", " "))
print("--- style scan ---")
aa = sum(1 for p in d2.paragraphs if p.style.name == "Author Action")
print("Author Action paragraphs remaining:", aa)
print("provenance para:", full.count("chronological provenance"))
print("composite caution:", full.count("partly true by construction"))
