#!/usr/bin/env python3
"""V17 amendments per reviewer feedback:
(1) 'Stop selling GPT' — AI tooling repositioned as enabling technology,
    not contribution (Sections 1.12, 5.3.9, 8.2, 9.8A; GPT brand reference removed);
(2) 'Write like an engineer' — targeted verb pass on self-claims
    (confirms/establishes/demonstrates/proved -> indicates/supports/documents/were);
(3) Unified architecture — FIAS (phenomenon) + adaptive-learning framework
    (principles) + dynamic game-state engine (implementation): new Section 7.5A,
    plus positioning sentences in 1.12, 8.1 and 9.12.
NOTE: V16 contains no LODF; the reviewer's 'LODF' maps to the Adaptive-Learning
Framework and their 'DSM' to the dynamic game-state architecture. No LODF term
is reintroduced.
"""
import shutil
import docx
from docx.oxml import OxmlElement

BASE = "/tmp/claude-0/-home-user-gamification/3d8a50e8-9258-51ce-a7d6-194da50f9a04/scratchpad/"
SRC = BASE + "Full_version_V16_EngD_Thesis_Submission_Draft.docx"
DST = BASE + "Full_version_V17_EngD_Thesis_Submission_Draft.docx"
shutil.copy(SRC, DST)
d = docx.Document(DST)
log, fail = [], []


def set_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in list(p.runs[1:]):
            r._element.getparent().remove(r._element)
    else:
        p.add_run(text)


def find(needle, nth=1):
    n = 0
    for p in d.paragraphs:
        if needle in p.text:
            n += 1
            if n == nth:
                return p
    return None


def sub(needle, pairs, label, nth=1):
    p = find(needle, nth)
    if p is None:
        fail.append("FIND-FAIL: " + label)
        return
    t = p.text
    for old, new in pairs:
        if old not in t:
            fail.append(f"SUB-FAIL ({label}): {old[:60]}")
        t = t.replace(old, new)
    set_text(p, t)
    log.append("edited: " + label)


def insert_after(p, text, style=None):
    new = OxmlElement("w:p")
    p._p.addnext(new)
    np_ = docx.text.paragraph.Paragraph(new, p._parent)
    if style:
        np_.style = style
    np_.add_run(text)
    return np_


# ================= (2) engineer's verb pass =================
sub("as Stage 3 confirms, until the gamified activity",
    [("as Stage 3 confirms,", "as the Stage-3 evidence indicates,")],
    "1.12 'confirms' -> 'indicates'")

sub("rescues gamification theory, establishes the coupling boundary condition",
    [("establishes the coupling boundary condition",
      "supports the coupling boundary condition")],
    "1.14 'establishes' -> 'supports'")

sub("Stage two of the transformative chain therefore establishes",
    [("therefore establishes that", "therefore indicates that")],
    "2.3 'establishes' -> 'indicates'")

sub("Second — the scope finding — it establishes a boundary condition",
    [("it establishes a boundary condition", "the evidence supports a boundary condition"),
     ("as Stage 3 confirms at literature-predicted acceptance levels",
      "as Stage 3 indicates at literature-predicted acceptance levels")],
    "8.1 'establishes/confirms' -> 'supports/indicates'")

sub("Stage-2 exits proved so absolute",
    [("proved so absolute", "were so absolute")],
    "8.6B 'proved' -> 'were'")

# ================= (1) de-sell the AI tooling =================
sub("Methodologically, the thesis contributes the null-hypothesis-plus-attribution architecture",
    [("the dual-purpose mechanic, the QFD evidence-to-design pipeline, a "
      "documented AI-assisted development workflow, and the commitment-gap "
      "measurement.",
      "the dual-purpose mechanic, the QFD evidence-to-design pipeline and the "
      "commitment-gap measurement; the AI-assisted development workflow is "
      "documented as the enabling method behind the artefact (§5.3.9), not "
      "claimed as a contribution — throughout the thesis, the AI tooling is "
      "an enabling technology, and the contribution is the behavioural "
      "design framework."),
     # (3) unified-architecture positioning sentence
     ("with the adaptive-learning framework as its engineering formalisation "
      "and zero-failure presentation as a supporting element.",
      "with the adaptive-learning framework as its engineering formalisation "
      "and zero-failure presentation as a supporting element. These outputs "
      "form one architecture: FIAS explains why users disengage after "
      "failure, the adaptive-learning framework states how systems should be "
      "redesigned to prevent that disengagement, and the dynamic game-state "
      "engine of §5.3.10 operationalises the framework in real time "
      "(§7.5A).")],
    "1.12 contributions reworded")

p = find("The Claude Code-assisted workflow is offered as a small methodological contribution")
if p is None:
    fail.append("FIND-FAIL: 5.3.9 closing")
else:
    set_text(p,
        "The Claude Code-assisted workflow is documented for reproducibility, "
        "not offered as a contribution in its own right: the AI tooling is an "
        "enabling technology, and the contribution it enabled is the "
        "behavioural design framework. Engineering doctoral programmes "
        "increasingly involve software components whose construction is "
        "itself an act of design. Documenting the vehicle by which the "
        "artefact was built, alongside the artefact itself, makes the action "
        "research more reproducible and clarifies which engineering decisions "
        "belong to the developer and which were accelerated by tooling.")
    log.append("rewrote: 5.3.9 workflow status")

sub("documentation of the Claude Code-assisted development workflow",
    [("documentation of the Claude Code-assisted development workflow",
      "documentation of the AI-assisted development workflow (recorded as "
      "enabling method, §5.3.9)")],
    "8.2 workflow demoted")

sub("a pre-trained large language model (GPT-class or Claude-class) is invoked",
    [("a pre-trained large language model (GPT-class or Claude-class) is invoked",
      "a pre-trained large language model is invoked")],
    "5.3.5 brand reference removed")

sub("it demonstrates how an AI-assisted development workflow can translate",
    [("The contribution is methodological as well as technical: it "
      "demonstrates how an AI-assisted development workflow can translate a "
      "behavioural and safety specification into candidate state logic, "
      "synthetic tests and auditable runtime rules.",
      "The workflow documents how a behavioural and safety specification was "
      "translated into candidate state logic, synthetic tests and auditable "
      "runtime rules; the AI tooling that accelerated this translation is an "
      "enabling technology, and the contribution remains the behavioural "
      "design framework the states implement.")],
    "9.8A workflow reworded")

# ================= (3) unified architecture: new Section 7.5A =================
anchor = find("The proposed staff coaching protocol (L6) operationalises trajectory framing")
if anchor is None:
    fail.append("FIND-FAIL: 7.5.3 anchor")
else:
    h = insert_after(anchor,
        "7.5A One Architecture: Phenomenon, Principles, Engine", style="Heading 2")
    p1 = insert_after(h,
        "The programme's outputs are best read as a single behavioural-"
        "engineering architecture rather than as three related ideas. FIAS is "
        "the phenomenon layer: it names and explains why older users "
        "disengage after failure — a post-contact behaviour driven by "
        "identity-level appraisal, characterised across Chapters 4–6. The "
        "adaptive-learning framework is the principle layer: it states how "
        "systems should be redesigned so that the disengagement is not "
        "produced, through identity disentanglement, trajectory installation "
        "and adaptive calibration (§7.4). The dynamic game-state architecture "
        "of §5.3.10 is the engine layer: it operationalises the principles in "
        "real time, monitoring the live performance-state vector and "
        "selecting among non-terminal states within pre-specified safety "
        "limits, so that the design response arrives at the moment the "
        "phenomenon would otherwise begin (§7.6A).")
    insert_after(p1,
        "The evidence supports each layer to a different degree, and the "
        "architecture inherits those gradations rather than smoothing them "
        "over. The phenomenon layer rests on the observed Stage-1 and "
        "Stage-2 record, and its frame-sensitivity is indicated — not "
        "isolated — by the Stage-3 comparison. The principle layer is "
        "derived from that record and from the learning-theory literature, "
        "and has not yet been evaluated as a package. The engine layer is an "
        "implemented specification whose behavioural effect is untested; its "
        "evaluation is pre-specified in §8.8. Stating the three as one "
        "architecture does not strengthen any single claim — it fixes where "
        "each claim lives, and it is what turns a set of findings into a "
        "design theory an engineering team can build against.")
    log.append("inserted: Section 7.5A (two paragraphs)")

toc = find("7.5.3 Coaching Protocol Specification\t")
if toc is not None:
    t = insert_after(toc, "7.5A One Architecture: Phenomenon, Principles, Engine\t69")
    t.style = toc._parent.paragraphs[0].style  # placeholder; fixed below
    # match the toc 2 style of the 7.5 line
    for p in d.paragraphs:
        if p.text.startswith("7.5 Prototype Stack") and p.style.name == "toc 2":
            t.style = p.style
            break
    log.append("inserted: ToC line for 7.5A")
else:
    fail.append("FIND-FAIL: ToC 7.5.3 line")

sub("Second — the scope finding — the evidence supports a boundary condition",
    [("(Table 2.1).",
      "(Table 2.1).")],
    "8.1 no-op guard")  # ensure paragraph still located after earlier edit

p = find("The thesis makes four principal theoretical contributions")
if p is None:
    fail.append("FIND-FAIL: 8.1 opening")
else:
    set_text(p, p.text.rstrip() + " Read together, the contributions form one "
        "behavioural-engineering architecture — the phenomenon (FIAS), the "
        "principles (the adaptive-learning framework) and the engine (the "
        "dynamic game-state architecture) — as set out in §7.5A.")
    log.append("appended: 8.1 unification sentence")

p = find("The contribution is therefore a calibrated chain")
if p is None:
    fail.append("FIND-FAIL: 9.12 chain paragraph")
else:
    set_text(p, p.text.rstrip() + " Read as one architecture, the chain "
        "closes: FIAS supplies the phenomenon, the adaptive-learning "
        "framework the design principles, and the dynamic game-state engine "
        "the real-time implementation (§7.5A).")
    log.append("appended: 9.12 unification sentence")

# ================= metadata =================
cp = d.core_properties
cp.title = "From FIDS to FIAS - Submission Draft v17"
cp.comments = ("Submission draft v17 (19 Jul 2026). Reviewer amendments: AI "
               "tooling repositioned as enabling technology; engineer-tone "
               "verb pass; new 7.5A unifies FIAS / adaptive-learning "
               "framework / dynamic game-state engine as one architecture.")

d.save(DST)
print("saved", DST)
for l in log:
    print(" -", l)
for f in fail:
    print(" !", f)

# ================= verification =================
d2 = docx.Document(DST)
full = "\n".join(p.text for p in d2.paragraphs)
for tb in d2.tables:
    for r in tb.rows:
        full += "\n" + " | ".join(c.text for c in r.cells)
print("\nLODF:", full.count("LODF"), "| DSM:", full.count("DSM"),
      "| GPT:", full.count("GPT"))
print("7.5A occurrences:", full.count("7.5A"))
print("'enabling technology':", full.count("enabling technology"))
for s in ("as Stage 3 confirms", "establishes the coupling",
          "proved so absolute", "it demonstrates how an AI-assisted",
          "GPT-class", "a documented AI-assisted development workflow"):
    print(f"stale '{s}':", full.count(s))
