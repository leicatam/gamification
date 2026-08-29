#!/usr/bin/env python3
"""V36 - HSESC approval received (application HSEARS20260805008, PI Prof.
C. F. Cheung, approved August 2026): (1) 3.4 review sentence replaced with
the concrete approval statement; (2) the Stage-2 consent procedure as
conducted (incomplete disclosure with scientific rationale, gatekeeper
consent, retrospective written consent) written into 3.4 - the account the
approved application itself contained."""
import shutil
import docx
from docx.oxml import OxmlElement

REPO = "/home/user/gamification/docs/thesis/"
SRC = REPO + "Full_version_V35_EngD_Thesis_Submission_Draft.docx"
DST = REPO + "Full_version_V36_EngD_Thesis_Submission_Draft.docx"
shutil.copy(SRC, DST)
d = docx.Document(DST)

anchor = None
for p in d.paragraphs:
    if "The programme was reviewed by the PolyU Human Subjects Ethics Sub-Committee" in p.text:
        anchor = p
        break
assert anchor is not None, "3.4 review sentence not found"

old = anchor.text
new = old.replace(
    "The programme was reviewed by the PolyU Human Subjects Ethics "
    "Sub-Committee; the approval reference and coverage statement are "
    "recorded in the ethics documentation held with the programme records.",
    "The programme was approved by the PolyU Human Subjects Ethics "
    "Sub-Committee (application reference HSEARS20260805008; Principal "
    "Investigator: Prof. Chi Fai Cheung, Department of Industrial and "
    "Systems Engineering; approved August 2026); the application described "
    "the procedures of all three stages and the expert consultations as "
    "conducted, and the approval memo is held with the programme records.")
assert new != old, "replacement failed"
for r in list(anchor.runs)[1:]:
    r._r.getparent().remove(r._r)
anchor.runs[0].text = new

newp = OxmlElement("w:p")
anchor._p.addnext(newp)
np_ = docx.text.paragraph.Paragraph(newp, anchor._parent)
np_.add_run(
    "Consent in Stage 2 is likewise stated exactly as it operated. The "
    "deployment was arranged with, and known to, the facility's operators, "
    "nursing team and general manager (gatekeeper consent, recorded at "
    "research start), and residents were invited by the facility to try the "
    "gamified system as a facility activity, voluntarily and free of "
    "charge. Participants were not told during the deployment that play "
    "formed part of a research project: the programme's subject matter is "
    "assessment fear, and announcing a fall-risk study would have "
    "re-imposed the clinical assessment frame whose behavioural effect was "
    "under study, so incomplete disclosure was integral to the design "
    "rather than incidental to it. Written consent was obtained "
    "retrospectively from the sustained participants at the close of the "
    "eight-week cycle, at the small recognition presentation for those who "
    "completed it; participants who had dropped out did not sign and were "
    "not interviewed. These arrangements - the incomplete disclosure and "
    "its rationale, the gatekeeper consent, and the retrospective timing "
    "of the written instruments - were stated in these terms in the "
    "approved ethics application, and their limitations are examined in "
    "8.7.")

cp = d.core_properties
cp.comments = "V36: HSESC approval HSEARS20260805008 inserted; Stage-2 consent-as-conducted written into 3.4."
d.save(DST)
print("V36 saved")
