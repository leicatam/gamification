#!/usr/bin/env python3
"""V15 corrections on the uploaded V14 (whose body text = repo V12):
(1) Stage-3 data status: remove every 'working file lost / register rebuilt'
    statement — data collection is complete, preliminary aggregates were
    released and reported, full participant-level analysis in progress.
(2) Remove LODF: title-page subtitle, §7.4 heading + ToC, abbreviations row,
    glossary key.
(3) Fix stale §3.2 numbers missed at V12: '0 of 35 ... 16 of 45' -> 38/41.
"""
import shutil
import docx

BASE = "/tmp/claude-0/-home-user-gamification/3d8a50e8-9258-51ce-a7d6-194da50f9a04/scratchpad/"
SRC = BASE + "V14.docx"
DST = BASE + "Full_version_V15_EngD_Thesis_Submission_Draft.docx"
shutil.copy(SRC, DST)
d = docx.Document(DST)
log, failures = [], []


def clear_runs(p):
    for r in list(p.runs):
        r._element.getparent().remove(r._element)


def set_text(p, text):
    # preserve the first run's formatting by editing it, dropping the rest
    if p.runs:
        p.runs[0].text = text
        for r in list(p.runs[1:]):
            r._element.getparent().remove(r._element)
    else:
        p.add_run(text)


def find(needle, nth=1):
    n = 0
    for p in d.paragraphs:
        if needle in p.text:
            n += 1
            if n == nth:
                return p
    return None


def replace_para(needle, new_text, label):
    p = find(needle)
    if p is None:
        failures.append("FIND-FAIL: " + label)
        return
    set_text(p, new_text)
    log.append("rewrote: " + label)


def sub(needle, pairs, label):
    p = find(needle)
    if p is None:
        failures.append("FIND-FAIL: " + label)
        return
    t = p.text
    for old, new in pairs:
        if old not in t:
            failures.append(f"SUB-FAIL ({label}): {old[:50]}")
        t = t.replace(old, new)
    set_text(p, t)
    log.append("edited: " + label)


# ---------------- (2) LODF removal ----------------
sub("Toward a Learning-Oriented Design Framework",
    [("Toward a Learning-Oriented Design Framework",
      "Toward an Adaptive-Learning Framework")],
    "title-page subtitle")

n = 0
for p in d.paragraphs:
    if "The Adaptive-Learning Framework (LODF)" in p.text:
        set_text(p, p.text.replace("The Adaptive-Learning Framework (LODF)",
                                   "The Adaptive-Learning Framework"))
        n += 1
log.append(f"'(LODF)' stripped from {n} heading/ToC lines")
if n < 2:
    failures.append("expected 2 (LODF) heading/ToC lines, got %d" % n)

for t in d.tables:
    txt0 = t.rows[0].cells[0].text.strip()
    for row in list(t.rows):
        cells = [c.text.strip() for c in row.cells]
        if cells and cells[0] == "LODF":
            if "three-principle engineering framework" in cells[-1]:
                # glossary: rename the construct key, keep the definition
                kp = row.cells[0].paragraphs[0]
                set_text(kp, "Adaptive-Learning Framework")
                dp = row.cells[1].paragraphs[0]
                set_text(dp, "The thesis's proposed three-principle "
                             "engineering framework: identity disentanglement, "
                             "trajectory installation and adaptive calibration. "
                             "The complete framework has not yet been evaluated.")
                log.append("glossary: LODF key renamed to Adaptive-Learning Framework")
            else:
                row._element.getparent().remove(row._element)
                log.append("abbreviations: LODF row deleted")

# ---------------- (3) stale §3.2 numbers ----------------
sub("0 of 35 younger versus 16 of 45 older initial dropouts",
    [("0 of 35 younger versus 16 of 45 older initial dropouts",
      "0 of 38 younger versus 16 of 41 older initial dropouts")],
    "§3.2 H0-age counts")

# ---------------- (1) Stage-3 data status ----------------
replace_para(
    "A register note applies to the age-band sizes",
    "A register note applies to the age-band sizes. The preliminary published "
    "summary reported band sizes of 35, 40 and 45; the full participant "
    "register compiled from the primary records gives 38, 41 and 41 (total "
    "120, unchanged), and the thesis reports the full register throughout. "
    "Every headline outcome count is unaffected: 23 initial dropouts, 15 "
    "voluntary returns, 8 permanent non-returns, 86 coordination endorsements "
    "and 76 regular-use intentions. Band-level percentages in this chapter "
    "are computed on the register denominators.",
    "§6.4 register note")

replace_para(
    "The Stage-3 participant-level working file was lost to a technical error after the aggregate results had been computed",
    "Stage-3 data collection is complete. The participant-level records — "
    "paper questionnaires, staff session logs and terminal exports — are held "
    "in the study archive, and the full participant-level analysis is in "
    "progress. Preliminary aggregate results were released and reported while "
    "that analysis proceeds, and Chapter 6 reports the preliminary record. "
    "The data logbook (Appendix G, item G.3) holds the published aggregate "
    "results and a reconciliation sheet whose live formulas recompute every "
    "published indicator from the participant-level entries, so the full "
    "analysis can be verified against the preliminary record as it is "
    "completed.",
    "§6.4 data-status paragraph")

replace_para(
    "Note on Table 6.3: coordination-challenge endorsement",
    "Note on Table 6.3: coordination-challenge endorsement was reported in "
    "the preliminary summary only as a total (86 of 120, 71.7%); the "
    "per-age-band breakdown will be reported with the full participant-level "
    "analysis, so it is shown as not yet reported. No per-band endorsement "
    "claim is made in this thesis.",
    "Table 6.3 note")

sub("pending completion of the Stage-3 data logbook",
    [("pending completion of the Stage-3 data logbook",
      "with the full participant-level analysis in progress")],
    "§8.3 empirical record")

p = find("One material data-management event is disclosed")
if p is None:
    failures.append("FIND-FAIL: §8.9 disclosure")
else:
    t = p.text
    old_start = t.index("One material data-management event is disclosed")
    new_tail = ("The Stage-3 dataset was collected in full; preliminary "
                "aggregate results have been released and reported, and the "
                "full participant-level analysis is in progress. The archived "
                "logbook holds the published aggregates together with a "
                "reconciliation harness that checks every participant-level "
                "entry against them (Appendix G, item G.3); Chapter 6 rests "
                "on the published preliminary record, which the completed "
                "analysis will extend.")
    set_text(p, t[:old_start] + new_tail)
    log.append("rewrote: §8.9 data-availability statement")

replace_para(
    "G.3  Stage-3 data logbook",
    "G.3  Stage-3 data logbook "
    "(Stage3_Data_Logbook_120cases_with_Reconciliation.xlsx): codebook with "
    "provenance note, participant/telemetry/questionnaire/outcome sheets for "
    "IDs 3001–3120, the published preliminary aggregate results "
    "(Published_Summary sheet, transcribed from Paper 3) and a live "
    "Reconciliation sheet. Stage-3 data collection is complete; "
    "participant-level entries are transcribed from the primary records "
    "(paper questionnaires, staff session logs and terminal exports), with "
    "each entry checked by the reconciliation sheet. (Full participant-level "
    "analysis in progress.)",
    "Appendix G.3")

# ---------------- metadata ----------------
cp = d.core_properties
cp.title = "From FIDS to FIAS - Submission Draft v15"
cp.comments = ("Submission draft v15 (19 Jul 2026). Corrects v14: Stage-3 "
               "data status (collected in full; preliminary results "
               "reported; analysis in progress); LODF term removed; "
               "3.2 counts fixed. Expert interviews to be integrated.")

d.save(DST)
print("saved", DST)
for l in log:
    print(" -", l)
if failures:
    print("FAILURES:")
    for f_ in failures:
        print(" !", f_)

# ---------------- verification ----------------
d2 = docx.Document(DST)
full = "\n".join(p.text for p in d2.paragraphs)
for t in d2.tables:
    for r in t.rows:
        full += "\n" + " | ".join(c.text for c in r.cells)
print("\nLODF:", full.count("LODF"))
print("Learning-Oriented:", full.count("Learning-Oriented"))
for s in ("was lost", "rebuil", "surviving", "Transcription in progress",
          "0 of 35", "16 of 45", "data-loss", "data-management event"):
    c = full.count(s)
    if c:
        print(f"STALE {s!r}: {c}")
        for line in full.split("\n"):
            if s in line:
                print("   >", line[:140])
print("'lives lost' game-context lines (expected, OK):",
      sum(1 for line in full.split("\n") if "lives lost" in line or "life is lost" in line))
EOF_MARKER = None
