#!/usr/bin/env python3
"""V31 — forward extension after the GrandPad comparator (Ch8): adaptive
learning as a later-life technology-onboarding pathway, supported by
verified literature (Park et al. 2014 Synapse; Czaja et al. 2018 PRISM;
Baltes & Baltes 1990 SOC; the An & Cheung lineage already in the list).
Framed explicitly as a design hypothesis. Three references added."""
import shutil
import docx
from docx.oxml import OxmlElement

REPO = "/home/user/gamification/docs/thesis/"
SRC = REPO + "Full_version_V30_EngD_Thesis_Submission_Draft.docx"
DST = REPO + "Full_version_V31_EngD_Thesis_Submission_Draft.docx"
shutil.copy(SRC, DST)
d = docx.Document(DST)
log, fail = [], []

anchor = None
for p in d.paragraphs:
    if "engagement with clinical purpose intact is the open problem the adaptive-learning framework addresses" in p.text:
        anchor = p
        break
if anchor is None:
    fail.append("GrandPad paragraph anchor")
else:
    new = OxmlElement("w:p")
    anchor._p.addnext(new)
    np_ = docx.text.paragraph.Paragraph(new, anchor._parent)
    np_.add_run(
        "The comparator also points to a forward extension of the "
        "adaptive-learning proposition itself: from sustaining engagement "
        "within one health technology to supporting the acquisition of "
        "new technological skills in later life. The learning literature "
        "supports the premise that older adults acquire demanding new "
        "skills when challenge is sustained and scaffolded: the Synapse "
        "Project found that three months of sustained engagement in "
        "cognitively demanding novel skills improved episodic memory "
        "relative to controls (D. C. Park et al., 2014); the PRISM "
        "randomised trial showed that a purpose-designed computer system "
        "with training reduced loneliness and increased perceived social "
        "support among older adults at risk of isolation (Czaja et al., "
        "2018); and within this programme's own research lineage, "
        "gamified learning systems have improved older adults' technology "
        "adoption and knowledge transfer in mobile payments and mobile "
        "navigation (An et al., 2024, 2025). Read with the model of "
        "selective optimisation with compensation (Baltes & Baltes, "
        "1990), these results frame later-life technology learning as an "
        "adaptive process that succeeds when tasks are selected, "
        "challenge is optimised and support compensates — the same regime "
        "the adaptive-learning framework specifies through calibrated "
        "challenge, a visible personal trajectory and a coaching scaffold "
        "(§7.5). A platform of GrandPad's class — subscription-funded, "
        "barrier-free and socially scaffolded — is therefore a plausible "
        "carrier for the framework beyond FRA: a curriculum of adaptively "
        "calibrated games through which an older user builds the generic "
        "capabilities (touch interaction, menu navigation, confidence "
        "under initial failure) that transfer to new health technologies. "
        "This is stated as a design hypothesis, not a finding: the "
        "framework's mechanism-level claims remain untested (§8.8), and "
        "any such deployment would need the same measurement-bearability "
        "discipline this thesis develops. But it locates the "
        "contribution's growth path — adaptive learning not only as a "
        "retention mechanism for one clinical device, but as a general "
        "onboarding pathway by which older adults learn their way into "
        "technologies that benefit their lives.")
    log.append("extension paragraph inserted")

REFS = [
 ("Baltes",
  "Baltes, P. B., & Baltes, M. M. (1990). Psychological perspectives on "
  "successful aging: The model of selective optimization with "
  "compensation. In P. B. Baltes & M. M. Baltes (Eds.), Successful "
  "aging: Perspectives from the behavioral sciences (pp. 1-34). "
  "Cambridge University Press."),
 ("Czaja",
  "Czaja, S. J., Boot, W. R., Charness, N., Rogers, W. A., & Sharit, J. "
  "(2018). Improving social support for older adults through "
  "technology: Findings from the PRISM randomized controlled trial. "
  "The Gerontologist, 58(3), 467-477. "
  "https://doi.org/10.1093/geront/gnw249"),
 ("Park, D",
  "Park, D. C., Lodi-Smith, J., Drew, L., Haber, S., Hebrank, A., "
  "Bischof, G. N., & Aamodt, W. (2014). The impact of sustained "
  "engagement on cognitive function in older adults: The Synapse "
  "Project. Psychological Science, 25(1), 103-112. "
  "https://doi.org/10.1177/0956797613499592"),
]
paras = d.paragraphs
ref_i = next(i for i, p in enumerate(paras)
             if p.text.strip() == "References"
             and p.style.name.startswith("Heading") and i > len(paras)//2)
for key, entry in REFS:
    paras = d.paragraphs
    end_i = next(j for j in range(ref_i+1, len(paras))
                 if paras[j].style.name.startswith("Heading"))
    target = None
    for j in range(ref_i+1, end_i):
        t = paras[j].text.strip()
        if t and t > key:
            target = j
            break
    prevp = paras[target-1] if target else paras[end_i-1]
    new = OxmlElement("w:p")
    prevp._p.addnext(new)
    np_ = docx.text.paragraph.Paragraph(new, prevp._parent)
    np_.style = paras[target].style if target else prevp.style
    np_.add_run(entry)
    log.append(f"ref added: {key}")

cp = d.core_properties
cp.comments = "V31: adaptive-learning onboarding-pathway extension with verified learning-literature citations."
d.save(DST)
print("LOG:", *log, sep="\n  ")
print("FAIL:", fail if fail else "none")
