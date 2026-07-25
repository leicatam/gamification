#!/usr/bin/env python3
"""V24 — Chapter 6 / §3.2 / Ch8 / Appendix G rewritten to describe the
actual Stage-3 deployment, per the author's account (25 Jul 2026):
no Hong Kong sites; the game was ported to an HTML build and distributed
by e-mail through the author's networks (work contacts, charity community
network, university support, working peers in medical centres); played
unsupervised on participants' own computers (trackpad, arrow keys, or a
foot-pad sensor where available); each participant e-mailed back the
build's logged results CSV (with score) and questionnaire responses,
from which the preliminary aggregates were compiled. Also removes a
duplicated editing fragment in the §6.9 scope paragraph and replaces the
false §6.4/§8.9/G.3 archive claims (paper questionnaires / staff session
logs / terminal exports) with the true record trail. Figure 6.0A remains
retired. Attribution language widened: platform, input device, setting
and supervision changed alongside the framing package."""
import shutil
import docx

REPO = "/home/user/gamification/docs/thesis/"
SRC = REPO + "Full_version_V23_EngD_Thesis_Submission_Draft.docx"
DST = REPO + "Full_version_V24_EngD_Thesis_Submission_Draft.docx"
shutil.copy(SRC, DST)
d = docx.Document(DST)
log, fail = [], []


def set_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in list(p.runs[1:]):
            r._element.getparent().remove(r._element)
    else:
        p.add_run(text)


def find(needle):
    for p in d.paragraphs:
        if needle in p.text:
            return p
    return None


def replace(tag, needle, new_text):
    p = find(needle)
    if p is None:
        fail.append(tag)
    else:
        set_text(p, new_text)
        log.append(tag)


# (1) §1 programme overview
replace("1-overview",
    "The same broad hardware platform, skiing mechanic, session length and five-attempt rule were retained.",
    "One hundred and twenty adults aged 25-80 were recruited through the "
    "author's Hong Kong-based networks — work contacts, a charity community "
    "network, university support channels and the author's working peers in "
    "medical centres. Stage 3 was a remote deployment: the game was ported "
    "to an HTML build and distributed by e-mail, and participants played on "
    "their own computers, using a trackpad or arrow keys, or a foot-pad "
    "sensor where one was available. The skiing mechanic, session structure "
    "and five-attempt rule were retained in the ported build; the "
    "pressure-panel apparatus, the clinical setting and in-person "
    "supervision were not. Stage 3 thereby changed a package of "
    "participant-facing clinical cues: no clinical pre-scan was conducted, "
    "clinical and fall-risk terminology was removed, participant-facing "
    "clinical outputs were withheld, recruitment described a coordination "
    "challenge, and the interface was re-skinned. Observed non-return was "
    "6.7% overall and 14.6% in the 65-80 band. Among 23 initial dropouts, "
    "15 returned voluntarily; in the 65-80 band, 10 of 16 returned. "
    "Coordination-challenge endorsement reached 71.7% and is treated as "
    "evidence that the intended frame was perceived, while regular-use "
    "intention reached 63.3% overall.")

# (2) §1 author position
p = find("Stage 3, by contrast, was conducted with a separate Hong Kong community cohort.")
if p is None:
    fail.append("1-position")
else:
    set_text(p, p.text.replace(
        "Stage 3, by contrast, was conducted with a separate Hong Kong community cohort.",
        "Stage 3, by contrast, was conducted remotely with a separate "
        "cohort recruited through the author's Hong Kong-based networks, "
        "playing an e-mailed HTML build of the game on their own computers."))
    log.append("1-position")

# (3) §1 scope
p = find("Stage 3 involved a Hong Kong community-recruited cohort.")
if p is None:
    fail.append("1-scope")
else:
    set_text(p, p.text.replace(
        "Stage 3 involved a Hong Kong community-recruited cohort. Both were "
        "Chinese-cultural in setting and language,",
        "Stage 3 involved a Hong Kong network-recruited cohort playing "
        "remotely on their own computers. Both were Chinese-cultural in "
        "language and cultural context,"))
    log.append("1-scope")

# (4) §2/§3 attribution status paragraph
replace("attribution-status",
    "Stage 3 retained the same broad pressure-panel apparatus, skiing mechanic, session length and five-attempt rule",
    "Stage 3 retained the skiing mechanic, session structure and "
    "five-attempt rule in an HTML port of the game, but not the "
    "pressure-panel apparatus or the supervised setting: the build was "
    "distributed by e-mail and played on participants' own computers with a "
    "trackpad, arrow keys or, where available, a foot-pad sensor, while the "
    "package of participant-facing clinical cues was changed and a separate "
    "Hong Kong network cohort was recruited. It is therefore a "
    "between-cohort quasi-experimental decoupling comparison, not a "
    "single-variable experiment. The increase from 0 of 21 observed "
    "Stage-2 re-engagements to 10 of 16 returns among initial dropouts in "
    "the Stage-3 65-80 group is a large directional signal. It supports "
    "the proposition that framing or related contextual features moderate "
    "post-failure return, but it does not isolate the effect of framing "
    "from all differences in cohort, setting, platform, input device, "
    "supervision, recruitment, prior exposure, staff interaction or "
    "observation window. Nor does it directly test the complete "
    "adaptive-learning framework.")

# (5) §3.2 Stage-3 method summary
replace("3.2-method",
    "Stage 3 was a between-cycle, between-cohort quasi-experimental field study",
    "Stage 3 was a between-cycle, between-cohort quasi-experimental field "
    "study with 120 adults aged 25-80 recruited through the author's Hong "
    "Kong-based networks (work contacts, a charity community network, "
    "university support channels and working peers in medical centres). "
    "The game was ported to an HTML build, distributed by e-mail, and "
    "played unsupervised on participants' own computers, with a trackpad "
    "or arrow keys, or a foot-pad sensor where one was available. The "
    "skiing mechanic, five-attempt rule and session structure were "
    "retained in the port, while participant-facing clinical framing was "
    "changed through multiple coordinated components: recruitment "
    "language, the absence of any pre-session clinical scan, removal of "
    "clinical labels and participant-facing clinical outputs, and a "
    "coordination-game interface. Because cohort, setting, platform, "
    "input device and supervision also differed, Stage 3 estimates the "
    "combined association of the decoupled condition and its context with "
    "behaviour; it evaluates a participant-facing redesign package in a "
    "separate cohort, and it does not isolate a single causal variable. "
    "Its methodological role is attribution: the Stage-2 rejection of H0 "
    "admits a dispositional reading (gamification does not transfer to "
    "older adults) and a coupling reading (the attachment to the clinical "
    "technology broke it), and these make opposite predictions for a "
    "decoupled game (§6.1). Stage 3 is therefore designed as the "
    "programme's positive control for gamification theory, while its "
    "multi-component, between-cohort character bounds the causal strength "
    "of that attribution. The cohort design serves this purpose "
    "deliberately: recruiting three age bands (25-44, 45-64, 65-80) "
    "builds the comparison group into the study itself, and licenses a "
    "second null hypothesis, H0-age — that under the decoupled condition "
    "there is no difference between age groups. Stage 3 nullifies it: "
    "initial engagement differs sharply by age (0 of 38 younger versus 16 "
    "of 41 older initial dropouts; Fisher's exact p < 0.0001), while "
    "non-return converges toward low levels in every band. The pattern is "
    "exactly what the identity-level account predicts — the burden that "
    "decoupling relieves is concentrated in the group whose identity the "
    "technology's purpose touches.")

# (6) §6.3 design paragraph
replace("6.3-design",
    "Stage 3 evaluates a participant-facing redesign package in a separate cohort. Stage 2 used a coupled clinical-game context",
    "Stage 3 evaluates a participant-facing redesign package in a separate "
    "cohort. Stage 2 used a coupled clinical-game context that included an "
    "InBody body-composition scan, clinical pass/fail language, parallel "
    "clinical recording and repeated use of the term \"assessment\". "
    "Stage 3 removed these participant-facing clinical cues by moving the "
    "game off the clinical apparatus altogether: the game was ported to an "
    "HTML build and sent to participants by e-mail, no pre-session InBody "
    "scan was conducted, recruitment and session language described a "
    "coordination challenge, clinical labels were removed, and the "
    "interface was re-skinned. The skiing mechanic, five-attempt rule and "
    "session structure were retained in the ported build; the "
    "pressure-panel apparatus, the clinical setting and in-person "
    "supervision were not. These changes constitute a multi-component "
    "decoupling package. The design retains the game mechanics but does "
    "not isolate lexical framing from procedural, architectural, platform "
    "and setting changes, nor from between-cohort differences.")

# (7) §6.4 cohort paragraph
replace("6.4-cohort",
    "Sessions used the same broad pressure-panel and skiing-game configuration as Stage 2",
    "One hundred and twenty adults aged 25-80 were recruited through the "
    "author's Hong Kong-based networks — work contacts, a charity "
    "community network, university support channels and working peers in "
    "medical centres — in three age bands: 38 aged 25-44, 41 aged 45-64 "
    "and 41 aged 65-80. Recruitment materials referred to a coordination "
    "challenge and did not use medical or fall-prevention language. Each "
    "participant received the HTML build by e-mail and played on their "
    "own computer, with a trackpad or arrow keys, or a foot-pad sensor "
    "where one was available; play was unsupervised. Of the 120 "
    "participants, 48 were male and 72 female, and all reported "
    "high-school education or above.")

# (8) §6.4 data-record paragraph
replace("6.4-records",
    "paper questionnaires, staff session logs and terminal exports",
    "Stage-3 data collection is complete. The participant-level records "
    "are the participants' returned materials: each participant e-mailed "
    "back the results file logged by the HTML build (a CSV including the "
    "session score and game-logged data) together with their "
    "questionnaire responses. The preliminary aggregates reported in this "
    "chapter were compiled from those returns as they arrived. Chapter 6 "
    "reports the preliminary aggregate record. The data logbook "
    "(Appendix G, item G.3) is structured to hold the participant-level "
    "compilation from the returned e-mails; its reconciliation sheet "
    "recomputes every preliminary indicator from the participant-level "
    "entries, so the compilation can be verified against the preliminary "
    "record as it is completed.")

# (9) §6.9 scope paragraph — also removes the duplicated fragment
replace("6.9-scope",
    "None of this dims the headline signal",
    "Stage 3's design choices set the scope of what it can say, and the "
    "thesis is comfortable inside that scope. It is a between-cycle, "
    "between-cohort comparison, so it speaks in strong associations "
    "rather than isolated causes. The decoupling intervention is a "
    "coordinated package delivered through a change of platform and "
    "setting — an e-mailed HTML build played unsupervised at home — which "
    "is how real deployments change frames, but which also means that "
    "platform, input device and supervision changed alongside framing. "
    "The proposed mechanisms (mental accounting, loss aversion, identity "
    "protection, stereotype threat) await their direct measures; "
    "volunteers chose to take part, as participants always do; follow-up "
    "is short, as first expeditions are; and the cohorts are "
    "Chinese-cultural, which makes the cross-cultural replication of "
    "§7.7 the natural next journey. The Stage-3 framing item measures "
    "frame recognition only, and the complete adaptive-learning framework "
    "was not yet deployed. None of this dims the headline signal — "
    "post-failure return, absent in Stage 2, became the majority "
    "behaviour — it simply names the confirmatory work that signal has "
    "earned.")

# (10) Ch8 objection: better-experience confound
replace("8-objection-better",
    "Concern: Stage 3 may have produced better outcomes",
    "Concern: Stage 3 may have produced better outcomes because the "
    "overall experience was simply better in ways not captured by the "
    "protocol. Response: the game mechanic, session structure and "
    "five-attempt rule were carried into the HTML port, but the platform, "
    "input device, setting, supervision, interface, pre-session procedure "
    "and recruitment all differed — Stage 3 was played remotely on "
    "participants' own computers rather than on the supervised "
    "pressure-panel station. \"Just better design\" and \"just a lighter "
    "context\" therefore remain viable mechanism-aspecific alternatives. "
    "Detailed implementation-fidelity records and within-cohort "
    "randomisation are required.")

# (11) Ch8 principal limitations
p = find("The principal limitations are as follows.")
if p is None:
    fail.append("8-limitations")
else:
    set_text(p, p.text.replace(
        "the Stage-3 decoupling package contains multiple changes;",
        "the Stage-3 decoupling package contains multiple changes, "
        "including a change of platform, input device, setting and "
        "supervision (a remotely played HTML build rather than the "
        "supervised pressure-panel station);"))
    log.append("8-limitations")

# (12) Ch8 conclusion summary
replace("8-conclusion",
    "Stage 3 examined a participant-facing decoupling package in a separate Hong Kong community cohort of 120 adults.",
    "Stage 3 examined a participant-facing decoupling package in a "
    "separate Hong Kong cohort of 120 adults recruited through the "
    "author's networks, playing an e-mailed HTML port of the game on "
    "their own computers. The skiing-game structure was retained in the "
    "port; the pressure-panel apparatus, clinical setting and in-person "
    "supervision were not, and clinical recruitment language, pre-session "
    "procedure, labels and participant-facing outputs were changed. "
    "Observed non-return was 6.7% overall and 14.6% in the 65-80 band. "
    "Fifteen of twenty-three initial dropouts returned overall; ten of "
    "sixteen returned in the 65-80 band. The result is a substantial "
    "behavioural difference and supports a frame-sensitive account, but "
    "the between-cohort design does not eliminate cohort and contextual "
    "alternatives. Coordination endorsement is interpreted as a "
    "manipulation check, not as a direct measure of latent learning "
    "orientation.")

# (13) Appendix D administration mode
p = find("The Stage-3 decoupling design used a short two-item post-session instrument")
if p is None:
    fail.append("appD-admin")
else:
    set_text(p, p.text.rstrip() + " In the Stage-3 remote deployment the "
        "instrument was distributed with the e-mailed game build, and "
        "participants returned their responses by e-mail together with "
        "the build's logged results file.")
    log.append("appD-admin")

# (14) Appendix G stack description
replace("appG-stack",
    "operational platform for both the in-frame skiing game and the decoupled Stage-3 condition",
    "The Stage-2 deployment ran on a five-layer architecture. The "
    "decoupled Stage-3 condition ran a ported HTML build of the L1 game "
    "layer on participants' own computers: L0 was replaced by trackpad or "
    "arrow-key input (or a foot-pad sensor where available), and the "
    "L3/L4 clinical pipeline was absent rather than merely suppressed. "
    "The full adaptive-learning reference stack adds three further "
    "layers (L5 User-Facing Trajectory, L6 Staff Coaching Protocol, "
    "L7 Adaptive-Learning Telemetry) without altering the underlying "
    "signal path.")

# (15) Appendix G L4 line
replace("appG-L4",
    "suppressed in the Stage-3 decoupled condition",
    "L4 · Clinical Output: Aggregated COP-derived balance index and "
    "category label, exposed to clinical staff in Stage 2; absent in the "
    "Stage-3 decoupled condition (the HTML build had no clinical "
    "pipeline).")

# (16) Appendix G item G.3
replace("appG-G3",
    "G.3  Stage-3 data logbook",
    "G.3  Stage-3 data logbook "
    "(Stage3_Data_Logbook_120cases_with_Reconciliation.xlsx): codebook "
    "with provenance note, participant/telemetry/questionnaire/outcome "
    "sheets for IDs 3001–3120, the preliminary aggregate results "
    "(Preliminary_Summary sheet) and a live Reconciliation sheet. "
    "Stage-3 data collection is complete; participant-level entries are "
    "compiled from the primary records of the remote deployment — the "
    "participants' returned e-mails, each carrying the HTML build's "
    "logged results file (CSV) and the questionnaire responses — with "
    "each entry checked by the reconciliation sheet.")

# (17) Appendix G item G.5
replace("appG-G5",
    "G.5  Photographs of the Hong Kong Stage-3 deployment",
    "G.5  The Stage-3 distribution and return e-mail record "
    "(de-identified index of dates, channels and returned files).")

cp = d.core_properties
cp.comments = "V24: Stage 3 rewritten as remote e-mailed HTML deployment; archive claims corrected; 6.9 duplicate removed."
d.save(DST)
print("LOG:", *log, sep="\n  ")
print("FAIL:", fail if fail else "none")
