#!/usr/bin/env python3
"""V32 — closing paragraph of the Ch8 GrandPad arc: what commercial
gerontechnology ignores (the avoided middle; the missing failure model;
static rather than adaptive personalisation), and the programme's
contribution to gerontology as AI-enabled health technology arrives.
Uses only already-cited references (Csikszentmihalyi, Vygotsky)."""
import shutil
import docx
from docx.oxml import OxmlElement

REPO = "/home/user/gamification/docs/thesis/"
SRC = REPO + "Full_version_V31_EngD_Thesis_Submission_Draft.docx"
DST = REPO + "Full_version_V32_EngD_Thesis_Submission_Draft.docx"
shutil.copy(SRC, DST)
d = docx.Document(DST)

anchor = None
for p in d.paragraphs:
    if "onboarding pathway by which older adults learn their way into technologies that benefit their lives" in p.text:
        anchor = p
        break
assert anchor is not None, "V31 extension paragraph not found"

new = OxmlElement("w:p")
anchor._p.addnext(new)
np_ = docx.text.paragraph.Paragraph(new, anchor._parent)
np_.add_run(
    "Read as a boundary case, the comparator also sharpens what "
    "commercial gerontechnology has so far left unaddressed, and locates "
    "this programme's contribution to gerontology as AI-enabled health "
    "technology arrives. Three omissions matter. First, the avoided "
    "middle: the market divides into clinical instruments that measure "
    "but are abandoned (the FRA of Chapter 4) and consumer platforms "
    "that engage but do not measure (the GrandPad class); the "
    "technologies that population ageing actually requires — instruments "
    "that assess fall risk, cognition and function, repeatedly and "
    "acceptably — sit in the middle that both poles avoid. Second, the "
    "missing failure model: GrandPad-class design treats disengagement "
    "as a usability and access problem and answers it by minimising "
    "occasions for failure; it carries no account of what an adverse "
    "result does to an older user's self-concept, because it never "
    "delivers one. Curated ease also caps challenge, and with it the "
    "mastery experiences on which learning and flow depend "
    "(Csikszentmihalyi, 1990; Vygotsky, 1978). Third, personalisation "
    "without adaptation: content is personalised — family, photographs, "
    "curated games — but challenge is not calibrated to the individual "
    "and does not move with them, so there is no rising trajectory for "
    "the user to read. The arrival of AI raises the stakes on all three "
    "counts at once. Adaptive models make individually calibrated "
    "challenge inexpensive to implement at scale — the enabling "
    "technology for the Chapter 7 framework — but the same wave is "
    "filling gerontechnology with continuous monitoring, risk scores and "
    "predictions: more verdicts about ageing bodies, delivered more "
    "often, to the population this thesis shows may respond to "
    "unfavourable verdicts by leaving permanently. Without a behavioural "
    "layer, AI-era gerontechnology risks industrialising the abandonment "
    "pattern documented here. The programme's contribution to "
    "gerontology is that layer: FIAS names the failure mode that "
    "verdict-bearing technology must design against; the Stage-3 "
    "frame-sensitivity evidence indicates where the leverage lies; and "
    "the adaptive-learning framework specifies how a measuring "
    "technology can remain bearable — challenge calibrated to the "
    "person, results presented as a trajectory, coaching absorbing the "
    "failure moment. The gamified FRA is the working demonstration that "
    "the middle ground is occupiable: one physical action that plays and "
    "measures, made returnable-to by design. In an AI decade, that is "
    "the difference between health technologies older adults live with "
    "and health technologies that fall silent (§8.6B).")

cp = d.core_properties
cp.comments = "V32: comparator-boundary argument — what gerontechnology misses; the contribution as AI arrives."
d.save(DST)
print("V32 saved")
