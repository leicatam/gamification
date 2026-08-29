#!/usr/bin/env python3
"""V23 — Figure 4.1B inserted: the author's photograph of the FRA balance
station in the Guangzhou facility rehabilitation room (supplied 25 Jul
2026; original archived in evidence/, EXIF absent in transit). Caption
describes the visible modular configuration and cross-references the
same platform under game control in Figure 5.1A. LoF entry added."""
import shutil
import docx
from docx.oxml import OxmlElement
from docx.shared import Inches

REPO = "/home/user/gamification/docs/thesis/"
SRC = REPO + "Full_version_V22_EngD_Thesis_Submission_Draft.docx"
DST = REPO + "Full_version_V23_EngD_Thesis_Submission_Draft.docx"
shutil.copy(SRC, DST)
d = docx.Document(DST)
log, fail = [], []


def find(needle, style_prefix=None):
    for p in d.paragraphs:
        if needle in p.text and (style_prefix is None
                                 or p.style.name.startswith(style_prefix)):
            return p
    return None


def insert_after(p, text="", style=None):
    new = OxmlElement("w:p")
    p._p.addnext(new)
    np_ = docx.text.paragraph.Paragraph(new, p._parent)
    if style:
        np_.style = style
    if text:
        np_.add_run(text)
    return np_


CAPTION = ("Figure 4.1B. The FRA balance station as deployed in the "
    "rehabilitation room of the Guangzhou facility (photograph by the "
    "author). The four-load-cell sensing surface — the dark glass panel "
    "marked with foot-placement guides — is set into an anti-slip base "
    "plate and flanked by support rails; the operator's monitor and "
    "keyboard stand on a mobile cart behind the platform. The "
    "configuration is the modular arrangement described in §4.1A "
    "(platform, support rails, operator computer); facility signage "
    "designates the occupational-therapy training area and the "
    "balance-test zone. The same platform is shown under game control "
    "in Figure 5.1A.")

cap41a = find("Figure 4.1A. Fall Risk Assessment (FRA) system architecture",
              "FigureCaption")
if cap41a is None:
    fail.append("Figure 4.1A caption")
else:
    imgp = insert_after(cap41a)
    imgp.add_run().add_picture(REPO + "figures/fig_4_1B.jpg",
                               width=Inches(3.4))
    capp = insert_after(imgp, CAPTION)
    try:
        capp.style = "FigureCaption"
        log.append("Figure 4.1B image + caption inserted")
    except Exception:
        fail.append("FigureCaption style")

lof41a = find("Figure 4.1A. Fall Risk Assessment (FRA) system architecture",
              "toc")
if lof41a is None:
    fail.append("LoF Figure 4.1A line")
else:
    t = insert_after(lof41a,
        "Figure 4.1B. The FRA balance station as deployed in the "
        "rehabilitation room of the Guangzhou facility.\t36")
    t.style = lof41a.style
    log.append("LoF entry for Figure 4.1B inserted")

cp = d.core_properties
cp.comments = "V23: Figure 4.1B (FRA station photograph) inserted with LoF entry."
d.save(DST)
print("LOG:", *log, sep="\n  ")
print("FAIL:", fail if fail else "none")
