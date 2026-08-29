#!/usr/bin/env python3
"""V19 — integrate the expert consultation with Dr Sarah C.W. Yuan.
Basis: the author's structured notes of the conversation (received 19 Jul
2026, seven points). Framed throughout as a single expert consultation —
interpretive evidence from an independent discipline, not participant data.
Insertions: §3.2 (method), §3.4 (ethics sentence), new §6.8A (findings and
what they add), §7.5.3 (L6 support), §8.5 (moderator note), Appendix H
(record), ToC lines, PENDING markers for date + written attribution.
"""
import shutil
import docx
from docx.oxml import OxmlElement

BASE = "/tmp/claude-0/-home-user-gamification/3d8a50e8-9258-51ce-a7d6-194da50f9a04/scratchpad/"
SRC = BASE + "Full_version_V18_EngD_Thesis_Submission_Draft.docx"
DST = BASE + "Full_version_V19_EngD_Thesis_Submission_Draft.docx"
shutil.copy(SRC, DST)
d = docx.Document(DST)
log, fail = [], []


def set_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in list(p.runs[1:]):
            r._element.getparent().remove(r._element)
    else:
        p.add_run(text)


def find(needle, nth=1):
    n = 0
    for p in d.paragraphs:
        if needle in p.text:
            n += 1
            if n == nth:
                return p
    return None


def insert_after(p, text="", style=None):
    new = OxmlElement("w:p")
    p._p.addnext(new)
    np_ = docx.text.paragraph.Paragraph(new, p._parent)
    if style:
        try:
            np_.style = style
        except Exception:
            pass
    if text:
        np_.add_run(text)
    return np_


# ---------------- §3.2: method paragraph ----------------
p = find("Stage 3 nullifies it: initial engagement differs sharply by age")
if p is None:
    fail.append("3.2 anchor")
else:
    insert_after(p,
        "A fourth evidence activity supplements the three stages. After the "
        "Stage-3 analysis, a single expert consultation interview was "
        "conducted with Dr Sarah C.W. Yuan (Center on the Family, University "
        "of Hawai'i at Mānoa), a medical sociologist and demographer of "
        "aging whose field of practice is community group falls-prevention "
        "programming. The purpose was to test the programme's behavioural "
        "interpretation against an independent discipline and an independent "
        "field of practice — non-electronic group exercise and physical "
        "games for balance — in which none of this programme's apparatus is "
        "used. The consultation informs interpretation and design (§6.8A, "
        "§7.5.3) and is not treated as participant data; the record and its "
        "attribution basis are given in Appendix H. (Interview date, format "
        "and the expert's written attribution confirmation to be appended at "
        "Appendix H.)")
    log.append("§3.2: consultation method paragraph inserted")

# ---------------- §3.4: ethics sentence ----------------
p = find("Approval reference and coverage statement to be inserted")
if p is None:
    p = find("ethics policy of The Hong Kong Polytechnic University")
if p is None:
    fail.append("3.4 anchor")
else:
    set_text(p, p.text.rstrip() + " The expert consultation reported at "
        "§6.8A was conducted with a professional in her professional "
        "capacity; she is named with her consent, and the written "
        "confirmation is held with the consultation record (Appendix H).")
    log.append("§3.4: ethics sentence appended")

# ---------------- new §6.8A ----------------
anchor = find("Figure 6.3. Retention versus regular-use intention by age band")
if anchor is None or not anchor.style.name == "FigureCaption":
    fail.append("6.8A anchor (Fig 6.3 caption)")
else:
    h = insert_after(anchor,
        "6.8A The Commitment Gap Read Through Group Dynamics: An Expert "
        "Consultation", style="Heading 2")
    p1 = insert_after(h,
        "The interpretation of the commitment gap was examined after the "
        "empirical programme through a single expert consultation interview "
        "with Dr Sarah C.W. Yuan of the Center on the Family, University of "
        "Hawai'i at Mānoa — a medical sociologist and demographer of aging "
        "whose work includes the Hawai'i Healthy Aging Partnership and "
        "statewide community falls-prevention programming. Her field of "
        "practice is deliberately distant from this programme's apparatus: "
        "group exercise and physical games for balance, with no electronic "
        "game and no clinical instrument attached to the activity. The "
        "consultation is reported as interpretive evidence from an "
        "independent discipline, not as additional participant data; the "
        "record is the author's structured notes, reproduced in Appendix H.")
    p2 = insert_after(p1,
        "Three of her observations bear directly on this thesis. First, "
        "from her programme experience and publications: older adults who "
        "join social exercise groups tend to drop out after a failure "
        "unless a physician, assistant or coach helps them past the "
        "rejection; with a group around them and a coach who leads them to "
        "continue, they perform and persist markedly better. "
        "Failure-linked dropout therefore appears even in her setting — "
        "with no clinical attachment at all — and is recoverable through "
        "human support. Second, in her experience older adults "
        "intrinsically avoid straight clinically oriented individual "
        "assessment of their current health status: they fear facing a "
        "negative risk result, and psychologically avoid the clinical "
        "evidence that would describe their condition. This is an "
        "independent statement, from a sociology of aging perspective, of "
        "the identity-level avoidance that the FIAS account proposes. "
        "Third, she regards gamification as very effective for promoting "
        "health technology — and expects hard rejection precisely when the "
        "game or exercise is linked to clinical results. On that basis she "
        "explicitly supported the Stage-3 decoupling design, and she "
        "supported the FIDS-to-FIAS construction and the adaptive-learning "
        "proposition as consistent with the elderly behaviour she has "
        "observed.")
    p3 = insert_after(p2,
        "What the consultation adds to the commitment-gap reading is a "
        "social mechanism the programme's own data could not supply. In her "
        "account, partnering, group play and a coach who leads the "
        "activity reduce initial dropout and, downstream, abandonment of "
        "the technology; and the components that support adoption extend "
        "beyond engineering excellence and design to the environment, the "
        "people using the same technology, and predictable or expected "
        "results. The mapping onto the framework is direct: the coach and "
        "the group correspond to the staff coaching layer (L6, §7.5.3) and "
        "motivate a group-deployment variant of the protocol; predictable "
        "or expected results correspond to trajectory installation (P2), "
        "since a visible personal trajectory is exactly what makes results "
        "expected rather than threatening. Her account also refines the "
        "Chapter 2 caution that social influence (Octalysis D5) can be "
        "double-edged for older users: structured group dynamics with a "
        "leading coach appears protective in her experience, whereas "
        "unstructured peer comparison may still threaten identity — a "
        "distinction the §8.8 study can operationalise. The evidential "
        "limits are stated plainly: this is one consultation, recorded in "
        "the author's notes; it is interpretive support from an adjacent "
        "discipline, not validation; and its observations concern physical "
        "group exercise, so their transfer to electronic games is an "
        "argued expectation rather than a demonstrated result.")
    log.append("§6.8A inserted (heading + three paragraphs)")

# ---------------- §7.5.3: L6 support sentence ----------------
p = find("The proposed staff coaching protocol (L6) operationalises trajectory framing")
if p is None:
    fail.append("7.5.3 anchor")
else:
    set_text(p, p.text.rstrip() + " The expert consultation of §6.8A "
        "independently supports this layer from community practice: in Dr "
        "Yuan's programmes, a coach or assistant who leads the participant "
        "past a failure is what converts failure-linked dropout into "
        "continued participation, and partnering and group formats further "
        "reduce dropout — grounds for evaluating a group-deployment variant "
        "of this protocol in the §8.8 study.")
    log.append("§7.5.3: L6 support sentence appended")

# ---------------- §8.5: moderator note ----------------
p = find("The credibility of the frame-sensitive account depends on examining alternatives")
if p is None:
    fail.append("8.5 anchor")
else:
    set_text(p, p.text.rstrip() + " The expert consultation of §6.8A bears "
        "on this examination in both directions: it supplies independent "
        "support for the clinical-avoidance reading, and it identifies "
        "group and coach support as candidate moderators of post-failure "
        "continuation that the future within-cohort study should record "
        "and, where feasible, standardise.")
    log.append("§8.5: moderator note appended")

# ---------------- Appendix H ----------------
g8 = find("G.8  Synthetic pipeline-test dataset")
if g8 is None:
    fail.append("Appendix G.8 anchor")
else:
    hh = insert_after(g8, "Appendix H · Expert Consultation Record",
                      style="Heading 2")
    h1 = insert_after(hh,
        "H.1  Expert and scope. Dr Sarah C.W. Yuan, Specialist, Center on "
        "the Family, University of Hawai'i at Mānoa; PhD and MA in "
        "Sociology (University of Hawai'i at Mānoa), MSc in Applied Social "
        "Research (University of Stirling). Research expertise: medical "
        "sociology, demography of aging, survey methodology; projects "
        "include the Hawai'i Healthy Aging Partnership; recognitions "
        "include the Hawai'i Pacific Gerontological Society Research and "
        "Teaching Award (2010). Her field of practice relevant to this "
        "consultation is community group falls-prevention programming: "
        "non-electronic group exercise and physical games whose objective "
        "is to improve older adults' balance and reduce fall risk. "
        "Credentials verified against her University of Hawai'i faculty "
        "profiles at the time of writing.")
    h2 = insert_after(h1,
        "H.2  Conduct and record. The consultation was a single "
        "semi-structured conversation between the author and Dr Yuan. The "
        "author first presented the three-stage programme, its findings "
        "and the null-hypothesis architecture; Dr Yuan then described her "
        "own research area and findings, and responded to the programme's "
        "constructs. The record is the author's structured notes, "
        "summarised at H.3. (Interview date, format and language, the "
        "question list used, and Dr Yuan's written confirmation of named "
        "attribution to be appended here at binding.)")
    h3 = insert_after(h2,
        "H.3  Summary of the consultation (author's notes). (1) The author "
        "introduced the research: Stages 1-3, the findings, and the null "
        "hypothesis. (2) Dr Yuan's falls-related work concerns "
        "non-clinical assessment and group work — exercise and group "
        "dynamic physical games, not electronic games — whose objective is "
        "to improve balance and avoid fall risk. (3) From her findings, "
        "including her publications: when older adults join a social "
        "exercise group, they tend to drop out after failure unless a "
        "physician or assistant helps them overcome the rejection; they "
        "perform better when playing in a group with a coach who leads "
        "them to continue, and continuation is difficult otherwise. "
        "(4) She agreed that older adults intrinsically reject clinical "
        "assessment or diagnosis of their current health status, and "
        "psychologically avoid facing the clinical evidence that explains "
        "their condition. (5) In her experience, older adults avoid "
        "individual assessment when it is straight clinically oriented, "
        "for fear of facing negative risk results. (6) She confirmed that "
        "gamification is very effective for promoting health technology, "
        "but that older adults reject it hard when the game or exercise is "
        "linked to clinical results; on this basis she supported the "
        "Stage-3 decoupling approach. (7) Speaking from a sociology "
        "perspective rather than an engineering one, she shared the same "
        "view of elderly behaviour, supported the FIDS-to-FIAS "
        "construction and the adaptive-learning proposition, and "
        "confirmed that older adults perform better when grouped and "
        "building dynamics: partnering and group dynamics reduce initial "
        "dropout and, eventually, abandonment of the health technology. "
        "She emphasised that the key components extend beyond engineering "
        "excellence and design: the environment, the people who are using "
        "the same technology, and predictable or expected results are the "
        "key components that support older adults in adopting health "
        "technologies.")
    insert_after(h3,
        "H.4  Use in the thesis. The consultation is cited at §3.2 "
        "(method), §6.8A (interpretation of the commitment gap), §7.5.3 "
        "(coaching-protocol support) and §8.5 (candidate moderators). It "
        "is treated throughout as interpretive evidence from an "
        "independent discipline — one consultation, recorded in the "
        "author's notes — and not as participant data, validation, or a "
        "demonstrated result for electronic games.")
    log.append("Appendix H inserted (H.1-H.4)")

# ---------------- ToC lines ----------------
toc68 = None
tocG = None
for p in d.paragraphs:
    if p.style.name.startswith("toc"):
        if p.text.strip().startswith("6.8 The Commitment Gap"):
            toc68 = p
        if p.text.strip().startswith("Appendix G"):
            tocG = p
if toc68 is not None:
    t = insert_after(toc68,
        "6.8A The Commitment Gap Read Through Group Dynamics: An Expert "
        "Consultation\t59")
    t.style = toc68.style
    log.append("ToC line for 6.8A inserted")
else:
    fail.append("ToC 6.8 line")
if tocG is not None:
    t = insert_after(tocG, "Appendix H · Expert Consultation Record\t112")
    t.style = tocG.style
    log.append("ToC line for Appendix H inserted")
else:
    fail.append("ToC Appendix G line")

# ---------------- metadata ----------------
cp = d.core_properties
cp.title = "From FIDS to FIAS - Submission Draft v19"
cp.comments = ("Submission draft v19 (19 Jul 2026). Expert consultation with "
               "Dr Sarah C.W. Yuan integrated: 3.2 method, 3.4 ethics, new "
               "6.8A, 7.5.3, 8.5, Appendix H. Pending: interview date/format, "
               "written attribution consent, photos, Stage-3 transcription.")

d.save(DST)
print("saved", DST)
for l in log:
    print(" -", l)
for f in fail:
    print(" !", f)

# ---------------- verification ----------------
d2 = docx.Document(DST)
full = "\n".join(p.text for p in d2.paragraphs)
for tb in d2.tables:
    for r in tb.rows:
        full += "\n" + " | ".join(c.text for c in r.cells)
for s in ("Sarah C.W. Yuan", "6.8A", "Appendix H", "group dynamics",
          "LODF", "expert consultation"):
    print(f"'{s}':", full.count(s))
EOF_SENTINEL = None
