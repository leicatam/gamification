#!/usr/bin/env python3
"""V37 REVIEW COPY - the integration pass: four-expert panel (3.2, Ch8,
H.6-H.8), assess-vs-train scoping defence (Ch4, 8.7), Stage 3B +
errorless/graceful positioning (8.8), Ch9 synthesis, family-channel
correction, three new references, YELLOW highlights on all pending
author inputs + front checklist."""
import shutil
import docx
from docx.oxml import OxmlElement
from docx.enum.text import WD_COLOR_INDEX

REPO = "docs/thesis/"
SRC = REPO + "Full_version_V36_EngD_Thesis_Submission_Draft.docx"
DST = REPO + "Full_version_V37_REVIEW_EngD_Thesis_Submission_Draft.docx"
shutil.copy(SRC, DST)
d = docx.Document(DST)
log, fail = [], []

def para_after(prev, text, yellow=False, style=None):
    new = OxmlElement("w:p")
    prev._p.addnext(new)
    np_ = docx.text.paragraph.Paragraph(new, prev._parent)
    if style: np_.style = style
    r = np_.add_run(text)
    if yellow: r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return np_

def find_para(substr):
    for p in d.paragraphs:
        if substr in p.text:
            return p
    return None

def find_heading(prefix):
    for p in d.paragraphs:
        if p.style.name.startswith("Heading") and p.text.strip().startswith(prefix):
            return p
    return None

def set_text(p, text):
    for r in list(p.runs)[1:]:
        r._r.getparent().remove(r._r)
    if p.runs: p.runs[0].text = text
    else: p.add_run(text)

# ---------- 0. FRONT YELLOW CHECKLIST ----------
first = d.paragraphs[0]
chk = para_after(first,
 "REVIEW COPY V37 - AUTHOR INPUT CHECKLIST (this page and every yellow mark "
 "must be resolved and removed before submission). Outstanding inputs: "
 "(1) exact HSESC approval date from the approval memo (s3.4 says 'August 2026'); "
 "(2) SS-293 camera-ready text for the consent-statement check; "
 "(3) the as-sent Mar-Apr 2026 participant invitation original (Appendix D/G.5); "
 "(4) photos of the two Stage-2 paper instruments (Appendix C); "
 "(5) verbatim Mandarin Stage-1 survey items + response CSV (Appendix A, G.1); "
 "(6) FIAS thirteen-attribute coding table + inter-coder kappa (Appendix B); "
 "(7) D1-D2 response scale (5 vs 7) + bilingual wording (Appendix D); "
 "(8) Stage-2 AI endpoint record - vendor/model/prompts/config (G.7) and whether "
 "session telemetry/AI-call logs were retained (G.2 or one disclosed sentence); "
 "(9) control person's written consent for Figure 5.1A (G.8); "
 "(10) Stage-3 participant-level logbook: distributor wrappers + questionnaire "
 "returns, or the disclosure sentence that compilation remains partial (s6.4); "
 "(11) Stage 3B results when the window closes (s8.8); "
 "(12) Dr Lau second consultation: game feedback, Mirror clearance, member-check "
 "(s8.5A, H.6); (13) member-checks for Dr Choy's and Dr Lau's summaries; "
 "(14) screenshots to re-send as files (dated invitations, Choy/Wong/Lau WhatsApp "
 "chains, HSEARS dashboard, ICMSS acceptance letter); "
 "(15) final Word field refresh (Ctrl+A, F9) + full verification sweep + proofread.",
 yellow=True)
log.append("front checklist")

# ---------- 1. s3.2 four-consultation method ----------
p = find_para("A fourth evidence activity supplements the three stages. After the Stage-3 analysis, a single expert consultation interview was conducted")
if p is None: fail.append("3.2 anchor")
else:
    set_text(p,
     "A fourth evidence activity supplements the three stages: a series of four "
     "expert consultations, conducted after the empirical work as a distinct "
     "data-collection exercise - not literature review - to test the programme's "
     "behavioural interpretation and its design framework against senior practice "
     "and research standpoints. The consultations were semi-structured interviews "
     "with: a senior physiotherapist and former chair of a Hong Kong professional "
     "physiotherapy body, pseudonymised as Expert P at his election (March 2026, "
     "telephone, approximately 25 minutes, Cantonese mixed with English; with a "
     "dated follow-up on 5 August 2026 in which he played the decoupled build); "
     "Dr Gary Kui-Kai Lau, Assistant Dean (Education Innovations) and Clinical "
     "Associate Professor, Division of Neurology, LKS Faculty of Medicine, The "
     "University of Hong Kong (June 2026, in person, approximately one hour, "
     "Cantonese mixed with English); Dr Sarah C.W. Yuan, Specialist, Center on "
     "the Family, University of Hawai'i at Manoa, a medical sociologist and "
     "demographer of aging (25 July 2026, telephone, English, approximately 35 "
     "minutes); and Dr Philip Choy, certified naturopathic physician (ANMCB) and "
     "President of the Hong Kong Academy of Naturopath, with over twenty-five "
     "years of clinical practice concentrated in elderly chronic disease "
     "(4 August 2026, telephone, approximately 30 minutes, Cantonese and "
     "English). Each expert gave written consent on a nine-statement form with "
     "an explicit identification election: three authorised named attribution "
     "and one elected pseudonymity; where an interview preceded its signed form, "
     "the record states both dates plainly. The researcher's notes are the "
     "primary records; each consultation summary is returned to its participant "
     "for validation (member checking). Analysis is thematic: expert statements "
     "are mapped onto the programme's constructs and reported as convergence, "
     "divergence or silence (Appendix H; discussed in Chapter 8). Positioning: "
     "expert input is interpretation and practice validation of the findings and "
     "a confirmatory-design review of the Chapter 7 framework; it is not "
     "additional empirical evidence and does not bear on the hypothesis tests.")
    log.append("3.2 rewritten")

# ---------- 2. Ch8 consultation analysis ----------
p = find_para("The interpretation of the commitment gap was examined after the empirical programme through a single expert consultation interview with Dr Sarah C.W. Yuan")
if p is None: fail.append("Ch8 Yuan anchor")
else:
    t = p.text
    set_text(p, t.replace(
      "through a single expert consultation interview with",
      "through expert consultation, beginning with"))
    p1 = para_after(p,
     "Three further consultations extend the validation across clinical-practice "
     "standpoints; their full records are in Appendix H.6-H.8, and the analysis "
     "below reports convergence, divergence and silence rather than agreement "
     "alone. Dr Philip Choy, drawing on a caseload of more than a thousand "
     "elderly chronic-disease patients, was asked the programme's definitional "
     "question directly - is the operative failure the game's, or the health "
     "result's? He agreed that the game itself does not cause the abandonment "
     "syndrome: it arises mainly from users' concerns about their own health "
     "performance, a pattern he described as very common among elderly people "
     "who have acknowledged a principal health condition. This is an "
     "independent, practice-based statement of the attribution claim at the "
     "centre of this thesis. He further located the adopting demographic - "
     "educated, financially independent adults from roughly 65-70 upward, "
     "engaging only once a concrete benefit is identified - and judged token "
     "reward economies insufficient for that segment.")
    p2 = para_after(p1,
     "Expert P, the panel's physiotherapist, contributed an abandonment window "
     "of two weeks to two months from post-surgical rehabilitation caseloads - "
     "an adjacent, clinically supervised population, cited here as analogy "
     "rather than corroboration of the FRA timing - together with three design "
     "observations: some assessments (balance, fall risk, cognition) cannot be "
     "simplified, supporting the occupied-middle argument of Figure 8.1; games "
     "pitched too simply read as childish to older players; and, most "
     "distinctively, that higher-educated patients who are observed failing "
     "will probably stop immediately - failure witnessed is a stronger stopping "
     "trigger than failure alone, which retrospectively supports Stage 3's "
     "choice of private, unsupervised play. In a dated follow-up he played the "
     "decoupled build itself and judged it suitable for fall-risk assessment, "
     "stating he would return to it regularly were it installed - the panel's "
     "first artifact-level validation, mirroring the study's intention "
     "construct behaviourally.")
    p3 = para_after(p2,
     "Dr Gary Lau's consultation bears on the framework's costliest element: "
     "coaching. In his account, coaching is always the best and most direct "
     "approach to sustaining rehabilitation - and it cannot scale: family "
     "involvement carries too high a social cost, and professional therapist "
     "resources are scarce. That pair of statements is the clinical case for "
     "the framework's coaching layer, and for AI's enabling role as the "
     "affordable way to give every user a coach. His diagnosis that current "
     "health technology offers no progressive, positive signs across a long, "
     "boring and lonely journey grounds the visible-trajectory element; and "
     "his stated position that technology should be embraced but not "
     "over-relied upon - care remaining people-centred - independently states "
     "the boundary this thesis draws for AI.")
    p3b = para_after(p3,
     "[AUTHOR INPUT NEEDED - Dr Lau second consultation: his played-the-game "
     "feedback and the Mirror-project convergence (avatar coaching; "
     "self-awareness and self-engagement as the key mechanism) await his "
     "explicit clearance and the author's record before this paragraph is "
     "completed.]", yellow=True)
    p4 = para_after(p3b,
     "The analysis records its own limits. The verdict-attribution "
     "corroboration rests on Dr Choy alone - the question was not put to the "
     "other experts. Dr Choy's and Expert P's observations share a "
     "private-practice sampling frame, and their convergence on the adopter "
     "demographic is weighted accordingly. Expert P's abandonment window "
     "derives from medically necessary, supervised rehabilitation rather than "
     "voluntary preventive assessment. And Dr Lau's honest uncertainty - "
     "whether games can change patients' minds and behaviour, he cannot yet "
     "be sure - is reported as his position, not resolved away. Expert views "
     "contextualise the behavioural evidence; they are not evidence for the "
     "hypotheses.")
    log.append("Ch8 consultation analysis inserted")

# ---------- 3. Ch4 scoping ----------
p = find_para("FRA is therefore the sensor-based artefact whose engagement profile is examined.")
if p is None: fail.append("Ch4 scoping anchor")
else:
    para_after(p,
     "A scoping statement belongs here. The commercial FRA proposition - in "
     "the InBody line as in its HUR antecedent - is not assessment alone but "
     "an assess-then-train model: measurement identifies fall risk, and a "
     "training programme on the same platform is intended to reduce it. This "
     "thesis studies the assessment gate only, and does so deliberately: the "
     "programme's central finding is that utilisation collapses at that gate, "
     "and when older adults refuse or abandon assessment, the training phase "
     "never begins. The assessment gate is thus the binding constraint on the "
     "entire model's preventive value, and its acceptance is the "
     "highest-leverage object of study. The training phase, and the extension "
     "of this thesis's framework to it, are addressed as an explicit "
     "limitation and as future work (s8.7, s8.8).")
    log.append("Ch4 scoping inserted")

# ---------- 4. s8.7 limitation ----------
h = find_heading("8.7")
if h is None: fail.append("8.7 heading")
else:
    para_after(h,
     "One scope limitation is stated before all others: this study covers the "
     "assessment phase of an assess-then-train product model and does not "
     "evaluate the training phase. Four considerations govern the boundary. "
     "First, the assessment gate is the binding constraint: abandonment there "
     "forecloses training altogether, and the non-utilisation figures that "
     "motivate the programme arise at this gate. Second, the mechanism "
     "documented here - verdict-driven abandonment - applies with equal force "
     "to the training phase, which delivers progress verdicts of its own; the "
     "framework is designed for, and extensible to, the full loop. Third, a "
     "supervised training intervention with elderly participants demands "
     "clinical resources and a longitudinal window beyond this part-time "
     "doctorate's scope. Fourth, the extension is specified rather than "
     "abandoned: the confirmatory programme of s8.8 carries the framework "
     "into the training phase.")
    log.append("8.7 limitation inserted")

# ---------- 5. s8.8 Stage 3B + positioning ----------
h = find_heading("8.8")
if h is None: fail.append("8.8 heading")
else:
    pA = para_after(h,
     "Since this section was first drafted, the first element of the "
     "confirmatory programme has been designed and pre-registered: Stage 3B, "
     "a 2x2 field experiment crossing adaptive against static presentation "
     "with zero-failure against failure-possible modes, in a new cohort of "
     "80-120 adults recruited in the same three age bands, with recorded "
     "in-build consent, a single pre-specified confirmatory contrast, and a "
     "trilingual instrument implementing the framework's calibration, "
     "trajectory and coaching elements. The design also brings into contact, "
     "for older adults in a health-technology context, two traditions the "
     "literature has kept apart: errorless learning, which holds that "
     "vulnerable learners acquire best when prevented from erring (Baddeley "
     "& Wilson, 1994), and the graceful-failure position of game-based "
     "learning, which treats low-consequence failure as productive (Plass, "
     "Homer, & Kinzer, 2015); dynamic difficulty adjustment supplies the "
     "adaptive arm's engineering lineage (Zohaib, 2018).")
    para_after(pA,
     "[AUTHOR INPUT NEEDED - Stage 3B results: this passage is completed when "
     "the collection window closes; report all pre-specified outcomes "
     "regardless of direction.]", yellow=True)
    log.append("8.8 Stage 3B inserted")

# ---------- 6. Ch9 synthesis before References ----------
paras = d.paragraphs
ref_i = None
for i, p in enumerate(paras):
    if p.text.strip() == "References" and p.style.name.startswith("Heading") and i > len(paras)//2:
        ref_i = i; break
if ref_i is None: fail.append("References heading")
else:
    prev = paras[ref_i-1]
    para_after(prev,
     "The programme's synthesis can now be stated with each claim at its "
     "evidential strength. It observed failure-induced disengagement (Stages "
     "1 and 2). It interprets the escalation into abandonment as "
     "verdict-driven - an interpretation corroborated in clinical practice. "
     "It demonstrated that removing the clinical frame restores acceptance "
     "(Stage 3), with the isolation of the zero-failure element assigned to "
     "the pre-registered Stage 3B. It hypothesises - supported by the "
     "learning literature and by behavioural precursors within the study "
     "windows - that sustained, adaptively calibrated, failure-protected "
     "practice is the pathway by which measurement-bearing technology "
     "becomes a learning tool older adults live with. And it proposes the "
     "adaptive-learning framework as the design discipline for that pathway, "
     "with AI as the enabling technology that makes calibration individual "
     "and scalable - enabling, not essential: the deployed system ran on a "
     "deterministic agent, and what AI adds is reach, language and "
     "personalisation at population scale. Habit itself, and the framework's "
     "mechanism-level claims, are assigned to the longitudinal programme of "
     "s8.8. The thesis claims what its evidence supports, specifies what "
     "would test the rest, and provides the instrument to do it.")
    log.append("Ch9 synthesis inserted")

# ---------- 7. family/personal channel correction ----------
n_ch = 0
for p in d.paragraphs:
    for old in [
      "work contacts, a charity community network, university support channels and the author's working peers in medical centres",
      "work contacts, a charity community network, university support channels and working peers in medical centres"]:
        if old in p.text:
            set_text(p, p.text.replace(old,
              "family and personal contacts, work contacts, a charity community "
              "network, university support channels and working peers in medical "
              "centres"))
            n_ch += 1
            break
log.append(f"channel lists corrected: {n_ch}")

# ---------- 8. Appendix H.6-H.8 ----------
lastH = None
for p in d.paragraphs:
    if p.text.strip().startswith("H.5"):
        lastH = p
if lastH is None:
    for p in d.paragraphs:
        if p.text.strip().startswith("H.4"):
            lastH = p
if lastH is None: fail.append("Appendix H anchor")
else:
    h6 = para_after(lastH,
     "H.6  Consultation record - Dr Gary Kui-Kai Lau. Standpoint: stroke "
     "neurologist; Assistant Dean (Education Innovations) and Clinical "
     "Associate Professor, Division of Neurology, LKS Faculty of Medicine, "
     "The University of Hong Kong; founder of the HKU STROKE Research Group "
     "(2018); developer of AI-enabled mobile applications for stroke "
     "recovery. Prior collaboration with the author on stroke-prevention "
     "education games is disclosed. Conduct: June 2026, in person, "
     "approximately one hour, Cantonese mixed with English; author's notes "
     "primary. Consent: completed 10 August 2026 (nine statements initialled "
     "GL; signature image), countersigned 11 August 2026; named attribution "
     "elected in writing ('Ok to be identified', WhatsApp, 11 Aug 2026) - "
     "the form's tick transcribes that written election, and both artifacts "
     "are archived. Substance: patients lack patience with slow-improving "
     "health technology; gamification contributes joyful participation; the "
     "long recovery journey is difficult, boring and lonely, and current "
     "technology shows no progressive positive signs; elderly patients give "
     "up easily and many abandon rehabilitation entirely - not a minority "
     "group; coaching is the best and most direct approach but cannot scale "
     "(family cost; therapist scarcity); whether games change minds and "
     "behaviour he cannot yet be sure - 'a good start to find out'.")
    h6y = para_after(h6,
     "[AUTHOR INPUT NEEDED - H.6 completion: second-consultation record "
     "(game feedback with date and build; Mirror-project paragraph subject "
     "to Dr Lau's clearance; member-check outcome).]", yellow=True)
    h7 = para_after(h6y,
     "H.7  Consultation record - Dr Philip Choy. Standpoint: certified "
     "naturopathic physician (ANMCB 06394) and Doctor of Natural Medicine; "
     "President of the Hong Kong Academy of Naturopath; over twenty-five "
     "years of clinical practice in Hong Kong with more than one thousand "
     "elderly patients, concentrated in chronic disease. He practises "
     "naturopathic and complementary medicine and is not a registered "
     "medical practitioner under the Medical Registration Ordinance; his "
     "contribution is practice-based insight into elderly health behaviour. "
     "Conduct: 4 August 2026, telephone, approximately 30 minutes, Cantonese "
     "and English; not audio-recorded; author's notes primary. Consent: "
     "signed the same day (all nine statements initialled; named attribution "
     "ticked on-screen, verified in the document's annotation layer; "
     "researcher countersigned) - the panel's only record needing no "
     "dual-dating; the WhatsApp transmittal chain is archived. Substance: "
     "the adopting demographic is educated, financially independent, 65-70 "
     "and upward, engaging on identified benefit; short-term visible benefit "
     "is critical; gamified health technology can sustain positive "
     "mentality where results are slow, with rewarding as the critical "
     "part; the failure-driven abandonment syndrome is very possible in his "
     "caseload; and - put the definitional question - the game does not "
     "cause it: it arises from concerns about one's own health performance. "
     "Token rewards do not motivate this educated segment.")
    h7y = para_after(h7,
     "[AUTHOR INPUT NEEDED - H.7: member-check outcome (send Dr Choy his "
     "summary for validation).]", yellow=True)
    h8 = para_after(h7y,
     "H.8  Consultation record - Expert P (pseudonymised at his election). "
     "Standpoint: senior physiotherapist, over twenty-five years of "
     "practice, trained locally; former chair of a Hong Kong professional "
     "physiotherapy body; operates his own practice serving clients from "
     "professional trainers to stroke-rehabilitation patients. Identity is "
     "held in the confidential consent record only. Conduct: March 2026 "
     "(the first consultation), telephone, approximately 25 minutes, "
     "Cantonese mixed with English; author's notes primary; consent signed "
     "and countersigned 4 August 2026 with anonymity elected (dual-dated); "
     "a dated follow-up on 5 August 2026 closed identified gaps. Substance: "
     "a two-week-to-two-month abandonment window in post-surgical "
     "rehabilitation; abandonment driven by reward scarcity, slow visible "
     "progress and loss of faith rather than technology difficulty as such, "
     "with difficulty acting at entry; necessary complexity of balance, "
     "fall-risk and cognitive assessment; an infantilisation warning; the "
     "observed-failure effect (higher-educated patients observed failing "
     "stop immediately); and the reward contradiction across educated and "
     "working-class elderly. In the follow-up he played the decoupled build "
     "and judged it suitable for FRA, with a stated regular-return "
     "intention; his summary was reviewed and endorsed by him in writing "
     "('Perfect', 5 Aug 2026) - the panel's first member-checked record.")
    log.append("H.6-H.8 inserted")

# ---------- 9. three new references ----------
REFS = [
 ("Baddeley, A",
  "Baddeley, A., & Wilson, B. A. (1994). When implicit learning fails: "
  "Amnesia and the problem of error elimination. Neuropsychologia, 32(1), "
  "53-68."),
 ("Plass, J",
  "Plass, J. L., Homer, B. D., & Kinzer, C. K. (2015). Foundations of "
  "game-based learning. Educational Psychologist, 50(4), 258-283."),
 ("Zohaib, M",
  "Zohaib, M. (2018). Dynamic difficulty adjustment (DDA) in computer "
  "games: A review. Advances in Human-Computer Interaction, 2018, "
  "5681652."),
]
paras = d.paragraphs
ref_i = None
for i, p in enumerate(paras):
    if p.text.strip() == "References" and p.style.name.startswith("Heading") and i > len(paras)//2:
        ref_i = i; break
if ref_i is None: fail.append("refs heading for additions")
else:
    for key, entry in REFS:
        paras = d.paragraphs
        end_i = next((j for j in range(ref_i+1, len(paras))
                      if paras[j].style.name.startswith("Heading")), len(paras))
        target = None
        for j in range(ref_i+1, end_i):
            t = paras[j].text.strip()
            if t and t > key:
                target = j; break
        prevp = paras[target-1] if target else paras[end_i-1]
        np_ = para_after(prevp, entry,
                         style=(paras[target].style if target else prevp.style))
        log.append(f"ref added: {key}")

# ---------- 10. approval-date yellow mark in 3.4 ----------
p = find_para("approved August 2026); the application described")
if p is None:
    p = find_para("approved August 2026")
if p is not None:
    r = p.add_run(" [AUTHOR INPUT NEEDED: exact approval date from the memo]")
    r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    log.append("3.4 approval-date mark")
else:
    fail.append("3.4 approval-date anchor")

cp = d.core_properties
cp.comments = ("V37 REVIEW COPY: panel integration (3.2, Ch8, H.6-H.8), "
               "assess-vs-train scoping (Ch4, 8.7), Stage 3B + learning-"
               "traditions positioning (8.8), Ch9 synthesis, family-channel "
               "correction, 3 refs, yellow author-input marks + front checklist.")
d.save(DST)
print("LOG:", *log, sep="\n  ")
print("FAIL:", fail if fail else "none")
