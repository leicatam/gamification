#!/usr/bin/env python3
"""V12 corrections per the author:
(a) Stage-3 raw data was NOT 'lost' outright — the working file was lost, the
    primary records survive and the register is being rebuilt (research stands);
(b) rebuilt register band sizes 38 / 41 / 41 (was 35 / 40 / 45), with the full
    cascade of band-level percentages, CIs and the commitment gap (35.6 -> 29.3 pp).
Outcome counts unchanged: dropouts 23, returns 15, permanent 8, endorsement 86,
intention 76; within-band counts unchanged (0/7/16, 0/5/10, 0/2/6, 25/28/23).
"""
import shutil
import docx
from docx.shared import Inches
from docx.oxml.ns import qn

BASE = "/tmp/claude-0/-home-user-gamification/3d8a50e8-9258-51ce-a7d6-194da50f9a04/scratchpad/"
SRC = BASE + "Tam_V11_EngD_Thesis_Final_Revision_Draft.docx"
DST = BASE + "Tam_V12_EngD_Thesis_Final_Revision_Draft.docx"
shutil.copy(SRC, DST)

d = docx.Document(DST)
log, failures = [], []


def clear_runs(p):
    for r in list(p.runs):
        r._element.getparent().remove(r._element)


def set_text(p, text):
    clear_runs(p)
    p.add_run(text)


def find(needle, nth=1):
    n = 0
    for p in d.paragraphs:
        if needle in p.text:
            n += 1
            if n == nth:
                return p
    return None


def sub(needle, pairs, nth=1, label=None):
    """Replace each (old, new) in the first paragraph containing `needle`."""
    p = find(needle, nth)
    lab = label or needle[:50]
    if p is None:
        failures.append("FIND-FAIL: " + lab)
        return
    t = p.text
    for old, new in pairs:
        if old not in t:
            failures.append(f"SUB-FAIL ({lab}): {old[:60]}")
        t = t.replace(old, new)
    set_text(p, t)
    log.append("edited: " + lab)


def sub_all(pairs_needle, pairs, label):
    """Apply pairs to every paragraph containing pairs_needle."""
    n = 0
    for p in d.paragraphs:
        if pairs_needle in p.text:
            t = p.text
            for old, new in pairs:
                t = t.replace(old, new)
            if t != p.text:
                set_text(p, t)
                n += 1
    log.append(f"edited {n} paragraphs: {label}")


def insert_after(p, text, style=None):
    new = docx.oxml.OxmlElement("w:p")
    p._p.addnext(new)
    np_ = docx.text.paragraph.Paragraph(new, p._parent)
    if style:
        np_.style = style
    np_.add_run(text)
    return np_


# ---------------- §6.4 band sizes + register note ----------------
p574 = find("35 aged 25-44, 40 aged 45-64 and 45 aged 65-80")
if p574 is None:
    failures.append("FIND-FAIL: band sizes sentence")
else:
    set_text(p574, p574.text.replace(
        "35 aged 25-44, 40 aged 45-64 and 45 aged 65-80",
        "38 aged 25-44, 41 aged 45-64 and 41 aged 65-80"))
    insert_after(p574,
        "A register note applies to the age-band sizes. The published interim "
        "summary reported band sizes of 35, 40 and 45; the register rebuilt "
        "from the surviving primary records gives 38, 41 and 41 (total 120, "
        "unchanged), and the thesis reports the rebuilt register throughout. "
        "Every headline outcome count is unaffected: 23 initial dropouts, 15 "
        "voluntary returns, 8 permanent non-returns, 86 coordination "
        "endorsements and 76 regular-use intentions. Band-level percentages "
        "in this chapter are computed on the rebuilt denominators.")
    log.append("edited: §6.4 band sizes + inserted register note")

# ---------------- §6.4 data-file framing ----------------
p = find("The Stage-3 participant-level records were lost to a technical error")
if p is None:
    failures.append("FIND-FAIL: §6.4 data-file paragraph")
else:
    set_text(p,
        "The Stage-3 participant-level working file was lost to a technical "
        "error after the aggregate results had been computed and reported. The "
        "underlying primary records — paper questionnaires, staff session logs "
        "and terminal exports — survive, and the participant register is being "
        "rebuilt from them; the research itself does not need to be redone. "
        "The archived data logbook (Appendix G, item G.3) preserves the "
        "published aggregate results and a reconciliation sheet whose live "
        "formulas recompute every published indicator from the rebuilt "
        "entries, so each transcribed record is verified as it is entered.")
    log.append("edited: §6.4 working-file/rebuild framing")

# ---------------- §6.5 headline paragraph ----------------
sub("Read against the two rival accounts of §6.1", [
    ("86.7% in-session retention in the 65–80 band",
     "85.4% in-session retention in the 65–80 band"),
    ("0.0%, 17.5%, 35.6%; chi-square across bands p = 0.0003",
     "0.0%, 17.1%, 39.0%; chi-square across bands p < 0.001"),
], label="§6.5 headline paragraph")

sub("Regular-use intention reached 63.3% overall and 51.1%", [
    ("63.3% overall and 51.1% in the 65-80 band",
     "63.3% overall and 56.1% in the 65-80 band"),
], label="§6.5 intention sentence")

p = find("Note on Table 6.3: coordination-challenge endorsement")
if p is None:
    failures.append("FIND-FAIL: Table 6.3 note")
else:
    set_text(p,
        "Note on Table 6.3: coordination-challenge endorsement was reported "
        "in the published summary only as a total (86 of 120, 71.7%); the "
        "per-age-band breakdown was not retained in that summary and awaits "
        "completion of the register rebuild, so it is shown as not preserved. "
        "No per-band endorsement claim is made in this thesis.")
    log.append("edited: Table 6.3 note")

# ---------------- §6.8 commitment gap ----------------
sub("a difference of 35.6 percentage points", [
    ("observed in-session retention was 86.7% while stated regular-use intention was 51.1%",
     "observed in-session retention was 85.4% while stated regular-use intention was 56.1%"),
    ("a difference of 35.6 percentage points", "a difference of 29.3 percentage points"),
], label="§6.8 commitment-gap paragraph")

# Figure 6.3 caption + LoF entry
sub_all("35.6-pp commitment gap",
        [("35.6-pp", "29.3-pp")], "Fig 6.3 caption/LoF")

# ---------------- §1.6 contribution qualifier ----------------
sub("The 35.6-percentage-point commitment gap does not identify", [
    ("The 35.6-percentage-point commitment gap",
     "The 29.3-percentage-point commitment gap"),
], label="§1.6 gap qualifier")

# ---------------- Table 6.5 note (§6.10) ----------------
sub("The confidence interval for the 35.6-percentage-point difference", [
    ("the 35.6-percentage-point difference", "the 29.3-percentage-point difference"),
], label="Table 6.5 note")

# ---------------- Ch7 opening ----------------
sub("as the 35.6-percentage-point gap between session retention", [
    ("the 35.6-percentage-point gap", "the 29.3-percentage-point gap"),
], label="Ch7 opening")

# ---------------- §8.3 empirical record ----------------
sub("and a descriptive 35.6-percentage-point difference", [
    ("a descriptive 35.6-percentage-point difference",
     "a descriptive 29.3-percentage-point difference"),
], label="§8.3 empirical record")

# ---------------- §8.9 reproducibility disclosure ----------------
p = find("One material data-loss event is disclosed")
if p is None:
    failures.append("FIND-FAIL: §8.9 disclosure")
else:
    t = p.text
    old = ("One material data-loss event is disclosed: the Stage-3 "
           "participant-level records were lost to a technical error after the "
           "aggregate results were computed and reported. The archived logbook "
           "preserves the published aggregates together with a reconciliation "
           "harness for any records that can be re-transcribed from surviving "
           "primary sources (Appendix G, item G.3); Chapter 6 rests on the "
           "published aggregate record unless and until the participant-level "
           "records are recovered.")
    new = ("One material data-management event is disclosed: the Stage-3 "
           "participant-level working file was lost to a technical error after "
           "the aggregate results were computed and reported. The primary "
           "records survive and the participant register is being rebuilt from "
           "them by verified transcription. The archived logbook preserves the "
           "published aggregates together with a reconciliation harness that "
           "checks every rebuilt entry against them (Appendix G, item G.3); "
           "Chapter 6 rests on the published aggregate record and the rebuilt "
           "register.")
    if old in t:
        set_text(p, t.replace(old, new))
        log.append("edited: §8.9 disclosure")
    else:
        failures.append("SUB-FAIL: §8.9 disclosure sentence mismatch")

# ---------------- §8.x Stage-3 summary ----------------
sub("Observed non-return was 6.7% overall and 13.3% in the 65-80 band", [
    ("13.3% in the 65-80 band", "14.6% in the 65-80 band"),
], label="§8 Stage-3 summary")

sub("observed session retention was 86.7% and stated regular-use intention was 51.1%", [
    ("observed session retention was 86.7% and stated regular-use intention was 51.1%",
     "observed session retention was 85.4% and stated regular-use intention was 56.1%"),
    ("The 35.6-percentage-point difference", "The 29.3-percentage-point difference"),
], label="§8 decoupling limits")

# ---------------- Appendix G.3 ----------------
p = find("G.3  Stage-3 data logbook")
if p is None:
    failures.append("FIND-FAIL: G.3")
else:
    set_text(p,
        "G.3  Stage-3 data logbook "
        "(Stage3_Data_Logbook_120cases_with_Reconciliation.xlsx): codebook "
        "with provenance note, participant/telemetry/questionnaire/outcome "
        "sheets for IDs 3001–3120, the published aggregate results "
        "(Published_Summary sheet, transcribed from Paper 3) and a live "
        "Reconciliation sheet. The original working file was lost to a "
        "technical error; the primary records survive, and the register is "
        "being rebuilt by transcription from those primary records only, with "
        "each entry checked by the reconciliation sheet. (Transcription in "
        "progress.)")
    log.append("edited: Appendix G.3")

# ---------------- tables ----------------
def cell_sub(t, pairs, label):
    hit = 0
    for row in t.rows:
        for c in row.cells:
            for p in c.paragraphs:
                txt = p.text
                new = txt
                for old, nw in pairs:
                    new = new.replace(old, nw)
                if new != txt:
                    clear_runs(p)
                    p.add_run(new)
                    hit += 1
    log.append(f"table {label}: {hit} cells edited")
    return hit

for t in d.tables:
    txt = "\n".join(" | ".join(c.text for c in r.cells) for r in t.rows)
    head = t.rows[0].cells[0].text.strip()
    if head == "Claim" and "H0-age" in txt:
        cell_sub(t, [("0.0% (25-44) vs 35.6% (65-80)",
                      "0.0% (25-44) vs 39.0% (65-80)")], "3.1 H0-age row")
    elif head == "Age band" and "Initial dropout" in txt and "Retention" in txt:
        cell_sub(t, [
            # order matters: do the composite strings first
            ("7 (17.5%)", "7 (17.1%)"),
            ("2 (5.0%)", "2 (4.9%)"),
            ("16 (35.6%)", "16 (39.0%)"),
            ("6 (13.3%)", "6 (14.6%)"),
            ("95.0%", "95.1%"),
            ("86.7%", "85.4%"),
        ], "6.2 stratified engagement")
        # n column: replace exact-cell values 35->38, 40->41, 45->41
        for row in t.rows[1:]:
            band = row.cells[0].text.strip()
            ncell = row.cells[1]
            v = ncell.text.strip()
            newv = {"25–44": "38", "45–64": "41", "65–80": "41"}.get(band)
            if newv and v != newv:
                clear_runs(ncell.paragraphs[0])
                ncell.paragraphs[0].add_run(newv)
        log.append("table 6.2: n column set to 38/41/41")
    elif head == "Item" and "Would use regularly" in txt:
        cell_sub(t, [("71.4%", "65.8%"), ("70.0%", "68.3%"),
                     ("51.1%", "56.1%")], "6.3 intention by band")
    elif head == "Contrast or quantity":
        cell_sub(t, [
            ("13.3%", "14.6%"), ("6.3–26.2%*", "6.9–28.4%*"),
            ("51.1%", "56.1%"), ("37.0–65.0%*", "41.0–70.1%*"),
            ("35.6 pp", "29.3 pp"),
            ("(86.7% retention and 51.1% stated intention)",
             "(85.4% retention and 56.1% stated intention)"),
        ], "6.5 statistical summary")
        # n cells 45 -> 41 on the three 65-80 Stage-3 rows
        for row in t.rows[1:]:
            first = row.cells[0].text
            if ("65-80" in first or "65–80" in first) and "Stage 3" in first:
                ncell = row.cells[3]
                if ncell.text.strip() == "45":
                    clear_runs(ncell.paragraphs[0])
                    ncell.paragraphs[0].add_run("41")
        log.append("table 6.5: 65-80 n cells set to 41")
    elif head == "Construct" and "Commitment Gap" in txt:
        cell_sub(t, [
            ("The 35.6-percentage-point difference", "The 29.3-percentage-point difference"),
            ("observed session retention (86.7%)", "observed session retention (85.4%)"),
            ("stated regular-use intention (51.1%)", "stated regular-use intention (56.1%)"),
        ], "glossary Commitment Gap")

# ---------------- figures 6.1 and 6.3 ----------------
def replace_figure(caption_start, png, width=5.8):
    prev = None
    for p in d.paragraphs:
        if p.text.strip().startswith(caption_start) and p.style.name == "FigureCaption":
            if prev is not None and prev._p.findall(".//" + qn("a:blip")):
                clear_runs(prev)
                prev.add_run().add_picture(BASE + png, width=Inches(width))
                log.append(f"figure replaced: {caption_start} <- {png}")
                return
            failures.append("FIG-FAIL no image before: " + caption_start)
            return
        prev = p
    failures.append("FIG-FAIL caption not found: " + caption_start)

replace_figure("Figure 6.1.", "figure_6_1_v12.png")
replace_figure("Figure 6.3.", "figure_6_3_v12.png")

d.save(DST)
print("saved", DST)
for l in log:
    print(" -", l)
if failures:
    print("FAILURES:")
    for f_ in failures:
        print(" !", f_)

# ---------------- verification sweep ----------------
d2 = docx.Document(DST)
full = "\n".join(p.text for p in d2.paragraphs)
for t in d2.tables:
    for r in t.rows:
        full += "\n" + " | ".join(c.text for c in r.cells)
stale = ["35.6", "86.7% in-session", "86.7% retention", "51.1", "17.5", "13.3",
         "35 aged", "40 aged", "45 aged", "6.3–26.2", "37.0–65.0",
         "p = 0.0003", "records were lost", "was lost with the participant-level"]
print("\nstale-string check:")
for s in stale:
    c = full.count(s)
    print(f"  {s!r}: {c}")
    if c:
        for line in full.split("\n"):
            if s in line:
                print("     >", line[:160])
