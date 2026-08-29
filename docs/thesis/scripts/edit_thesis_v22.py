#!/usr/bin/env python3
"""V22 — device lineage, development lineage, and Figure 5.1A, from the
author's 25 Jul 2026 clarifications and archived evidence (evidence/):
(1) §4.1A: platform lineage — InBody FRA line derives from the HUR
    SmartBalance modular balance solution (Asia market rights; game mode),
    stated as the author's account with the brochure archived.
(2) §5.3.9: development lineage — the input coupling was first validated
    on the open-source Tux Racer game (Aug 2024, control persons, videos
    archived) before the custom 2D Pygame build; no code carried over.
(3) §5.3.2: Figure 5.1A inserted between Figures 5.1 and 5.1B — frame from
    the 27 Aug 2024 demo video; caption records the control-person setting
    and the facility privacy rule that precludes participant imagery.
    LoF entry added.
"""
import shutil
import docx
from docx.oxml import OxmlElement
from docx.shared import Inches

REPO = "/home/user/gamification/docs/thesis/"
SRC = REPO + "Full_version_V21_EngD_Thesis_Submission_Draft.docx"
DST = REPO + "Full_version_V22_EngD_Thesis_Submission_Draft.docx"
shutil.copy(SRC, DST)
d = docx.Document(DST)
log, fail = [], []


def find(needle, style_prefix=None, nth=1):
    n = 0
    for p in d.paragraphs:
        if needle in p.text and (style_prefix is None
                                 or p.style.name.startswith(style_prefix)):
            n += 1
            if n == nth:
                return p
    return None


def insert_after(p, text="", style=None):
    new = OxmlElement("w:p")
    p._p.addnext(new)
    np_ = docx.text.paragraph.Paragraph(new, p._parent)
    if style:
        np_.style = style
    if text:
        np_.add_run(text)
    return np_


# ---------- (1) §4.1A platform lineage ----------
p = find("refers to the multilateral body-balance ability analysis system "
         "manufactured by InBody")
if p is None:
    fail.append("4.1A anchor")
else:
    insert_after(p,
        "The platform lineage is European. According to the author's "
        "industry documentation, the InBody FRA line derives from the HUR "
        "SmartBalance modular balance solution (HUR Oy, Finland) — balance "
        "platform, support rail and touchscreen computer — for which InBody "
        "holds the Asian market rights; the Asia-market units used in this "
        "programme are modified versions whose software adds a game mode. "
        "The vendor brochure supporting this account is held in the study "
        "archive (Appendix G, item G.8). Device naming throughout the "
        "thesis refers to the InBody-branded instruments as deployed, with "
        "the HUR platform as their design origin.")
    log.append("4.1A lineage paragraph inserted")

# ---------- (2) §5.3.9 development lineage ----------
p = find("The skiing game was developed using Anthropic Claude Code as the "
         "primary engineering assistant.")
if p is None:
    fail.append("5.3.9 anchor")
else:
    insert_after(p,
        "Before any custom code was written, the platform-to-game input "
        "coupling was validated on an off-the-shelf scaffold. In the "
        "earliest development phase the balance platform was connected to "
        "the open-source Tux Racer skiing game — a 3D title in which a "
        "penguin slides belly-down over a snow slope collecting fish — so "
        "that weight-shift steering could be demonstrated end-to-end before "
        "the study game existed. Two contemporaneous videos, recorded at "
        "the development office on 27 August 2024 and held in the study "
        "archive (Appendix G, item G.8), show one of two control persons "
        "steering this prototype by weight shift alone. The deployed "
        "Stage-2 game described in this chapter is a separate "
        "implementation: a 2D Pygame build, with the avatar aesthetic moved "
        "from the prototype's penguin to a human skier during development. "
        "The open-source scaffold contributed no code to the deployed "
        "build; its role — proving the input coupling before engineering "
        "effort was committed — is recorded here so that the development "
        "lineage is complete.")
    log.append("5.3.9 lineage paragraph inserted")

# ---------- (3) Figure 5.1A between Figures 5.1 and 5.1B ----------
cap51 = find("Figure 5.1. Skiing game on-screen display", "FigureCaption")
if cap51 is None:
    fail.append("Figure 5.1 caption")
else:
    imgp = insert_after(cap51)
    imgp.add_run().add_picture(
        REPO + "figures/fig_5_1A_candidate_platform_demo.jpg",
        width=Inches(2.9))
    capp = insert_after(imgp,
        "Figure 5.1A. The platform-game coupling demonstrated: one of two "
        "control persons steers the skiing-game prototype by weight shift "
        "on the balance platform (development office, 27 August 2024; "
        "frame from the archived video, Appendix G, item G.8). Photography "
        "of study participants was not permitted under the residential "
        "facility's privacy rules, so no participant play imagery exists; "
        "this development-phase demonstration stands in its place.",
        style=None)
    try:
        capp.style = "FigureCaption"
        log.append("Figure 5.1A image + caption inserted")
    except Exception:
        fail.append("FigureCaption style")

lof51 = find("Figure 5.1. Skiing game on-screen display", "toc")
if lof51 is None:
    fail.append("LoF Figure 5.1 line")
else:
    t = insert_after(lof51,
        "Figure 5.1A. The platform-game coupling demonstrated by a control "
        "person (development office, August 2024).\t40")
    t.style = lof51.style
    log.append("LoF entry for Figure 5.1A inserted")

cp = d.core_properties
cp.comments = "V22: HUR lineage (4.1A); Tux Racer prototype lineage (5.3.9); Figure 5.1A inserted."
d.save(DST)
print("LOG:", *log, sep="\n  ")
print("FAIL:", fail if fail else "none")
