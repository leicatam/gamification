#!/usr/bin/env python3
"""Build the consolidated Synapep v0.4 Business Plan & Valuation Word (.docx) document."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"; normal.font.size = Pt(10.5)
for lvl, sz in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)]:
    st = doc.styles[lvl]; st.font.name = "Calibri"; st.font.size = Pt(sz)
    st.font.color.rgb = NAVY; st.font.bold = True


def para(text="", size=10.5, bold=False, italic=False, color=None, align=None, space_after=6):
    p = doc.add_paragraph()
    if align: p.alignment = align
    r = p.add_run(text); r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True; p.add_run(text)
    else:
        p.add_run(text)
    return p


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        run = t.rows[0].cells[i].paragraphs[0].add_run(h); run.bold = True; run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            run = cells[i].paragraphs[0].add_run(str(val)); run.font.size = Pt(9.5)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows: row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# Title page
doc.add_paragraph().paragraph_format.space_after = Pt(60)
para("SYNAPEP", size=34, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Business Plan & Valuation", size=18, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para("AI-personalised regenerative aesthetic medicine — physician-configured peptide treatments",
     size=12, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Powered by CodeLife.AI (proprietary model) · IT-EXO® · SynExo",
     size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)
para("Joint Venture: WBI (Wellbiz International) 70%  ·  Eyesel (manufacturing) 30%",
     size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Draft v0.4  ·  Prepared 6 June 2026", size=10, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)
para("CONFIDENTIAL — for internal planning and discussion only. Illustrative projections, not an "
     "offer of securities and not investment, legal, or tax advice.",
     size=8.5, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# 1 Executive summary
doc.add_heading("1. Executive Summary", level=1)
para("Synapep Labs is a physician co-development lab for regenerative aesthetic medicine. Elite, "
     "globally competitive Korean aesthetic doctors bring their clinical experience into the lab and "
     "co-design purpose-built products — personalised to the physician's clinical case, not the "
     "individual patient — powered by CodeLife.AI (our proprietary aesthetic-medicine model), "
     "IT-EXO® (immune-tolerant exosomes) and SynExo (synthetic recombinant exosomes), delivered via "
     "an applicator device plus physician-personalised treatment kits.")
para("The company runs on two co-equal moats: (A) an AI engine + compounding data flywheel — "
     "peptides engineered to outperform generalised products, validated by proprietary AI "
     "quality/analytics; and (B) a physician co-development engine — elite Korean KOLs, a fast/cheap "
     "regulatory 'modification engine' that turns each doctor's request into a registered SKU, and a "
     "compliance-first benefit-sharing flywheel in which the inventing physician champions the "
     "product to peers. The wedge is razor-and-blades (applicator device + recurring kits, ~$30–35 "
     "ex-factory); the long-run vision is a data-science company sitting on a portfolio of "
     "doctor-originated SKUs.")
para("Why now / why this works.", bold=True, space_after=2)
bullet(" medical aesthetics ≈ $17–18.5B (2024), ~10–13% CAGR → ~$56B by 2033; injectables >40% of "
       "revenue. Korea is a global aesthetic hub; Hong Kong a gateway.", "Large, growing market:")
bullet(" a modification engine — minor changes to approved products on Korea's improved-device / "
       "negative-list regime and functional-cosmetic/quasi-drug routes (target <3% change, "
       "~3-month cycle, <$30K per SKU). ~6 months → ~20 registered SKUs across KFDA-recognising "
       "Asia. (Working assumptions — to be validated.)", "Regulatory speed:")
bullet(" the inventing doctor champions each SKU to peers; benefit-shared as an FMV royalty for "
       "genuine IP co-invention, decoupled from the doctor's own usage (K-Sunshine compliant).",
       "Doctor-driven flywheel:")
bullet(" regenerative, doctor-personalised devices/kits sit beside (not replace) HA fillers and "
       "botox.", "New category:")
para("Structure & scope.", bold=True, space_after=2)
para("JV — WBI 70% (IP: CodeLife.AI / IT-EXO / SynExo + R&D) and Eyesel 30% (manufacturing). Valued "
     "standalone at the ex-factory line; distributors are separate, non-consolidated.")
para("Economics & ask.", bold=True, space_after=2)
para("~65% blended gross margin; revenue $5.3M (2027) → $20.3M (2028) → $30.2M (2029) → $39.9M "
     "(2030); EBITDA-positive in the launch year. Raising $1.0M to kick off at $5–7M pre-money "
     "($6M midpoint → $1M ≈ 14.3%). The AI/data-platform thesis underpins a premium, "
     "institutional-AI Series A in 2028.")

# 2 Technology & AI moat
doc.add_heading("2. Technology & the AI Moat", level=1)
para("Three pillars — the AI pillar is the spine; devices and kits are the wedge and the "
     "data-capture layer.")
table(["Pillar", "What it is / why it matters"],
      [["1 — AI engine & data flywheel", "CodeLife.AI designs personalised peptides engineered to "
        "OUTPERFORM generalised products (differentiator is performance, not just personalisation). "
        "A proprietary AI quality-&-analytics suite validates every batch/variant and improves the "
        "model from each treatment — a compounding data flywheel. Built as physician decision-support "
        "within a cleared envelope to manage SaMD exposure."],
       ["2 — Devices & kits (revenue wedge)", "An applicator device placed/sold to clinics (also the "
        "data-capture endpoint) plus recurring personalised treatment kits (~$30–35 ex-factory), "
        "CodeLife-configured within an approved device family."],
       ["3 — Data monetisation (future upside)", "Licensing the model/insights, physician outcome "
        "analytics, and R&D partnerships — the institutional-AI-investor story. Not in the base P&L."]],
      widths=[1.9, 4.6])
table(["Core technology", "What it is"],
      [["CodeLife.AI", "Proprietary AI platform/model designing peptides & proteins and optimising "
        "delivery within an approved envelope; basis of the data-science-company vision."],
       ["IT-EXO®", "Immune-tolerant layer: surface HLA-G suppresses NK/T-cell attack, so efficacy "
        "compounds rather than declines."],
       ["SynExo", "Synthetic recombinant exosome (salmon-derived) carrying growth-factor/peptide "
        "cargo; ~20B particles/mL, 30–150 nm, >95% purity; 12-month stability at −80 °C."]],
      widths=[1.7, 4.8])

# 3 Vision
doc.add_heading("3. Vision — From Devices to a Data-Science Company", level=1)
para("Devices and kits are the wedge that puts CodeLife into clinics and starts the data flywheel. "
     "As personalised treatments and outcomes accumulate, Synapep builds a proprietary "
     "aesthetic-medicine dataset and its own large model — the durable, compounding asset. The "
     "endpoint is a data-science company in aesthetic medicine: AI-designed, outcome-validated, "
     "continuously improving — and fundable by institutional AI investors.")

# 4 Regulatory
doc.add_heading("4. Regulatory Strategy (Korea → Hong Kong)", level=1)
bullet(" register the microneedling/topical applicator + kit as a device family; CodeLife variants "
       "register as \"modified devices\" (same intended use / mechanism / raw materials) on the fast "
       "path (Class I ≈ 15 business days; Class II \"modified\" via conformity bodies; ~3 months "
       "realistic).", "Korea (MFDS) first —")
bullet(" light, voluntary MDACS listing; move while the voluntary window is open (a statutory "
       "framework is coming).", "Hong Kong next —")
bullet(" personalise via device parameters (delivery, depth, concentration within an approved "
       "range, a menu of pre-cleared options), NOT by changing the active molecule, which would "
       "break the \"same raw materials\" test.", "Keep the fast path valid —")
bullet(" microneedling/topical line carries the device fast-path; the IM injectable line runs on a "
       "separate drug/cosmetic pathway (see Risks).", "Two tracks —")

# 5 GTM
doc.add_heading("5. Go-to-Market", level=1)
bullet(" aesthetic/derm physicians & clinics — Korea first, then Hong Kong.", "Customers —")
bullet(" place/sell the applicator device, then earn recurring revenue on personalised kits; "
       "CodeLife + outcome data make the relationship sticky.", "Model —")
bullet(" a focused clinical/KOL salesforce in two concentrated markets; physician advocates and "
       "head-to-head efficacy data drive adoption.", "Build —")

# 6 AI & data-science build plan
doc.add_heading("6. AI & Data-Science Build Plan & Costs", level=1)
para("How Pillar 1 (the moat) gets built — and what it costs. All AI/data spend below is a subset "
     "of the OpEx \"R&D incl. AI\" line, not additional.")
table(["Phase", "Window", "Deliverables"],
      [["1 — Foundation", "2026", "Data schema + device-side capture pipeline (consented inputs, "
        "settings, images); CodeLife v1 configuration engine within the approved envelope; AI-QA "
        "tooling (batch/variant validation, anomaly detection)."],
       ["2 — Data flywheel", "2027–28", "Live outcome capture from clinics; models linking "
        "formulation/device parameters → outcomes; proprietary dataset compounds; physician "
        "dashboards (beta)."],
       ["3 — Own large model", "2029–30", "Train a proprietary aesthetic-medicine model on "
        "multimodal data (images + formulations + outcomes); ship physician outcome-analytics "
        "products (Pillar 3 upside)."]],
      widths=[1.5, 1.0, 4.0])
doc.add_heading("6.1 Team & cost (AI/data subset of OpEx, $000s)", level=2)
table(["Year", "AI/data FTE", "Personnel", "Compute & data ops", "Total AI spend", "(Total OpEx)"],
      [["2026", "2", "250", "50", "300", "(700)"],
       ["2027", "4", "600", "100", "700", "(2,600)"],
       ["2028", "7", "1,050", "250", "1,300", "(6,000)"],
       ["2029", "10", "1,600", "400", "2,000", "(8,200)"],
       ["2030", "12", "2,000", "600", "2,600", "(10,000)"]],
      widths=[0.7, 1.0, 1.0, 1.4, 1.2, 1.1])
para("Core roles: ML/AI lead (peptide design + model strategy), data/ML engineers, "
     "data-platform/MLOps, clinical-data/annotation, regulatory-aware ML (SaMD). Recruit the "
     "ML/AI lead and a data engineer first (2026); scale annotation and MLOps as clinic data arrives.")
para("Data moat: every CodeLife-configured treatment generates consented inputs, device parameters "
     "and (where permitted) outcome data via the applicator device. More treatments → better models "
     "→ peptides that outperform generalised products → more clinics → more data. The dataset, not "
     "the code, is the durable moat — and the line item that turns Synapep from a medtech story "
     "(Lens A) into an AI/data-company story (Lens B). Governance: consent, de-identification, "
     "Korea/HK data-protection compliance, physician in the loop.")

# 7 Financials
doc.add_heading("7. Financial Model", level=1)
para("Bottom-up by clinics (~100–125 kits/clinic/month at maturity). Kit unit economics: "
     "COGS = 20%×EXW + $4.50 material = $11.00 at $32.50 → GP $21.50, 66.2%. Device ASP $3,000 at "
     "40% GM. All figures USD; illustrative.")
table(["Driver", "2026 prep", "2027", "2028", "2029", "2030"],
      [["Clinics (cumulative)", "0", "150", "400", "700", "1,000"],
       ["New clinics (devices sold)", "0", "150", "250", "300", "300"],
       ["Treatment kits (000s)", "0", "148", "600", "900", "1,200"]],
      widths=[2.0, 0.9, 0.8, 0.8, 0.8, 0.8])
doc.add_heading("7.1 P&L ($ thousands)", level=2)
table(["Line", "2026", "2027", "2028", "2029", "2030"],
      [["Kit revenue ($32.50)", "0", "4,810", "19,500", "29,250", "39,000"],
       ["Device revenue ($3k ASP)", "0", "450", "750", "900", "900"],
       ["Total revenue", "0", "5,260", "20,250", "30,150", "39,900"],
       ["Kit gross profit (66.2%)", "0", "3,182", "12,900", "19,350", "25,800"],
       ["Device gross profit (40%)", "0", "180", "300", "360", "360"],
       ["Total gross profit", "0", "3,362", "13,200", "19,710", "26,160"],
       ["Blended gross margin", "—", "63.9%", "65.2%", "65.4%", "65.6%"],
       ["OpEx (sales, reg, R&D incl. AI, G&A)", "700", "2,600", "6,000", "8,200", "10,000"],
       ["EBITDA", "(700)", "762", "7,200", "11,510", "16,160"]],
      widths=[2.1, 0.8, 0.8, 0.85, 0.85, 0.85])
doc.add_heading("7.2 Scenarios (2029)", level=2)
table(["Scenario", "Clinics", "EXW", "Revenue", "Gross profit"],
      [["Conservative", "~500", "$30", "~$22M", "~$14M"],
       ["Base", "700", "$32.50", "$30.2M", "$19.7M"],
       ["Upside", "~850", "$35", "~$38M", "~$25M"]],
      widths=[1.3, 1.0, 0.9, 1.1, 1.1])
para("Pillar-3 data monetisation (model/insight licensing, outcome analytics) is excluded from the "
     "base P&L and treated as upside.", italic=True, size=9.5)

# 8 Valuation
doc.add_heading("8. Valuation — Two Lenses", level=1)
para("Lens A — device + consumables business (the floor; sets the kickoff price).", bold=True, space_after=2)
para("Recommend $5–7M pre-money / $6–8M post, $6M midpoint → $1.0M buys ~14.3%. Three methods "
     "reconcile: seed convention for an IP-owning JV ($5–7M); DCF (40% discount, 8× EBITDA exit) "
     "≈ $35M un-risk-adjusted EV → ~$5–7M risk-adjusted; forward multiple. Recurring medical "
     "consumables + regulatory moat support the upper half and a higher revenue multiple (3–6×).")
para("Lens B — AI / data-science company (the prize; attracts institutional AI capital).", bold=True, space_after=2)
para("With a proprietary model, an AI-QA/analytics suite, and a compounding data flywheel, the "
     "comparables shift from medtech to AI/data platforms — much higher multiples. This does not "
     "change the $1M kickoff price (priced on Lens A) but raises the valuation ceiling and reframes "
     "the 2028 Series A as a premium, institutional AI round.")
doc.add_heading("8.1 Cap table (post-raise, $6.0M pre-money)", level=2)
table(["Holder", "Contribution", "Founder %", "% post-raise"],
      [["WBI", "IP (CodeLife.AI/IT-EXO/SynExo) + R&D", "70.0%", "60.0%"],
       ["Eyesel", "Manufacturing (production, fill-finish, QA)", "30.0%", "25.7%"],
       ["Kickoff investor(s)", "$1.0M cash", "—", "14.3%"],
       ["Total", "", "100%", "100%"]],
      widths=[1.5, 3.0, 1.0, 1.0])
para("Recommendation: price the kickoff on Lens A ($6M midpoint); pitch Lens B as the prize and the "
     "Series A thesis. Optionally test institutional-AI appetite for a larger round.", italic=True, size=9.5)

# 9 What a technical investor will challenge
doc.add_heading("9. What a Technical Investor Will Challenge (and Our Answer)", level=1)
para("A chemical-engineer investor entering aesthetics will press these; pre-empting them:")
table(["Challenge", "Our answer / action"],
      [["Cold chain & COGS are asserted, not built (−80 °C is unshippable to clinics; COGS is "
        "top-down with no BOM).", "Resolve stability to 2–8 °C / ambient (or cost the cold chain) "
        "and publish a bottom-up BOM incl. per-batch QC-release cost; treat early-year margins as "
        "lower until scale."],
       ["Personalisation vs \"same raw materials\" is a contradiction — real personalisation "
        "changes the active and breaks the device fast-path.", "Personalisation lives in a "
        "pre-cleared configuration envelope (delivery, depth, concentration range, option menu) — "
        "defended as both regulatorily clean and clinically differentiating."],
       ["Exosomes are a regulatory/characterisation minefield (Korea ad ban; EU human-exosome ban; "
        "US FDA warnings).", "Lead with non-human (salmon/synthetic) origin; provide "
        "COA/particle/potency/purity/stability/batch data; market on function, not medical claims."],
       ["\"$25M exposure\" is weasel-phrased — it aggregates affiliated entities.", "Re-present at "
        "entity/product/channel level with reorder data; weight as market signal, not Synapep revenue."],
       ["\"Own large model\" overreaches for a 2→12 FTE team.", "Reframe as a structured, consented "
        "outcomes dataset + decision-support analytics; moat is data + clinic data-rights clauses."],
       ["Clinic ramp is aggressive (0→150 in year one; plan itself calls 150 'upside').",
        "Rebuild the base off anchor-clinic actuals (real kits/clinic/month + reorder cohorts) "
        "before scaling capital."]],
      widths=[2.6, 3.9])

# 10 International expansion
doc.add_heading("10. International Expansion (KFDA as an Export Springboard)", level=1)
para("Synapep is not a Korea-only play: a Korean MFDS registration anchors a reliance-led export "
     "strategy (full detail in export-strategy.md). Two-speed thesis — the applicator device "
     "travels widely on the Korean technical file; the peptide/exosome kit is wrapped per market "
     "(cosmetic vs device vs drug), leaning on salmon/synthetic origin and strict claim discipline.")
table(["Tier", "Markets & KFDA leverage"],
      [["Tier 1 — ASEAN + Hong Kong (fastest)", "Hong Kong MDACS recognises Korea MFDS as a "
        "reference regulator (pair Korea + Singapore for the ≥2-reference expedited route); Vietnam "
        "accepts MFDS certifications; Singapore HSA is the gateway; one ASEAN AMDD/CSDT dossier "
        "cascades to Thailand, Malaysia, Indonesia, Philippines."],
       ["Tier 2 — GCC + Greater China", "Saudi SFDA (GCC regional reference) + UAE MOHAP; Taiwan "
        "TFDA accepts Korean data; China NMPA via Hainan/cross-border first."]],
      widths=[2.2, 4.3])
para("Precedent: Korean salmon-PDRN boosters (e.g., Rejuran) exported across SEA/Middle East/Europe "
     "off a Korean approval — a near-exact analog. Sequence: Korea anchor → HK/Vietnam/Singapore "
     "(2027–28) → ASEAN + GCC (2028–29) → Taiwan/China (2029–30); each market needs a local "
     "responsible person/distributor. Cold-chain caveat: −80 °C cross-border logistics is even "
     "harder than domestic — solve stability (§9) before Wave 1. Effect: the clinic ramp becomes a "
     "multi-market TAM, strengthening Lens A and the Series A.")

# 11 Risks
doc.add_heading("11. Key Risks", level=1)
for lead, rest in [
    ("Device-vs-drug classification (#1) — ", "an IM injectable peptide is very likely a "
     "drug/biologic, not a device, in Korea and HK. Lead the device fast-path with the "
     "microneedling/topical line + applicator system; run the IM/injectable line on a separate track."),
    ("\"Same raw materials\" constraint — ", "personalise via device parameters within an approved "
     "envelope; changing the active molecule can break the fast path."),
    ("AI substantiation — ", "the \"outperforms generalised products\" and data-moat claims need "
     "head-to-head efficacy evidence and a real ML capability/dataset behind the \"own large model.\""),
    ("SaMD exposure — ", "CodeLife may be regulated as Software as a Medical Device; keep the "
     "physician in the loop within a cleared envelope."),
    ("Incumbents — ", "HA/botox (~$5–6.5B, still growing) are not displaced — position as a new "
     "adjacent category."),
    ("Capital adequacy — ", "$1M kicks off; multi-market registration + clinical salesforce + AI "
     "team likely need a 2028 Series A."),
]:
    bullet(rest, lead)

# 12 Next steps
doc.add_heading("12. What Would Make This Investment-Grade", level=1)
for lead, rest in [
    ("Confirm the regulatory split ", "— microneedling/topical (device fast-path) vs IM injectable (drug)."),
    ("Produce head-to-head efficacy data ", "(AI-personalised vs generalised) and define the AI/data "
     "team and dataset plan behind the \"own large model.\""),
    ("Document the IP-contribution scope ", "WBI assigns into the JV."),
    ("Confirm device ASP/GM, clinic ramp, kits/clinic/month; ", "name the core team and an "
     "anchor-clinic pipeline in Korea/HK."),
]:
    bullet(rest, lead)
doc.add_paragraph()
para("Sources — medical aesthetics market: MarketsandMarkets, P&S Intelligence, DataM. Korea MFDS "
     "modified-device path: Emergo by UL, ElendiLabs, MedDeviceGuide. Hong Kong MDACS: mdd.gov.hk, "
     "Asia Actual. Botox/HA: Future Market Insights, Allied Market Research, Grand View Research.",
     italic=True, size=9, color=GREY)

out = "/home/user/gamification/synexo/Synapep_Business_Plan_and_Valuation.docx"
doc.save(out)
print("Saved:", out)
