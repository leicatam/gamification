#!/usr/bin/env python3
"""GenApep operating financial model (bottom-up, founder inputs Jun 2026).

Vial-based razor-and-blades consumable model. 5-year horizon: H2-2026 prep, production from Feb 2027.
Outputs: operating-financials.csv, operating-financial-model.md, GenApep_Financial_Model.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

DIR = "/home/user/gamification/synexo/"
NAVY = RGBColor(0x1B, 0x33, 0x55); GREY = RGBColor(0x55, 0x55, 0x55); ACCENT = RGBColor(0x2E, 0x86, 0xC1)

# ---------------- founder inputs ----------------
PRICE = 32.5            # ex-factory $/vial (range 30–35; same for Korea direct & overseas distributor)
PRICE_LO, PRICE_HI = 30.0, 35.0
COST_EYESEL = 4.0       # Eyesel bottoming/fill-finish per vial
COST_PROD = 3.0         # exosome + peptide production incl. raw material per vial
COST = COST_EYESEL + COST_PROD          # $7 total variable cost/vial
GP_VIAL = PRICE - COST                   # $25.5 base
VIALS_PER_CLINIC_MO = 100               # market requirement per clinic
CAP_PER_MO = 60000                      # production capacity (vials/month) with the 4+2 team
MOQ = 300                               # vials per order
ENDUSER_LO, ENDUSER_HI = 60.0, 70.0     # clinic's end-user price

YEARS = [2026, 2027, 2028, 2029, 2030]

# 2027 monthly active-clinic ramp (production ready Feb; waves per founder)
MONTHS_2027 = ["Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]
CLINICS_2027 = [0, 10, 20, 30, 40, 50, 70, 85, 100, 105, 108, 110]   # Korea wave1 → +HK/VN/MY → wave2 (100)

# end-of-year clinic counts (2027 from ramp; 2028–30 are ASSUMPTIONS pending founder input)
CLINICS_END = {2027: 110, 2028: 220, 2029: 350, 2030: 480}

# OpEx by year ($000s). Lab team (4) + 2 shared support is FIXED to ~60k/mo capacity;
# testing/facilities shared by Eyesel. Growth is mainly registration + sales + G&A.
OPEX = {
    2026: {"Personnel (4 lab + 2 shared, partial)": 108, "Regulatory / cGMP / ISO / KR registration": 200,
           "Shared testing & facilities (Eyesel)": 40, "Sales, marketing & KOL": 20, "G&A / management": 92},
    2027: {"Personnel (4 lab + 2 shared)": 215, "Regulatory & registration (HK/VN/MY)": 150,
           "Shared testing & facilities (Eyesel)": 80, "Sales, marketing & KOL": 255, "G&A / management": 200},
    2028: {"Personnel (4 lab + 2 shared + QA)": 240, "Regulatory & registration": 120,
           "Shared testing & facilities (Eyesel)": 100, "Sales, marketing & KOL": 540, "G&A / management": 300},
    2029: {"Personnel": 260, "Regulatory & registration": 110, "Shared testing & facilities (Eyesel)": 120,
           "Sales, marketing & KOL": 710, "G&A / management": 400},
    2030: {"Personnel": 280, "Regulatory & registration": 100, "Shared testing & facilities (Eyesel)": 140,
           "Sales, marketing & KOL": 880, "G&A / management": 500},
}


def avg_active(year):
    if year == 2026:
        return 0.0
    if year == 2027:
        return sum(CLINICS_2027) / 12.0
    start = CLINICS_END[year - 1]; end = CLINICS_END[year]
    return (start + end) / 2.0


def vials(year):
    return round(avg_active(year) * 12 * VIALS_PER_CLINIC_MO)


def opex_total(year):
    return sum(OPEX[year].values())


def pnl(year, price=PRICE):
    v = vials(year); rev = v * price; cogs = v * COST; gp = rev - cogs
    ox = opex_total(year) * 1000; ebitda = gp - ox
    return dict(vials=v, rev=rev, cogs=cogs, gp=gp, gm=(gp / rev if rev else 0), opex=ox, ebitda=ebitda)


P = {y: pnl(y) for y in YEARS}


def k(x):
    return round(x / 1000.0)


def m(x):
    return x / 1e6


# ------------------------------------------------------------------ CSV
with open(DIR + "operating-financials.csv", "w") as f:
    f.write("GenApep Operating Financial Model (USD). Vial model; price $%.2f, cost $%.2f.\n\n" % (PRICE, COST))
    f.write("2027 monthly ramp,," + ",".join(MONTHS_2027) + ",TOTAL\n")
    f.write("Active clinics,," + ",".join(str(c) for c in CLINICS_2027) + ",-\n")
    vm = [c * VIALS_PER_CLINIC_MO for c in CLINICS_2027]
    f.write("Vials,," + ",".join(str(x) for x in vm) + ",%d\n\n" % sum(vm))
    f.write("Annual P&L ($000s)," + ",".join(str(y) for y in YEARS) + "\n")
    f.write("Avg active clinics," + ",".join("%.0f" % avg_active(y) for y in YEARS) + "\n")
    f.write("Vials," + ",".join(str(P[y]["vials"]) for y in YEARS) + "\n")
    f.write("Revenue," + ",".join(str(k(P[y]["rev"])) for y in YEARS) + "\n")
    f.write("COGS," + ",".join(str(k(P[y]["cogs"])) for y in YEARS) + "\n")
    f.write("Gross profit," + ",".join(str(k(P[y]["gp"])) for y in YEARS) + "\n")
    f.write("Gross margin %," + ",".join("%.1f" % (P[y]["gm"] * 100) for y in YEARS) + "\n")
    f.write("OpEx," + ",".join(str(opex_total(y)) for y in YEARS) + "\n")
    f.write("EBITDA," + ",".join(str(k(P[y]["ebitda"])) for y in YEARS) + "\n")

# ------------------------------------------------------------------ Markdown
def row(label, vals):
    return "| " + label + " | " + " | ".join(vals) + " |\n"

with open(DIR + "operating-financial-model.md", "w") as f:
    f.write("# GenApep — Operating Financial Model (bottom-up)\n\n")
    f.write("> 5-year plan from **July 2026** (prep, training, regulatory applications, cGMP/ISO). "
            "**First production readiness February 2027.** Vial-based razor-and-blades consumable "
            "model. Built from founder operating inputs; 2028–2030 clinic counts are assumptions. "
            "Illustrative; USD.\n\n")
    f.write("## Assumptions\n\n")
    f.write(row("Item", ["Value"])); f.write("|---|---|\n")
    f.write(row("Ex-factory price / vial", ["$%.0f–%.0f (base $%.2f); same Korea-direct & overseas distributor" % (PRICE_LO, PRICE_HI, PRICE)]))
    f.write(row("Cost / vial", ["$%.0f = Eyesel bottoming $%.0f + exosome/peptide+raw material $%.0f" % (COST, COST_EYESEL, COST_PROD)]))
    f.write(row("Gross profit / vial", ["$%.2f (%.1f%% margin)" % (GP_VIAL, GP_VIAL / PRICE * 100)]))
    f.write(row("Vials / clinic / month", ["%d (MOQ %d/order → ~quarterly stocking)" % (VIALS_PER_CLINIC_MO, MOQ)]))
    f.write(row("Production capacity", ["%s vials/month (= %s/yr → ~%d clinics) with 4 lab + 2 shared staff" % (f"{CAP_PER_MO:,}", f"{CAP_PER_MO*12:,}", CAP_PER_MO // VIALS_PER_CLINIC_MO)]))
    f.write(row("End-user price / vial", ["$%.0f–%.0f (clinic ~50%% margin)" % (ENDUSER_LO, ENDUSER_HI)]))
    f.write(row("Team", ["4 Exosome Lab + 2 support (shared by Eyesel); testing & facilities shared by Eyesel"]))
    f.write("\n## 2027 monthly ramp (production from Feb)\n\n")
    f.write(row("Month", MONTHS_2027 + ["FY27"]))
    f.write("|" + "---|" * 14 + "\n")
    f.write(row("Clinics", [str(c) for c in CLINICS_2027] + ["~110"]))
    f.write(row("Vials", [str(c * VIALS_PER_CLINIC_MO) for c in CLINICS_2027] + [f"{sum(vm):,}"]))
    f.write("\n## Annual P&L ($000s)\n\n")
    f.write(row("Line", [str(y) for y in YEARS])); f.write("|---|---|---|---|---|---|\n")
    f.write(row("Avg active clinics", ["%.0f" % avg_active(y) for y in YEARS]))
    f.write(row("Vials (000s)", ["%.1f" % (P[y]["vials"] / 1000.0) for y in YEARS]))
    f.write(row("**Revenue**", ["**%s**" % f"{k(P[y]['rev']):,}" for y in YEARS]))
    f.write(row("COGS ($7/vial)", [f"{k(P[y]['cogs']):,}" for y in YEARS]))
    f.write(row("Gross profit", [f"{k(P[y]['gp']):,}" for y in YEARS]))
    f.write(row("Gross margin", ["%.1f%%" % (P[y]["gm"] * 100) if P[y]["rev"] else "—" for y in YEARS]))
    f.write(row("OpEx", [f"{opex_total(y):,}" for y in YEARS]))
    f.write(row("**EBITDA**", ["**%s**" % f"{k(P[y]['ebitda']):,}" for y in YEARS]))
    f.write("\n## Unit economics\n\n")
    f.write("- **Per vial (GenApep):** price $%.2f − cost $%.0f = **$%.2f GP (%.1f%%)**.\n" % (PRICE, COST, GP_VIAL, GP_VIAL / PRICE * 100))
    f.write("- **Per clinic / year (GenApep):** %d vials × $%.2f = **$%s revenue / $%s GP**.\n" % (VIALS_PER_CLINIC_MO * 12, PRICE, f"{int(VIALS_PER_CLINIC_MO*12*PRICE):,}", f"{int(VIALS_PER_CLINIC_MO*12*GP_VIAL):,}"))
    f.write("- **Channel (clinic):** buys $%.0f–%.0f, sells end-user $%.0f–%.0f → ~50%% clinic margin (~$%s/month gross per clinic).\n" % (PRICE_LO, PRICE_HI, ENDUSER_LO, ENDUSER_HI, f"{int(VIALS_PER_CLINIC_MO*PRICE):,}"))
    f.write("\n## Price sensitivity (2030 & 5-yr cumulative)\n\n")
    f.write(row("Ex-factory price", ["$30.00", "$32.50", "$35.00"])); f.write("|---|---|---|---|\n")
    cumrev = {}; cumeb = {}
    for pr in (30.0, 32.5, 35.0):
        cumrev[pr] = sum(pnl(y, pr)["rev"] for y in YEARS)
        cumeb[pr] = sum(pnl(y, pr)["ebitda"] for y in YEARS)
    f.write(row("2030 revenue", ["$%.1fM" % m(pnl(2030, pr)["rev"]) for pr in (30, 32.5, 35)]))
    f.write(row("2030 EBITDA", ["$%.1fM" % m(pnl(2030, pr)["ebitda"]) for pr in (30, 32.5, 35)]))
    f.write(row("5-yr cum. revenue", ["$%.1fM" % m(cumrev[pr]) for pr in (30, 32.5, 35)]))
    f.write(row("5-yr cum. EBITDA", ["$%.1fM" % m(cumeb[pr]) for pr in (30, 32.5, 35)]))
    f.write("\n## Capacity, funding & notes\n\n")
    f.write("- **Capacity headroom:** 2030 at ~%.0f%% of the 60K/month line (%s vials); beyond ~600 clinics needs a second line.\n" % (P[2030]["vials"] / (CAP_PER_MO * 12) * 100, f"{P[2030]['vials']:,}"))
    f.write("- **Capital efficiency:** **EBITDA-positive from 2027** (first production year). Cumulative EBITDA turns positive during 2027, so the **$1.0M kickoff funds prep + working capital** to self-sustaining — no large survival round required.\n")
    f.write("- **Operating leverage:** the 4-lab + 2-shared team is fixed to 60K/month, so margin expands as clinics grow (lab headcount flat; growth is sales/registration).\n")
    f.write("- **Working capital:** MOQ 300/order + distributor terms imply inventory + receivables; size a WC buffer in the raise.\n")
    f.write("- **Scope:** vial (consumable) revenue only — applicator device placement is excluded (placed/loaned or priced separately; confirm).\n")
    f.write("- **Assumptions to confirm:** 2028–2030 clinic counts (220/350/480), base price, OpEx (Korea-loaded salaries), and shared-services charge from Eyesel.\n")

# ------------------------------------------------------------------ Word
def nd():
    doc = Document()
    n = doc.styles["Normal"]; n.font.name = "Calibri"; n.font.size = Pt(10.5)
    for lvl, sz in [("Heading 1", 16), ("Heading 2", 13)]:
        st = doc.styles[lvl]; st.font.name = "Calibri"; st.font.size = Pt(sz)
        st.font.color.rgb = NAVY; st.font.bold = True
    return doc


def para(doc, t="", size=10.5, bold=False, italic=False, color=None, align=None, after=6):
    p = doc.add_paragraph()
    if align: p.alignment = align
    r = p.add_run(t); r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after); return p


def bullet(doc, t, lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if lead:
        r = p.add_run(lead); r.bold = True
    p.add_run(t); p.paragraph_format.space_after = Pt(3); return p


def tbl(doc, headers, rows, widths=None, fs=9.5, hl_rows=()):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        r = t.rows[0].cells[i].paragraphs[0].add_run(str(h)); r.bold = True; r.font.size = Pt(fs)
    for ri, row in enumerate(rows):
        c = t.add_row().cells
        for i, v in enumerate(row):
            rr = c[i].paragraphs[0].add_run(str(v)); rr.font.size = Pt(fs)
            if ri in hl_rows or i == 0: rr.bold = True
    if widths:
        for i, w in enumerate(widths):
            for r in t.rows: r.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2); return t


doc = nd()
para(doc, "GENAPEP", size=26, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
para(doc, "Operating Financial Model (5-year, bottom-up)", size=14, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
para(doc, "From July 2026 (prep) · first production Feb 2027 · vial consumable model. Illustrative — not investment advice.",
     size=9, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)

doc.add_heading("1. Assumptions", level=2)
tbl(doc, ["Item", "Value"],
    [["Ex-factory price / vial", "$%.0f–%.0f (base $%.2f) — same Korea-direct & overseas distributor" % (PRICE_LO, PRICE_HI, PRICE)],
     ["Cost / vial", "$%.0f = Eyesel bottoming $%.0f + exosome/peptide incl. raw material $%.0f" % (COST, COST_EYESEL, COST_PROD)],
     ["Gross profit / vial", "$%.2f (%.1f%% margin)" % (GP_VIAL, GP_VIAL / PRICE * 100)],
     ["Vials / clinic / month", "%d (MOQ %d per order)" % (VIALS_PER_CLINIC_MO, MOQ)],
     ["Production capacity", "%s vials/month (≈ %d clinics) — 4 lab + 2 shared staff" % (f"{CAP_PER_MO:,}", CAP_PER_MO // VIALS_PER_CLINIC_MO)],
     ["End-user price / vial", "$%.0f–%.0f (clinic ~50%% margin)" % (ENDUSER_LO, ENDUSER_HI)],
     ["Team & facilities", "4 Exosome Lab + 2 support (shared by Eyesel); testing & facilities shared by Eyesel"]],
    widths=[2.3, 4.2])

doc.add_heading("2. 2027 monthly ramp (production from February)", level=2)
tbl(doc, ["Month"] + MONTHS_2027 + ["FY27"],
    [["Clinics"] + [str(c) for c in CLINICS_2027] + ["~110"],
     ["Vials"] + [str(c * VIALS_PER_CLINIC_MO) for c in CLINICS_2027] + [f"{sum(vm):,}"]],
    widths=[0.8] + [0.42] * 12 + [0.7], fs=8)

doc.add_heading("3. Annual P&L ($000s)", level=2)
tbl(doc, ["Line"] + [str(y) for y in YEARS],
    [["Avg active clinics"] + ["%.0f" % avg_active(y) for y in YEARS],
     ["Vials (000s)"] + ["%.1f" % (P[y]["vials"] / 1000.0) for y in YEARS],
     ["Revenue"] + [f"{k(P[y]['rev']):,}" for y in YEARS],
     ["COGS ($7/vial)"] + [f"{k(P[y]['cogs']):,}" for y in YEARS],
     ["Gross profit"] + [f"{k(P[y]['gp']):,}" for y in YEARS],
     ["Gross margin"] + [("%.1f%%" % (P[y]["gm"] * 100)) if P[y]["rev"] else "—" for y in YEARS],
     ["OpEx"] + [f"{opex_total(y):,}" for y in YEARS],
     ["EBITDA"] + [f"{k(P[y]['ebitda']):,}" for y in YEARS]],
    widths=[1.9, 1.0, 1.0, 1.0, 1.0, 1.1])

doc.add_heading("4. OpEx detail ($000s)", level=2)
lines = list(OPEX[2030].keys())
# normalise line labels across years by index
keylist = [list(OPEX[y].keys()) for y in YEARS]
rows = []
for idx in range(5):
    label = ["Personnel (4 lab + 2 shared)", "Regulatory / cGMP / ISO / registration",
             "Shared testing & facilities (Eyesel)", "Sales, marketing & KOL", "G&A / management"][idx]
    rows.append([label] + [str(list(OPEX[y].values())[idx]) for y in YEARS])
rows.append(["Total OpEx"] + [str(opex_total(y)) for y in YEARS])
tbl(doc, ["OpEx line"] + [str(y) for y in YEARS], rows, widths=[2.6, 0.85, 0.85, 0.85, 0.85, 0.85], hl_rows=(5,))

doc.add_heading("5. Unit economics", level=2)
bullet(doc, "price $%.2f − cost $%.0f = $%.2f GP (%.1f%%)." % (PRICE, COST, GP_VIAL, GP_VIAL / PRICE * 100), "Per vial — ")
bullet(doc, "%d vials × $%.2f = $%s revenue / $%s GP per year." % (VIALS_PER_CLINIC_MO * 12, PRICE, f"{int(VIALS_PER_CLINIC_MO*12*PRICE):,}", f"{int(VIALS_PER_CLINIC_MO*12*GP_VIAL):,}"), "Per clinic — ")
bullet(doc, "buys $%.0f–%.0f, sells end-user $%.0f–%.0f → ~50%% margin (~$%s/month gross per clinic)." % (PRICE_LO, PRICE_HI, ENDUSER_LO, ENDUSER_HI, f"{int(VIALS_PER_CLINIC_MO*PRICE):,}"), "Channel — ")

doc.add_heading("6. Price sensitivity", level=2)
tbl(doc, ["Metric", "$30.00", "$32.50", "$35.00"],
    [["2030 revenue"] + ["$%.1fM" % m(pnl(2030, pr)["rev"]) for pr in (30, 32.5, 35)],
     ["2030 EBITDA"] + ["$%.1fM" % m(pnl(2030, pr)["ebitda"]) for pr in (30, 32.5, 35)],
     ["5-yr cum. revenue"] + ["$%.1fM" % m(sum(pnl(y, pr)["rev"] for y in YEARS)) for pr in (30, 32.5, 35)],
     ["5-yr cum. EBITDA"] + ["$%.1fM" % m(sum(pnl(y, pr)["ebitda"] for y in YEARS)) for pr in (30, 32.5, 35)]],
    widths=[2.2, 1.4, 1.4, 1.4])

doc.add_heading("7. Capacity, funding & notes", level=2)
bullet(doc, "2030 runs at ~%.0f%% of the 60K/month line; beyond ~600 clinics needs a second line." % (P[2030]["vials"] / (CAP_PER_MO * 12) * 100), "Capacity — ")
bullet(doc, "EBITDA-positive from 2027; cumulative cash turns positive during 2027, so the $1.0M kickoff funds prep + working capital to self-sustaining — no large survival round required.", "Capital efficiency — ")
bullet(doc, "the 4-lab + 2-shared team is fixed to 60K/month; margin expands as clinics grow (production headcount flat).", "Operating leverage — ")
bullet(doc, "MOQ 300/order + distributor terms imply inventory + receivables; size a working-capital buffer.", "Working capital — ")
bullet(doc, "vial (consumable) revenue only — applicator device placement excluded (confirm).", "Scope — ")
bullet(doc, "2028–2030 clinic counts (220/350/480), base price, and Korea-loaded OpEx are assumptions to confirm.", "To confirm — ")
para(doc, "Founder inputs: Jul-2026 start (prep/training/regulatory/cGMP/ISO); first production Feb-2027; "
          "4 lab + 2 shared staff; Eyesel bottoming $4 + production $3 = $7/vial; 60K/month capacity; "
          "100 vials/clinic/month; MOQ 300; price $30–35; end-user $60–70; first wave 20–40 Korea + 20 "
          "(HK/VN/MY); second wave Q3-2027 → 100 clinics.", size=9, italic=True, color=GREY)
doc.save(DIR + "GenApep_Financial_Model.docx")

for y in YEARS:
    print("%d: clinics %.0f, vials %s, rev $%.2fM, GP $%.2fM (%.1f%%), OpEx $%.2fM, EBITDA $%.2fM"
          % (y, avg_active(y), f"{P[y]['vials']:,}", m(P[y]["rev"]), m(P[y]["gp"]),
             P[y]["gm"] * 100 if P[y]["rev"] else 0, opex_total(y) / 1000.0, m(P[y]["ebitda"])))
print("Saved operating-financials.csv, operating-financial-model.md, GenApep_Financial_Model.docx")
