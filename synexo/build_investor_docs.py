#!/usr/bin/env python3
"""Build standalone investor Word docs: DD scorecard, milestone-gated SAFE term sheet, export strategy."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
GREY = RGBColor(0x55, 0x55, 0x55)
DIR = "/home/user/gamification/synexo/"


def new_doc():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"; normal.font.size = Pt(10.5)
    for lvl, sz in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)]:
        st = doc.styles[lvl]; st.font.name = "Calibri"; st.font.size = Pt(sz)
        st.font.color.rgb = NAVY; st.font.bold = True
    return doc


def para(doc, text="", size=10.5, bold=False, italic=False, color=None, align=None, space_after=6):
    p = doc.add_paragraph()
    if align: p.alignment = align
    r = p.add_run(text); r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True; p.add_run(text)
    else:
        p.add_run(text)
    return p


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        run = t.rows[0].cells[i].paragraphs[0].add_run(h); run.bold = True; run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            run = cells[i].paragraphs[0].add_run(str(val)); run.font.size = Pt(9.5)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows: row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def title(doc, line1, line2):
    para(doc, "GENAPEP", size=26, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, line1, size=15, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, line2, size=9, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)


# ---------------------------------------------------------------- DD scorecard
def build_scorecard():
    doc = new_doc()
    title(doc, "Investor DD Scorecard (technical-investor lens)",
          "Chemical-engineer / aesthetic-medicine investor read. Illustrative diligence aid, not investment advice.")
    para(doc, "Legend: GREEN = credible/de-risked · AMBER = plausible but unproven, needs evidence · "
              "RED = material gap to resolve before scaling/closing.", italic=True, size=9.5)
    table(doc, ["#", "Category", "RAG", "Rationale", "Gating DD question"],
          [["1", "Commercial validation", "AMBER", "\"$25M+ exposure\" unverified & aggregates "
            "affiliates; testimony = usability not efficacy.", "Show the $25M by entity/product/channel "
            "+ reorder data — how much is THIS product's demand?"],
           ["2", "Chemistry / formulation", "AMBER", "SynExo/IT-EXO treated as settled; the "
            "personalisation-vs-'same raw materials' envelope is asserted.", "Define the pre-cleared "
            "envelope — is variation clean AND clinically differentiating?"],
           ["3", "Manufacturing & supply chain", "AMBER", "Stabilisation tech removes the −80 °C cold-chain "
            "dependency (a distribution/export advantage); COGS still top-down, no BOM.", "Show the stability/shelf-life dossier "
            "+ BOM + per-batch QC release cost."],
           ["4", "Regulatory", "AMBER", "Strong topical-vs-IM segmentation (green-leaning); exosome "
            "scrutiny + SaMD pull to amber.", "Classification memo: does the device path survive the "
            "envelope; how are exosome claims controlled?"],
           ["5", "AI / data moat", "AMBER", "\"Own large model\" overreaches for 2→12 FTE; real asset "
            "is consented outcomes data, gated on clinic rights.", "Are data-rights in the clinic "
            "contract, and what does CodeLife v1 do today (demo)?"],
           ["6", "Financials", "AMBER", "Ramp 0→150 yr1 aggressive (plan calls 150 'upside'); early "
            "margins likely overstated.", "Rebuild base off anchor-clinic actuals: kits/clinic/month "
            "+ reorder cohorts."],
           ["7", "Valuation / deal", "AMBER", "$7M ask leans on validation 'still to be packaged'; "
            "floor supports ~$5–6M today.", "What price holds before the validation pack exists; what "
            "re-rates it at Series A?"],
           ["8", "IP & JV", "AMBER", "Most patents are applications; microneedling device granted "
            "(good); JV IP contribution undocumented.", "Show the WBI IP assignment/exclusive-license "
            "schedule into the JV."],
           ["9", "Team / capital adequacy", "AMBER", "$1M must cover reg + study + mfg/QC + AI v1 + "
            "launch — thin; high Series A dependency.", "Which milestones can $1M actually hit before a "
            "financing gap?"],
           ["10", "Market reach & export", "AMBER", "KFDA is a real export springboard, but the kit's "
            "class varies per market and exosome claims are restricted.", "Per market: cosmetic, device "
            "or drug — does salmon/synthetic origin clear EU/cosmetic bans?"]],
          widths=[0.3, 1.4, 0.6, 2.3, 2.3])
    doc.add_heading("Overall verdict", level=2)
    para(doc, "Conditional interest — diligence yes, their price no. The narrative is fundable and "
              "unusually disciplined (claim control, regulatory segmentation, validation-pack thinking). "
              "But incremental value rests on unverified commercial claims, and the parts a chemical "
              "engineer judges best — COGS BOM, formulation/personalisation, exosome status, royalty compliance "
              "— are least evidenced.")
    bullet(doc, " a validated-ecosystem, regulatory-segmented consumables business with a credible "
                "data-upside option.", "What I'd underwrite today —")
    bullet(doc, " $5–6M pre, milestone-gated (not $7M, not the AI/data-company multiple yet).", "Price —")
    bullet(doc, " produce the validation pack (items 1, 2, 8) plus the stability dossier + COGS BOM "
                "(item 3). Do that and amber→green across the board.", "Single biggest unlock —")
    doc.save(DIR + "GenApep_DD_Scorecard.docx")
    print("Saved GenApep_DD_Scorecard.docx")


# ---------------------------------------------------------------- Term sheet
def build_termsheet():
    doc = new_doc()
    title(doc, "Indicative Term Sheet — Milestone-Gated SAFE",
          "NON-BINDING; for discussion only. Not an offer of securities or legal/tax/investment advice.")
    para(doc, "Except for No-Shop and Confidentiality (binding when countersigned), no obligation "
              "arises until definitive documents are signed. Figures illustrative, consistent with the "
              "v0.4 financials.", italic=True, size=9.5)
    doc.add_heading("1. Headline terms", level=2)
    table(doc, ["Term", "Proposed"],
          [["Company", "GenApep (JV — WBI 70% / Eyesel 30%; IP-owning entity)"],
           ["Investor", "[Lead investor / SPV]"],
           ["Instrument", "Post-money SAFE (YC post-money form, modified for milestone tranches)"],
           ["Total amount", "US$1,000,000, funded in two tranches"],
           ["Valuation cap", "US$6,000,000 post-money (range $5–6M; $6M if validation pack complete at close)"],
           ["Discount", "20% to the next priced round"],
           ["MFN", "Yes — best terms granted to any other SAFE in this round"],
           ["Conversion", "Converts to preferred at the next Qualified Financing (priced round ≥ US$3,000,000), better of cap or discount"],
           ["Pro-rata rights", "Yes — maintain ownership % in the Series A"],
           ["Information rights", "Quarterly management accounts, annual budget, cap table on request"],
           ["Governance", "Board observer while the SAFE is outstanding (no board control)"]],
          widths=[1.7, 4.8])
    doc.add_heading("2. Tranche structure", level=2)
    table(doc, ["Tranche", "Amount", "Release trigger"],
          [["Tranche 1 — Close", "US$400,000", "On signing + satisfaction of Conditions Precedent (§3)"],
           ["Tranche 2 — Milestone", "US$600,000", "On Investor's written confirmation Milestones (§4) met, within 12 months of close"]],
          widths=[1.6, 1.2, 3.7])
    para(doc, "If Milestones are not met within 12 months, Tranche 2 is at the Investor's discretion "
              "or renegotiated. Both tranches share the same cap/discount.", italic=True, size=9.5)
    doc.add_heading("3. Conditions precedent (before Tranche 1)", level=2)
    bullet(doc, " executed assignment or exclusive license of WBI IP (CodeLife.AI, IT-EXO®, SynExo, "
                "know-how, patents/trademarks) into GenApep, with a contribution schedule.", "IP into the JV —")
    bullet(doc, " independent verification of a representative sample of the \"US$25M+\" sales at "
                "entity/product/channel level.", "Commercial verification —")
    bullet(doc, " a clinic-contract template embedding consent, de-identification and GenApep "
                "data-ownership/usage rights.", "Data rights —")
    bullet(doc, " clean cap table, JV agreement, good standing, customary reps & warranties.", "Corporate —")
    doc.add_heading("4. Milestones (release Tranche 2)", level=2)
    bullet(doc, " the stability/shelf-life dossier confirming the cold-chain-free (ambient/2–8 °C, lyophilised) "
                "clinic format and batch reproducibility, plus a bottom-up COGS BOM with per-batch "
                "QC-release costs.", "Stability dossier + COGS —")
    bullet(doc, " 5–10 anchor clinics live with documented kits/clinic/month and reorder evidence.", "Anchor clinics —")
    bullet(doc, " one prospective GenApep protocol readout on the chosen beachhead indication.", "Clinical readout —")
    bullet(doc, " a classification memo (Korea + first export market) confirming the topical/microneedling "
                "device path holds for the personalisation envelope; IM/injectable on a separate pathway.", "Regulatory —")
    doc.add_heading("5. Use of funds (tied to milestones)", level=2)
    table(doc, ["Bucket", "Purpose"],
          [["Regulatory & legal", "Korea/HK + first-export classification memo; JV/IP documentation; claims review"],
           ["Clinical / KOL", "Anchor-clinic protocol, prospective beachhead study, testimony dossier, safety monitoring"],
           ["Manufacturing / chemistry", "Stabilisation/stability validation (cold-chain-free), COGS BOM, QA/release validation, initial inventory"],
           ["AI / data", "CodeLife.AI v1, consent workflow, device-side capture, QA analytics"],
           ["Launch", "KOL education, clinic onboarding, first-export distributor readiness"]],
          widths=[1.9, 4.6])
    doc.add_heading("6. Market / kit-classification risk (export note)", level=2)
    para(doc, "KFDA registration anchors export via reliance (Hong Kong MDACS reference, Vietnam, ASEAN "
              "cascade, GCC/SFDA). However the treatment-kit's classification varies by market (cosmetic "
              "vs device vs drug) and exosome claims are restricted (Korea ad ban; EU human-exosome ban; "
              "US FDA warnings). Investor relies on Company maintaining claim discipline and using "
              "salmon/synthetic (non-human) origin where it eases classification; a per-market regulatory "
              "map is a Tranche-2 expectation.")
    doc.add_heading("7. Other terms", level=2)
    bullet(doc, " 45 days exclusivity from countersignature (binding).", "No-Shop —")
    bullet(doc, " binding; mutual.", "Confidentiality —")
    bullet(doc, " each party bears its own (or capped reimbursement at Investor's option).", "Expenses —")
    bullet(doc, " lapses if not accepted within [14] days.", "Expiry —")
    para(doc, "Prepared as a planning artifact reflecting the independent technical-investor evaluation. "
              "Final terms subject to definitive documentation and the Company's counsel review.",
         italic=True, size=9, color=GREY)
    doc.save(DIR + "GenApep_Term_Sheet.docx")
    print("Saved GenApep_Term_Sheet.docx")


# ---------------------------------------------------------------- Export strategy
def build_export():
    doc = new_doc()
    title(doc, "International Expansion / Export Strategy",
          "KFDA/MFDS registration as an export springboard. Illustrative; verify each pathway with local counsel.")
    doc.add_heading("1. Thesis — two-speed export", level=2)
    bullet(doc, " built on the Korean technical file, the microneedling/topical device rides "
                "reliance/recognition into reference-friendly markets.", "The device travels widely —")
    bullet(doc, " the peptide/exosome kit's legal class (cosmetic vs device vs drug) varies by country; "
                "salmon/synthetic (non-human) origin sidesteps EU/UK human-exosome cosmetic bans.", "The kit is wrapped per market —")
    bullet(doc, " Korea banned the word \"exosome\" in cosmetic ads (Jan 2025); US FDA = 12 warning "
                "letters, none approved. Market on function/observed benefit.", "Claims are the recurring risk —")
    para(doc, "Precedent: Korean salmon-PDRN skin boosters (e.g., Rejuran, MFDS-approved 2014) exported "
              "across SEA/Middle East/Europe off a Korean approval — a near-exact analog.")
    doc.add_heading("2. Tier 1 — ASEAN + Hong Kong (reliance-led, fastest)", level=2)
    table(doc, ["Market", "How KFDA helps", "Kit approach", "Notes"],
          [["Hong Kong", "MDACS recognises Korea MFDS as reference; Expedited favours ≥2 references",
            "MDACS device listing / cosmetic", "Pair Korea + Singapore for the expedited route"],
           ["Vietnam", "Explicitly accepts MFDS certifications", "Device via MFDS reliance", "Among the most direct routes"],
           ["Singapore", "HSA approval is a regional gateway (onward-recognised TH/PH/HK)", "HSA device / cosmetic notification", "Strategic 2nd reference approval"],
           ["Thailand", "Reliance via Singapore HSA; Korean data supports", "Aesthetic devices ARE medical devices", "Large market; K-beauty halo"],
           ["Malaysia", "ASEAN AMDD/CSDT dossier on Korean file", "Aesthetic devices designated medical devices (Jun 2026)", "MDA registration"],
           ["Indonesia / Philippines", "ASEAN AMDD/CSDT cascade", "Device or cosmetic per claim", "One CSDT dossier → multiple filings"]],
          widths=[1.3, 2.0, 1.7, 1.5])
    doc.add_heading("3. Tier 2 — GCC + Greater China", level=2)
    table(doc, ["Market", "How KFDA helps", "Kit approach", "Notes"],
          [["Saudi Arabia (SFDA)", "GCC regional reference; KFDA+CE accelerate the TFA route", "Medical device (MDMA)", "SFDA approval references onward across GCC"],
           ["UAE (MOHAP)", "Reliance weights recognised approvals (CE strongest)", "Medical device", "High-spend market; Gulf gateway"],
           ["Taiwan (TFDA)", "Accepts Korean clinical data", "Medical device", "Mature aesthetics market"],
           ["China (NMPA)", "Large but slow; use faster on-ramps", "Hainan cross-border / topical cosmetic first", "Biggest prize, longest timeline"]],
          widths=[1.4, 2.0, 1.8, 1.3])
    doc.add_heading("4. Sequence & operating requirements", level=2)
    table(doc, ["Phase", "Markets", "Trigger"],
          [["Anchor", "Korea (MFDS)", "Device-family + first kit registered; anchor clinics live"],
           ["Wave 1 (2027–28)", "Hong Kong, Vietnam, Singapore", "Korean approval + Singapore reference secured"],
           ["Wave 2 (2028–29)", "Thailand, Malaysia, Indonesia, Philippines; Saudi, UAE", "ASEAN CSDT dossier; SFDA TFA"],
           ["Wave 3 (2029–30)", "Taiwan, China (Hainan → NMPA)", "Regional revenue base + Series A capital"]],
          widths=[1.4, 2.8, 2.3])
    para(doc, "Each market needs a local responsible person / authorised distributor, a per-market claims "
              "map, and label/registration localisation. Cold-chain-free advantage: stabilisation gives an "
              "ambient/2–8 °C format, so cross-border distribution is far simpler than for typical "
              "before Wave 1.")
    doc.add_heading("5. Why this strengthens the investment case", level=2)
    bullet(doc, " converts a single-country plan into a multi-market clinic TAM (the 150→1,000 ramp is "
                "spread across geographies).")
    bullet(doc, " strengthens Lens A market size and the Series A narrative (regional platform, not a Korea SKU).")
    bullet(doc, " keeps the company honest: device scales on reliance; kit scales on disciplined per-market classification.")
    doc.add_heading("6. Key risks", level=2)
    bullet(doc, " a market may treat the kit as a drug/biologic (slower/heavier).", "Kit classification drift —")
    bullet(doc, " manage with non-human origin + claim discipline.", "Exosome scrutiny / ad limits —")
    bullet(doc, " recognition speeds review but local registration, fees, RP and post-market duties still apply.", "Reliance ≠ automatic —")
    bullet(doc, " provide the shelf-life dossier confirming the cold-chain-free (ambient/2–8 °C) format.", "Stability substantiation —")
    para(doc, "Sources: HK MDACS reference recognition (Asia Actual, MedDeviceGuide); Vietnam MFDS "
              "reliance & ASEAN AMDD/CSDT (Cisema, Pacific Bridge); Thailand/Singapore reliance (Asia "
              "Actual); SFDA/UAE reliance (Emergo by UL, GSL); exosome regulation (FDA; EU Annex II; "
              "Korea MFDS); PDRN/Rejuran precedent (industry press).", italic=True, size=9, color=GREY)
    doc.save(DIR + "GenApep_Export_Strategy.docx")
    print("Saved GenApep_Export_Strategy.docx")


if __name__ == "__main__":
    build_scorecard()
    build_termsheet()
    build_export()
