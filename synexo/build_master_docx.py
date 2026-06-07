#!/usr/bin/env python3
"""Merge all Synapep source documents into ONE Word file -> Synapep_Master_Dossier.docx.

Renders the markdown sources (single source of truth) so the dossier never drifts from the
standalone docs. Run build_multimarket_model.py / build_docx.py first if sources changed.
"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

DIR = "/home/user/gamification/synexo/"
NAVY = RGBColor(0x1F, 0x3A, 0x5F); GREY = RGBColor(0x55, 0x55, 0x55); ACCENT = RGBColor(0x2E, 0x86, 0xC1)
BODY = RGBColor(0x22, 0x22, 0x22); MONO = RGBColor(0x80, 0x3A, 0x00)

# Ordered parts: (file, Part title shown on divider)
PARTS = [
    ("business-plan.md",      "Business Plan"),
    ("financial-model.md",    "Financial Model"),
    ("multimarket-model.md",  "Multi-Market Revenue Model"),
    ("valuation.md",          "Valuation"),
    ("export-strategy.md",    "International Expansion / Export Strategy"),
    ("dd-scorecard.md",       "Investor Diligence Scorecard"),
    ("term-sheet.md",         "Indicative Term Sheet"),
]

INLINE = re.compile(r'(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))')


def init():
    d = Document()
    n = d.styles["Normal"]; n.font.name = "Calibri"; n.font.size = Pt(10.5); n.font.color.rgb = BODY
    for lvl, sz in [("Heading 1", 17), ("Heading 2", 13.5), ("Heading 3", 11.5)]:
        st = d.styles[lvl]; st.font.name = "Calibri"; st.font.size = Pt(sz)
        st.font.color.rgb = NAVY; st.font.bold = True
    return d


def add_runs(p, text, size=10.5):
    """Render inline **bold**, `code`, [text](url) into a paragraph."""
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = p.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = p.add_run(tok[1:-1]); r.font.name = "Consolas"; r.font.color.rgb = MONO
        elif tok.startswith("[") and "](" in tok:
            label = tok[1:tok.index("](")]; r = p.add_run(label); r.italic = True
        else:
            r = p.add_run(tok)
        r.font.size = Pt(size)


def is_sep_row(cells):
    return all(re.fullmatch(r':?-{2,}:?', c.strip()) for c in cells if c.strip() != "") and any(cells)


def split_row(line):
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def render_table(doc, block):
    rows = [split_row(l) for l in block]
    rows = [r for r in rows if not is_sep_row(r)]
    if not rows:
        return
    ncol = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=ncol); t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.autofit = True
    for ri, r in enumerate(rows):
        cells = t.add_row().cells
        for ci in range(ncol):
            val = r[ci] if ci < len(r) else ""
            para = cells[ci].paragraphs[0]
            add_runs(para, val, size=9.5)
            if ri == 0:
                for run in para.runs:
                    run.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def is_quote(line):
    s = line.strip()
    return s.startswith("> ") or s == ">"


def is_block_start(line):
    """True if the line begins a new structural block (so paragraph/list coalescing stops)."""
    s = line.strip()
    if s == "":
        return True
    if re.match(r'#{1,6}\s', s):
        return True
    if s.startswith("|"):
        return True
    if is_quote(line):
        return True
    if re.fullmatch(r'-{3,}', s):
        return True
    if re.match(r'^\s*[-*]\s+', line):
        return True
    if re.match(r'^\s*\d+\.\s+', line):
        return True
    return False


def render_markdown(doc, text):
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # table block
        if stripped.startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i]); i += 1
            render_table(doc, block); continue

        # blockquote (collapse consecutive "> " lines into one paragraph)
        if is_quote(line):
            buf = []
            while i < len(lines) and is_quote(lines[i]):
                buf.append(lines[i].strip()[1:].strip()); i += 1
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(8)
            add_runs(p, " ".join(buf), size=10)
            for r in p.runs:
                r.italic = True; r.font.color.rgb = GREY
            continue

        # headings
        m = re.match(r'(#{1,6})\s+(.*)', stripped)
        if m:
            level = min(len(m.group(1)), 3)
            doc.add_heading("", level=level).add_run(m.group(2))
            i += 1; continue

        # horizontal rule -> skip (part dividers handled separately)
        if re.fullmatch(r'-{3,}', stripped):
            i += 1; continue

        # bullets (one nesting level via indentation) + wrapped continuations
        mb = re.match(r'^(\s*)[-*]\s+(.*)', line)
        if mb:
            indent = len(mb.group(1)); buf = [mb.group(2)]; i += 1
            while i < len(lines) and not is_block_start(lines[i]):
                buf.append(lines[i].strip()); i += 1
            style = "List Bullet 2" if indent >= 2 else "List Bullet"
            p = doc.add_paragraph(style=style); add_runs(p, " ".join(buf))
            p.paragraph_format.space_after = Pt(3); continue

        # numbered list + wrapped continuations
        mn = re.match(r'^(\s*)\d+\.\s+(.*)', line)
        if mn:
            buf = [mn.group(2)]; i += 1
            while i < len(lines) and not is_block_start(lines[i]):
                buf.append(lines[i].strip()); i += 1
            p = doc.add_paragraph(style="List Number"); add_runs(p, " ".join(buf))
            p.paragraph_format.space_after = Pt(3); continue

        # blank
        if stripped == "":
            i += 1; continue

        # paragraph: join wrapped continuation lines
        buf = [stripped]; i += 1
        while i < len(lines) and not is_block_start(lines[i]):
            buf.append(lines[i].strip()); i += 1
        p = doc.add_paragraph(); add_runs(p, " ".join(buf))
        p.paragraph_format.space_after = Pt(6)


def divider(doc, n, title):
    doc.add_page_break()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PART %d" % n); r.font.size = Pt(12); r.bold = True; r.font.color.rgb = ACCENT
    p.paragraph_format.space_after = Pt(2)
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(title); r2.font.size = Pt(22); r2.bold = True; r2.font.color.rgb = NAVY
    p2.paragraph_format.space_after = Pt(10)


def main():
    d = init()
    # Cover
    d.add_paragraph().paragraph_format.space_after = Pt(60)
    c = d.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run("SYNAPEP"); r.font.size = Pt(40); r.bold = True; r.font.color.rgb = NAVY
    s = d.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = s.add_run("Master Investor Dossier"); rs.font.size = Pt(20); rs.font.color.rgb = GREY
    s2 = d.add_paragraph(); s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs2 = s2.add_run("AI-Personalised Aesthetic Medical Devices · Korea-anchored, export-led")
    rs2.font.size = Pt(12); rs2.italic = True; rs2.font.color.rgb = GREY
    d.add_paragraph().paragraph_format.space_after = Pt(30)
    disc = d.add_paragraph(); disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rd = disc.add_run("Consolidates the business plan, financial model, multi-market scenarios, "
                      "valuation, export strategy, diligence scorecard and indicative term sheet "
                      "into one document. Illustrative — not investment, legal or tax advice.")
    rd.font.size = Pt(10); rd.italic = True; rd.font.color.rgb = GREY
    # Contents
    d.add_page_break(); d.add_heading("Contents", level=1)
    for n, (_, title) in enumerate(PARTS, 1):
        p = d.add_paragraph(style="List Number"); p.add_run(title).font.size = Pt(11)
    # Parts
    for n, (fname, title) in enumerate(PARTS, 1):
        divider(d, n, title)
        with open(DIR + fname, encoding="utf-8") as f:
            text = f.read()
        # drop the file's own first H1 (we show the Part title divider instead)
        text = re.sub(r'^\s*#\s+.*\n', '', text, count=1)
        render_markdown(d, text)
    d.save(DIR + "Synapep_Master_Dossier.docx")
    print("Saved Synapep_Master_Dossier.docx (%d parts)" % len(PARTS))


if __name__ == "__main__":
    main()
