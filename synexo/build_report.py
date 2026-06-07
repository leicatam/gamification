#!/usr/bin/env python3
"""Build the comprehensive investor report -> Synapep_Comprehensive_Report.docx.

Single-sources all financials from build_multimarket_model (RESULTS) so the report, the model
doc and the deck never drift.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
import build_multimarket_model as mm

DIR = "/home/user/gamification/synexo/"
NAVY = RGBColor(0x1F, 0x3A, 0x5F); GREY = RGBColor(0x55, 0x55, 0x55); RED = RGBColor(0xB1, 0x2A, 0x2A)
Y = mm.YEARS; R = mm.RESULTS


def doc_init():
    d = Document()
    n = d.styles["Normal"]; n.font.name = "Calibri"; n.font.size = Pt(10.5)
    for lvl, sz in [("Heading 1", 15), ("Heading 2", 12.5), ("Heading 3", 11)]:
        st = d.styles[lvl]; st.font.name = "Calibri"; st.font.size = Pt(sz)
        st.font.color.rgb = NAVY; st.font.bold = True
    return d


def P(d, text="", size=10.5, bold=False, italic=False, color=None, align=None, after=6):
    p = d.add_paragraph()
    if align: p.alignment = align
    r = p.add_run(text); r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after); return p


def B(d, text, lead=None):
    p = d.add_paragraph(style="List Bullet")
    if lead:
        r = p.add_run(lead); r.bold = True
    p.add_run(text); return p


def NB(d, text, n):
    p = d.add_paragraph(style="List Number"); p.add_run(text); return p


def T(d, headers, rows, widths=None, fs=9.5):
    t = d.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        r = t.rows[0].cells[i].paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(fs)
    for row in rows:
        c = t.add_row().cells
        for i, v in enumerate(row):
            rr = c[i].paragraphs[0].add_run(str(v)); rr.font.size = Pt(fs)
    if widths:
        for i, w in enumerate(widths):
            for rr in t.rows: rr.cells[i].width = Inches(w)
    d.add_paragraph().paragraph_format.space_after = Pt(2); return t


def H1(d, t): d.add_heading(t, level=1)
def H2(d, t): d.add_heading(t, level=2)


def usd_m(scn, year, key): return "$%.1fM" % (R[scn][0][year][key] / 1e6)


d = doc_init()

# ---------------- Cover
P(d, "SYNAPEP", size=30, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
P(d, "Comprehensive Investor Report", size=17, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
P(d, "AI-Personalised Aesthetic Medical Devices — Korea-anchored, export-led",
  size=12, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=20)
P(d, "Prepared as an independent operator/investor read (chemical-engineering + aesthetic-medicine "
     "lens). Illustrative planning document — not investment, legal or tax advice. Financials are "
     "scenario estimates consistent with the v0.4 model.", size=9, italic=True, color=GREY,
  align=WD_ALIGN_PARAGRAPH.CENTER)
d.add_page_break()

# ---------------- 1 Executive summary
H1(d, "1. Executive summary")
P(d, "Synapep Labs is a physician co-development lab for regenerative aesthetic medicine. Elite "
     "Korean aesthetic doctors co-design purpose-built products (personalised to the physician's "
     "clinical case, not the patient): an applicator device plus recurring CodeLife-configured "
     "peptide/exosome kits. Each doctor request becomes a registered SKU via a fast/cheap regulatory "
     "'modification engine'; the inventing physician then champions it to peers under a "
     "compliance-first benefit-sharing royalty. The company is defensible on two co-equal moats — an "
     "AI/data flywheel and a physician-network + regulatory-process engine — and registered in Korea "
     "(MFDS), then exported across Asia and the Gulf on reliance pathways.")
H2(d, "The numbers (multi-market, same product)")
T(d, ["2030 outcome", "Conservative", "Base", "Upside"],
  [["Markets", "11 (China excl.)", "12", "12"],
   ["Clinics", "{:,}".format(R["Conservative"][0][2030]["cum"]),
    "{:,}".format(R["Base"][0][2030]["cum"]), "{:,}".format(R["Upside"][0][2030]["cum"])],
   ["Revenue", usd_m("Conservative", 2030, "tot_rev"), usd_m("Base", 2030, "tot_rev"),
    usd_m("Upside", 2030, "tot_rev")],
   ["Gross profit", usd_m("Conservative", 2030, "tot_gp"), usd_m("Base", 2030, "tot_gp"),
    usd_m("Upside", 2030, "tot_gp")],
   ["EBITDA", usd_m("Conservative", 2030, "ebitda"), usd_m("Base", 2030, "ebitda"),
    usd_m("Upside", 2030, "ebitda")]],
  widths=[1.7, 1.5, 1.4, 1.4])
H2(d, "The ask")
B(d, " US$1.0M kickoff on a milestone-gated post-money SAFE, $6M cap / 20% discount, in two "
     "tranches ($400k close, $600k on milestones).", "Round —")
B(d, " priced today on the device/consumables floor ($5–6M pre); the AI/data-company upside is the "
     "Series A thesis, not the seed price.", "Valuation —")
B(d, " resolve cold-chain/COGS and produce the validation pack; stand up Korea + Wave-1 export "
     "(HK, Vietnam, Singapore); ship CodeLife v1 + the data-rights clause.", "What $1M buys —")

# ---------------- 2 Market & opportunity
H1(d, "2. Company & market opportunity")
B(d, " medical aesthetics ≈ $17–18.5B (2024), ~10–13% CAGR, → ~$56B by 2033; injectables >40% of "
     "revenue. Korea is a global aesthetics hub; Hong Kong a gateway.", "Market —")
B(d, " commodity injectables (HA, botox) are not displaced (~$5–6.5B, still growing). Synapep is a "
     "new adjacent regenerative category, AI-personalised by physician purpose.", "Positioning —")
B(d, " JV — WBI 70% (IP: CodeLife.AI, IT-EXO®/SynExo, R&D) / Eyesel 30% (manufacturing); valued "
     "standalone at ex-factory; distributors arm's-length / non-consolidated.", "Structure —")

# ---------------- 3 Evaluation
H1(d, "3. Independent evaluation — strengths and corrections")
H2(d, "Validated strengths")
B(d, " Korea MFDS 'modified device' route + Hong Kong MDACS voluntary listing enable a fast, "
     "low-cost first registration and an export springboard.", "Regulatory arbitrage —")
B(d, " device placement + recurring kits is stickier and higher-value than commodity OEM vials.", "Razor-and-blades —")
B(d, " designing peptides that outperform generalised products, plus an AI-QA/analytics loop, "
     "creates a compounding data advantage.", "AI efficacy + flywheel —")
H2(d, "Corrections carried into the plan")
NB(d, "Device-vs-drug: an IM injectable is likely a drug/biologic — lead the device fast-path with "
      "the microneedling/topical line; run IM separately.", 1)
NB(d, "'Same raw materials': personalise via device parameters within a pre-cleared envelope, not "
      "new actives.", 2)
NB(d, "Exosomes: lead with non-human (salmon/synthetic) origin and claim discipline (Korea ad ban; "
      "EU human-exosome ban; US FDA warnings).", 3)
NB(d, "Cold chain & COGS: reformulate to 2–8 °C / ambient or cost the chain; publish a bottom-up "
      "BOM. (This is the single biggest unlock.)", 4)
NB(d, "'$25M exposure': re-present at entity/product/channel level; treat as market signal, not "
      "company revenue.", 5)

# ---------------- 4 Business model
H1(d, "4. Business model — two co-equal moats")
P(d, "Synapep Labs is a physician co-development lab, defensible on two fronts at once. Either "
     "moat alone is copyable; together they compound. Devices/kits are the revenue wedge and the "
     "data-capture layer.")
B(d, " CodeLife.AI designs personalised peptides engineered to outperform generalised products; "
     "an AI-QA suite validates every batch and feeds a compounding outcomes dataset (the durable "
     "asset). Physician decision-support within a cleared envelope (SaMD-aware).",
  "Moat A — AI engine & data flywheel —")
B(d, " elite Korean KOLs co-develop purpose-built SKUs (doctor-personalised, not patient), "
     "registered fast/cheap via a regulatory 'modification engine', and spread through a "
     "compliance-first benefit-shared KOL flywheel. A relationship + process asset, hard to copy.",
  "Moat B — Physician co-development engine —")
H2(d, "Moat B — the three engines")
NB(d, "KOL co-development lab — Korea's world-class physicians co-design products from the doctor's "
      "clinical perspective; CodeLife accelerates each design.", 1)
NB(d, "Modification engine — each variant is a minor change under Korea's improved-device / "
      "negative-list regime and functional-cosmetic/quasi-drug routes (company target <3% change, "
      "~3-month cycle, <$30K/SKU → ~20 SKUs in ~6 months; <$600K total — to be validated).", 2)
NB(d, "Benefit-shared KOL flywheel (compliance-first) — the inventing physician champions the SKU "
      "to peers and shares economics as an FMV royalty for genuine IP co-invention, DECOUPLED from "
      "their own usage; K-Sunshine / anti-rebate compliant. The 'double-edged' risk is managed by "
      "design.", 3)
P(d, "See the Model Clarification annex for how this maps to the exosome / synthetic-biotech "
     "aesthetics landscape (ExoCoBio, Brexogen, designer exosomes) and its regulatory reality.",
  size=9.5, italic=True, color=GREY)

# ---------------- 5 Technology & AI moat
H1(d, "5. Technology & the AI moat")
P(d, "The differentiator is performance, not merely personalisation. The moat compounds as the same "
     "product, used across many markets, returns structured, consented outcome data that improves "
     "the model — which improves the peptides — which deepens clinic lock-in.")
B(d, " AI peptide design (CodeLife), AI-QA/batch analytics, consented outcomes capture at the "
     "device, and a closed-loop retraining cadence.", "Components —")
B(d, " 'own large model' is a direction, not a day-one claim for a small team; the defensible asset "
     "now is the dataset + data-rights + analytics, not a foundation model.", "Honest framing —")

# ---------------- 6 Regulatory
H1(d, "6. Regulatory strategy + the modification engine")
T(d, ["Regulated object", "Class / pathway", "Speed / note"],
  [["Applicator (microneedling/ion)", "Medical device — improved-device / negative-list change",
    "Fast; KR 10-2654857 anchors the family"],
   ["Doctor-personalised topical/exosome kit", "Functional cosmetic / quasi-drug (mostly outside device class)",
    "Same-active notification; <3% deltas; the SKU engine"],
   ["IM / injectable peptide", "Drug / biologic (separate track)", "Slower; must not gate launch"],
   ["CodeLife software", "Config/decision-support in a cleared envelope", "Manages SaMD exposure"]],
  widths=[2.4, 2.3, 1.8])

# ---------------- 7 Export
H1(d, "7. International expansion — KFDA as an export springboard")
P(d, "Two-speed: the device travels on the Korean technical file via reliance/recognition; the kit "
     "is wrapped per market (cosmetic vs device vs drug), leaning on salmon/synthetic origin. "
     "Precedent: Korean salmon-PDRN boosters (e.g., Rejuran) exported across SEA/Middle East/Europe "
     "off a Korean approval.")
T(d, ["Tier", "Markets & KFDA leverage"],
  [["Tier 1 — ASEAN + Hong Kong", "HK MDACS recognises Korea MFDS as a reference (pair Korea + "
    "Singapore for the ≥2-reference expedited route); Vietnam accepts MFDS; Singapore HSA gateway; "
    "ASEAN AMDD/CSDT cascade to Thailand, Malaysia, Indonesia, Philippines."],
   ["Tier 2 — GCC + Greater China", "Saudi SFDA (GCC regional reference) + UAE MOHAP; Taiwan TFDA "
    "accepts Korean data; China NMPA via Hainan/cross-border first."]],
  widths=[1.9, 4.6])
P(d, "Sequence: Korea anchor (2027) → HK/Vietnam/Singapore (2027–28) → ASEAN + GCC (2028–29) → "
     "Taiwan/China (2029–30). Each market needs a local responsible person/distributor.", size=9.5,
  italic=True, color=GREY)

# ---------------- 8 Financials
H1(d, "8. Multi-market financial model (3 scenarios)")
P(d, "Same Korean-developed product across all markets — scenarios flex only the clinic ramp, price "
     "and utilisation, never the product, so margins are stable and R&D is not duplicated.")
H2(d, "Total revenue ($000s)")
T(d, ["Scenario"] + [str(y) for y in Y],
  [[s] + [mm.fmt(R[s][0][y]["tot_rev"]) for y in Y] for s in mm.ORDER],
  widths=[1.7, 1.1, 1.1, 1.1, 1.1])
H2(d, "EBITDA ($000s)")
T(d, ["Scenario"] + [str(y) for y in Y],
  [[s] + [mm.fmt(R[s][0][y]["ebitda"]) for y in Y] for s in mm.ORDER],
  widths=[1.7, 1.1, 1.1, 1.1, 1.1])
P(d, "Unit economics (identical everywhere): kit EXW $30–35, kit GM ~64–67%%; device ASP $3,000, "
     "GM 35–42%%. Distributors arm's-length / non-consolidated; revenue booked at ex-factory. Full "
     "per-market clinic build in the Multi-Market Model document.", size=9.5, italic=True, color=GREY)

# ---------------- 9 Valuation
H1(d, "9. Valuation — two lenses")
B(d, " device/consumables business — $5–7M pre / $6–8M post; recommend $6M midpoint for the $1M "
     "kickoff.", "Lens A (floor, priced today) —")
B(d, " AI/data-science company — shifts comps from medtech to AI/data multiples; reframes the 2028 "
     "Series A as a premium, institutional AI round. The prize, not the seed price.", "Lens B (upside) —")
H2(d, "Cap table — $6.0M pre / $1.0M (illustrative)")
T(d, ["Holder", "Pre", "Post-raise"],
  [["WBI", "70.0%", "60.0%"], ["Eyesel", "30.0%", "25.7%"], ["Investor (SAFE-converted)", "—", "14.3%"]],
  widths=[2.6, 1.5, 1.5])

# ---------------- 10 DD scorecard
H1(d, "10. Diligence scorecard (technical-investor RAG)")
T(d, ["Category", "RAG", "Gating question"],
  [["Commercial validation", "AMBER", "$25M by entity/product/channel + reorder data?"],
   ["Chemistry / formulation", "AMBER", "Is the personalisation envelope clean AND differentiating?"],
   ["Manufacturing & cold chain", "RED", "2–8 °C reformulation or costed cold chain + BOM?"],
   ["Regulatory", "AMBER", "Does the device path survive the envelope; claims controlled?"],
   ["AI / data moat", "AMBER", "Data-rights in the clinic contract; CodeLife v1 demo?"],
   ["Financials", "AMBER", "Rebuild base off anchor-clinic actuals?"],
   ["Valuation / deal", "AMBER", "What price holds before the validation pack exists?"],
   ["IP & JV", "AMBER", "WBI IP assignment/licence schedule into the JV?"],
   ["Team / capital adequacy", "AMBER", "Which milestones can $1M actually hit?"],
   ["Market reach & export", "AMBER", "Per-market kit class; salmon/synthetic clears bans?"],
   ["Physician benefit-sharing", "ORANGE", "FMV co-invention royalty, decoupled from usage; K-Sunshine?"]],
  widths=[2.0, 0.9, 3.6], fs=9)
P(d, "Verdict: conditional interest — diligence yes, their price no. Price $5–6M, milestone-gated. "
     "Resolve cold chain + COGS (RED) and the validation pack and amber→green across the board.",
  italic=True, size=9.5)

# ---------------- 11 Proposed deal
H1(d, "11. Proposed deal — milestone-gated SAFE")
T(d, ["Term", "Proposed"],
  [["Instrument", "Post-money SAFE (milestone tranches)"],
   ["Amount", "US$1.0M — $400k close / $600k milestone"],
   ["Cap / discount", "$6.0M / 20% / MFN"],
   ["Conditions precedent", "IP into JV; $25M sample verification; clinic data-rights clause; "
    "physician co-development/IP-royalty framework + K-Sunshine compliance opinion; clean cap table"],
   ["Milestones (Tranche 2)", "Cold-chain spec + COGS BOM; 5–10 anchor clinics w/ reorders; one "
    "prospective readout; classification memo"],
   ["Governance", "Board observer; pro-rata; information rights"]],
  widths=[1.9, 4.6])

# ---------------- 12 Risks
H1(d, "12. Key risks & mitigations")
T(d, ["Risk", "Mitigation"],
  [["Device-vs-drug reclassification", "Lead device fast-path; IM on separate track; system registration"],
   ["Cold-chain / COGS unproven", "Reformulate to 2–8 °C; bottom-up BOM; QC release costing"],
   ["Exosome claims restricted", "Non-human origin; function-based claims; per-market claims map"],
   ["Aggressive clinic ramp", "Rebuild base off anchor-clinic reorder cohorts"],
   ["Kit class varies by market", "Per-market regulatory wrapper; reliance where available"],
   ["Capital adequacy / Series A risk", "Milestone gating; lean export via distributors"]],
  widths=[2.3, 4.2])

# ---------------- 13 Roadmap / asks
H1(d, "13. Roadmap, use of funds & asks")
B(d, " Korea registration + anchor clinics; cold-chain/COGS resolution; CodeLife v1 + data-rights; "
     "Wave-1 export prep (HK, Vietnam, Singapore).", "0–12 months ($1M) —")
B(d, " ASEAN cascade + GCC; prospective beachhead readout; Series A (AI round) to fund multi-market "
     "registration + clinical salesforce.", "12–30 months —")
H2(d, "Open items to confirm")
for t in ["Verify the $25M at entity/product/channel level (validation pack).",
          "Confirm device ASP/GM and per-market clinic ramps vs clinic counts.",
          "Head-to-head efficacy evidence that AI peptides outperform generalised products.",
          "Current ML capability/data assets vs the 'own large model' build plan.",
          "Round strategy: lean $1M seed (Lens A) vs a larger institutional-AI round (Lens B)."]:
    B(d, " " + t)

d.save(DIR + "Synapep_Comprehensive_Report.docx")
print("Saved Synapep_Comprehensive_Report.docx")
