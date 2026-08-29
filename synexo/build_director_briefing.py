#!/usr/bin/env python3
"""Director heads-up briefing -> director-briefing.md + GenApep_Director_Briefing.docx.

Heads-up for the directors of WBI and Eyesel on the GenApep IPO project:
structure, implementation path, pros and cons, recommended shareholding
structure, and what it means for WBI shareholders.

Sources (figures cited from): GenApep_JV_Financial_Plan.xlsx (rev 29 Aug 2026),
GLP-1 proforma (fvh-glp1-proforma.xlsx), WBI x Eyesel MOU 12 Aug 2026 +
Amendment No. 1 + Addendum A, GenApep x First Vital Stage 2 Integration
Framework (29 Aug 2026), GenApep-DPW Licensing Framework (28 Aug 2026).
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

DIR = "/home/user/gamification/synexo/"
NAVY = RGBColor(0x1B, 0x33, 0x55); GREY = RGBColor(0x55, 0x55, 0x55); RED = RGBColor(0xC0, 0x39, 0x2B)

# ---------------- figures (JV Financial Plan, rev 29 Aug 2026) ----------------
YEARS = ["FY2025A", "2026P", "2027P", "2028P", "2029P", "2030P"]
REV = [11.43, 16.39, 25.31, 31.70, 43.82, 58.91]
EBIT = [0.76, 2.40, 4.54, 5.95, 8.86, 12.51]
POOL = 0.075                       # deadlock-fallback midpoint (Amendment No.1, A4)
EY_IN, WBI_IN = 0.51, 0.49         # internal JV split
EY_POST = EY_IN * (1 - POOL)       # 47.2%
WBI_POST = WBI_IN * (1 - POOL)     # 45.3%
GA_MID = 35.0                      # combined GenApep value, mid of $30-40M planning range
FVH_SCEN = [5.0, 10.0, 13.0]       # FVH verified-value scenarios ($M)

def block(fvh_v, ga_v=GA_MID):
    return ga_v / (ga_v + fvh_v)

CONF = ("PRIVATE & CONFIDENTIAL — for the directors of WBI and Eyesel only. "
        "Not for release to First Vital or any third party.")
DISCLAIMER = ("Heads-up briefing prepared to open board discussion. Non-binding; not legal, tax, "
              "accounting, securities, healthcare-regulatory or clinical advice. Figures are "
              "planning figures from the parties' own workbooks and are illustrative.")

PROJECT_BRIEF = (
    "The project takes WBI and Eyesel to a US public market in two stages. Stage 1 forms the "
    "Joint Venture: GenApep Holdings (Cayman) owns GenApep Korea, which acquires 100% of WBI and "
    "100% of Eyesel; Eyesel shareholders hold 51%, WBI shareholders 49%, both diluted pro rata by "
    "a 5–10% technology sweat pool. Stage 2 is the listing route, decided later and separately: "
    "Route A (default) — GenApep's own US offering, confidential SEC draft by end-2027, listing "
    "2028; or Route B — combination with First Vital Health and Wellness Inc (FVH), a US company "
    "with a stated SEC filing record and an operating GLP-1 patient-management programme. "
    "Signing Amendment No. 1 starts Stage 1 now; Addendum A creates the machinery to consider "
    "Route B without committing to it. The PCAOB audit of both companies is the critical path "
    "under every route, so forming the JV first costs no calendar time.")

STEPS = [
    ("1. Sign Amendment No. 1", "Immediate",
     "Fixes nine defects in the 12-Aug MOU that matter under every scenario: mandatory 51:49 "
     "equalisation once a >15% variance is shown, consideration to shareholder bodies pro rata "
     "(never to individuals), Third-Director disclosure and deadlock fallback, technology "
     "chain-of-title as a condition precedent."),
    ("2. Execute Stage 1", "Q3-2026 → Q2-2027",
     "Incorporate GenApep Holdings (Cayman) and GenApep Korea; deliver Eyesel's updated financial "
     "report (30 days); appoint the PCAOB-registered auditor in Q4-2026 — the critical path; "
     "definitive SHA + sweat-pool KPI schedule; share exchange Q1–Q2 2027."),
    ("3. Sign Addendum A when ready to consider Route B", "When FVH discussion is live",
     "Makes the route a Reserved Matter decision: Route A remains the default; GenApep Korea is "
     "the only permitted acquisition vehicle; any Stage 2 transaction distributes consideration "
     "strictly pro rata through the block."),
    ("4. Verify First Vital before any terms", "This week onward",
     "Gate zero: obtain FVH's CIK and complete EDGAR filing index. Request July–August GLP-1 "
     "actual enrolments, claims submitted/paid and denial rates. No terms before verification."),
    ("5. If Route B is chosen: one block, one ratio", "After gates pass",
     "Independent valuations of both sides + a written fairness opinion; FVH issues shares to "
     "GenApep Holdings' shareholders pro rata for 100% of GenApep Holdings, simultaneously at "
     "closing; the sweat pool carries through on equivalent terms; then reporting obligations and, "
     "separately, an exchange-listing application when the group qualifies."),
]

ROUTE_AB = [
    ("Speed to being public",
     "Own registration: confidential draft end-2027, listing 2028; market-dependent.",
     "Potentially faster to reporting status via the combination — but reporting status is not an "
     "exchange listing; a separate listing application follows either way."),
    ("Capital raised",
     "Primary raise at the offering (illustrative US$20M at US$80M pre-money).",
     "A combination raises no money. A financing plan must be built separately."),
    ("Issuer profile",
     "Foreign private issuer assumed: Form F-1, IFRS, 20-F/6-K reporting.",
     "US parent likely means domestic issuer: S-1/S-4, US GAAP, 10-K/10-Q/8-K and full proxy "
     "rules — more cost and time than currently budgeted. Written counsel opinion required."),
    ("What we get besides listing",
     "Nothing external — the group's own story only.",
     "FVH's GLP-1 programme (US$2.2M 2026 → US$16.7M 2027 proforma; recurring MRR US$1.59M/month "
     "exiting 2027) and a US clinical channel for GenApep peptide products."),
    ("Counterparty risk",
     "None — no external counterparty.",
     "FVH's SEC status is unverified; the GLP-1 programme has ~2 months of trading history; "
     "possible shell-company restrictions; inherited liabilities; related-party optics."),
    ("Control of the process",
     "Fully in the JV's hands.",
     "Shared with FVH's board and existing holders; exclusivity and governance must be negotiated."),
]

FVH_PROS = [
    ("Existing SEC filing record", "If verified, cuts the time and execution risk of becoming a reporting company."),
    ("Operating US clinical channel", "The GLP-1 patient base is a natural route to market for GenApep peptides — GI/metabolic support, muscle preservation, GLP-1-associated hair loss."),
    ("Real recurring-revenue programme", "CPT-reimbursed care pathway; proforma US$2.2M (H2-2026) → US$16.7M (2027); ARPU US$168.92/patient/month."),
    ("US infrastructure & shareholder base", "Corporate platform and float that the JV would otherwise build from zero."),
]
FVH_CONS = [
    ("Filing status unverified", "A public search did not surface an identifiable EDGAR record; if the Dec-2024 filing was a Reg-A offering statement, FVH may not be a reporting company at all. Gate zero."),
    ("Two months of trading history", "The GLP-1 programme began enrolling 1 July 2026; the 10,000-patient ramp is a plan, not a track record. July–August actuals are the cheapest diligence available."),
    ("Proforma is revenue-only", "No clinician cost, billing cost, denials, collections. The 20% EBITDA margin used in the draft IPO MOU is an assumption; a cost build is needed."),
    ("Issuer/accounting consequences", "US-parent structure likely forfeits foreign-private-issuer status: US GAAP + domestic reporting — unbudgeted cost and time."),
    ("Shell risk & no capital", "Possible shell-company resale restrictions; a combination raises no money; the group could emerge public but undercapitalised."),
    ("Related-party exposure", "Principals could sit on both sides; only independent valuations, a fairness opinion and disinterested approval make the transaction defensible."),
]

REC_STRUCTURE = (
    "First Vital's share structure is undetermined, which is an opportunity: the parties can set "
    "it rationally rather than inherit it. The recommendation is a single value-based exchange "
    "ratio: independent valuations of GenApep (planning range US$30–40M before technology "
    "upside) and of FVH (verified value — to be established in diligence), a written fairness "
    "opinion on the ratio, and entry as ONE block: FVH issues shares to GenApep Holdings' "
    "shareholders pro rata. The GenApep block percentage then divides internally at 51:49 less "
    "the sweat pool, automatically. No shares to individuals; no separate negotiations inside "
    "the FVH cap table; the draft IPO MOU's 40/35/25 allocation to two individuals is replaced.")

# cap table at mid: GA $35M, FVH verified $10M -> block 77.8%
CAP_ROWS = []
for fvh_v in FVH_SCEN:
    b = block(fvh_v)
    CAP_ROWS.append(["FVH verified at US$%.0fM" % fvh_v,
                     "%.0f%%" % (b * 100),
                     "%.1f%%" % (EY_POST * b * 100),
                     "%.1f%%" % (WBI_POST * b * 100),
                     "%.1f%%" % (POOL * b * 100),
                     "%.0f%%" % ((1 - b) * 100)])

DPW = (
    "A commercial licensing framework with DPW (an entity under the listed Richards Group) was "
    "circulated on 28 August 2026. It is a non-binding modular structure — a Master Technology "
    "Commercialization Agreement with asset-specific Rights Schedules choosing, per asset, between "
    "(1) an exclusive field-of-use licence, (2) acquisition of defined AI-peptide patent assets "
    "with included development work, or (3) non-exclusive access with first-notice rights. "
    "Economics are architecture only at this stage: access fee + development/regulatory "
    "milestones + running royalty + annual minimums. Why it matters to valuation: an arm's-length "
    "licence with a listed group's subsidiary is third-party evidence of the technology's value — "
    "it supports the AI.pep value scenarios (US$5–10M illustrative) that the JV plan currently "
    "carries at US$0 — and adds a royalty income line. Two cautions: no grant should be signed "
    "before the chain-of-title condition (Clause 5.7) is satisfied, and field-of-use exclusivity "
    "must be drawn narrowly so the platform is not stranded ahead of the IPO story.")

WBI_LENS = [
    ("Your starting position",
     "WBI shareholders hold 49% of GenApep, diluting to ~45.3% at the 7.5% pool midpoint — worth "
     "≈ US$15.9M at the US$35M combined planning mid, before any AI.pep/DPW technology credit."),
    ("Route A outcome (workbook illustration)",
     "At an US$80M pre-money IPO with a US$20M raise, WBI shareholders ≈ US$36.3M post-IPO "
     "(JV plan, 7.5% pool). MM Studio — WBI's own growth engine (US$2.9M → US$36.3M by 2030) — "
     "is the single biggest swing factor in that valuation."),
    ("Route B outcome (recommended structure)",
     "At FVH verified US$10M, WBI shareholders ≈ 35.3% of the listed company — ≈ US$15.9M at the "
     "US$45M combined value, and ≈ US$28–35M if the market values the group at US$80–100M. The "
     "GLP-1 channel adds a US revenue engine the JV does not otherwise have."),
    ("What the draft IPO MOU would have done",
     "It allocated 40% of FVH to one individual and 35% to another, with nothing to the wider "
     "WBI and Eyesel shareholder bodies at that step, and produced ~53:47 — reversing the agreed "
     "51:49. Amendment A2 and Addendum clauses B4/B5 exist precisely to prevent this: every "
     "shareholder rides through pro rata or the transaction does not proceed."),
    ("Technology upside accrues to the JV",
     "AI.pep and the DPW licence sit inside GenApep, so WBI shareholders participate through "
     "their block stake; the sweat pool (Keith · Sidney · Charles + team) is the incentive "
     "mechanism and survives any change of vehicle (Clause 7.8)."),
]

INTERIM_INTRO = (
    "Eyesel's updated financial data is not yet available, and the 51:49 split rests on it. That "
    "does not block signing: fairness comes from mechanism, not from guessing the missing numbers. "
    "The four devices below let the parties commit now and correct automatically when the data "
    "arrives — the first is already drafted into Amendment No. 1.")

INTERIM_MECH = [
    ["1. Peg + mandatory true-up (Amendment A1 — already drafted)",
     "Sign on a stated basis recorded on the face of the document (Eyesel revenue US$8.0M, value "
     "US$15–20M). If verified figures differ by >15%, equalisation is MANDATORY — cash or share "
     "adjustment — with an independent valuer deciding if boards cannot agree in 30 days. "
     "Symmetric: it also protects Eyesel if its real numbers are higher."],
    ["2. Verified financials as a condition precedent to the share exchange",
     "Not to signing. The exchange is Q1–Q2 2027 and the PCAOB audit produces verified numbers on "
     "the critical path anyway — nothing is lost by signing the MOU and Amendment now."],
    ["3. Escrow holdback (fallback)",
     "If the 30-day report deadline slips: hold back 5–8 points of the Eyesel allocation in "
     "escrow, released or reallocated on verification."],
    ["4. Parity of method, not of number",
     "The same independent valuer applies the same valuation method and multiples to both "
     "companies (both are manufacturing-economics businesses). The parties agree the METHOD "
     "today; the numbers follow from the data."],
]

INTERIM_NUMS = [
    ["Eyesel revenue", "US$8.0M", "The documented management-memo basis — the LOWER of the two circulating figures; conservative, and the true-up protects Eyesel if ~US$15M proves right"],
    ["Gross margin", "40%", "Mirrors WBI's NTS-certified FY2025 actual — same manufacturing economics"],
    ["EBITDA margin", "~12%", "Typical GMP-manufacturer profile; placeholder"],
    ["Growth", "~10%/yr", "JV plan placeholder"],
    ["Value band", "US$15–20M", "Same band as WBI — symmetric valuation-parity; prejudices neither side"],
]

INTERIM_CHECK = (
    "The check to show both boards: on a pure value-proportional basis Eyesel needs a verified "
    "value of ≥ ~US$18M (against WBI's US$17.5M mid) to arithmetically support 51%. At US$8.0M "
    "revenue that likely requires counting the GMP facility, licences and quality systems as "
    "asset value — a legitimate argument, but one that must be documented, or equalisation "
    "applies. At ~US$15M revenue, 51:49 is comfortable.")

INTERIM_FVH = (
    "Structurally, the First Vital discussion never waits on Eyesel's data: under block entry the "
    "FVH ratio is GenApep-block vs FVH-verified-value, and the internal 51:49 trues up inside the "
    "block without touching the listed-company cap table. Only the internal split is gated — and "
    "it self-corrects.")

RESOLVE = [
    "Eyesel revenue: US$8.0M (JV plan placeholder, basis of the 51% stake) vs ~US$15M (draft IPO MOU). Mandatory equalisation applies beyond 15% variance — this figure moves the split.",
    "Eyesel principal: Mr. Park (12-Aug MOU) vs Mr. Kim (draft IPO MOU). Written confirmation needed.",
    "Technology chain of title for AI.pep, CodeLife.AI, IT-EXO, SynExo — condition precedent (Clause 5.7); also DPW's first diligence gap.",
    "One operating model: WBI-core/MM-Studio/Eyesel margins (JV plan) vs the vial model at 78.5% GM (draft IPO MOU). One architecture must be chosen before any valuation.",
    "FVH verification: CIK + full EDGAR index; shell status; July–August GLP-1 actuals.",
    "The draft IPO MOU itself: not signable as drafted (individual allocations, FVH Korea vehicle, no sweat pool). To be revised to the block-entry structure above.",
]

DECISIONS = [
    "Sign Amendment No. 1 and start Stage 1 now (both boards).",
    "Treat Addendum A as the gate for any First Vital discussion; sign it when ready to consider Route B.",
    "Instruct the FA to obtain FVH's CIK and complete EDGAR index, and July–August GLP-1 actuals — this week.",
    "Eyesel: deliver the updated financial report and confirm the principal — within 30 days.",
    "Appoint the PCAOB-registered auditor in Q4-2026 (critical path), scoped for either route.",
    "Do not sign the draft IPO MOU in its current form; instruct revision to the value-based block-entry structure.",
    "Progress the DPW framework to asset-level term sheets only after chain-of-title is documented.",
]

DOCMAP = [
    ("WBI × Eyesel MOU (12 Aug 2026)", "The Joint Venture itself — structure, 51:49, contributions, governance."),
    ("Amendment No. 1", "Nine fixes needed under every scenario; read Part 1 table first."),
    ("Addendum A", "Route A/B machinery and the protections for any external listing counterparty."),
    ("GenApep JV Financial Plan (xlsx)", "The boards' own numbers: 5-year plan, valuation & dilution, IPO roadmap, sensitivity."),
    ("Stage 2 Integration Framework", "The full First Vital analysis: sequencing, conflicts register, diligence programme, go/no-go."),
    ("GLP-1 proforma (xlsx)", "FVH's programme model — revenue-only; note the MRR-sheet labelling issue."),
    ("GenApep × DPW Licensing Framework", "The technology-licence structure with Richards Group's DPW; scenario menu and diligence gaps."),
    ("Draft IPO Framework MOU", "The earlier FVH draft — superseded in approach by this briefing; do not sign as-is."),
]

# ------------------------------------------------------------------ Markdown
def wmd():
    L = []; a = L.append
    a("# GenApep — Director Briefing: the IPO project, First Vital and the road ahead\n")
    a("> **%s**\n>\n> %s\n" % (CONF, DISCLAIMER))
    a("\n## 1. The project in brief\n\n%s\n" % PROJECT_BRIEF)
    a("\n## 2. How we implement\n")
    for t, w, d in STEPS:
        a("- **%s** (*%s*) — %s\n" % (t, w, d))
    a("\n## 3. Route A vs Route B — pros and cons\n\n")
    a("| Dimension | Route A — own IPO (default) | Route B — combine with First Vital |\n|---|---|---|\n")
    for dim, ra, rb in ROUTE_AB:
        a("| %s | %s | %s |\n" % (dim, ra, rb))
    a("\n### What First Vital specifically brings\n")
    for t, d in FVH_PROS:
        a("- **%s** — %s\n" % (t, d))
    a("\n### And the risks it carries\n")
    for t, d in FVH_CONS:
        a("- **%s** — %s\n" % (t, d))
    a("\n## 4. Recommended shareholding structure (FVH structure is ours to set)\n\n%s\n" % REC_STRUCTURE)
    a("\n**Illustrative combined cap table** (GenApep at US$%.0fM mid; 7.5%% pool; value-based ratio):\n\n" % GA_MID)
    a("| Scenario | GenApep block | Eyesel shareholders | WBI shareholders | Tech pool | FVH holders |\n|---|---|---|---|---|---|\n")
    for r in CAP_ROWS:
        a("| " + " | ".join(r) + " |\n")
    a("\n*The block percentage is set by the valuations; the internal split never needs renegotiating.*\n")
    a("\n### The DPW licence — technology credit to the valuation\n\n%s\n" % DPW)
    a("\n## 5. What this means for WBI shareholders\n")
    for t, d in WBI_LENS:
        a("- **%s** — %s\n" % (t, d))
    a("\n## 6. Fair without Eyesel's numbers — interim assumptions\n\n%s\n" % INTERIM_INTRO)
    a("\n### The four mechanisms\n")
    for t, d in INTERIM_MECH:
        a("- **%s** — %s\n" % (t, d))
    a("\n### Planning numbers until the data arrives\n\n| Item | Assumption | Rationale |\n|---|---|---|\n")
    for r in INTERIM_NUMS:
        a("| " + " | ".join(r) + " |\n")
    a("\n%s\n\n%s\n" % (INTERIM_CHECK, INTERIM_FVH))
    a("\n## 7. What must be resolved before definitive documents\n")
    for i, r in enumerate(RESOLVE, 1):
        a("%d. %s\n" % (i, r))
    a("\n## 8. The financial picture (JV plan, planning basis)\n\n")
    a("| US$M | " + " | ".join(YEARS) + " |\n|---|" + "---|" * 6 + "\n")
    a("| Revenue | " + " | ".join("%.1f" % v for v in REV) + " |\n")
    a("| Operating profit | " + " | ".join("%.1f" % v for v in EBIT) + " |\n")
    a("\nCombined pre-JV valuation planning range US$30–40M (before AI.pep/DPW credit); illustrative "
      "IPO US$80M pre-money + US$20M raise. First Vital adds the GLP-1 line (US$2.2M → US$16.7M "
      "proforma) under Route B.\n")
    a("\n## 9. Decisions requested\n")
    for i, d in enumerate(DECISIONS, 1):
        a("%d. %s\n" % (i, d))
    a("\n## 10. Document map\n\n| Document | Read it for |\n|---|---|\n")
    for t, d in DOCMAP:
        a("| %s | %s |\n" % (t, d))
    open(DIR + "director-briefing.md", "w").write("".join(L))

wmd()

# ------------------------------------------------------------------ Word
def nd():
    doc = Document()
    n = doc.styles["Normal"]; n.font.name = "Calibri"; n.font.size = Pt(10.5)
    for lvl, sz in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)]:
        st = doc.styles[lvl]; st.font.name = "Calibri"; st.font.size = Pt(sz)
        st.font.color.rgb = NAVY; st.font.bold = True
    return doc

def para(doc, t="", size=10.5, bold=False, italic=False, color=None, align=None, after=6):
    p = doc.add_paragraph()
    if align: p.alignment = align
    r = p.add_run(t); r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after); return p

def bullet(doc, t, lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if lead:
        r = p.add_run(lead); r.bold = True
    p.add_run(t); p.paragraph_format.space_after = Pt(3); return p

def numbered(doc, t):
    p = doc.add_paragraph(style="List Number")
    p.add_run(t); p.paragraph_format.space_after = Pt(3); return p

def tbl(doc, headers, rows, widths=None, fs=9.5, hl_rows=()):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        r = t.rows[0].cells[i].paragraphs[0].add_run(str(h)); r.bold = True; r.font.size = Pt(fs)
    for ri, row in enumerate(rows):
        c = t.add_row().cells
        for i, v in enumerate(row):
            rr = c[i].paragraphs[0].add_run(str(v)); rr.font.size = Pt(fs)
            if ri in hl_rows or i == 0: rr.bold = True
    if widths:
        for i, w in enumerate(widths):
            for r in t.rows: r.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2); return t

doc = nd()
para(doc, "GENAPEP", size=26, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
para(doc, "Director Briefing — the IPO project, First Vital and the road ahead",
     size=14, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
para(doc, CONF, size=9, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
para(doc, DISCLAIMER, size=8.5, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)

doc.add_heading("1. The project in brief", level=2)
para(doc, PROJECT_BRIEF)

doc.add_heading("2. How we implement", level=2)
tbl(doc, ["Step", "Window", "What happens"], [[t, w, d] for t, w, d in STEPS], widths=[1.7, 1.2, 3.6])

doc.add_heading("3. Route A vs Route B — pros and cons", level=2)
tbl(doc, ["Dimension", "Route A — own IPO (default)", "Route B — combine with First Vital"],
    [[d, a_, b_] for d, a_, b_ in ROUTE_AB], widths=[1.4, 2.5, 2.6])
doc.add_heading("What First Vital specifically brings", level=3)
for t, d in FVH_PROS:
    bullet(doc, d, t + " — ")
doc.add_heading("And the risks it carries", level=3)
for t, d in FVH_CONS:
    bullet(doc, d, t + " — ")

doc.add_heading("4. Recommended shareholding structure", level=2)
para(doc, REC_STRUCTURE)
para(doc, "Illustrative combined cap table (GenApep at US$35M mid; 7.5% pool; value-based ratio):",
     bold=True, after=3)
tbl(doc, ["Scenario", "GenApep block", "Eyesel sh.", "WBI sh.", "Tech pool", "FVH holders"],
    CAP_ROWS, widths=[1.9, 1.1, 1.0, 1.0, 0.9, 1.0], hl_rows=(1,))
para(doc, "The block percentage is set by the valuations; the internal 51:49-less-pool split never "
          "needs renegotiating.", size=9, italic=True, color=GREY)
doc.add_heading("The DPW licence — technology credit to the valuation", level=3)
para(doc, DPW)

doc.add_heading("5. What this means for WBI shareholders", level=2)
for t, d in WBI_LENS:
    bullet(doc, d, t + " — ")

doc.add_heading("6. Fair without Eyesel's numbers — interim assumptions", level=2)
para(doc, INTERIM_INTRO)
doc.add_heading("The four mechanisms", level=3)
for t, d in INTERIM_MECH:
    bullet(doc, d, t + " — ")
doc.add_heading("Planning numbers until the data arrives", level=3)
tbl(doc, ["Item", "Assumption", "Rationale"], INTERIM_NUMS, widths=[1.4, 1.1, 4.0])
para(doc, INTERIM_CHECK)
para(doc, INTERIM_FVH, bold=True)

doc.add_heading("7. What must be resolved before definitive documents", level=2)
for r in RESOLVE:
    numbered(doc, r)

doc.add_heading("8. The financial picture (JV plan, planning basis)", level=2)
tbl(doc, ["US$M"] + YEARS,
    [["Revenue"] + ["%.1f" % v for v in REV],
     ["Operating profit"] + ["%.1f" % v for v in EBIT]],
    widths=[1.3] + [0.87] * 6, hl_rows=())
para(doc, "Combined pre-JV valuation planning range US$30–40M (before AI.pep/DPW credit); "
          "illustrative IPO US$80M pre-money + US$20M primary raise. First Vital adds the GLP-1 "
          "line (US$2.2M → US$16.7M proforma) under Route B.", size=9, italic=True, color=GREY)

doc.add_heading("9. Decisions requested", level=2)
for d_ in DECISIONS:
    numbered(doc, d_)

doc.add_heading("10. Document map", level=2)
tbl(doc, ["Document", "Read it for"], [[t, d] for t, d in DOCMAP], widths=[2.6, 3.9])
para(doc, "Prepared to open discussion at the WBI and Eyesel boards. All terms subject to "
          "diligence, valuations and definitive agreements.",
     size=8.5, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.save(DIR + "GenApep_Director_Briefing.docx")
print("Saved director-briefing.md, GenApep_Director_Briefing.docx")
for r in CAP_ROWS:
    print("  cap:", r)
