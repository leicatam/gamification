#!/usr/bin/env python3
"""Word version of the investor deck -> Synapep_Investor_Deck.docx (one section per slide).
Numbers single-sourced from build_multimarket_model (RESULTS)."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import build_multimarket_model as mm

DIR = "/home/user/gamification/synexo/"
NAVY = RGBColor(0x1F, 0x3A, 0x5F); GREY = RGBColor(0x55, 0x55, 0x55); ACCENT = RGBColor(0x2E, 0x86, 0xC1)
Y = mm.YEARS; R = mm.RESULTS


def init():
    d = Document()
    n = d.styles["Normal"]; n.font.name = "Calibri"; n.font.size = Pt(11)
    for lvl, sz in [("Heading 1", 15), ("Heading 2", 12)]:
        st = d.styles[lvl]; st.font.name = "Calibri"; st.font.size = Pt(sz)
        st.font.color.rgb = NAVY; st.font.bold = True
    return d


def P(d, t="", size=11, bold=False, italic=False, color=None, align=None, after=6):
    p = d.add_paragraph()
    if align: p.alignment = align
    r = p.add_run(t); r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after); return p


def slide(d, n, title):
    h = d.add_heading("", level=1)
    r1 = h.add_run("Slide %d  " % n); r1.font.color.rgb = ACCENT
    r2 = h.add_run("· " + title)
    return h


def B(d, t, sub=False):
    p = d.add_paragraph(style="List Bullet 2" if sub else "List Bullet")
    p.add_run(t); p.paragraph_format.space_after = Pt(3); return p


def LB(d, lead, rest):
    p = d.add_paragraph(style="List Bullet")
    r = p.add_run(lead); r.bold = True; p.add_run(rest)
    p.paragraph_format.space_after = Pt(3); return p


def TBL(d, headers, rows, widths=None):
    t = d.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hh in enumerate(headers):
        r = t.rows[0].cells[i].paragraphs[0].add_run(hh); r.bold = True; r.font.size = Pt(10)
    for row in rows:
        c = t.add_row().cells
        for i, v in enumerate(row):
            rr = c[i].paragraphs[0].add_run(str(v)); rr.font.size = Pt(10)
    if widths:
        for i, w in enumerate(widths):
            for rr in t.rows: rr.cells[i].width = Inches(w)
    d.add_paragraph().paragraph_format.space_after = Pt(2); return t


d = init()
# Cover
P(d, "SYNAPEP", size=28, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
P(d, "Investor Deck — Word version", size=15, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
P(d, "AI-Personalised Aesthetic Medical Devices · Korea-anchored, export-led", size=11.5,
  italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
P(d, "Illustrative — not investment advice. Figures are v0.4 scenario estimates.", size=9,
  italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)

slide(d, 2, "The opportunity")
B(d, "Medical aesthetics ≈ $17–18.5B (2024), ~10–13% CAGR → ~$56B by 2033")
B(d, "Injectables >40% of revenue; Korea a global hub, Hong Kong a gateway")
B(d, "Commodity injectables (HA, botox) still grow but commoditise — buyers want differentiated, "
     "regenerative, personalised outcomes", sub=True)
B(d, "Opening: an AI-personalised, physician-configured regenerative category beside HA/botox")

slide(d, 3, "The solution — device + recurring kits + AI")
B(d, "Applicator device in clinics + recurring, CodeLife-configured peptide/exosome kits, "
     "registered as a medical-device family")
B(d, "Personalised by physician purpose within an approved envelope — fast 'modified-device' "
     "registration, not bespoke per-patient devices", sub=True)
B(d, "Razor-and-blades: device is the wedge and data-capture endpoint; kits are recurring revenue")
B(d, "Kit ~$30–35 ex-factory; ~64–67% gross margin; device ASP $3,000", sub=True)

slide(d, 4, "The moat — AI efficacy + data flywheel")
B(d, "Differentiator is performance, not merely personalisation: peptides engineered to outperform "
     "generalised products")
B(d, "Proprietary AI-QA/analytics validate every batch and feed continuous improvement")
B(d, "Closed loop: inputs/outcomes → better model → better peptides → deeper clinic lock-in", sub=True)
B(d, "Vision: a data-science company in aesthetic medicine — devices/kits are the data layer")
B(d, "Honest framing: day-one asset is the consented dataset + data-rights + analytics, not a "
     "foundation model", sub=True)

slide(d, 5, "Same product, more markets")
LB(d, "One product family — ", "the Korean-developed device + CodeLife kit is exported as-is")
LB(d, "No per-market R&D — ", "personalisation is by physician purpose within the same approved envelope")
LB(d, "Identical unit economics everywhere — ", "the only variables are clinic base and timing")
LB(d, "Scale + data benefits — ", "one production/QC/stability spec; one comparable global dataset")
LB(d, "Regulatory — ", "the Korean technical file is the master dossier for every reliance filing")

slide(d, 6, "Regulatory strategy (Korea → Hong Kong → export)")
TBL(d, ["Product line", "Pathway", "Speed / note"],
    [["Microneedling/topical kit + applicator", "Korea MFDS modified-device; HK MDACS listing",
      "Fast (~months) — the lead family"],
     ["IM / injectable peptide", "Drug/biologic (separate) or cosmetic", "Slower — doesn't gate launch"],
     ["CodeLife software", "Config/decision-support in cleared envelope", "Manages SaMD exposure"]],
    widths=[2.4, 2.3, 1.8])

slide(d, 7, "Export springboard — KFDA reliance")
TBL(d, ["Tier", "Markets", "KFDA leverage"],
    [["Tier 1: ASEAN + Hong Kong", "HK, Vietnam, Singapore → Thailand, Malaysia, Indonesia, Philippines",
      "MDACS reference; VN accepts MFDS; ASEAN CSDT cascade"],
     ["Tier 2: GCC + Greater China", "Saudi, UAE; Taiwan; China",
      "SFDA regional reference; TFDA uses Korean data; China via Hainan"]],
    widths=[1.7, 2.7, 2.1])

slide(d, 8, "Multi-market revenue — 3 scenarios")
TBL(d, ["$M (2030)", "Conservative", "Base", "Upside"],
    [["Revenue", "$%.1f" % (R["Conservative"][0][2030]["tot_rev"] / 1e6),
      "$%.1f" % (R["Base"][0][2030]["tot_rev"] / 1e6), "$%.1f" % (R["Upside"][0][2030]["tot_rev"] / 1e6)],
     ["EBITDA", "$%.1f" % (R["Conservative"][0][2030]["ebitda"] / 1e6),
      "$%.1f" % (R["Base"][0][2030]["ebitda"] / 1e6), "$%.1f" % (R["Upside"][0][2030]["ebitda"] / 1e6)],
     ["Clinics", "{:,}".format(R["Conservative"][0][2030]["cum"]),
      "{:,}".format(R["Base"][0][2030]["cum"]), "{:,}".format(R["Upside"][0][2030]["cum"])]],
    widths=[1.8, 1.6, 1.5, 1.5])
P(d, "Revenue trajectory ($M): " + "; ".join(
    "%s %s" % (s, "/".join("%.0f" % (R[s][0][y]["tot_rev"] / 1e6) for y in Y)) for s in mm.ORDER)
  + "  (years %s)" % "/".join(str(y) for y in Y), size=9.5, italic=True, color=GREY)

slide(d, 9, "Scenario assumptions (product unchanged)")
TBL(d, ["Lever", "Conservative", "Base", "Upside"],
    [["Markets", "11 (China excl.)", "12", "12"],
     ["Clinic ramp vs base", "0.62x", "1.00x", "1.30x"],
     ["Kit EXW", "$30.00", "$32.50", "$35.00"],
     ["Kit gross margin", "%.0f%%" % (mm.kit_gm(30) * 100), "%.0f%%" % (mm.kit_gm(32.5) * 100),
      "%.0f%%" % (mm.kit_gm(35) * 100)],
     ["Device GM", "35%", "40%", "42%"],
     ["Kits/clinic/yr", "1,200", "1,400", "1,500"]],
    widths=[1.9, 1.6, 1.5, 1.5])

slide(d, 10, "Valuation — two lenses")
LB(d, "Lens A (priced today) — ", "device/consumables floor: $5–7M pre; $6M midpoint for the $1M kickoff")
LB(d, "Lens B (upside) — ", "AI/data-science company: medtech → AI multiples; premium Series A in 2028")
TBL(d, ["Holder", "Pre", "Post-raise"],
    [["WBI", "70.0%", "60.0%"], ["Eyesel", "30.0%", "25.7%"], ["Investor (SAFE)", "—", "14.3%"]],
    widths=[2.6, 1.5, 1.5])

slide(d, 11, "The ask — milestone-gated SAFE")
B(d, "US$1.0M kickoff — post-money SAFE, $6.0M cap, 20% discount, MFN")
B(d, "Two tranches: $400k at close · $600k on milestones", sub=True)
B(d, "Conditions precedent: WBI IP into the JV; verify the $25M sample; clinic data-rights clause")
B(d, "Milestones: stability dossier + COGS BOM; 5–10 anchor clinics w/ reorders; one prospective "
     "readout; classification memo")
B(d, "Use of funds: Korea launch + Wave-1 export (HK/Vietnam/Singapore); CodeLife v1; QC/stability")

slide(d, 12, "What we will prove (diligence-ready)")
LB(d, "Cold-chain-free: ", "stabilisation removes the −80 °C dependency; document the shelf-life dossier + bottom-up COGS BOM")
LB(d, "Validation pack: ", "the $25M at entity/product/channel level + reorder cohorts")
LB(d, "Formulation envelope: ", "clean AND clinically differentiating within the approved family")
LB(d, "Exosome discipline: ", "non-human (salmon/synthetic) origin; function-based claims")
LB(d, "Data moat: ", "CodeLife v1 demo + data-rights in the clinic contract")

slide(d, 13, "Roadmap")
TBL(d, ["Phase", "Markets", "Milestones"],
    [["2027 — Anchor", "Korea", "Registration; anchor clinics; CodeLife v1"],
     ["2027–28 — Wave 1", "Hong Kong, Vietnam, Singapore", "Reliance filings; reorder evidence"],
     ["2028–29 — Wave 2", "Thailand, Malaysia, Indonesia, Philippines; Saudi, UAE",
      "ASEAN CSDT cascade; SFDA; Series A"],
     ["2029–30 — Wave 3", "Taiwan, China (Hainan→NMPA)", "Greater-China entry"]],
    widths=[1.7, 2.7, 2.1])

slide(d, 14, "Closing")
P(d, "Same product. More markets. A data moat.", size=16, bold=True, color=NAVY, after=4)
P(d, "Base case ~$%.0fM revenue / ~$%.0fM EBITDA by 2030 across 12 markets on one Korean-developed "
     "product. $1M milestone-gated kickoff." % (R["Base"][0][2030]["tot_rev"] / 1e6,
                                                R["Base"][0][2030]["ebitda"] / 1e6),
  size=12, color=GREY)

d.save(DIR + "Synapep_Investor_Deck.docx")
print("Saved Synapep_Investor_Deck.docx")
