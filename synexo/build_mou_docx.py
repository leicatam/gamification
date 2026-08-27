#!/usr/bin/env python3
"""GenApep IPO framework MOU -> mou-ipo-framework.md + GenApep_IPO_MOU.docx.

Draft non-binding MOU among FVH / WBI / Eyesel outlining the US-listing structure,
consolidated forecast, valuation-management framework and business plan.
GenApep vial-model numbers single-sourced from build_financials (P).
FVH GLP-1 Bridge program numbers from fvh-glp1-proforma.xlsx (company proforma,
Jul-2026 -> Dec-2027); constants below cite that file.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import build_financials as fin

DIR = "/home/user/gamification/synexo/"
NAVY = RGBColor(0x1B, 0x33, 0x55); GREY = RGBColor(0x55, 0x55, 0x55); ACCENT = RGBColor(0x2E, 0x86, 0xC1)

YEARS = fin.YEARS                      # 2026..2030
P = fin.P                              # GenApep vial-model P&L (dollars)

# ---------------- FVH GLP-1 Bridge program (source: fvh-glp1-proforma.xlsx) ----------------
FVH_REV = {2026: 2.224, 2027: 16.714}          # $M — proforma Annual Summary (Jul-26 start)
FVH_REV_EXTRAP = {2028: 25.0, 2029: 30.0, 2030: 34.0}   # $M [extrapolation — ~2027 enrollment pace, 10% attrition; to be confirmed by FVH]
FVH_EBITDA_MARGIN = 0.20               # [assumption] proforma is gross reimbursement only (no OpEx given)
FVH_ARPU = 168.92                      # $/patient/month recurring (99426+99454+99457)
FVH_ENROLL_2026, FVH_ENROLL_TOTAL = 2500, 10000
FVH_PEAK_ACTIVE = 9414
FVH_EXIT_MRR = 1.590                   # $M/month, Dec-2027
FVH_LTV = 3788                         # $/patient (proforma avg revenue per patient)

def fvh_rev(y):
    return FVH_REV.get(y) or FVH_REV_EXTRAP.get(y, 0.0)

# ---------------- consolidation assumptions (illustrative; to be confirmed) ----------------
EYESEL_REV_2026 = 15.0                 # $M — existing Eyesel revenue contributed to the group
EYESEL_GROWTH = 0.06                   # 6%/yr organic (assumption range 5-8%)
EYESEL_EBITDA_MARGIN = 0.12            # [assumption] typical GMP manufacturer margin
HAIR_REV = {2026: 0.0, 2027: 0.0, 2028: 1.0, 2029: 3.0, 2030: 6.0}   # $M — hair-loss + other new lines
HAIR_EBITDA_MARGIN = 0.30              # [assumption] contribution after launch costs

def eyesel_rev(y):
    return EYESEL_REV_2026 * (1 + EYESEL_GROWTH) ** (y - 2026)

ROWS = {}
for y in YEARS:
    gen_rev = P[y]["rev"] / 1e6; gen_eb = P[y]["ebitda"] / 1e6
    ey_rev = eyesel_rev(y); ey_eb = ey_rev * EYESEL_EBITDA_MARGIN
    fv_rev = fvh_rev(y); fv_eb = fv_rev * FVH_EBITDA_MARGIN
    ha_rev = HAIR_REV[y]; ha_eb = ha_rev * HAIR_EBITDA_MARGIN
    ROWS[y] = dict(gen_rev=gen_rev, gen_eb=gen_eb, ey_rev=ey_rev, ey_eb=ey_eb,
                   fv_rev=fv_rev, fv_eb=fv_eb, ha_rev=ha_rev, ha_eb=ha_eb,
                   tot_rev=gen_rev + ey_rev + fv_rev + ha_rev,
                   tot_eb=gen_eb + ey_eb + fv_eb + ha_eb)

M = lambda v: "%.1f" % v

# ---------------- shared content ----------------
TITLE = "GENAPEP"
SUBTITLE = "Memorandum of Understanding — IPO Framework (Draft for Discussion)"
DISCLAIMER = ("DRAFT — NON-BINDING. This memorandum outlines a proposed framework to kick-start "
              "discussion among the parties and to brief IPO sponsors, potential investors and the "
              "current shareholders of WBI. It is not an offer or solicitation of securities, creates "
              "no legal obligation except as expressly stated, and all figures are illustrative and "
              "subject to due diligence, audit and definitive agreements.")

PARTIES = [
    ("First Vital Health and Wellness Inc (“FVH”)",
     "U.S. digital-health and clinical-services company — remote patient monitoring (RPM) and "
     "chronic-care management — and operator of the GLP-1 Bridge patient-management program "
     "(Section 5). Made its SEC filing in December 2024 (filing history to be confirmed against "
     "FVH's EDGAR records). Proposed listed holding vehicle. Referenced existing shareholder: "
     "Mr. Ernie Lee."),
    ("WBI", "Technology and IP owner — CodeLife.AI peptide-design engine, IT-EXO/SynExo exosome and "
            "peptide science. Its current shareholders are an intended audience of this MOU."),
    ("Eyesel", "GMP manufacturer contributing certified production capacity, AI-peptide manufacturing "
               "capability and approximately US$15M of existing annual revenue. Principal: Mr. Kim."),
    ("GenApep principals", "Ms. Theresa Jang (CEO) and Mr. Kim (Eyesel principal), the proposed "
                           "controlling shareholders of FVH after the share issuance."),
]

STEPS = [
    ("Step 1 — SEC status (condition precedent)", "Now → 2027",
     "FVH maintains its SEC filing/reporting status and extends it into 2027. The framework in this "
     "MOU carries forward ONLY if this status is extended; otherwise the MOU lapses with no "
     "obligation on any party."),
    ("Step 2 — Share issuance to GenApep principals", "Before Dec 2026",
     "From/through FVH shareholder Mr. Ernie Lee, new FVH shares are issued such that (indicatively, "
     "assuming 75% of FVH) Ms. Theresa Jang holds 40% and Mr. Kim holds 35%. Existing FVH "
     "shareholders (including Mr. Lee) retain approximately 25% (working assumption). Final "
     "percentages, instruments and pricing to be defined in definitive agreements."),
    ("Step 3 — Korean subsidiary & acquisitions", "H2 2026 → Jan 2027",
     "FVH establishes a wholly-owned Korean subsidiary (“FVH Korea”). FVH Korea acquires "
     "100% of WBI and 100% of Eyesel — together the GenApep business. Consideration structure "
     "(share swap / cash mix) and valuations [TBD], supported by independent valuations."),
    ("Step 4 — Restructure complete", "January 2027",
     "The consolidated group operates as GenApep under listed FVH: FVH (US listed) → FVH Korea "
     "(100%) → WBI + Eyesel. Consolidated PCAOB-auditable financials from FY2027."),
]

SYNERGY = [
    ("FVH contribution — GLP-1 Bridge in GI (US)", "The operating GLP-1 patient-management program "
     "detailed in Section 5 — a CPT-reimbursed, recurring-revenue care platform (assessment → care "
     "plan → monthly RPM/PCM monitoring) targeting 10,000 enrolled patients by end-2027 — plus the "
     "listed vehicle, SEC reporting history and a U.S. clinical channel for the group's products."),
    ("GLP-1 × GenApep synergy", "FVH's GLP-1 patient base is the natural U.S. channel for GenApep "
     "peptide programs: GI and metabolic support on- and off-therapy, muscle-preservation "
     "formulations, and GLP-1-associated hair-loss management — designed by CodeLife.AI, "
     "manufactured by Eyesel under GMP."),
    ("Hair-loss management", "AI-designed peptide/exosome programs for hair-loss management — an "
     "adjacent high-recurrence consumer-medical category, including GLP-1-associated hair loss."),
    ("Eyesel contribution", "GMP-certified manufacturing, AI-peptide production capability and "
     "~US$15M existing revenue — the consolidated group's revenue and cash-flow base from day one."),
    ("WBI contribution", "CodeLife.AI design engine, IT-EXO/SynExo science and the consented "
     "outcomes dataset — the group's differentiation and pipeline engine."),
    ("GenApep operating engine", "The vial-based razor-and-blades consumable model (US$7/vial cost, "
     "~78.5% gross margin, EBITDA-positive from first production year 2027) already documented in "
     "the GenApep operating financial model."),
]

# FVH care pathway (CPT code, when, rate, note) — rates implied by the proforma (LA County NP rates)
FVH_PATHWAY = [
    ("99204", "Day 1", "$150.76", "Comprehensive GI & metabolic assessment"),
    ("99213", "Week 2", "$89.16", "Labs review, assessment, treatment initiation"),
    ("G0506", "Month 1", "$62.44", "Care-plan establishment"),
    ("99453", "Month 1", "$22.00", "RPM device (scale) set-up"),
    ("99426 + 99454 + 99457", "Monthly (recurring)", "$168.92 / month",
     "PCM $67.80 + RPM data $52.11 + RPM weight management $49.01 — the MRR engine"),
    ("99213", "Quarterly", "$89.16", "Recurring quarterly E/M follow-up"),
]

FVH_METRICS = [
    ("Model period", "July 1, 2026 → December 31, 2027 (18 months); accelerating enrollment ramp"),
    ("Enrollment", "2,500 patients by Dec-2026; 10,000 cumulative by Dec-2027; ~10% annual attrition"),
    ("Active patients", "2,463 avg (2026) → 9,414 peak (Dec-2027)"),
    ("Revenue", "US$2.2M (H2-2026) → US$16.7M (2027) — gross reimbursement, per proforma"),
    ("Recurring revenue", "MRR US$1.59M/month exiting Dec-2027 (≈ US$19M annualised); ARPU $168.92/patient/month"),
    ("Unit value", "≈ US$3,788 average revenue per enrolled patient (proforma LTV basis)"),
]

VAL_POINTS = [
    ("Entry reference points", "GenApep JV previously framed at US$6M pre-money for a US$1M kickoff "
     "(device/consumables lens); Eyesel contributes ~US$15M revenue; FVH contributes an operating "
     "GLP-1 program at a US$19M annualised MRR run-rate exiting 2027 (proforma). Acquisition "
     "valuations for WBI and Eyesel [TBD] with independent support."),
    ("Method for the listed company", "Sponsor-led comparables: specialty peptide / medical-aesthetics "
     "/ CDMO peers for the Korean business, and recurring-revenue digital-health / care-management "
     "peers for the FVH GLP-1 platform (indicative placeholder ranges: 3–6× revenue or 10–15× EBITDA "
     "— to be validated by the IPO sponsor), cross-checked with DCF on the consolidated plan."),
    ("Value-creation levers", "Consolidated growth (US$15M Eyesel base + FVH GLP-1 ramp + GenApep "
     "ramp + hair-loss optionality), gross-margin expansion toward the ~78.5% consumable model, "
     "recurring-revenue mix (care-management MRR + consumable reorders), and liquidity/uplisting "
     "premium as reporting history builds."),
    ("Governance & fairness", "Ms. Jang and Mr. Kim would stand on both sides of the acquisitions "
     "(as FVH shareholders and as WBI/Eyesel principals). The parties will use independent "
     "valuations / fairness opinions, disinterested-director or shareholder approval where "
     "applicable, and full related-party disclosure in SEC filings."),
]

CONDITIONS = [
    "FVH successfully maintaining and extending its SEC listing/reporting status into 2027 — the governing condition of this MOU.",
    "Satisfactory mutual due diligence (corporate, financial, IP, regulatory, tax) on FVH, WBI and Eyesel — including confirmation of FVH's EDGAR filing history and the assumptions of the GLP-1 proforma (reimbursement rates, enrollment, attrition).",
    "Definitive agreements: share-issuance agreement, share-purchase/exchange agreements, shareholder agreements.",
    "PCAOB-standard audits of Eyesel and WBI sufficient for SEC consolidated reporting.",
    "Korean regulatory, foreign-direct-investment and tax clearances for FVH Korea and the acquisitions.",
    "U.S. healthcare-regulatory compliance of the GLP-1 program (billing/coding, telehealth and RPM/PCM rules, state scope-of-practice).",
    "Board and shareholder approvals of each party, including the current shareholders of WBI.",
    "Independent valuations / fairness opinions for the related-party acquisitions.",
    "No assurance of financing, listing venue, uplisting or valuation is given or implied.",
]

TIMELINE = [
    ("Q3–Q4 2026", "Three-party discussion on this MOU; due diligence; definitive agreements; PCAOB audit preparation; FVH SEC status maintenance. FVH GLP-1 program enrolling from Jul-2026 (2,500 patients by Dec-2026)."),
    ("Before Dec 2026", "FVH share issuance completed — Theresa Jang 40% / Mr. Kim 35% (indicative)."),
    ("H2 2026", "FVH Korea incorporated; acquisition agreements for WBI and Eyesel signed."),
    ("Jan 2027", "Restructure complete — WBI and Eyesel consolidated under FVH Korea; group operates as GenApep."),
    ("FY2027", "First consolidated year: ~US$35M revenue (illustrative); GenApep production ramp from Feb-2027; FVH GLP-1 program scaling to 10,000 patients; GLP-1-bridge and hair-loss synergy programs formalised."),
    ("2028+", "Compounded growth with synergy lines; reporting history and scale toward an uplisting path [venue TBD]."),
]

# ------------------------------------------------------------------ Markdown
def wmd():
    L = []
    a = L.append
    a("# GenApep — Memorandum of Understanding: IPO Framework (Draft)\n")
    a("> %s\n" % DISCLAIMER)
    a("\n## 1. Parties\n")
    for n, d in PARTIES:
        a("- **%s** — %s\n" % (n, d))
    a("\n## 2. Purpose\n")
    a("Outline the proposed structure and roadmap for taking the GenApep business to the US public "
      "market via FVH; kick-start discussion among the three parties; and provide a framework "
      "document for IPO sponsors, potential investors and the current shareholders of WBI.\n")
    a("\n## 3. Proposed structure & steps\n")
    for t, w, d in STEPS:
        a("- **%s** (*%s*) — %s\n" % (t, w, d))
    a("\n**Resulting structure:** FVH (US listed) → FVH Korea (100%) → WBI + Eyesel (= GenApep).\n")
    a("\n**Indicative FVH ownership after Step 2:** Theresa Jang 40% · Mr. Kim 35% · existing FVH "
      "shareholders (incl. Ernie Lee) ~25% — final percentages [to be defined].\n")
    a("\n## 4. Post-restructure business plan (consolidation & synergy)\n")
    for t, d in SYNERGY:
        a("- **%s** — %s\n" % (t, d))
    a("\n## 5. FVH — the GLP-1 Bridge program (US)\n")
    a("\nFVH's operating content: a GLP-1 patient-management program (gastro-intestinal & metabolic "
      "care around GLP-1 therapy) billed under U.S. CPT/HCPCS codes at Nurse-Practitioner rates "
      "(proforma basis: LA County). Source: FVH GLP-1 proforma, Jul-2026 → Dec-2027 "
      "(`fvh-glp1-proforma.xlsx`).\n")
    a("\n### 5.1 Care pathway & reimbursement\n\n")
    a("| Code | When | Rate | Service |\n|---|---|---|---|\n")
    for c, w, r, n in FVH_PATHWAY:
        a("| %s | %s | %s | %s |\n" % (c, w, r, n))
    a("\n### 5.2 Program metrics (per proforma)\n\n")
    a("| Metric | Value |\n|---|---|\n")
    for k, v in FVH_METRICS:
        a("| %s | %s |\n" % (k, v))
    a("\n## 6. Consolidated revenue forecast (illustrative, US$M)\n\n")
    hdr = "| Line | " + " | ".join(str(y) for y in YEARS) + " |\n"
    a(hdr); a("|---|" + "---|" * len(YEARS) + "\n")
    a("| Eyesel (existing, ~6%/yr) | " + " | ".join(M(ROWS[y]["ey_rev"]) for y in YEARS) + " |\n")
    a("| FVH GLP-1 Bridge (proforma → extrapolated) | " + " | ".join(M(ROWS[y]["fv_rev"]) for y in YEARS) + " |\n")
    a("| GenApep vial model | " + " | ".join(M(ROWS[y]["gen_rev"]) for y in YEARS) + " |\n")
    a("| Hair-loss + other new lines | " + " | ".join(M(ROWS[y]["ha_rev"]) for y in YEARS) + " |\n")
    a("| **Consolidated revenue** | " + " | ".join("**%s**" % M(ROWS[y]["tot_rev"]) for y in YEARS) + " |\n")
    a("| **Consolidated EBITDA** | " + " | ".join("**%s**" % M(ROWS[y]["tot_eb"]) for y in YEARS) + " |\n")
    a("\nAssumptions: Eyesel base US$15M growing ~6%/yr at ~12% EBITDA margin [assumption]; FVH "
      "GLP-1 per company proforma for 2026–2027, extrapolated for 2028–2030 at roughly the 2027 "
      "enrollment pace with 10% attrition [to be confirmed by FVH], at ~20% EBITDA margin "
      "[assumption — proforma is gross reimbursement only]; GenApep vial model per the GenApep "
      "operating financial model (78.5% GM, EBITDA-positive 2027); hair-loss and other new lines "
      "ramp from 2028 at ~30% contribution [illustrative placeholders].\n")
    a("\n## 7. Valuation management\n")
    for t, d in VAL_POINTS:
        a("- **%s** — %s\n" % (t, d))
    a("\n## 8. Conditions precedent & key risks\n")
    for i, c in enumerate(CONDITIONS, 1):
        a("%d. %s\n" % (i, c))
    a("\n## 9. Indicative timeline\n\n| Window | Milestones |\n|---|---|\n")
    for w, d in TIMELINE:
        a("| %s | %s |\n" % (w, d))
    a("\n## 10. Non-binding nature & confidentiality\n")
    a("This MOU is non-binding except for confidentiality and governing-law clauses in the executed "
      "version. Each party bears its own costs. The parties intend to negotiate definitive "
      "agreements in good faith. Contents are confidential to the parties, their advisers, IPO "
      "sponsors, potential investors and WBI shareholders under duty of confidence.\n")
    a("\n## 11. Signatures (draft)\n")
    a("- For FVH: ______________________  Name/Title: [Ernie Lee, TBD]  Date: ______\n")
    a("- For WBI: ______________________  Name/Title: [TBD]  Date: ______\n")
    a("- For Eyesel: ______________________  Name/Title: [Mr. Kim, TBD]  Date: ______\n")
    open(DIR + "mou-ipo-framework.md", "w").write("".join(L))

wmd()

# ------------------------------------------------------------------ Word
def nd():
    doc = Document()
    n = doc.styles["Normal"]; n.font.name = "Calibri"; n.font.size = Pt(10.5)
    for lvl, sz in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)]:
        st = doc.styles[lvl]; st.font.name = "Calibri"; st.font.size = Pt(sz)
        st.font.color.rgb = NAVY; st.font.bold = True
    return doc


def para(doc, t="", size=10.5, bold=False, italic=False, color=None, align=None, after=6):
    p = doc.add_paragraph()
    if align: p.alignment = align
    r = p.add_run(t); r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after); return p


def bullet(doc, t, lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if lead:
        r = p.add_run(lead); r.bold = True
    p.add_run(t); p.paragraph_format.space_after = Pt(3); return p


def numbered(doc, t):
    p = doc.add_paragraph(style="List Number")
    p.add_run(t); p.paragraph_format.space_after = Pt(3); return p


def tbl(doc, headers, rows, widths=None, fs=9.5, hl_rows=()):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        r = t.rows[0].cells[i].paragraphs[0].add_run(str(h)); r.bold = True; r.font.size = Pt(fs)
    for ri, row in enumerate(rows):
        c = t.add_row().cells
        for i, v in enumerate(row):
            rr = c[i].paragraphs[0].add_run(str(v)); rr.font.size = Pt(fs)
            if ri in hl_rows or i == 0: rr.bold = True
    if widths:
        for i, w in enumerate(widths):
            for r in t.rows: r.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2); return t


doc = nd()
para(doc, TITLE, size=26, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
para(doc, SUBTITLE, size=14, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
para(doc, DISCLAIMER, size=8.5, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)

doc.add_heading("1. Parties", level=2)
tbl(doc, ["Party", "Description"], [[n, d] for n, d in PARTIES], widths=[2.1, 4.4])

doc.add_heading("2. Purpose", level=2)
para(doc, "Outline the proposed structure and roadmap for taking the GenApep business to the US "
          "public market via FVH; kick-start discussion among the three parties; and provide a "
          "framework document for IPO sponsors, potential investors and the current shareholders "
          "of WBI.")

doc.add_heading("3. Proposed structure & steps", level=2)
tbl(doc, ["Step", "Window", "Description"], [[t, w, d] for t, w, d in STEPS],
    widths=[1.7, 1.0, 3.8])
para(doc, "Resulting structure:  FVH (US listed)  →  FVH Korea (100%)  →  WBI + Eyesel "
          "(together, GenApep).", bold=True)
tbl(doc, ["Indicative FVH ownership after Step 2", "%"],
    [["Ms. Theresa Jang", "40%"], ["Mr. Kim", "35%"],
     ["Existing FVH shareholders (incl. Mr. Ernie Lee)", "~25%"],
     ["Total", "100%"]], widths=[4.2, 1.2], hl_rows=(3,))
para(doc, "Percentages are indicative (“assuming 75% of FVH”) and will be defined in "
          "definitive agreements.", size=9, italic=True, color=GREY)

doc.add_heading("4. Post-restructure business plan — consolidation & synergy", level=2)
for t, d in SYNERGY:
    bullet(doc, d, t + " — ")

doc.add_heading("5. FVH — the GLP-1 Bridge program (US)", level=2)
para(doc, "FVH's operating content: a GLP-1 patient-management program (gastro-intestinal & "
          "metabolic care around GLP-1 therapy) billed under U.S. CPT/HCPCS codes at "
          "Nurse-Practitioner rates (proforma basis: LA County). Source: FVH GLP-1 proforma, "
          "July 2026 → December 2027 (fvh-glp1-proforma.xlsx).")
doc.add_heading("5.1 Care pathway & reimbursement", level=3)
tbl(doc, ["Code", "When", "Rate", "Service"], [list(r) for r in FVH_PATHWAY],
    widths=[1.5, 1.2, 1.1, 2.7])
doc.add_heading("5.2 Program metrics (per proforma)", level=3)
tbl(doc, ["Metric", "Value"], [[k, v] for k, v in FVH_METRICS], widths=[1.6, 4.9])

doc.add_heading("6. Consolidated revenue forecast (illustrative, US$M)", level=2)
tbl(doc, ["Line"] + [str(y) for y in YEARS],
    [["Eyesel (existing, ~6%/yr)"] + [M(ROWS[y]["ey_rev"]) for y in YEARS],
     ["FVH GLP-1 Bridge (proforma → extrap.)"] + [M(ROWS[y]["fv_rev"]) for y in YEARS],
     ["GenApep vial model"] + [M(ROWS[y]["gen_rev"]) for y in YEARS],
     ["Hair-loss + other new lines"] + [M(ROWS[y]["ha_rev"]) for y in YEARS],
     ["Consolidated revenue"] + [M(ROWS[y]["tot_rev"]) for y in YEARS],
     ["Consolidated EBITDA"] + [M(ROWS[y]["tot_eb"]) for y in YEARS]],
    widths=[2.5, 0.8, 0.8, 0.8, 0.8, 0.8], hl_rows=(4, 5))
para(doc, "Assumptions: Eyesel base US$15M growing ~6%/yr at ~12% EBITDA margin [assumption]; FVH "
          "GLP-1 per company proforma for 2026–2027, extrapolated 2028–2030 at roughly the 2027 "
          "enrollment pace with 10% attrition [to be confirmed by FVH], at ~20% EBITDA margin "
          "[assumption — proforma is gross reimbursement only]; GenApep vial model per the GenApep "
          "operating financial model (78.5% GM, EBITDA-positive 2027); hair-loss and other new "
          "lines ramp from 2028 at ~30% contribution [illustrative placeholders].",
     size=9, italic=True, color=GREY)

doc.add_heading("7. Valuation management", level=2)
for t, d in VAL_POINTS:
    bullet(doc, d, t + " — ")

doc.add_heading("8. Conditions precedent & key risks", level=2)
for c in CONDITIONS:
    numbered(doc, c)

doc.add_heading("9. Indicative timeline", level=2)
tbl(doc, ["Window", "Milestones"], [[w, d] for w, d in TIMELINE], widths=[1.4, 5.1])

doc.add_heading("10. Non-binding nature & confidentiality", level=2)
para(doc, "This MOU is non-binding except for confidentiality and governing-law clauses in the "
          "executed version. Each party bears its own costs. The parties intend to negotiate "
          "definitive agreements in good faith. Contents are confidential to the parties, their "
          "advisers, IPO sponsors, potential investors and the current shareholders of WBI under "
          "duty of confidence.")

doc.add_heading("11. Signatures (draft)", level=2)
tbl(doc, ["Party", "Signature", "Name / Title", "Date"],
    [["FVH", "", "[Mr. Ernie Lee — TBD]", ""],
     ["WBI", "", "[TBD]", ""],
     ["Eyesel", "", "[Mr. Kim — TBD]", ""]], widths=[1.0, 1.8, 2.5, 1.0])
para(doc, "Prepared as a draft framework to open discussion. All terms subject to change.",
     size=8.5, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.save(DIR + "GenApep_IPO_MOU.docx")
print("Saved mou-ipo-framework.md, GenApep_IPO_MOU.docx")
for y in YEARS:
    print("%d: consolidated rev $%sM (FVH $%sM), EBITDA $%sM"
          % (y, M(ROWS[y]["tot_rev"]), M(ROWS[y]["fv_rev"]), M(ROWS[y]["tot_eb"])))
