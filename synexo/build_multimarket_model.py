#!/usr/bin/env python3
"""Multi-market revenue model — SAME Korean-developed product family across all markets.

One product (applicator device + CodeLife-configured peptide/exosome kit) developed for Korean
physicians is exported as-is. Personalisation = physician purpose within the same approved
envelope; NO per-market product development. Unit economics are identical in every market; only
the clinic base and timing differ.

Three scenarios via a single switch:
  - Conservative : slower ASEAN/GCC ramp, China EXCLUDED, lower price/utilisation
  - Base         : full 12-market plan
  - Upside       : faster ramp, higher price/utilisation

Importable: SCENARIOS, ORDER, YEARS, BASE_CLINICS, RESULTS, kit_gm(), k(), fmt().
Run as a script to (re)generate multimarket-model.csv/.md and GenApep_MultiMarket_Model.docx.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

DIR = "/home/user/gamification/synexo/"
NAVY = RGBColor(0x1F, 0x3A, 0x5F); GREY = RGBColor(0x55, 0x55, 0x55)

DEVICE_ASP = 3000.0
YEARS = [2027, 2028, 2029, 2030]

# Base-case cumulative clinics per market per year (same product sold into each).
# wave: Korea anchor 2027; Wave1 HK/VN/SG 2027-28; Wave2 ASEAN+GCC 2028-29; Wave3 TW/CN 2029-30
BASE_CLINICS = {
    "Korea (anchor)":        {2027: 120, 2028: 250, 2029: 400, 2030: 550},
    "Hong Kong":             {2027: 20,  2028: 60,  2029: 100, 2030: 140},
    "Vietnam":               {2027: 10,  2028: 50,  2029: 110, 2030: 180},
    "Singapore":             {2027: 10,  2028: 35,  2029: 60,  2030: 85},
    "Thailand":              {2027: 0,   2028: 40,  2029: 120, 2030: 220},
    "Malaysia":              {2027: 0,   2028: 20,  2029: 70,  2030: 130},
    "Indonesia":             {2027: 0,   2028: 15,  2029: 70,  2030: 150},
    "Philippines":           {2027: 0,   2028: 10,  2029: 45,  2030: 100},
    "Saudi Arabia (SFDA)":   {2027: 0,   2028: 20,  2029: 70,  2030: 140},
    "UAE (MOHAP)":           {2027: 0,   2028: 15,  2029: 50,  2030: 95},
    "Taiwan":                {2027: 0,   2028: 0,   2029: 35,  2030: 90},
    "China (Hainan->NMPA)":  {2027: 0,   2028: 0,   2029: 40,  2030: 140},
}

SCENARIOS = {
    "Conservative": dict(clinic_mult=0.62, exclude=["China (Hainan->NMPA)"], exw=30.0,
                         device_gm=0.35, kits_yr=1200,
                         opex={2027: 2800, 2028: 6000, 2029: 9000, 2030: 12000}),
    "Base":         dict(clinic_mult=1.00, exclude=[], exw=32.5,
                         device_gm=0.40, kits_yr=1400,
                         opex={2027: 3000, 2028: 7000, 2029: 11000, 2030: 16000}),
    "Upside":       dict(clinic_mult=1.30, exclude=[], exw=35.0,
                         device_gm=0.42, kits_yr=1500,
                         opex={2027: 3500, 2028: 8500, 2029: 13500, 2030: 19000}),
}
ORDER = ["Conservative", "Base", "Upside"]


def kit_gm(exw):
    """COGS = 20% x EXW + $4.50 material; GP margin derived (same product, same recipe)."""
    return (exw - (0.20 * exw + 4.50)) / exw


def scaled_clinics(scn):
    s = SCENARIOS[scn]; out = {}
    for m, yrs in BASE_CLINICS.items():
        if m in s["exclude"]:
            out[m] = {y: 0 for y in YEARS}
        else:
            out[m] = {y: round(yrs[y] * s["clinic_mult"]) for y in YEARS}
    return out


def compute(scn):
    s = SCENARIOS[scn]; clinics = scaled_clinics(scn); gm_kit = kit_gm(s["exw"])
    res = {}
    for i, y in enumerate(YEARS):
        cum_total = sum(clinics[m][y] for m in clinics)
        prev_total = sum(clinics[m][YEARS[i - 1]] for m in clinics) if i > 0 else 0
        new_clinics = cum_total - prev_total
        avg_clinics = (prev_total + cum_total) / 2.0
        kits = avg_clinics * s["kits_yr"]
        kit_rev = kits * s["exw"]; dev_rev = new_clinics * DEVICE_ASP
        tot_rev = kit_rev + dev_rev
        tot_gp = kit_rev * gm_kit + dev_rev * s["device_gm"]
        ebitda = tot_gp - s["opex"][y] * 1000
        res[y] = dict(cum=cum_total, new=new_clinics, kits=kits, kit_rev=kit_rev, dev_rev=dev_rev,
                      tot_rev=tot_rev, tot_gp=tot_gp, gm=tot_gp / tot_rev, ebitda=ebitda)
    return res, clinics, gm_kit


RESULTS = {scn: compute(scn) for scn in ORDER}  # {scn: (res, clinics, gm_kit)}


def k(x):
    return round(x / 1000.0)


def fmt(x):
    return "{:,}".format(k(x))


# ------------------------------------------------------------------ docx helpers
def _new_doc():
    doc = Document()
    n = doc.styles["Normal"]; n.font.name = "Calibri"; n.font.size = Pt(10.5)
    for lvl, sz in [("Heading 1", 16), ("Heading 2", 13)]:
        st = doc.styles[lvl]; st.font.name = "Calibri"; st.font.size = Pt(sz)
        st.font.color.rgb = NAVY; st.font.bold = True
    return doc


def _para(doc, text="", size=10.5, bold=False, italic=False, color=None, align=None, after=6):
    p = doc.add_paragraph()
    if align: p.alignment = align
    r = p.add_run(text); r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after); return p


def _table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        r = t.rows[0].cells[i].paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(9.5)
    for rowv in rows:
        c = t.add_row().cells
        for i, v in enumerate(rowv):
            rr = c[i].paragraphs[0].add_run(str(v)); rr.font.size = Pt(9.5)
    if widths:
        for i, w in enumerate(widths):
            for rr in t.rows: rr.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2); return t


# ------------------------------------------------------------------ outputs
def write_csv():
    with open(DIR + "multimarket-model.csv", "w") as f:
        f.write("GenApep Multi-Market Model — SAME Korean-developed product across all markets (USD)\n")
        f.write("Three scenarios. Distributors arm's-length / non-consolidated; revenue at ex-factory.\n\n")
        for scn in ORDER:
            s = SCENARIOS[scn]; res = RESULTS[scn][0]; gmk = RESULTS[scn][2]
            f.write("=== %s (EXW $%.2f, kitGM %.1f%%, deviceGM %.0f%%, kits/clinic/yr %d, China %s) ===\n"
                    % (scn, s["exw"], gmk * 100, s["device_gm"] * 100, s["kits_yr"],
                       "excluded" if s["exclude"] else "included"))
            f.write("Line," + ",".join(str(y) for y in YEARS) + "\n")
            f.write("Total clinics," + ",".join(str(res[y]["cum"]) for y in YEARS) + "\n")
            f.write("Kits (000s)," + ",".join(str(k(res[y]["kits"])) for y in YEARS) + "\n")
            f.write("Total revenue ($000s)," + ",".join(str(k(res[y]["tot_rev"])) for y in YEARS) + "\n")
            f.write("Gross profit ($000s)," + ",".join(str(k(res[y]["tot_gp"])) for y in YEARS) + "\n")
            f.write("EBITDA ($000s)," + ",".join(str(k(res[y]["ebitda"])) for y in YEARS) + "\n\n")
        f.write("BASE per-market cumulative clinics," + ",".join(str(y) for y in YEARS) + "\n")
        base_clinics = RESULTS["Base"][1]
        for m in BASE_CLINICS:
            f.write(m + "," + ",".join(str(base_clinics[m][y]) for y in YEARS) + "\n")


def _row(label, vals):
    return "| " + label + " | " + " | ".join(vals) + " |\n"


def write_md():
    with open(DIR + "multimarket-model.md", "w") as f:
        f.write("# GenApep — Multi-Market Revenue Model (3 scenarios)\n\n")
        f.write("> **Same product, more markets.** One product family — the applicator **device** + "
                "**CodeLife-configured peptide/exosome kit** developed for **Korean physicians** — is "
                "exported **as-is**. Personalisation is by *physician purpose within the same approved "
                "envelope*; there is **no per-market product development**, so unit economics are "
                "**identical in every market**. Scenarios flex only the **clinic ramp, price and "
                "utilisation** — not the product. Illustrative; USD.\n\n")
        f.write("## Scenario assumptions\n\n")
        f.write(_row("Lever", ["Conservative", "Base", "Upside"]))
        f.write("|---|---|---|---|\n")
        f.write(_row("Markets", ["11 (China excluded)", "12", "12"]))
        f.write(_row("Clinic ramp vs base", ["0.62×", "1.00×", "1.30×"]))
        f.write(_row("Kit EXW", ["$30.00", "$32.50", "$35.00"]))
        f.write(_row("Kit gross margin", ["%.1f%%" % (kit_gm(30) * 100),
                                          "%.1f%%" % (kit_gm(32.5) * 100), "%.1f%%" % (kit_gm(35) * 100)]))
        f.write(_row("Device GM", ["35%", "40%", "42%"]))
        f.write(_row("Kits/clinic/yr", ["1,200", "1,400", "1,500"]))
        f.write("\n## Scenario comparison — Total revenue ($000s)\n\n")
        f.write(_row("Scenario", [str(y) for y in YEARS])); f.write("|---|---|---|---|---|\n")
        for scn in ORDER:
            f.write(_row(scn, [fmt(RESULTS[scn][0][y]["tot_rev"]) for y in YEARS]))
        f.write("\n## Scenario comparison — EBITDA ($000s)\n\n")
        f.write(_row("Scenario", [str(y) for y in YEARS])); f.write("|---|---|---|---|---|\n")
        for scn in ORDER:
            f.write(_row(scn, [fmt(RESULTS[scn][0][y]["ebitda"]) for y in YEARS]))
        f.write("\n## 2030 snapshot\n\n")
        f.write(_row("Scenario", ["Clinics", "Revenue", "Blended GM", "Gross profit", "EBITDA"]))
        f.write("|---|---|---|---|---|---|\n")
        for scn in ORDER:
            r = RESULTS[scn][0][2030]
            f.write(_row(scn, ["{:,}".format(r["cum"]), "$%.1fM" % (r["tot_rev"] / 1e6),
                              "%.1f%%" % (r["gm"] * 100), "$%.1fM" % (r["tot_gp"] / 1e6),
                              "$%.1fM" % (r["ebitda"] / 1e6)]))
        f.write("\n## Base case — cumulative clinics by market\n\n")
        base_clinics = RESULTS["Base"][1]
        f.write(_row("Market", [str(y) for y in YEARS])); f.write("|---|---|---|---|---|\n")
        for m in BASE_CLINICS:
            f.write(_row(m, [str(base_clinics[m][y]) for y in YEARS]))
        f.write(_row("**Total**", ["**%d**" % RESULTS["Base"][0][y]["cum"] for y in YEARS]))
        f.write("\n*Kits = average active clinics in year × kits/clinic/yr (first-year clinics ramp, "
                "approximated by averaging opening/closing counts). Devices = new clinics × $3,000. "
                "Distributors are arm's-length and non-consolidated; revenue is booked at ex-factory.*\n\n")
        f.write("**Read-through:** the **Conservative** case (China excluded, slower ramp, lower price) "
                "lands near the Korea-only base (~$35M vs $39.9M/2030) but spread across 11 markets "
                "(less single-country risk); **Base** ~doubles the Korea-only base; **Upside** roughly "
                "triples it — all on the **same product**, so the swing is distribution reach, not R&D.\n\n")
        f.write("**Caveats:** per-market ramps are assumptions pending anchor-clinic evidence; the "
                "kit's regulatory class varies by market and exosome claims are restricted (see "
                "`export-strategy.md`, business-plan §9). Distribution is **cold-chain-free** "
                "(ambient/2–8 °C stabilised format) — substantiate with the shelf-life dossier.\n")


def write_docx():
    doc = _new_doc()
    _para(doc, "GENAPEP", size=26, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    _para(doc, "Multi-Market Revenue Model — Conservative / Base / Upside", size=14, color=GREY,
          align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    _para(doc, "Same Korean-developed personalised product family, exported across markets. Illustrative.",
          size=9, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    _para(doc, "Same product, more markets.", bold=True, after=2)
    _para(doc, "One product family — the applicator device + CodeLife-configured peptide/exosome kit "
               "developed for Korean physicians — is exported as-is (personalisation by physician "
               "purpose within the same approved envelope; no per-market product development). Unit "
               "economics are identical everywhere; scenarios flex only the clinic ramp, price and "
               "utilisation.")
    doc.add_heading("Scenario assumptions", level=2)
    _table(doc, ["Lever", "Conservative", "Base", "Upside"],
           [["Markets", "11 (China excluded)", "12", "12"],
            ["Clinic ramp vs base", "0.62x", "1.00x", "1.30x"],
            ["Kit EXW", "$30.00", "$32.50", "$35.00"],
            ["Kit gross margin", "%.1f%%" % (kit_gm(30) * 100), "%.1f%%" % (kit_gm(32.5) * 100),
             "%.1f%%" % (kit_gm(35) * 100)],
            ["Device GM", "35%", "40%", "42%"],
            ["Kits/clinic/yr", "1,200", "1,400", "1,500"]],
           widths=[1.8, 1.6, 1.6, 1.5])
    doc.add_heading("Scenario comparison — Total revenue ($000s)", level=2)
    _table(doc, ["Scenario"] + [str(y) for y in YEARS],
           [[scn] + [fmt(RESULTS[scn][0][y]["tot_rev"]) for y in YEARS] for scn in ORDER],
           widths=[1.8, 1.1, 1.1, 1.1, 1.1])
    doc.add_heading("Scenario comparison — EBITDA ($000s)", level=2)
    _table(doc, ["Scenario"] + [str(y) for y in YEARS],
           [[scn] + [fmt(RESULTS[scn][0][y]["ebitda"]) for y in YEARS] for scn in ORDER],
           widths=[1.8, 1.1, 1.1, 1.1, 1.1])
    doc.add_heading("2030 snapshot", level=2)
    _table(doc, ["Scenario", "Clinics", "Revenue", "Blended GM", "Gross profit", "EBITDA"],
           [[scn, "{:,}".format(RESULTS[scn][0][2030]["cum"]),
             "$%.1fM" % (RESULTS[scn][0][2030]["tot_rev"] / 1e6),
             "%.1f%%" % (RESULTS[scn][0][2030]["gm"] * 100),
             "$%.1fM" % (RESULTS[scn][0][2030]["tot_gp"] / 1e6),
             "$%.1fM" % (RESULTS[scn][0][2030]["ebitda"] / 1e6)] for scn in ORDER],
           widths=[1.5, 1.0, 1.0, 1.1, 1.1, 1.0])
    doc.add_heading("Base case — cumulative clinics by market", level=2)
    base_clinics = RESULTS["Base"][1]
    _table(doc, ["Market"] + [str(y) for y in YEARS],
           [[m] + [str(base_clinics[m][y]) for y in YEARS] for m in BASE_CLINICS]
           + [["TOTAL"] + [str(RESULTS["Base"][0][y]["cum"]) for y in YEARS]],
           widths=[2.2, 0.9, 0.9, 0.9, 0.9])
    _para(doc, "Read-through: the Conservative case (China excluded, slower ramp, lower price) lands "
               "near the Korea-only base (~$35M vs $39.9M/2030) but across 11 markets (less "
               "single-country risk); Base ~doubles it; Upside roughly triples it — all on the same "
               "product, so the swing is distribution reach, not R&D. Per-market ramps are assumptions "
               "pending anchor-clinic evidence; kit regulatory class varies by market; distribution is "
               "cold-chain-free (ambient/2–8 °C stabilised format).", italic=True, size=9, color=GREY)
    doc.save(DIR + "GenApep_MultiMarket_Model.docx")


def main():
    write_csv(); write_md(); write_docx()
    for scn in ORDER:
        r = RESULTS[scn][0][2030]
        print("%-12s 2030: clinics %d, rev $%.1fM, GP $%.1fM, EBITDA $%.1fM, GM %.1f%%"
              % (scn, r["cum"], r["tot_rev"] / 1e6, r["tot_gp"] / 1e6, r["ebitda"] / 1e6, r["gm"] * 100))
    print("Saved multimarket-model.csv/.md, GenApep_MultiMarket_Model.docx")


if __name__ == "__main__":
    main()
